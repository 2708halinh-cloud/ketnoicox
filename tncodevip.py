import os
import logging
import asyncio
import time
import threading
import sqlite3
import hashlib
import secrets
import json
import re
import shutil
import atexit
import signal
import zipfile
import io
import csv
import html
import tempfile
import random
import unicodedata
import subprocess
from datetime import datetime
import urllib.parse
import urllib.request
import urllib.error
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, BotCommand
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes

# Cấu hình theo dõi lỗi
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("httpcore").setLevel(logging.WARNING)


class RedactTelegramTokenFilter(logging.Filter):
    def filter(self, record):
        message = record.getMessage()
        redacted = re.sub(r"bot\d+:[A-Za-z0-9_-]+", "bot<redacted>", message)
        if redacted != message:
            record.msg = redacted
            record.args = ()
        return True


for handler in logging.getLogger().handlers:
    handler.addFilter(RedactTelegramTokenFilter())


def env_int(name: str, default: int = 0) -> int:
    raw = os.getenv(name, str(default)).strip()
    try:
        return int(raw or default)
    except Exception:
        return int(default)


def env_bool(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None or str(raw).strip() == "":
        return default
    return str(raw).strip().lower() not in {"0", "false", "no", "off"}


def load_local_env(path: str = ".env"):
    if not os.path.isabs(path):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), path)
    if not os.path.exists(path):
        return
    with open(path, "r", encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value


load_local_env()


# ==================== CẤU HÌNH HỆ THỐNG MẶC ĐỊNH ====================
BOT_TOKEN_FILE = os.getenv("BOT_TOKEN_FILE", r"D:\ZALO_MOVE\bot_token.txt").strip()
TOKEN_BOT = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()
if (not TOKEN_BOT) and BOT_TOKEN_FILE and os.path.exists(BOT_TOKEN_FILE):
    try:
        with open(BOT_TOKEN_FILE, "r", encoding="utf-8") as _f:
            TOKEN_BOT = (_f.read() or "").strip()
    except Exception:
        TOKEN_BOT = ""
ADMIN_ID = env_int("TELEGRAM_ADMIN_ID", 7959686962)
ADMIN_USERNAME = os.getenv("TELEGRAM_ADMIN_USERNAME", "@gifhub2708").strip()
SUPREME_ADMIN_ID = env_int("TELEGRAM_SUPREME_ADMIN_ID", ADMIN_ID)
SUPREME_ADMIN_USERNAME = os.getenv("TELEGRAM_SUPREME_ADMIN_USERNAME", ADMIN_USERNAME).strip().lstrip("@")
BOT_ATTENDANCE_CHAT = os.getenv("BOT_ATTENDANCE_CHAT", "").strip()
BOT_ATTENDANCE_ENABLED = env_bool("BOT_ATTENDANCE_ENABLED", True)
SUPPER_ADMIN_GROUP_LINK = os.getenv("SUPPER_ADMIN_GROUP_LINK", "").strip()
SUPPER_ADMIN_GROUP_ID = os.getenv("SUPPER_ADMIN_GROUP_ID", "").strip()
SUPPER_ADMIN_ENABLED = env_bool("SUPPER_ADMIN_ENABLED", True)
ADMIN_MANAGEMENT_GROUP_LINK = os.getenv("ADMIN_MANAGEMENT_GROUP_LINK", "").strip()
ADMIN_MANAGEMENT_GROUP_ID = os.getenv("ADMIN_MANAGEMENT_GROUP_ID", "").strip()
ADMIN_MANAGEMENT_ENABLED = env_bool("ADMIN_MANAGEMENT_ENABLED", True)
BOT_ATTENDANCE_GREETINGS = [
    "Co mat, admin check giup em nhe.",
    "Da len ca, san sang nhan lenh.",
    "Em vua khoi dong xong.",
    "Diem danh dau ca day admin.",
    "Online roi, cho lenh quet cua admin.",
    "Bao cao co mat.",
]
TARGET_CHAT = os.getenv("TARGET_CHAT", "").strip()
WEB_URL = os.getenv("WEB_URL", "https://TXX88.NET").strip()
LINK_NHAP_CODE = os.getenv("LINK_NHAP_CODE", "xx88code.com").strip()
CAMPAIGN_DEFAULT_GROUP_LINK = os.getenv("CAMPAIGN_START_GROUP_LINK", "https://t.me/+Iv1UDpbsB8w0ODA1").strip()
CAMPAIGN_REGISTER_LINK = os.getenv("CAMPAIGN_REGISTER_LINK", WEB_URL).strip()
CAMPAIGN_SUPPORT_USERNAME = os.getenv("CAMPAIGN_SUPPORT_USERNAME", ADMIN_USERNAME).strip().lstrip("@") or "gifhub2708"
CAMPAIGN_START_TEMPLATE = os.getenv("CAMPAIGN_START_TEMPLATE", "").strip()
CAMPAIGN_ALLOWED_GROUP_IDS = os.getenv("CAMPAIGN_ALLOWED_GROUP_IDS", "").strip()
CAMPAIGN_ASSETS_MANIFEST = os.getenv("CAMPAIGN_ASSETS_MANIFEST", "").strip()
CAMPAIGN_CONSENT_REQUIRED = env_bool("CAMPAIGN_CONSENT_REQUIRED", True)
CAMPAIGN_START_TEMPLATE_DEFAULT = (
    "Xin chào {first_name}!\n\n"
    "Cảm ơn bạn đã quan tâm. Bạn có thể dùng các liên kết chính thức bên dưới:\n\n"
    "- Nhóm chính: {group_link}\n"
    "- Đăng ký/thông tin: {register_link}\n"
    "- Hỗ trợ admin: https://t.me/{support_username}\n\n"
    "Nếu bạn không muốn nhận tin nữa, hãy nhắn /stop."
)
ZALO_IMEI = os.getenv("ZALO_IMEI", "").strip()
COOKIE_RAW = os.getenv("ZALO_COOKIE_RAW", "").strip()
ZALO_AUTO_REPLY = (
    "🇧🇷 👑 XX88 BRAZIL CODE VIP XIN CHÀO QUÝ KHÁCH 👑 🇧🇷\n\n"
    "✨✨✨          🌸🌸🌸🌸🌸          ✨✨✨\n"
    "🎁 MÃ ƯU ĐÃI: TN58 🎁\n"
    "🎈🎈🎈🎈🎈🎈🎈🎈🎈\n\n"
    "🥰 Thành viên mới chỉ cần:\n"
    "🤣 Điền đầy đủ thông tin đăng ký\n"
    "🤣 Xác minh số điện thoại qua Trung tâm khuyến mãi\n"
    "😌 Ngay lập tức nhận CODE VIP trải nghiệm miễn phí\n\n"
    "🔗 Đăng ký nhanh: https://TXX88.NET\n"
    "🤖 Nhận code tự động: t.me/xx88_code_bot\n\n"
    "⚠️ Lưu ý:\n"
    "💥 Chỉ được rút tiền khi đã tối thiểu 3 điểm.\n"
    "⚠️ Vi phạm điều kiện sẽ bị khấu trừ toàn bộ tiền thưởng và tiền thắng.\n\n"
    "😮 Nhanh tay tham gia để không bỏ lỡ ưu đãi tân thủ!"
)
ZALO_ADMIN_UID = os.getenv("ZALO_ADMIN_UID", "").strip()
ZALO_ADMIN_PHONE = os.getenv("ZALO_ADMIN_PHONE", "").strip()
ZALO_VERIFY_CMD = os.getenv("ZALO_VERIFY_CMD", "").strip()
ZALO_ADMIN_AUTO_FIND = True
ZALO_ADMIN_REP_MODE = "text"
ZALO_ADMIN_REP_IMAGE = ""
ZALO_AUTO_JOIN = False
TELE_AUTO_JOIN = False
ZALO_AUTO_REP_ENABLED = True
ZALO_AUTO_REP_ALL = True
ZALO_AUTO_REP_NGUOILA = True
ZALO_AUTO_REP_BANBE = True
ZALO_AUTO_REP_NEW = True
ZALO_PENDING_CMDS = {}
ZALO_2FA_ADMIN_CODE = os.getenv("ZALO_2FA_ADMIN_CODE", "").strip()
ZALO_FULL_HELP_UNLOCKED = set()
ZALO_GROUP_ADMINS = set()
ZALO_BANBE_BROADCAST_ENABLED = True
ZALO_BANBE_BROADCAST_INTERVAL = 120
ZALO_BANBE_BROADCAST_MSG = ""
ZALO_LAST_BANBE_BROADCAST_TS = 0.0
ZALO_BANBE_MIN_DELAY_SEC = 60
ZALO_BANBE_MAX_DELAY_SEC = 180
PENDING_APPROVAL_BY_MAIN = {}
GROUP_PRO_AI_NAME = "GROUP PRO AI"
OPENAI_INALL_ZALO = False
OPENAI_INALL_TELE = False
ZALO_BLESS_TIMERS = {}
OPENAI_INALL_ZALO_GROUPS = set()
OPENAI_INALL_TELE_GROUPS = set()
ADMIN_CHAT_NOTIFY = False
TELEGRAM_GROUP_SILENT_MODE = os.getenv("TELEGRAM_GROUP_SILENT_MODE", "1").strip().lower() not in {"0", "false", "no", "off"}
ZALO_SILENT_MODE = os.getenv("ZALO_SILENT_MODE", "1").strip().lower() not in {"0", "false", "no", "off"}
SCAN_REPORT_INTERVAL_SEC = 1800
PHONE_EVENT_SUMMARY_INTERVAL_SEC = 1800
BUSINESS_PRO_ZALO_GROUPS = set()
SESSION_BACKUP_INTERVAL_SEC = 3600
LIVE_CAPTURE_ZALO_GROUPS = set()
BACKUP_REMIND_INTERVAL_SEC = 3600
LAST_BACKUP_REMIND_TS = 0.0
BACKUP_PACKAGE_PASSWORD = os.getenv("BACKUP_PACKAGE_PASSWORD", "").strip()
CONTACT_EXPORT_INTERVAL_SEC = 1800
ZALO_BLESS_ENABLED = True
HELP_IMAGE_URLS = []
ZALO_2FA_IMAGE_URL = ""
CONTACT_BACKUP_DIR = r"D:\ZALO_MOVE\AI_FAMILY_BACKUP"
CONTACT_BACKUP_DAILY_CSV = os.path.join(CONTACT_BACKUP_DIR, "contacts_daily.csv")
CONTACT_BACKUP_DAILY_JSON = os.path.join(CONTACT_BACKUP_DIR, "contacts_daily.json")
CONTACT_BACKUP_TOTAL_CSV = os.path.join(CONTACT_BACKUP_DIR, "contacts_total.csv")
CONTACT_BACKUP_TOTAL_JSON = os.path.join(CONTACT_BACKUP_DIR, "contacts_total.json")
SCAN_PHONE_RE = re.compile(
    r"(?<!\d)(?:(?:\+?84|0084|0)(?:[\s.\-\u00A0\u2007\u202F]*\d){8,10}|[35789](?:[\s.\-\u00A0\u2007\u202F]*\d){8})(?!\d)"
)
SCAN_UID_RE = re.compile(r"(?<![\w+])(?:-100\d{6,20}|-?\d{6,25})(?![\w])")
SCAN_HANDLE_RE = re.compile(
    r"(?i)(?:^|[^\w])@([a-z0-9_][a-z0-9_.]{2,63})"
    r"|(?:https?://)?(?:t\.me|telegram\.me|facebook\.com|fb\.com)/(?!c/)([a-z0-9_][a-z0-9_.]{2,63})"
    r"|(?:https?://)?(?:www\.)?tiktok\.com/@([a-z0-9_][a-z0-9_.]{2,63})"
    r"|\b(t_[a-z0-9_]{5,63})\b"
    r"|\b(user\d{5,25})\b"
)
SCAN_FILE_MAX_BYTES = int(os.getenv("SCAN_FILE_MAX_BYTES", str(80 * 1024 * 1024)))
SCAN_ARCHIVE_MAX_FILES = int(os.getenv("SCAN_ARCHIVE_MAX_FILES", "300"))
SCAN_SQLITE_MAX_ROWS_PER_TABLE = int(os.getenv("SCAN_SQLITE_MAX_ROWS_PER_TABLE", "20000"))

BOT_ACTIVE = True
CHUC_NANG_AUTO_DUYET = True
THONG_BAO_START_ACTIVE = True
HIEN_THI_ANH_DONG = True

KEYWORDS_ROUTING = {
    "dangky": "https://TXX88.NET",
    "nhancode": "https://xx88code.com",
    "hotro": "https://t.me/gifhub2708"
}
BANNED_USERS = set()
EXTRA_ADMINS = set()
VIP_USERS = set()

QUANG_CAO_TEXT = (
    "🇧🇷🔥 BẢNG ƯU ĐÃI VIP XX88 BRAZIL 🔥🇧🇷\n\n"
    "✨✨✨          🌸🌸🌸🌸🌸          ✨✨✨\n"
    "🎁 CODE TÂN THỦ: TN58\n\n"
    "💎 Nạp đầu nhận thưởng 100% + CODE VIP theo từng mốc.\n"
    "👇 Quý khách chọn đúng mốc để hệ thống ghi nhận:\n\n"
    "💰 50K nhận 100% + Code 18K\n"
    "💰 100K nhận 100% + Code 38K\n"
    "💰 200K nhận 100% + Code 58K\n"
    "💎 500K nhận 100% + Code 128K\n"
    "🔥 1 Triệu nhận 100% + Code 288K\n\n"
    "⚠️ Chỉ áp dụng cho tài khoản đủ điều kiện nhận ưu đãi."
)
LOI_CHAO_MAC_DINH = (
    "🇧🇷 👑 TỔNG XX88 BRAZIL XIN CHÀO QUÝ KHÁCH 👑 🇧🇷\n\n"
    "✨✨✨          🌸🌸🌸🌸🌸          ✨✨✨\n"
    "🎁 MÃ ƯU ĐÃI: TN58 🎁\n"
    "🎈🎈🎈🎈🎈🎈🎈🎈🎈\n\n"
    "🥰 Rất hân hạnh được phục vụ quý khách!\n"
    "🔐 Đây là HỆ THỐNG TỰ ĐỘNG gửi link đăng ký và hỗ trợ nhận CODE VIP theo khu vực Brazil.\n\n"
    "🔥 Quý khách vui lòng chọn 1 trong 2 hình thức bên dưới để tiếp tục:"
)
TELE_REPLY_TEMPLATES_DEFAULT = {
    "register_prompt": (
        "🇧🇷🚀 KÍCH HOẠT TÀI KHOẢN XX88 BRAZIL 🚀🇧🇷\n\n"
        "✨✨✨          🌸🌸🌸🌸🌸          ✨✨✨\n"
        "🎁 MÃ ƯU ĐÃI: TN58\n\n"
        "🔗 Link đăng ký chính thức:\n"
        "{WEB_URL}\n\n"
        "📌 Sau khi đăng ký xong, quý khách nhập TÊN TÀI KHOẢN GAME xuống đây để hệ thống ghi nhận."
    ),
    "old_member_phone_prompt": (
        "🔐 XÁC THỰC THÀNH VIÊN CŨ 🔐\n\n"
        "📱 Quý khách vui lòng nhập SỐ ĐIỆN THOẠI đã đăng ký tài khoản game.\n"
        "✅ Hệ thống sẽ đối soát và chuyển admin hỗ trợ nhận CODE."
    ),
    "moc_intro": (
        "🇧🇷🔥 BẢNG MỐC KHUYẾN MÃI VIP XX88 BRAZIL 🔥🇧🇷\n\n"
        "💎 Nạp đầu nhận thưởng 100% + CODE VIP theo từng mốc.\n"
        "👇 Quý khách vui lòng chọn đúng mốc nạp để hệ thống ghi nhận:"
    ),
    "join_not_verified": (
        "⚠️ Hệ thống chưa xác nhận quý khách đã tham gia kênh.\n\n"
        "👉 Vui lòng vào {TARGET_CHAT}\n"
        "✅ Sau đó quay lại bấm nút xác nhận thêm một lần."
    ),
    "choose_reward_mode": (
        "🎁 CHỌN HÌNH THỨC NHẬN THƯỞNG 🎁\n\n"
        "👨‍💼 Liên hệ admin nếu cần hỗ trợ nhanh.\n"
        "🤖 Chọn hệ thống tự động nếu quý khách muốn gửi yêu cầu duyệt CODE ngay."
    ),
    "auto_review_disabled": (
        "⚠️ Hệ thống tự động đang tạm khóa.\n"
        "👨‍💼 Quý khách vui lòng liên hệ admin để được hỗ trợ trực tiếp."
    ),
    "auto_review_waiting": (
        "🔄 Hệ thống đang kiểm tra yêu cầu.\n\n"
        "✅ Hồ sơ đã được gửi tới admin xử lý.\n"
        "⏳ Quý khách vui lòng chờ trong giây lát."
    ),
    "promo_other": (
        "🎁 ƯU ĐÃI KHÁC 🎁\n\n"
        "🔥 Thưởng nạp ngày vàng lần 2 nhận 20%.\n"
        "📌 Chi tiết sẽ hiển thị trực tiếp trên trang chủ khi đủ điều kiện."
    ),
    "old_member_phone_saved": (
        "✅ ĐÃ GHI NHẬN SỐ ĐIỆN THOẠI\n\n"
        "📱 Hệ thống đã đồng bộ thông tin thành viên cũ.\n"
        "👨‍💼 Quý khách vui lòng nhắn admin để được hỗ trợ nhận CODE:\n"
        "🔗 t.me/{ADMIN_USERNAME_NOSTRIP}"
    ),
    "ask_phone_after_tk": (
        "🇧🇷✅ ĐÃ GHI NHẬN TÀI KHOẢN GAME:\n"
        "{ACCOUNT}\n\n"
        "📱 Quý khách vui lòng nhập tiếp SỐ ĐIỆN THOẠI CHÍNH CHỦ.\n"
        "🔐 Hệ thống dùng số này để đối soát và bảo mật hồ sơ nhận CODE VIP."
    ),
    "join_required": (
        "🇧🇷📢 BƯỚC BẮT BUỘC ĐỂ MỞ CODE VIP 📢🇧🇷\n\n"
        "👉 Quý khách vui lòng tham gia kênh ưu đãi chính thức:\n"
        "{TARGET_CHAT}\n\n"
        "🎁 Sau khi tham gia, bấm nút xác nhận bên dưới để hệ thống mở bảng mốc thưởng."
    ),
}
TELE_REPLY_TEMPLATES = dict(TELE_REPLY_TEMPLATES_DEFAULT)
TELE_REPLY_MEDIA = {}
USER_DATA = {}
ADMIN_CONTENT_WIZARD = {}
ZALO_FRIEND_IDS = set()
ZALO_SEEN_USER_IDS = set()
ZALO_LAST_AUTO_REPLY_TS = {}
ZALO_REPLY_COOLDOWN_SEC = 180
DB_PATH = os.getenv("CODEX_DB", r"D:\ZALO_MOVE\codevip_merge.db")
DB_BACKUP_PATH = os.getenv("CODEX_DB_BACKUP", r"D:\ZALO_MOVE\codevip_merge.db.bak")
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DATA_DIR = os.getenv("ZALO_MOVE_DIR", r"D:\ZALO_MOVE").strip() or BASE_DIR
if not os.path.isdir(PROJECT_DATA_DIR):
    PROJECT_DATA_DIR = BASE_DIR
DEFAULT_USERBOT_SHARD_ROOT = os.path.join(PROJECT_DATA_DIR, "userbot_shards")
DEFAULT_USERBOT_SHARD_ROOTS = ";".join(
    os.path.join(DEFAULT_USERBOT_SHARD_ROOT, f"botuser_{idx:02d}") for idx in range(1, 6)
)
SUPPER_ADMIN_SHARED_CONFIG = os.getenv("SUPPER_ADMIN_SHARED_CONFIG", os.path.join(BASE_DIR, "supper_admin_group.json")).strip()
ADMIN_MANAGEMENT_SHARED_CONFIG = os.getenv("ADMIN_MANAGEMENT_SHARED_CONFIG", os.path.join(BASE_DIR, "admin_quanly_group.json")).strip()
AI_SECRETS_CSV = os.getenv("BOT_AI_SECRETS_CSV", os.path.join(BASE_DIR, "bot_ai_secrets.csv"))
AI_SECRETS_ENC = os.getenv("BOT_AI_SECRETS_ENC", os.path.join(BASE_DIR, "bot_ai_secrets.csv.enc"))
USERBOT_SCAN_JOBS_DIR = os.getenv("USERBOT_SCAN_JOBS_DIR", os.path.join(PROJECT_DATA_DIR, "userbot_scan_jobs"))
try:
    USERBOT_JOB_PROFILE_LIMIT = max(1, int(os.getenv("USERBOT_JOB_PROFILE_LIMIT", "12").strip() or "12"))
except Exception:
    USERBOT_JOB_PROFILE_LIMIT = 12
CONTACT_IMPORT_PENDING_JSONL = os.path.join(USERBOT_SCAN_JOBS_DIR, "pending.jsonl")
CONTACT_IMPORT_CLAIMED_JSON = os.path.join(USERBOT_SCAN_JOBS_DIR, "claimed_contact_import_ids.json")
CONTACT_IMPORT_PROCESSED_JSON = os.path.join(USERBOT_SCAN_JOBS_DIR, "processed_contact_import_ids.json")
CONTACT_IMPORT_WORKER_LOG = os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_import_contacts_worker.log")
CONTACT_IMPORT_RESULTS_JSONL = os.path.join(USERBOT_SCAN_JOBS_DIR, "contact_import_results.jsonl")
CONTACT_IMPORT_LAST_SUMMARY = os.path.join(USERBOT_SCAN_JOBS_DIR, "contact_import_last_summary.txt")
CONTACT_IMPORT_LAST_USERS_CSV = os.path.join(USERBOT_SCAN_JOBS_DIR, "contact_import_last_users.csv")
CONTACT_IMPORT_USERS_JSONL = os.path.join(USERBOT_SCAN_JOBS_DIR, "contact_import_users.jsonl")
USERBOT_SEND_LIVE_STATUS_JSON = os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_send_live_status.json")
USERBOT_SEND_LAST_SUMMARY = os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_send_last_summary.txt")
USERBOT_SEND_RESULTS_JSONL = os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_send_results.jsonl")
USERBOT_SEND_RECIPIENTS_CSV = os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_send_last_recipients.csv")
USERBOT_STOP_ALL_JSON = os.path.join(USERBOT_SCAN_JOBS_DIR, "stop_all_jobs.json")
USERBOT_STOP_ALL_LAST_SUMMARY = os.path.join(USERBOT_SCAN_JOBS_DIR, "stop_all_jobs_last_summary.txt")
AI_API_ENDPOINT = "https://models.github.ai/inference/chat/completions"
AI_MODEL = "meta/Llama-4-Scout-17B-16E-Instruct"
AI_API_TOKEN_ENV = "GH_MODELS_TOKEN"
OPENAI_API_ENDPOINT = "https://api.openai.com/v1/chat/completions"
OPENAI_MODEL = "gpt-4o"
AI_PROVIDER_ENV_HINTS = {
    "openai": "OPENAI_API_KEY",
    "github": AI_API_TOKEN_ENV,
    "git": AI_API_TOKEN_ENV,
    "meta": "META_API_KEY",
    "alo": "ALO_API_KEY",
}
OPENAI_GITHUB_BASE_URL = "https://models.inference.ai.azure.com"
GIT_EXE_PATH = r"D:\Git\cmd\git.exe"
os.environ["GIT_PYTHON_GIT_EXECUTABLE"] = GIT_EXE_PATH
ADMIN_NOTIFY_LAST_TS = {}
LAST_ZALO_BOT = None

# ===== ONE-FILE USERBOT ACCOUNT BOOTSTRAP =====
USERBOT_SYNC_ENABLED = os.getenv("USERBOT_SYNC_ENABLED", "1") == "1"
USERBOT_ACCOUNTS_ROOT = os.getenv("USERBOT_ACCOUNTS_ROOT", os.path.join(DEFAULT_USERBOT_SHARD_ROOT, "botuser_01"))
USERBOT_20_ACCOUNTS_ROOT = os.getenv("USERBOT_20_ACCOUNTS_ROOT", r"D:\ZALO_MOVE\20ac xinhan 2 6")
USERBOT_12_LO_ROOT = os.getenv("USERBOT_12_LO_ROOT", r"D:\ZALO_MOVE\10ac xinhan 265 dm")
USERBOT_20_XINHAN_ROOT = os.getenv("USERBOT_20_XINHAN_ROOT", r"D:\ZALO_MOVE\20ac xinhan 2 6")
USERBOT_20_NEW_ROOT = os.getenv("USERBOT_20_NEW_ROOT", r"D:\ZALO_MOVE\20ac_moi_20260603")
USERBOT_32_LO_ROOTS_RAW = os.getenv(
    "USERBOT_32_LO_ROOTS",
    r"D:\ZALO_MOVE\10ac xinhan 265 dm;D:\ZALO_MOVE\20ac xinhan 2 6",
).strip()
USERBOT_52_ROOT = os.getenv("USERBOT_52_ROOT", r"D:\ZALO_MOVE\52acc_tong_hop_20260603")
GUITN_USERBOT_ACCOUNTS_ROOT = os.getenv("GUITN_USERBOT_ACCOUNTS_ROOT", "").strip()
USERBOT_SHARD_ROOTS_RAW = os.getenv("USERBOT_SHARD_ROOTS", DEFAULT_USERBOT_SHARD_ROOTS).strip()
USERBOT_ACCOUNT_COOLDOWN_JSON = os.getenv(
    "USERBOT_ACCOUNT_COOLDOWN_JSON",
    os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_account_cooldown.json"),
).strip()
USERBOT_API_ID_RAW = os.getenv("USERBOT_API_ID", "0").strip()
USERBOT_API_ID = int(USERBOT_API_ID_RAW or "0")
USERBOT_API_HASH = os.getenv("USERBOT_API_HASH", "").strip()
USERBOT_ROLE = os.getenv("USERBOT_ROLE", "admin-main").strip().lower()


def load_machine_identity_code() -> str:
    path = os.path.join(BASE_DIR, "machine_identity.json")
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8-sig") as f:
                data = json.load(f)
            return str(data.get("machine_code") or "").strip()
    except Exception:
        return ""
    return ""


DEFAULT_MACHINE_DEVICE = load_machine_identity_code() or "VGAH510"
USERBOT_ADMIN_DEVICE = os.getenv("USERBOT_ADMIN_DEVICE", DEFAULT_MACHINE_DEVICE).strip()
USERBOT_MANAGER_DEVICE = os.getenv("USERBOT_MANAGER_DEVICE", DEFAULT_MACHINE_DEVICE).strip()


def sync_userbot_account_jsons() -> dict:
    report = {"total": 0, "updated": 0, "created": 0, "failed": 0}
    if not USERBOT_SYNC_ENABLED:
        return report
    if not os.path.isdir(USERBOT_ACCOUNTS_ROOT):
        logging.warning("Userbot accounts folder not found: %s", USERBOT_ACCOUNTS_ROOT)
        return report

    for name in sorted(os.listdir(USERBOT_ACCOUNTS_ROOT)):
        acc_dir = os.path.join(USERBOT_ACCOUNTS_ROOT, name)
        if not os.path.isdir(acc_dir):
            continue
        report["total"] += 1
        json_path = os.path.join(acc_dir, f"{name}.json")
        created = not os.path.exists(json_path)
        try:
            if created:
                data = {}
            else:
                with open(json_path, "r", encoding="utf-8-sig") as f:
                    data = json.load(f)
                if not isinstance(data, dict):
                    data = {}

            data["session_file"] = str(data.get("session_file") or name)
            data["phone"] = str(data.get("phone") or name)
            if USERBOT_API_ID:
                data["app_id"] = USERBOT_API_ID
            elif created and not data.get("app_id"):
                data["app_id"] = 0
            if USERBOT_API_HASH:
                data["app_hash"] = USERBOT_API_HASH
            elif created and "app_hash" not in data:
                data["app_hash"] = ""
            data["sdk"] = str(data.get("sdk") or "Windows 10")
            data["app_version"] = str(data.get("app_version") or "6.6.2 x64")
            data["device"] = USERBOT_ADMIN_DEVICE if USERBOT_ROLE == "admin-main" else USERBOT_MANAGER_DEVICE
            data["lang_pack"] = str(data.get("lang_pack") or "tdesktop")
            data["system_lang_pack"] = str(data.get("system_lang_pack") or "en-US")
            data.setdefault("username", None)
            data["ipv6"] = bool(data.get("ipv6", False))
            data.setdefault("first_name", None)
            data.setdefault("last_name", None)
            data.setdefault("register_time", None)
            data.setdefault("sex", None)
            data.setdefault("last_check_time", None)
            data["lang_code"] = str(data.get("lang_code") or "en")
            data["avatar"] = str(data.get("avatar") or "img/default.png")
            data.setdefault("proxy", None)
            data.setdefault("twoFA", None)
            data["block"] = bool(data.get("block", False))
            data["system_lang_code"] = str(data.get("system_lang_code") or "en-US")
            data.setdefault("id", None)

            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=4)

            report["created" if created else "updated"] += 1
        except Exception as exc:
            report["failed"] += 1
            logging.warning("Cannot sync account JSON %s: %s", json_path, exc)
    return report


def validate_telegram_bot_token() -> bool:
    if not TOKEN_BOT or TOKEN_BOT.startswith("<") or ":" not in TOKEN_BOT:
        print("[BOOT][ERROR] TELEGRAM_BOT_TOKEN missing/invalid.")
        print(f"[BOOT][HINT] Put the BotFather token into: {BOT_TOKEN_FILE}")
        return False
    try:
        req = urllib.request.Request(f"https://api.telegram.org/bot{TOKEN_BOT}/getMe", method="GET")
        with urllib.request.urlopen(req, timeout=15) as resp:
            payload = json.loads(resp.read().decode("utf-8", errors="replace"))
        if not payload.get("ok"):
            print("[BOOT][ERROR] Telegram token rejected.")
            return False
        bot = payload.get("result", {}) or {}
        print(f"[BOOT] Telegram bot OK: @{bot.get('username', 'unknown')} id={bot.get('id', '?')}")
        return True
    except Exception as exc:
        print(f"[BOOT][ERROR] Telegram token check failed: {exc}")
        return False


def one_file_bootstrap() -> bool:
    rep = sync_userbot_account_jsons()
    if USERBOT_SYNC_ENABLED:
        print(
            f"[BOOT] Userbot JSON sync: total={rep['total']} updated={rep['updated']} "
            f"created={rep['created']} failed={rep['failed']}"
        )
    return validate_telegram_bot_token()
CMD_EXPLAIN = {
    "zai": "Tinh nang nay cho phep Admin goi AI tu moi truong Github de tra loi cau hoi.",
    "zaddadmiai": "Lenh nay cap quyen AI rieng cho admin phu, chi duoc dung lenh zai.",
    "zaddadminall": "Lenh nay cap toan quyen cau hinh bot cho admin phu.",
    "zautorep": "Bat/tat tu dong tra loi theo bo tu khoa da cai dat.",
    "zautojoin": "Bat/tat che do tu dong xu ly link nhom.",
    "zsetrep": "Dat noi dung tra loi theo tu khoa khach gui den.",
    "zchecksetrep": "Kiem tra danh sach tu khoa auto-rep hien co.",
    "zcheckiudall": "Tong hop toan bo UID da quet kem ten.",
    "zbanbe": "Bat/tat gui tin tu dong den danh ba theo chu ky thoi gian.",
    "zsettn": "Dat noi dung tin nhan dung cho zbanbe.",
    "zaddmingr": "Cap them admin nhom zalo.",
    "zcapnhat": "Goi AI phan tich yeu cau cap nhat va de xuat cach lam.",
    "zaitokens": "Quan ly da moi truong AI: zaitokens add <api_nen_tang> <api_token> <ten_ai> | use | del | list.",
    "quetnhanh": "Quet nhanh SDT trong text/reply, map voi user/uid trong danh ba bot.",
    "quetfile": "Quet file txt/csv gui len Telegram, map phone -> user/uid/ten tu du lieu bot.",
    "quetfileepath": "Quet file local theo duong dan tren may dang chay bot.",
    "quetfilepath": "Alias cu cua /quetfileepath.",
    "quetnow": "Quet nhanh hang loat so dien thoai bang text admin gui, map user/uid/link va luu vao danh ba bot.",
    "quetuser": "Quet danh sach SDT va map user/uid/link hien co; so chua co map van duoc luu vao danh ba bot.",
    "setsuppergroup": "Luu/bat/tat nhom supper admin: /setsuppergroup on|off hoac go trong nhom de set chat id.",
    "setsupergroup": "Alias cua /setsuppergroup.",
    "adminquanly": "Luu/bat/tat nhom /adminquanly: /adminquanly on|off de nhan report van hanh.",
    "keoall": "Reply file uid|user roi dry-run/chay that keo nguoi vao nhom bang 1 userbot; ket qua luu tung folder rieng.",
    "keoallstatus": "Xem bao cao /keoall gan nhat bang tieng Viet.",
    "quetstatus": "Xem tien trinh quet/import SDT -> UID/user theo tong job va tung acc userbot.",
    "guitnstatus": "Xem tien trinh gui tin /guitn theo tong so va tung acc userbot.",
    "guitnlai": "Lay file khach /guitn gui chua thanh cong gan nhat, hoac file reply, de gui lai.",
    "stopall": "Dung tat ca job userbot dang cho/dang chay; neu co job dang chay se tat worker de cat job ngay.",
    "stop": "Dung rieng bot father duoc tag bang cu phap /stop@ten_bot.",
    "suarepsdt": "Sua mau tin bot yeu cau khach nhap so dien thoai, giu nguyen xuong dong.",
    "suarep": "Sua tung mau tin khach theo ma: /suarep <ma> <noi_dung>.",
    "xemrep": "Xem danh sach ma mau tin khach co the sua.",
}

GIF_CHAO_HOI = "https://media.giphy.com/media/v1.Y2lkPTc5MGI3NjExM2I4YTM0NmRiaWVmYTg3Y2Z0ZzB6bXN4NTR6b3g1YW90Y2N4ZHp6OCZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/l0ExhcMymdL6vY9aM/giphy.gif"
GIF_DANG_KY = "https://media.giphy.com/media/v1.Y2lkPTc5MGI3NjExbXByZndkMjVvOTU1Y3VlMWxlYWVnOXcyYTVoZGQwaW1oZjBwazl0byZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/3o6gDWzmAzrpi5DQU8/giphy.gif"
GIF_EP_JOIN = "https://media.giphy.com/media/v1.Y2lkPTc5MGI3NjExZ3NjcTR3dXRwbWtwb3Y5dmRxeTVod2V1Y2Z2N204MXg4czBwd2hyeSZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/26AHPxxnSw1L9sh1u/giphy.gif"
GIF_MOC_NAP = "https://media.giphy.com/media/v1.Y2lkPTc5MGI3NjExbnkzamNidDRwZDR4NDNqd3Fhc3Rnb3FhcWw1N3F4cm50YWw0ZHp6OCZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/3oz8xAFtqoOUUrsh7W/giphy.gif"
GIF_CHO_DUYET = "https://media.giphy.com/media/v1.Y2lkPTc5MGI3NjEx3ZidTh0YThqdGptMnN2ZHBidmN0NDNsd3pka3R2YWJ0Yzh2bHpxOCZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/tXL4FHPSnVJ0A/giphy.gif"
GIF_QUANG_CAO = "https://media.giphy.com/media/v1.Y2lkPTc5MGI3NjExaXpneDB4b3hmdG8zbmt4dDJndWR5bmY3bjUxeWVzOHJ1aDk2ZnB3NyZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/L1QMTl9gYOf3TH762E/giphy.gif"
USERBOT_GUITN_GIF = ""

def db_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_store():
    with db_conn() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS config (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS install_tokens (
                token_hash TEXT PRIMARY KEY,
                raw_token TEXT NOT NULL,
                label TEXT DEFAULT '',
                created_at REAL NOT NULL,
                expires_at REAL NOT NULL,
                redeemed INTEGER NOT NULL DEFAULT 0
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS audit (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                actor_platform TEXT NOT NULL,
                actor_uid TEXT NOT NULL,
                action TEXT NOT NULL,
                detail TEXT DEFAULT '',
                created_at REAL NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS ai_permissions (
                platform TEXT NOT NULL,
                uid TEXT NOT NULL,
                enabled INTEGER NOT NULL DEFAULT 1,
                daily_quota INTEGER NOT NULL DEFAULT 30,
                used_today INTEGER NOT NULL DEFAULT 0,
                day_key TEXT DEFAULT '',
                created_at REAL NOT NULL,
                PRIMARY KEY (platform, uid)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS outbox (
                platform TEXT NOT NULL,
                target_uid TEXT NOT NULL,
                msg_hash TEXT NOT NULL,
                created_at REAL NOT NULL,
                PRIMARY KEY(platform, target_uid, msg_hash)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS customer_contacts (
                uid TEXT PRIMARY KEY,
                platform TEXT NOT NULL,
                name TEXT DEFAULT '',
                username TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                last_message TEXT DEFAULT '',
                updated_at REAL NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS campaign_target_consent (
                platform TEXT NOT NULL DEFAULT 'telegram',
                target_uid TEXT NOT NULL,
                username TEXT DEFAULT '',
                phone TEXT DEFAULT '',
                source TEXT DEFAULT '',
                status TEXT NOT NULL DEFAULT 'opt_in',
                updated_at REAL NOT NULL,
                PRIMARY KEY(platform, target_uid)
            )
            """
        )
        conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_campaign_target_consent_username ON campaign_target_consent(platform, username)"
        )
        conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_campaign_target_consent_phone ON campaign_target_consent(platform, phone)"
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS keyword_replies (
                keyword TEXT PRIMARY KEY,
                content TEXT NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS scanned_phone_leads (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                phone TEXT NOT NULL,
                username TEXT DEFAULT '',
                full_name TEXT DEFAULT '',
                source_uid TEXT DEFAULT '',
                source_chat TEXT DEFAULT '',
                source_type TEXT DEFAULT '',
                created_at REAL NOT NULL,
                UNIQUE(phone, source_uid, source_chat)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS bot_phonebook (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                phone TEXT NOT NULL,
                platform TEXT NOT NULL,
                source_uid TEXT DEFAULT '',
                username TEXT DEFAULT '',
                full_name TEXT DEFAULT '',
                source_chat TEXT DEFAULT '',
                first_seen REAL NOT NULL,
                last_seen REAL NOT NULL,
                hit_count INTEGER NOT NULL DEFAULT 1,
                last_message_snippet TEXT DEFAULT '',
                UNIQUE(phone, platform, source_uid, source_chat)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS phone_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                phone TEXT NOT NULL,
                platform TEXT NOT NULL,
                source_uid TEXT DEFAULT '',
                username TEXT DEFAULT '',
                full_name TEXT DEFAULT '',
                source_chat TEXT DEFAULT '',
                is_new_phone INTEGER NOT NULL DEFAULT 0,
                is_new_uid INTEGER NOT NULL DEFAULT 0,
                created_at REAL NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS telegram_group_members (
                chat_id TEXT NOT NULL,
                user_id TEXT NOT NULL,
                name TEXT DEFAULT '',
                username TEXT DEFAULT '',
                last_seen REAL NOT NULL,
                hit_count INTEGER NOT NULL DEFAULT 1,
                PRIMARY KEY(chat_id, user_id)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS userbot_scan_jobs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                command TEXT NOT NULL,
                chat_id TEXT DEFAULT '',
                chat_title TEXT DEFAULT '',
                requested_by TEXT DEFAULT '',
                dry_run INTEGER NOT NULL DEFAULT 1,
                status TEXT NOT NULL DEFAULT 'queued',
                detail TEXT DEFAULT '',
                created_at REAL NOT NULL,
                updated_at REAL NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                uid INTEGER PRIMARY KEY,
                step TEXT DEFAULT '',
                ten_game TEXT DEFAULT '',
                sdt TEXT DEFAULT '',
                goi_nap TEXT DEFAULT '',
                is_banned INTEGER NOT NULL DEFAULT 0,
                is_vip INTEGER NOT NULL DEFAULT 0,
                created_at REAL NOT NULL,
                updated_at REAL NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS bot_admins (
                uid INTEGER PRIMARY KEY,
                role TEXT NOT NULL DEFAULT 'extra',
                created_at REAL NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS broadcast_queue (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                platform TEXT NOT NULL DEFAULT 'telegram',
                target_uid TEXT NOT NULL,
                message TEXT NOT NULL,
                msg_hash TEXT NOT NULL,
                status TEXT NOT NULL DEFAULT 'queued',
                attempts INTEGER NOT NULL DEFAULT 0,
                last_error TEXT DEFAULT '',
                requested_by TEXT DEFAULT '',
                created_at REAL NOT NULL,
                updated_at REAL NOT NULL,
                sent_at REAL
            )
            """
        )
        conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_broadcast_queue_status ON broadcast_queue(status, updated_at)"
        )
        conn.execute(
            "CREATE UNIQUE INDEX IF NOT EXISTS idx_broadcast_queue_dedupe ON broadcast_queue(platform, target_uid, msg_hash)"
        )
        conn.commit()


def backup_db(reason: str = "manual"):
    try:
        if os.path.exists(DB_PATH):
            shutil.copy2(DB_PATH, DB_BACKUP_PATH)
            logging.info("[BACKUP] DB backup ok (%s) -> %s", reason, DB_BACKUP_PATH)
    except Exception as exc:
        logging.warning("Backup DB that bai: %s", exc)


def cleanup_legacy_session_backup_files(out_dir: str) -> None:
    if not os.path.isdir(out_dir):
        return
    keep = {"backup_download_latest.zip", "backup_password_latest.txt", "latest_runtime.json"}
    timestamp_backup_re = re.compile(r"^\d{8}_\d{6}_.+\.(db|json|sqlite|bak)$", re.IGNORECASE)
    for fn in os.listdir(out_dir):
        lower = fn.lower()
        if lower in keep or lower.startswith("latest_"):
            continue
        if not timestamp_backup_re.match(fn):
            continue
        path = os.path.join(out_dir, fn)
        try:
            if os.path.isfile(path):
                os.remove(path)
        except Exception as exc:
            logging.warning("Xoa session backup cu that bai %s: %s", path, exc)


def backup_session_snapshot(reason: str = "manual") -> str:
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        out_dir = os.path.join(base_dir, "session_backups")
        os.makedirs(out_dir, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        db_src = os.path.abspath(DB_PATH)
        db_name = os.path.basename(db_src)
        db_out = os.path.join(out_dir, f"latest_{db_name}")
        if os.path.exists(db_src):
            shutil.copy2(db_src, db_out)
        # Save runtime config snapshot
        cfg = {
            "reason": reason,
            "created_at": ts,
            "zalo_admin_uid": ZALO_ADMIN_UID,
            "zalo_admin_phone": ZALO_ADMIN_PHONE,
            "bot_active": BOT_ACTIVE,
            "openai_inall_zalo_groups": sorted(list(OPENAI_INALL_ZALO_GROUPS)),
            "openai_inall_tele_groups": sorted(list(OPENAI_INALL_TELE_GROUPS)),
            "business_pro_zalo_groups": sorted(list(BUSINESS_PRO_ZALO_GROUPS)),
        }
        with open(os.path.join(out_dir, "latest_runtime.json"), "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
        cleanup_legacy_session_backup_files(out_dir)
        return out_dir
    except Exception as exc:
        logging.warning("Backup session snapshot that bai: %s", exc)
        return ""


def auto_restore_from_latest_backup():
    try:
        # Neu DB ton tai va mo duoc thi bo qua restore
        if os.path.exists(DB_PATH):
            try:
                conn = sqlite3.connect(DB_PATH)
                conn.execute("SELECT 1")
                conn.close()
                return
            except Exception:
                pass

        base_dir = os.path.dirname(os.path.abspath(__file__))
        backup_dir = os.path.join(base_dir, "session_backups")
        candidates = []
        if os.path.isdir(backup_dir):
            for fn in os.listdir(backup_dir):
                if fn.endswith(".db") or fn.endswith(".bak") or fn.endswith(".sqlite") or "codevip_merge.db" in fn:
                    candidates.append(os.path.join(backup_dir, fn))
        if os.path.exists(DB_BACKUP_PATH):
            candidates.append(os.path.abspath(DB_BACKUP_PATH))
        if not candidates:
            return
        latest = max(candidates, key=lambda p: os.path.getmtime(p))
        shutil.copy2(latest, DB_PATH)
        logging.info("[RESTORE] Da khoi phuc DB tu backup: %s", latest)
    except Exception as exc:
        logging.warning("Auto restore backup that bai: %s", exc)


def make_backup_download_package() -> str:
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        out_dir = os.path.join(base_dir, "session_backups")
        os.makedirs(out_dir, exist_ok=True)
        zip_path = os.path.join(out_dir, "backup_download_latest.zip")
        note_path = os.path.join(out_dir, "backup_password_latest.txt")
        with open(note_path, "w", encoding="utf-8") as f:
            f.write(f"Backup password: {BACKUP_PACKAGE_PASSWORD}\n")
            f.write("Use this password when restoring backup package.\n")
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            if os.path.exists(DB_PATH):
                zf.write(DB_PATH, arcname=os.path.basename(DB_PATH))
            if os.path.exists(DB_BACKUP_PATH):
                zf.write(DB_BACKUP_PATH, arcname=os.path.basename(DB_BACKUP_PATH))
            zf.write(note_path, arcname=os.path.basename(note_path))
        return zip_path
    except Exception as exc:
        logging.warning("Tao backup download package that bai: %s", exc)
        return ""


def _contact_backup_key(row: dict) -> tuple[str, str, str, str]:
    platform = str(row.get("platform") or "").strip().lower()
    uid = str(row.get("uid") or "").strip()
    username = str(row.get("username") or "").strip().lower().lstrip("@")
    phone = "".join(ch for ch in str(row.get("phone") or "") if ch.isdigit())
    return platform, uid, username, phone


def _merge_contact_backup_rows(existing: list[dict], incoming: list[dict]) -> list[dict]:
    merged: dict[tuple[str, str, str, str], dict] = {}
    for row in existing + incoming:
        item = {
            "uid": str(row.get("uid") or "").strip(),
            "platform": str(row.get("platform") or "").strip(),
            "name": str(row.get("name") or "").strip(),
            "username": str(row.get("username") or "").strip(),
            "phone": str(row.get("phone") or "").strip(),
            "updated_at": float(row.get("updated_at") or 0),
        }
        key = _contact_backup_key(item)
        if not any(key):
            continue
        old = merged.get(key)
        if not old or item["updated_at"] >= float(old.get("updated_at") or 0):
            merged[key] = item
    return sorted(merged.values(), key=lambda x: float(x.get("updated_at") or 0), reverse=True)


def _read_contact_backup_json(path: str) -> list[dict]:
    if not os.path.exists(path):
        return []
    try:
        with open(path, "r", encoding="utf-8-sig") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _write_contact_backup_files(rows: list[dict], json_path: str, csv_path: str) -> None:
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(rows, f, ensure_ascii=False, indent=2)
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["uid", "platform", "name", "username", "phone", "updated_at"])
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in ["uid", "platform", "name", "username", "phone", "updated_at"]})


def cleanup_legacy_contact_backup_files() -> None:
    keep = {
        os.path.basename(CONTACT_BACKUP_DAILY_CSV).lower(),
        os.path.basename(CONTACT_BACKUP_DAILY_JSON).lower(),
        os.path.basename(CONTACT_BACKUP_TOTAL_CSV).lower(),
        os.path.basename(CONTACT_BACKUP_TOTAL_JSON).lower(),
    }
    if not os.path.isdir(CONTACT_BACKUP_DIR):
        return
    for fn in os.listdir(CONTACT_BACKUP_DIR):
        lower = fn.lower()
        if lower in keep:
            continue
        if not (lower.startswith("contacts_") and lower.endswith((".csv", ".json"))):
            continue
        path = os.path.join(CONTACT_BACKUP_DIR, fn)
        try:
            if os.path.isfile(path):
                os.remove(path)
        except Exception as exc:
            logging.warning("Xoa backup cu that bai %s: %s", path, exc)


def export_contacts_to_my_documents(reason: str = "auto") -> str:
    try:
        out_dir = CONTACT_BACKUP_DIR
        os.makedirs(out_dir, exist_ok=True)
        day_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0).timestamp()
        with db_conn() as conn:
            db_rows = conn.execute(
                "SELECT uid, platform, name, username, phone, updated_at FROM customer_contacts ORDER BY updated_at DESC"
            ).fetchall()
        incoming = [
            {
                "uid": str(r["uid"] or ""),
                "platform": str(r["platform"] or ""),
                "name": str(r["name"] or ""),
                "username": str(r["username"] or ""),
                "phone": str(r["phone"] or ""),
                "updated_at": float(r["updated_at"] or 0),
            }
            for r in db_rows
        ]
        total_rows = _merge_contact_backup_rows(_read_contact_backup_json(CONTACT_BACKUP_TOTAL_JSON), incoming)
        daily_incoming = [row for row in incoming if float(row.get("updated_at") or 0) >= day_start]
        daily_rows = _merge_contact_backup_rows(_read_contact_backup_json(CONTACT_BACKUP_DAILY_JSON), daily_incoming)
        daily_rows = [row for row in daily_rows if float(row.get("updated_at") or 0) >= day_start]
        _write_contact_backup_files(total_rows, CONTACT_BACKUP_TOTAL_JSON, CONTACT_BACKUP_TOTAL_CSV)
        _write_contact_backup_files(daily_rows, CONTACT_BACKUP_DAILY_JSON, CONTACT_BACKUP_DAILY_CSV)
        cleanup_legacy_contact_backup_files()
        logging.info("[EXPORT] Contacts backup updated reason=%s daily=%s total=%s dir=%s", reason, len(daily_rows), len(total_rows), out_dir)
        return out_dir
    except Exception as exc:
        logging.warning("Export contacts backup that bai: %s", exc)
        return ""


def load_runtime_state():
    global USER_DATA, EXTRA_ADMINS, ZALO_ADMIN_REP_MODE, ZALO_ADMIN_REP_IMAGE
    global ZALO_ADMIN_UID, ZALO_ADMIN_PHONE, ZALO_VERIFY_CMD, ZALO_ADMIN_AUTO_FIND, ZALO_IMEI, COOKIE_RAW
    global SUPPER_ADMIN_GROUP_ID, SUPPER_ADMIN_GROUP_LINK, SUPPER_ADMIN_ENABLED
    global ADMIN_MANAGEMENT_GROUP_ID, ADMIN_MANAGEMENT_GROUP_LINK, ADMIN_MANAGEMENT_ENABLED
    global TARGET_CHAT, WEB_URL, LINK_NHAP_CODE
    global CAMPAIGN_DEFAULT_GROUP_LINK, CAMPAIGN_REGISTER_LINK, CAMPAIGN_SUPPORT_USERNAME
    global CAMPAIGN_START_TEMPLATE, CAMPAIGN_ALLOWED_GROUP_IDS, CAMPAIGN_ASSETS_MANIFEST, CAMPAIGN_CONSENT_REQUIRED
    global GIF_CHAO_HOI, GIF_DANG_KY, GIF_EP_JOIN, GIF_MOC_NAP, GIF_CHO_DUYET, GIF_QUANG_CAO, USERBOT_GUITN_GIF
    global ZALO_AUTO_REPLY, LOI_CHAO_MAC_DINH, QUANG_CAO_TEXT
    global ZALO_AUTO_JOIN, TELE_AUTO_JOIN, ZALO_AUTO_REP_ENABLED, ZALO_AUTO_REP_ALL
    global ZALO_AUTO_REP_NGUOILA, ZALO_AUTO_REP_BANBE, ZALO_AUTO_REP_NEW
    global ZALO_GROUP_ADMINS, ZALO_BANBE_BROADCAST_ENABLED, ZALO_BANBE_BROADCAST_INTERVAL, ZALO_BANBE_BROADCAST_MSG
    global OPENAI_INALL_ZALO, OPENAI_INALL_TELE, OPENAI_INALL_ZALO_GROUPS, OPENAI_INALL_TELE_GROUPS, ADMIN_CHAT_NOTIFY, TELEGRAM_GROUP_SILENT_MODE, ZALO_SILENT_MODE, USERBOT_JOB_PROFILE_LIMIT, BUSINESS_PRO_ZALO_GROUPS, LIVE_CAPTURE_ZALO_GROUPS, ZALO_BLESS_ENABLED, HELP_IMAGE_URLS, ZALO_2FA_IMAGE_URL, TELE_REPLY_TEMPLATES, TELE_REPLY_MEDIA
    user_data_snapshot = {}
    try:
        user_data_snapshot = json.loads(get_config("user_data_json", "{}") or "{}")
        USER_DATA = {int(k): v for k, v in user_data_snapshot.items()}
    except Exception:
        USER_DATA = {}
        user_data_snapshot = {}
    extra_admins_snapshot = set()
    try:
        EXTRA_ADMINS.clear()
        for x in json.loads(get_config("extra_admins_json", "[]") or "[]"):
            uid = int(x)
            EXTRA_ADMINS.add(uid)
            extra_admins_snapshot.add(uid)
    except Exception:
        pass
    migrate_user_state_tables(user_data_snapshot, extra_admins_snapshot)
    load_user_state_tables()
    ZALO_ADMIN_UID = (get_config("zalo_admin_uid", ZALO_ADMIN_UID) or ZALO_ADMIN_UID).strip()
    ZALO_ADMIN_PHONE = (get_config("zalo_admin_phone", ZALO_ADMIN_PHONE) or ZALO_ADMIN_PHONE).strip()
    ZALO_VERIFY_CMD = (get_config("zalo_verify_cmd", ZALO_VERIFY_CMD) or ZALO_VERIFY_CMD).strip()
    ZALO_ADMIN_AUTO_FIND = get_config("zalo_admin_auto_find", "1" if ZALO_ADMIN_AUTO_FIND else "0") == "1"
    shared_supper_config = _read_shared_config(SUPPER_ADMIN_SHARED_CONFIG)
    SUPPER_ADMIN_GROUP_ID = (
        get_config("supper_admin_group_id", SUPPER_ADMIN_GROUP_ID or str(shared_supper_config.get("chat_id") or ""))
        or SUPPER_ADMIN_GROUP_ID
        or str(shared_supper_config.get("chat_id") or "")
    ).strip()
    SUPPER_ADMIN_GROUP_LINK = (
        get_config("supper_admin_group_link", SUPPER_ADMIN_GROUP_LINK or str(shared_supper_config.get("link") or ""))
        or SUPPER_ADMIN_GROUP_LINK
        or str(shared_supper_config.get("link") or "")
    ).strip()
    shared_supper_enabled = _read_shared_enabled(SUPPER_ADMIN_SHARED_CONFIG, SUPPER_ADMIN_ENABLED)
    SUPPER_ADMIN_ENABLED = _env_bool_override(
        "SUPPER_ADMIN_ENABLED",
        get_config("supper_admin_enabled", "1" if shared_supper_enabled else "0") == "1",
    )
    shared_admin_management_config = _read_shared_config(ADMIN_MANAGEMENT_SHARED_CONFIG)
    ADMIN_MANAGEMENT_GROUP_ID = (
        get_config("admin_management_group_id", ADMIN_MANAGEMENT_GROUP_ID or str(shared_admin_management_config.get("chat_id") or ""))
        or ADMIN_MANAGEMENT_GROUP_ID
        or str(shared_admin_management_config.get("chat_id") or "")
    ).strip()
    ADMIN_MANAGEMENT_GROUP_LINK = (
        get_config("admin_management_group_link", ADMIN_MANAGEMENT_GROUP_LINK or str(shared_admin_management_config.get("link") or ""))
        or ADMIN_MANAGEMENT_GROUP_LINK
        or str(shared_admin_management_config.get("link") or "")
    ).strip()
    shared_admin_management_enabled = _read_shared_enabled(ADMIN_MANAGEMENT_SHARED_CONFIG, ADMIN_MANAGEMENT_ENABLED)
    ADMIN_MANAGEMENT_ENABLED = _env_bool_override(
        "ADMIN_MANAGEMENT_ENABLED",
        get_config("admin_management_enabled", "1" if shared_admin_management_enabled else "0") == "1",
    )
    TARGET_CHAT = (get_config("target_chat", TARGET_CHAT) or TARGET_CHAT).strip()
    WEB_URL = (get_config("web_url", WEB_URL) or WEB_URL).strip()
    LINK_NHAP_CODE = (get_config("link_nhap_code", LINK_NHAP_CODE) or LINK_NHAP_CODE).strip()
    CAMPAIGN_DEFAULT_GROUP_LINK = (
        get_config("campaign_start_group_link", CAMPAIGN_DEFAULT_GROUP_LINK)
        or CAMPAIGN_DEFAULT_GROUP_LINK
        or TARGET_CHAT
    ).strip()
    CAMPAIGN_REGISTER_LINK = (get_config("campaign_register_link", CAMPAIGN_REGISTER_LINK) or CAMPAIGN_REGISTER_LINK or WEB_URL).strip()
    CAMPAIGN_SUPPORT_USERNAME = (
        get_config("campaign_support_username", CAMPAIGN_SUPPORT_USERNAME)
        or CAMPAIGN_SUPPORT_USERNAME
        or ADMIN_USERNAME
    ).strip().lstrip("@")
    CAMPAIGN_START_TEMPLATE = (get_config("campaign_start_template", CAMPAIGN_START_TEMPLATE) or CAMPAIGN_START_TEMPLATE).strip()
    CAMPAIGN_ALLOWED_GROUP_IDS = (get_config("campaign_allowed_group_ids", CAMPAIGN_ALLOWED_GROUP_IDS) or CAMPAIGN_ALLOWED_GROUP_IDS).strip()
    CAMPAIGN_ASSETS_MANIFEST = (get_config("campaign_assets_manifest", CAMPAIGN_ASSETS_MANIFEST) or CAMPAIGN_ASSETS_MANIFEST).strip()
    CAMPAIGN_CONSENT_REQUIRED = get_config("campaign_consent_required", "1" if CAMPAIGN_CONSENT_REQUIRED else "0") == "1"
    GIF_CHAO_HOI = (get_config("gif_chao_hoi", GIF_CHAO_HOI) or GIF_CHAO_HOI).strip()
    GIF_DANG_KY = (get_config("gif_dang_ky", GIF_DANG_KY) or GIF_DANG_KY).strip()
    GIF_EP_JOIN = (get_config("gif_ep_join", GIF_EP_JOIN) or GIF_EP_JOIN).strip()
    GIF_MOC_NAP = (get_config("gif_moc_nap", GIF_MOC_NAP) or GIF_MOC_NAP).strip()
    GIF_CHO_DUYET = (get_config("gif_cho_duyet", GIF_CHO_DUYET) or GIF_CHO_DUYET).strip()
    GIF_QUANG_CAO = (get_config("gif_quang_cao", GIF_QUANG_CAO) or GIF_QUANG_CAO).strip()
    USERBOT_GUITN_GIF = (get_config("userbot_guitn_gif", USERBOT_GUITN_GIF) or "").strip()
    ZALO_IMEI = (get_config("zalo_imei", ZALO_IMEI) or ZALO_IMEI).strip()
    COOKIE_RAW = (get_config("zalo_cookie_raw", COOKIE_RAW) or COOKIE_RAW).strip()
    ZALO_AUTO_REPLY = get_config("zalo_auto_reply_text", ZALO_AUTO_REPLY) or ZALO_AUTO_REPLY
    LOI_CHAO_MAC_DINH = get_config("telegram_welcome_text", LOI_CHAO_MAC_DINH) or LOI_CHAO_MAC_DINH
    QUANG_CAO_TEXT = get_config("telegram_qc_text", QUANG_CAO_TEXT) or QUANG_CAO_TEXT
    ZALO_ADMIN_REP_MODE = get_config("zalo_admin_rep_mode", "text") or "text"
    ZALO_ADMIN_REP_IMAGE = get_config("zalo_admin_rep_image", "")
    ZALO_AUTO_JOIN = get_config("zalo_auto_join", "0") == "1"
    TELE_AUTO_JOIN = get_config("tele_auto_join", "0") == "1"
    ZALO_AUTO_REP_ENABLED = get_config("zalo_auto_rep_enabled", "0") == "1"
    ZALO_AUTO_REP_ALL = get_config("zalo_auto_rep_all", "0") == "1"
    ZALO_AUTO_REP_NGUOILA = get_config("zalo_auto_rep_nguoila", "0") == "1"
    ZALO_AUTO_REP_BANBE = get_config("zalo_auto_rep_banbe", "0") == "1"
    ZALO_AUTO_REP_NEW = get_config("zalo_auto_rep_new", "0") == "1"
    try:
        ZALO_GROUP_ADMINS = set(json.loads(get_config("zalo_group_admins_json", "[]") or "[]"))
    except Exception:
        ZALO_GROUP_ADMINS = set()
    ZALO_BANBE_BROADCAST_ENABLED = get_config("zalo_banbe_broadcast_enabled", "0") == "1"
    ZALO_BANBE_BROADCAST_INTERVAL = int(get_config("zalo_banbe_broadcast_interval", "120") or "120")
    ZALO_BANBE_BROADCAST_MSG = get_config("zalo_banbe_broadcast_msg", "")
    OPENAI_INALL_ZALO = get_config("openai_inall_zalo", "0") == "1"
    OPENAI_INALL_TELE = get_config("openai_inall_tele", "0") == "1"
    try:
        OPENAI_INALL_ZALO_GROUPS = set(json.loads(get_config("openai_inall_zalo_groups_json", "[]") or "[]"))
    except Exception:
        OPENAI_INALL_ZALO_GROUPS = set()
    try:
        OPENAI_INALL_TELE_GROUPS = set(json.loads(get_config("openai_inall_tele_groups_json", "[]") or "[]"))
    except Exception:
        OPENAI_INALL_TELE_GROUPS = set()
    try:
        BUSINESS_PRO_ZALO_GROUPS = set(json.loads(get_config("business_pro_zalo_groups_json", "[]") or "[]"))
    except Exception:
        BUSINESS_PRO_ZALO_GROUPS = set()
    try:
        LIVE_CAPTURE_ZALO_GROUPS = set(json.loads(get_config("live_capture_zalo_groups_json", "[]") or "[]"))
    except Exception:
        LIVE_CAPTURE_ZALO_GROUPS = set()
    try:
        HELP_IMAGE_URLS = [str(x).strip() for x in json.loads(get_config("help_image_urls_json", "[]") or "[]") if str(x).strip()]
    except Exception:
        HELP_IMAGE_URLS = []
    ZALO_2FA_IMAGE_URL = (get_config("zalo_2fa_image_url", "") or "").strip()
    ZALO_BLESS_ENABLED = get_config("zalo_bless_enabled", "1") == "1"
    TELEGRAM_GROUP_SILENT_MODE = _env_bool_override(
        "TELEGRAM_GROUP_SILENT_MODE",
        get_config("telegram_group_silent_mode", "1" if TELEGRAM_GROUP_SILENT_MODE else "0") == "1",
    )
    ZALO_SILENT_MODE = _env_bool_override(
        "ZALO_SILENT_MODE",
        get_config("zalo_silent_mode", "1" if ZALO_SILENT_MODE else "0") == "1",
    )
    try:
        env_profile_limit = os.getenv("USERBOT_JOB_PROFILE_LIMIT", "").strip()
        raw_profile_limit = env_profile_limit or get_config("userbot_job_profile_limit", str(USERBOT_JOB_PROFILE_LIMIT)) or "1"
        USERBOT_JOB_PROFILE_LIMIT = max(1, int(raw_profile_limit))
    except Exception:
        USERBOT_JOB_PROFILE_LIMIT = 1
    ADMIN_CHAT_NOTIFY = get_config("admin_chat_notify", "0") == "1"
    if ZALO_SILENT_MODE:
        ADMIN_CHAT_NOTIFY = False
        ZALO_AUTO_JOIN = False
        ZALO_AUTO_REP_ENABLED = False
        ZALO_AUTO_REP_ALL = False
        ZALO_AUTO_REP_NGUOILA = False
        ZALO_AUTO_REP_BANBE = False
        ZALO_AUTO_REP_NEW = False
        ZALO_BANBE_BROADCAST_ENABLED = False
        OPENAI_INALL_ZALO = False
    TELE_REPLY_TEMPLATES = dict(TELE_REPLY_TEMPLATES_DEFAULT)
    try:
        raw_templates = json.loads(get_config("tele_reply_templates_json", "{}") or "{}")
        for k in TELE_REPLY_TEMPLATES_DEFAULT:
            if k in raw_templates and str(raw_templates[k]).strip() != "":
                TELE_REPLY_TEMPLATES[k] = str(raw_templates[k])
    except Exception:
        pass
    try:
        raw_media = json.loads(get_config("tele_reply_media_json", "{}") or "{}")
        TELE_REPLY_MEDIA = {
            str(k): str(v).strip()
            for k, v in raw_media.items()
            if str(k).strip() and str(v).strip()
        }
    except Exception:
        TELE_REPLY_MEDIA = {}
    write_supper_admin_shared_config()
    write_admin_management_shared_config()


def persist_runtime_state():
    persist_user_state_tables()
    refresh_group_enabled_from_shared()
    set_config("user_data_json", json.dumps(USER_DATA, ensure_ascii=False))
    set_config("extra_admins_json", json.dumps(sorted(list(EXTRA_ADMINS))))
    set_config("zalo_admin_uid", str(ZALO_ADMIN_UID))
    set_config("zalo_admin_phone", str(ZALO_ADMIN_PHONE))
    set_config("zalo_verify_cmd", str(ZALO_VERIFY_CMD))
    set_config("zalo_admin_auto_find", "1" if ZALO_ADMIN_AUTO_FIND else "0")
    set_config("supper_admin_group_id", SUPPER_ADMIN_GROUP_ID or "")
    set_config("supper_admin_group_link", SUPPER_ADMIN_GROUP_LINK or "")
    set_config("supper_admin_enabled", "1" if SUPPER_ADMIN_ENABLED else "0")
    set_config("admin_management_group_id", ADMIN_MANAGEMENT_GROUP_ID or "")
    set_config("admin_management_group_link", ADMIN_MANAGEMENT_GROUP_LINK or "")
    set_config("admin_management_enabled", "1" if ADMIN_MANAGEMENT_ENABLED else "0")
    set_config("target_chat", TARGET_CHAT or "")
    set_config("web_url", WEB_URL or "")
    set_config("link_nhap_code", LINK_NHAP_CODE or "")
    set_config("campaign_start_group_link", CAMPAIGN_DEFAULT_GROUP_LINK or TARGET_CHAT or "")
    set_config("campaign_register_link", CAMPAIGN_REGISTER_LINK or WEB_URL or "")
    set_config("campaign_support_username", CAMPAIGN_SUPPORT_USERNAME or ADMIN_USERNAME.replace("@", ""))
    set_config("campaign_start_template", CAMPAIGN_START_TEMPLATE or CAMPAIGN_START_TEMPLATE_DEFAULT)
    set_config("campaign_allowed_group_ids", CAMPAIGN_ALLOWED_GROUP_IDS or "")
    set_config("campaign_assets_manifest", CAMPAIGN_ASSETS_MANIFEST or "")
    set_config("campaign_consent_required", "1" if CAMPAIGN_CONSENT_REQUIRED else "0")
    set_config("gif_chao_hoi", GIF_CHAO_HOI or "")
    set_config("gif_dang_ky", GIF_DANG_KY or "")
    set_config("gif_ep_join", GIF_EP_JOIN or "")
    set_config("gif_moc_nap", GIF_MOC_NAP or "")
    set_config("gif_cho_duyet", GIF_CHO_DUYET or "")
    set_config("gif_quang_cao", GIF_QUANG_CAO or "")
    set_config("userbot_guitn_gif", USERBOT_GUITN_GIF or "")
    write_supper_admin_shared_config()
    write_admin_management_shared_config()
    if ZALO_IMEI:
        set_config("zalo_imei", ZALO_IMEI)
    if COOKIE_RAW:
        set_config("zalo_cookie_raw", COOKIE_RAW)
    set_config("zalo_auto_reply_text", ZALO_AUTO_REPLY or "")
    set_config("telegram_welcome_text", LOI_CHAO_MAC_DINH or "")
    set_config("telegram_qc_text", QUANG_CAO_TEXT or "")
    set_config("zalo_admin_rep_mode", ZALO_ADMIN_REP_MODE)
    set_config("zalo_admin_rep_image", ZALO_ADMIN_REP_IMAGE)
    set_config("zalo_auto_join", "1" if ZALO_AUTO_JOIN else "0")
    set_config("tele_auto_join", "1" if TELE_AUTO_JOIN else "0")
    set_config("zalo_auto_rep_enabled", "1" if ZALO_AUTO_REP_ENABLED else "0")
    set_config("zalo_auto_rep_all", "1" if ZALO_AUTO_REP_ALL else "0")
    set_config("zalo_auto_rep_nguoila", "1" if ZALO_AUTO_REP_NGUOILA else "0")
    set_config("zalo_auto_rep_banbe", "1" if ZALO_AUTO_REP_BANBE else "0")
    set_config("zalo_auto_rep_new", "1" if ZALO_AUTO_REP_NEW else "0")
    set_config("zalo_group_admins_json", json.dumps(sorted(list(ZALO_GROUP_ADMINS))))
    set_config("zalo_banbe_broadcast_enabled", "1" if ZALO_BANBE_BROADCAST_ENABLED else "0")
    set_config("zalo_banbe_broadcast_interval", str(max(ZALO_BANBE_MIN_DELAY_SEC, int(ZALO_BANBE_BROADCAST_INTERVAL))))
    set_config("zalo_banbe_broadcast_msg", ZALO_BANBE_BROADCAST_MSG or "")
    set_config("openai_inall_zalo", "1" if OPENAI_INALL_ZALO else "0")
    set_config("openai_inall_tele", "1" if OPENAI_INALL_TELE else "0")
    set_config("openai_inall_zalo_groups_json", json.dumps(sorted(list(OPENAI_INALL_ZALO_GROUPS))))
    set_config("openai_inall_tele_groups_json", json.dumps(sorted(list(OPENAI_INALL_TELE_GROUPS))))
    set_config("business_pro_zalo_groups_json", json.dumps(sorted(list(BUSINESS_PRO_ZALO_GROUPS))))
    set_config("live_capture_zalo_groups_json", json.dumps(sorted(list(LIVE_CAPTURE_ZALO_GROUPS))))
    set_config("help_image_urls_json", json.dumps(HELP_IMAGE_URLS, ensure_ascii=False))
    set_config("zalo_2fa_image_url", ZALO_2FA_IMAGE_URL or "")
    set_config("admin_chat_notify", "1" if ADMIN_CHAT_NOTIFY else "0")
    set_config("telegram_group_silent_mode", "1" if TELEGRAM_GROUP_SILENT_MODE else "0")
    set_config("zalo_silent_mode", "1" if ZALO_SILENT_MODE else "0")
    set_config("userbot_job_profile_limit", str(max(1, int(USERBOT_JOB_PROFILE_LIMIT or 1))))
    set_config("zalo_bless_enabled", "1" if ZALO_BLESS_ENABLED else "0")
    set_config("tele_reply_templates_json", json.dumps(TELE_REPLY_TEMPLATES, ensure_ascii=False))
    set_config("tele_reply_media_json", json.dumps(TELE_REPLY_MEDIA, ensure_ascii=False))


def upsert_contact(platform: str, uid: str, name: str = "", username: str = "", phone: str = "", last_message: str = ""):
    with db_conn() as conn:
        conn.execute(
            """
            INSERT INTO customer_contacts(uid, platform, name, username, phone, last_message, updated_at)
            VALUES(?,?,?,?,?,?,?)
            ON CONFLICT(uid) DO UPDATE SET
                platform=excluded.platform,
                name=CASE WHEN excluded.name!='' THEN excluded.name ELSE customer_contacts.name END,
                username=CASE WHEN excluded.username!='' THEN excluded.username ELSE customer_contacts.username END,
                phone=CASE WHEN excluded.phone!='' THEN excluded.phone ELSE customer_contacts.phone END,
                last_message=CASE WHEN excluded.last_message!='' THEN excluded.last_message ELSE customer_contacts.last_message END,
                updated_at=excluded.updated_at
            """,
            (str(uid), platform, name[:150], username[:150], phone[:50], (last_message or "")[:1000], time.time()),
        )
        conn.commit()


def is_known_zalo_group_uid(uid: str) -> bool:
    target = str(uid or "").strip()
    if not target:
        return False
    with db_conn() as conn:
        row = conn.execute(
            "SELECT 1 FROM customer_contacts WHERE uid=? AND platform='zalo_group' LIMIT 1",
            (target,),
        ).fetchone()
    return bool(row)


def telegram_display_name(user) -> str:
    if not user:
        return ""
    first = str(getattr(user, "first_name", "") or "").strip()
    last = str(getattr(user, "last_name", "") or "").strip()
    username = str(getattr(user, "username", "") or "").strip()
    name = " ".join(x for x in (first, last) if x).strip()
    return name or username or str(getattr(user, "id", "") or "")


def cache_telegram_group_member(update: Update):
    user = getattr(update, "effective_user", None)
    chat = getattr(update, "effective_chat", None)
    if not user or not chat:
        return
    chat_type = str(getattr(chat, "type", "") or "").lower()
    if chat_type not in ("group", "supergroup"):
        return
    chat_id = str(getattr(chat, "id", "") or "")
    user_id = str(getattr(user, "id", "") or "")
    if not chat_id or not user_id:
        return
    username = str(getattr(user, "username", "") or "").strip()
    name = telegram_display_name(user)
    chat_title = str(getattr(chat, "title", "") or "").strip()
    now_ts = time.time()
    upsert_contact("telegram_user", user_id, name=name, username=username)
    upsert_contact("telegram_group", chat_id, name=chat_title)
    with db_conn() as conn:
        conn.execute(
            """
            INSERT INTO telegram_group_members(chat_id, user_id, name, username, last_seen, hit_count)
            VALUES(?,?,?,?,?,1)
            ON CONFLICT(chat_id, user_id) DO UPDATE SET
                name=CASE WHEN excluded.name!='' THEN excluded.name ELSE telegram_group_members.name END,
                username=CASE WHEN excluded.username!='' THEN excluded.username ELSE telegram_group_members.username END,
                last_seen=excluded.last_seen,
                hit_count=telegram_group_members.hit_count + 1
            """,
            (chat_id, user_id, name[:150], username[:150], now_ts),
        )
        conn.commit()


def list_cached_group_members(chat_id: str, limit: int = 5000):
    with db_conn() as conn:
        return conn.execute(
            """
            SELECT chat_id, user_id, name, username, last_seen, hit_count
            FROM telegram_group_members
            WHERE chat_id=?
            ORDER BY lower(COALESCE(NULLIF(name,''), username, user_id))
            LIMIT ?
            """,
            (str(chat_id), int(limit)),
        ).fetchall()


def list_cached_telegram_groups(limit: int = 500):
    with db_conn() as conn:
        return conn.execute(
            """
            SELECT uid AS chat_id, name AS chat_title, updated_at
            FROM customer_contacts
            WHERE platform='telegram_group'
            ORDER BY updated_at DESC
            LIMIT ?
            """,
            (int(limit),),
        ).fetchall()


def list_userbot_profiles(limit: int = 50, root: str | None = None) -> list[str]:
    accounts_root = root or USERBOT_ACCOUNTS_ROOT
    if not os.path.isdir(accounts_root):
        return []
    cooldowns = load_userbot_account_cooldowns()
    profiles = []
    for name in sorted(os.listdir(accounts_root)):
        path = os.path.join(accounts_root, name)
        if os.path.isdir(path) and not userbot_account_is_cooldown(name, cooldowns):
            profiles.append(name)
        if len(profiles) >= limit:
            break
    return profiles


def load_userbot_account_cooldowns() -> dict:
    try:
        with open(USERBOT_ACCOUNT_COOLDOWN_JSON, "r", encoding="utf-8-sig") as f:
            data = json.load(f)
        accounts = data.get("accounts") if isinstance(data, dict) else {}
        return accounts if isinstance(accounts, dict) else {}
    except Exception:
        return {}


def userbot_account_is_cooldown(phone: str, cooldowns: dict | None = None) -> bool:
    cooldowns = cooldowns if isinstance(cooldowns, dict) else load_userbot_account_cooldowns()
    entry = cooldowns.get(str(phone or "").strip()) or {}
    if not isinstance(entry, dict):
        return False
    try:
        until = float(entry.get("until") or 0)
    except Exception:
        until = 0
    if until > time.time():
        return True
    if until <= 0:
        status = str(entry.get("status") or "").strip().lower()
        reason = str(entry.get("reason") or "").strip().lower()
        permanent_markers = (
            "disabled",
            "relogin",
            "missing_session",
            "missing_tdata",
            "missing_key",
            "unauthorized",
            "authkey",
            "deactivated",
            "peer_flood",
            "flood",
        )
        return any(marker in status or marker in reason for marker in permanent_markers)
    return False


def _split_path_list(raw: str) -> list[str]:
    parts = re.split(r"[;\n\r]+", str(raw or ""))
    roots = []
    seen = set()
    for part in parts:
        p = part.strip().strip('"')
        if not p:
            continue
        try:
            p = os.path.abspath(p)
        except Exception:
            pass
        key = os.path.normcase(os.path.normpath(p))
        if key in seen:
            continue
        seen.add(key)
        roots.append(p)
    return roots


def get_userbot_shard_roots() -> list[str]:
    roots = []
    for root in _split_path_list(os.getenv("USERBOT_SHARD_ROOTS", USERBOT_SHARD_ROOTS_RAW)):
        if os.path.isdir(root):
            roots.append(root)
    return roots


def get_userbot_job_roots(default_root: str | None = None) -> list[str]:
    shard_roots = get_userbot_shard_roots()
    if shard_roots:
        return shard_roots
    root = default_root or USERBOT_ACCOUNTS_ROOT
    return [root] if root else []


def _existing_userbot_dir(path: str | None) -> str:
    path = str(path or "").strip().strip('"')
    if not path:
        return ""
    try:
        path = os.path.abspath(path)
    except Exception:
        pass
    return path if os.path.isdir(path) else ""


def _dedupe_existing_userbot_dirs(paths: list[str]) -> list[str]:
    roots = []
    seen = set()
    for path in paths or []:
        root = _existing_userbot_dir(path)
        if not root:
            continue
        key = os.path.normcase(os.path.normpath(root))
        if key in seen:
            continue
        seen.add(key)
        roots.append(root)
    return roots


def _first_existing_userbot_dir(*paths: str | None) -> str:
    for path in paths:
        root = _existing_userbot_dir(path)
        if root:
            return root
    return ""


def get_userbot_12_lo_root() -> str:
    return _first_existing_userbot_dir(
        USERBOT_12_LO_ROOT,
        os.path.join(USERBOT_52_ROOT, "10ac xinhan 265 dm"),
    )


def get_userbot_20_xinhan_root() -> str:
    return _first_existing_userbot_dir(
        USERBOT_20_XINHAN_ROOT,
        USERBOT_20_ACCOUNTS_ROOT,
        os.path.join(USERBOT_52_ROOT, "20ac xinhan 2 6"),
    )


def get_userbot_20_new_root() -> str:
    return _first_existing_userbot_dir(
        USERBOT_20_NEW_ROOT,
        os.path.join(USERBOT_52_ROOT, "20ac_moi_20260603"),
    )


def get_userbot_40_lot_roots() -> list[str]:
    return _dedupe_existing_userbot_dirs(
        [
            get_userbot_20_xinhan_root(),
            get_userbot_20_new_root(),
        ]
    )


def get_userbot_52_lot_roots() -> list[str]:
    roots = _dedupe_existing_userbot_dirs(
        [
            get_userbot_12_lo_root(),
            get_userbot_20_xinhan_root(),
            get_userbot_20_new_root(),
        ]
    )
    if roots:
        return roots
    return _dedupe_existing_userbot_dirs([USERBOT_52_ROOT])


def split_rows_by_roots(rows: list, roots: list[str]) -> list[list]:
    if not roots:
        return []
    buckets = [[] for _ in roots]
    for idx, row in enumerate(rows or []):
        buckets[idx % len(roots)].append(row)
    return buckets


def get_userbot_job_profile_limit() -> int:
    try:
        limit = max(1, int(USERBOT_JOB_PROFILE_LIMIT or 1))
    except Exception:
        limit = 1
    try:
        if is_baokybcr_bot_instance() and not get_userbot_shard_roots():
            return max(20, limit)
    except Exception:
        pass
    return limit


def summarize_userbot_roots(limit: int | None = None) -> tuple[list[dict], int]:
    try:
        profile_limit = max(1, int(limit or get_userbot_job_profile_limit() or 1))
    except Exception:
        profile_limit = 1
    stats = []
    total_profiles = 0
    for idx, root in enumerate(get_userbot_job_roots(USERBOT_ACCOUNTS_ROOT), start=1):
        profiles = list_userbot_profiles(limit=profile_limit, root=root)
        total_profiles += len(profiles)
        stats.append(
            {
                "index": idx,
                "root": root,
                "name": os.path.basename(os.path.normpath(root)) or f"shard_{idx}",
                "profiles": profiles,
                "profile_count": len(profiles),
            }
        )
    return stats, total_profiles


def active_userbot_root_profiles(roots: list[str], profile_limit: int | None = None) -> list[tuple[str, list[str]]]:
    try:
        limit = max(1, int(profile_limit or get_userbot_job_profile_limit() or 1))
    except Exception:
        limit = 1
    active = []
    for root in roots or []:
        if not root or not os.path.isdir(root):
            continue
        profiles = list_userbot_profiles(limit=limit, root=root)
        if profiles:
            active.append((root, profiles))
    return active


def attach_userbot_report_origin(detail: dict, update: Update) -> dict:
    detail = dict(detail or {})
    chat = getattr(update, "effective_chat", None)
    user = getattr(update, "effective_user", None)
    detail.setdefault("origin_bot_slot", os.getenv("BOT_INSTANCE_SLOT", "").strip().lower() or "default")
    detail.setdefault("origin_chat_id", str(getattr(chat, "id", "") or ""))
    detail.setdefault("origin_chat_title", str(getattr(chat, "title", "") or getattr(chat, "first_name", "") or ""))
    detail.setdefault("origin_chat_type", str(getattr(chat, "type", "") or ""))
    detail.setdefault("origin_user_id", str(getattr(user, "id", "") or ""))
    return detail


def queue_userbot_scan_job(command: str, chat_id: str, chat_title: str, requested_by: str, dry_run: bool, detail: dict) -> int:
    now_ts = time.time()
    detail = dict(detail or {})
    detail.setdefault("origin_bot_slot", os.getenv("BOT_INSTANCE_SLOT", "").strip().lower() or "default")
    detail_text = json.dumps(detail, ensure_ascii=False)
    with db_conn() as conn:
        cur = conn.execute(
            """
            INSERT INTO userbot_scan_jobs(command, chat_id, chat_title, requested_by, dry_run, status, detail, created_at, updated_at)
            VALUES(?,?,?,?,?,?,?,?,?)
            """,
            (command, str(chat_id), str(chat_title), str(requested_by), 1 if dry_run else 0, "dry_run" if dry_run else "queued", detail_text, now_ts, now_ts),
        )
        conn.commit()
        job_id = int(cur.lastrowid)
    try:
        os.makedirs(USERBOT_SCAN_JOBS_DIR, exist_ok=True)
        job_path = os.path.join(USERBOT_SCAN_JOBS_DIR, "pending.jsonl")
        slot_name = os.getenv("BOT_INSTANCE_SLOT", "default").strip() or "default"
        queue_id = f"{slot_name}-{int(now_ts * 1000)}-{job_id}"
        payload = {
            "id": queue_id,
            "db_id": job_id,
            "command": command,
            "chat_id": str(chat_id),
            "chat_title": str(chat_title),
            "requested_by": str(requested_by),
            "dry_run": bool(dry_run),
            "detail": detail or {},
            "created_at": now_ts,
        }
        with open(job_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(payload, ensure_ascii=False) + "\n")
    except Exception as exc:
        logging.warning("Ghi queue userbot_scan_jobs that bai: %s", exc)
    return job_id


def _clean_import_username(raw: str) -> str:
    return str(raw or "").strip().lstrip("@")


def _fold_contact_import_key(raw: str) -> str:
    text = str(raw or "").strip().lstrip("\ufeff").lower().replace("đ", "d")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    return text


def _row_get_any(row: dict, *keys: str) -> str:
    if not isinstance(row, dict):
        return ""
    wanted = {_fold_contact_import_key(k) for k in keys}
    for key, value in row.items():
        if _fold_contact_import_key(key) in wanted:
            return str(value or "").strip()
    return ""


def _normalize_contact_import_row(row: dict) -> dict:
    return {
        "job_id": _row_get_any(row, "job_id", "job"),
        "account": _row_get_any(row, "account", "acc", "profile", "session", "tai_khoan"),
        "phone": normalize_vn_phone(
            _row_get_any(row, "phone", "sdt", "so_dien_thoai", "dien_thoai", "mobile", "tel")
        ),
        "uid": _row_get_any(row, "uid", "user_id", "userid", "telegram_id", "id"),
        "name": _row_get_any(row, "name", "ten", "tên", "full_name", "ho_ten", "display_name"),
        "username": _clean_import_username(
            _row_get_any(row, "username", "user", "ten_user", "telegram_user", "nick", "nickname")
        ),
        "type": _row_get_any(row, "type", "loai"),
        "status": _row_get_any(row, "status", "trang_thai"),
        "source": _row_get_any(row, "source", "nguon", "nguồn"),
        "consent": normalize_consent_status(_row_get_any(row, "consent", "consent_status", "opt_in", "dong_y", "đồng_ý")),
        "platform": _row_get_any(row, "platform", "nen_tang"),
        "link": _row_get_any(row, "link", "url"),
    }


def _dedupe_contact_import_rows(rows: list[dict], valid_only: bool = True) -> list[dict]:
    seen = set()
    out = []
    for row in rows:
        item = _normalize_contact_import_row(row)
        if valid_only and not (item["username"] or item["uid"]):
            continue
        if not (item["uid"] or item["username"] or item["phone"]):
            continue
        key = (item.get("uid") or "", item.get("username") or "", item.get("phone") or "")
        if key in seen:
            continue
        seen.add(key)
        out.append(item)
    return out


def _read_contact_import_csv(path: str, valid_only: bool = True) -> list[dict]:
    if not os.path.exists(path):
        return []
    rows = []
    for enc in ("utf-8-sig", "utf-8", "cp1258", "cp1252", "latin1"):
        try:
            with open(path, "r", encoding=enc, newline="") as f:
                for row in csv.DictReader(f):
                    rows.append(row)
            break
        except UnicodeDecodeError:
            rows = []
            continue
        except Exception:
            return []
    return _dedupe_contact_import_rows(rows, valid_only)


_CONTACT_IMPORT_HEADER_KEYS = {
    _fold_contact_import_key(k)
    for k in (
        "job_id",
        "account",
        "acc",
        "profile",
        "phone",
        "sdt",
        "so_dien_thoai",
        "uid",
        "user_id",
        "userid",
        "telegram_id",
        "id",
        "name",
        "ten",
        "tên",
        "full_name",
        "username",
        "user",
        "ten_user",
        "telegram_user",
        "source",
        "nguon",
        "consent",
        "consent_status",
        "opt_in",
        "dong_y",
        "đồng_ý",
    )
}


def _contact_import_has_header(fields: list[str]) -> bool:
    folded = {_fold_contact_import_key(field) for field in fields}
    return bool(folded & _CONTACT_IMPORT_HEADER_KEYS)


def _read_contact_import_delimited_text(text: str, delimiter: str, valid_only: bool = True) -> list[dict]:
    raw_lines = []
    for line in str(text or "").splitlines():
        clean = line.strip()
        if not clean:
            continue
        low = clean.lower()
        if low.startswith("--- sheet:") or low.startswith("total="):
            continue
        raw_lines.append(line)
    if not raw_lines:
        return []
    for idx, line in enumerate(raw_lines[:80]):
        try:
            fields = next(csv.reader([line], delimiter=delimiter))
        except Exception:
            continue
        if not _contact_import_has_header(fields):
            continue
        try:
            reader = csv.DictReader(io.StringIO("\n".join(raw_lines[idx:])), delimiter=delimiter)
            return _dedupe_contact_import_rows(list(reader), valid_only)
        except Exception:
            return []
    return []


def _read_contact_import_text(text: str, valid_only: bool = True) -> list[dict]:
    for delimiter in (",", "\t", ";", "|"):
        rows = _read_contact_import_delimited_text(text, delimiter, valid_only)
        if rows:
            return rows
    return _dedupe_contact_import_rows(parse_guitn_recipients_text(text), valid_only)


def _read_contact_import_text_file(path: str, valid_only: bool = True) -> list[dict]:
    if not os.path.exists(path):
        return []
    for enc in ("utf-8-sig", "utf-8", "cp1258", "cp1252", "latin1"):
        try:
            with open(path, "r", encoding=enc, errors="strict") as f:
                return _read_contact_import_text(f.read(), valid_only)
        except UnicodeDecodeError:
            continue
        except Exception:
            return []
    return []


def _read_latest_contact_import_export(valid_only: bool = True) -> tuple[list[dict], str]:
    if not os.path.isdir(USERBOT_SCAN_JOBS_DIR):
        return [], USERBOT_SCAN_JOBS_DIR
    candidates = []
    prefixes = ("quetall_now_", "quetall_now_job_", "uid_now_", "user_now_", "tt_now_", "sdt_now_", "sdtnow_", "scan_entities_map_")
    for name in os.listdir(USERBOT_SCAN_JOBS_DIR):
        low = name.lower()
        if not low.endswith((".txt", ".csv")):
            continue
        if not low.startswith(prefixes):
            continue
        path = os.path.join(USERBOT_SCAN_JOBS_DIR, name)
        try:
            candidates.append((os.path.getmtime(path), path))
        except OSError:
            continue
    for _, path in sorted(candidates, reverse=True)[:40]:
        rows = _read_contact_import_text_file(path, valid_only)
        if rows:
            return rows, path
    return [], CONTACT_IMPORT_LAST_USERS_CSV


def _read_best_contact_import_jsonl(valid_only: bool = True) -> tuple[list[dict], str]:
    if not os.path.exists(CONTACT_IMPORT_USERS_JSONL):
        return [], CONTACT_IMPORT_USERS_JSONL
    grouped: dict[str, list[dict]] = {}
    try:
        with open(CONTACT_IMPORT_USERS_JSONL, "r", encoding="utf-8-sig") as f:
            for line in f:
                try:
                    row = json.loads(line)
                except Exception:
                    continue
                job_id = str(row.get("job_id") or "unknown").strip() or "unknown"
                grouped.setdefault(job_id, []).append(row)
    except Exception:
        return [], CONTACT_IMPORT_USERS_JSONL
    best_rows: list[dict] = []
    best_job = ""
    for job_id, raw_rows in grouped.items():
        rows = _dedupe_contact_import_rows(raw_rows, valid_only)
        if len(rows) > len(best_rows):
            best_rows = rows
            best_job = job_id
    source = CONTACT_IMPORT_USERS_JSONL + (f"#job={best_job}" if best_job else "")
    return best_rows, source


def load_contact_import_users(valid_only: bool = True) -> tuple[list[dict], str]:
    last_rows = _read_contact_import_csv(CONTACT_IMPORT_LAST_USERS_CSV, valid_only)
    best_rows, best_source = _read_best_contact_import_jsonl(valid_only)
    export_rows, export_source = _read_latest_contact_import_export(valid_only)
    candidates = [
        (last_rows, CONTACT_IMPORT_LAST_USERS_CSV),
        (best_rows, best_source),
        (export_rows, export_source),
    ]
    rows, source = max(candidates, key=lambda item: len(item[0]))
    return rows, source


def _contact_import_row_ok_for_mode(row: dict, mode: str) -> bool:
    username = _clean_import_username(row.get("username") or "")
    uid = str(row.get("uid") or "").strip()
    phone = normalize_vn_phone(row.get("phone") or "")
    status = str(row.get("status") or "").strip().lower()
    platform = str(row.get("platform") or "").strip().lower()
    row_type = str(row.get("type") or "").strip().lower()
    raw_unverified_uid = status in {"raw_uid", "raw"} or (row_type == "uid" and not platform and not username)
    verified_uid = bool(uid and uid.isdigit() and not raw_unverified_uid)
    if mode == "user":
        return bool(username)
    if mode == "uid":
        return bool(verified_uid or username)
    if mode == "sdt":
        return bool(phone or username)
    if mode == "tt":
        return bool(phone or uid or username or str(row.get("name") or "").strip())
    if mode == "all":
        return bool(verified_uid or username)
    return bool(username)


def _filter_contact_import_rows_for_mode(rows: list[dict], mode: str) -> list[dict]:
    return [row for row in rows if _contact_import_row_ok_for_mode(row, mode)]


def _contact_import_condition_text(mode: str) -> str:
    if mode == "user":
        return "Điều kiện hợp lệ: dòng phải có username/user."
    if mode == "uid":
        return "Điều kiện hợp lệ: dòng phải có UID hoặc username/user."
    if mode == "sdt":
        return "Điều kiện hợp lệ: dòng phải có SĐT hoặc username/user."
    if mode == "tt":
        return "Điều kiện hợp lệ: dòng phải có SĐT, UID hoặc username/user."
    if mode == "all":
        return "Điều kiện hợp lệ: dòng phải có UID hoặc username/user; SĐT chỉ là cột phụ."
    return "Điều kiện hợp lệ: dòng phải có dữ liệu khách."


def load_contact_import_users_for_mode(mode: str) -> tuple[list[dict], str]:
    candidates = []
    for rows, source in (
        (_read_contact_import_csv(CONTACT_IMPORT_LAST_USERS_CSV, valid_only=False), CONTACT_IMPORT_LAST_USERS_CSV),
        _read_best_contact_import_jsonl(valid_only=False),
        _read_latest_contact_import_export(valid_only=False),
    ):
        candidates.append((_filter_contact_import_rows_for_mode(rows, mode), source))
    rows, source = max(candidates, key=lambda item: len(item[0]))
    return rows, source


def _read_json_file_safe(path: str, default):
    try:
        with open(path, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    except Exception:
        return default


def _env_bool_override(name: str, current: bool) -> bool:
    raw = os.getenv(name)
    if raw is None or str(raw).strip() == "":
        return current
    return str(raw).strip().lower() not in {"0", "false", "no", "off"}


def _read_contact_import_pending_jobs() -> list[dict]:
    jobs = []
    if not os.path.exists(CONTACT_IMPORT_PENDING_JSONL):
        return jobs
    try:
        with open(CONTACT_IMPORT_PENDING_JSONL, "r", encoding="utf-8-sig") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    job = json.loads(line)
                except Exception:
                    continue
                detail = job.get("detail") or {}
                if str(job.get("command") or "") == "phat500k_import_contacts" or detail.get("job_type") == "import_contacts":
                    jobs.append(job)
    except Exception:
        pass
    return jobs


def _read_userbot_queue_jobs() -> list[dict]:
    jobs = []
    if not os.path.exists(CONTACT_IMPORT_PENDING_JSONL):
        return jobs
    try:
        with open(CONTACT_IMPORT_PENDING_JSONL, "r", encoding="utf-8-sig") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    job = json.loads(line)
                except Exception:
                    continue
                if isinstance(job, dict):
                    jobs.append(job)
    except Exception:
        pass
    return jobs


def _is_supported_userbot_queue_job(job: dict) -> bool:
    detail = job.get("detail") or {}
    command = str(job.get("command") or "")
    job_type = str(detail.get("job_type") or "")
    return command in {
        "phat500k_import_contacts",
        "import_contacts",
        "userbot_send_message",
        "guitn",
        "keoall",
        "invite_to_group",
    } or job_type in {"import_contacts", "send_message", "invite_to_group"}


def _userbot_queue_job_kind(job: dict) -> str:
    detail = job.get("detail") or {}
    command = str(job.get("command") or "")
    job_type = str(detail.get("job_type") or "")
    if command in {"userbot_send_message", "guitn"} or job_type == "send_message":
        return "gửi tin"
    if command in {"keoall", "invite_to_group"} or job_type == "invite_to_group":
        return "kéo nhóm"
    if command in {"phat500k_import_contacts", "import_contacts"} or job_type == "import_contacts":
        return "quét SĐT"
    return "job"


def _write_json_file_safe(path: str, data) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _terminate_userbot_worker_processes() -> dict:
    script = r"""
$items = Get-CimInstance Win32_Process | Where-Object {
  $_.CommandLine -match 'userbot_import_contacts\.py'
}
foreach ($p in $items) {
  try {
    Stop-Process -Id $p.ProcessId -Force -ErrorAction Stop
    "stopped:$($p.ProcessId)"
  } catch {
    "failed:$($p.ProcessId):$($_.Exception.Message)"
  }
}
"""
    try:
        proc = subprocess.run(
            ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
            capture_output=True,
            text=True,
            timeout=20,
        )
        out = (proc.stdout or "") + ("\n" + proc.stderr if proc.stderr else "")
        stopped = re.findall(r"stopped:(\d+)", out)
        failed = re.findall(r"failed:(\d+):([^\r\n]+)", out)
        return {"stopped": stopped, "failed": failed, "raw": out.strip(), "returncode": proc.returncode}
    except Exception as exc:
        return {"stopped": [], "failed": [], "raw": f"{type(exc).__name__}: {exc}", "returncode": -1}


def stop_all_userbot_jobs(requested_by: str, stop_workers: bool = True) -> dict:
    stop_at = time.time()
    jobs = [job for job in _read_userbot_queue_jobs() if _is_supported_userbot_queue_job(job)]
    processed_raw = _read_json_file_safe(CONTACT_IMPORT_PROCESSED_JSON, [])
    processed = {str(x) for x in processed_raw} if isinstance(processed_raw, list) else set()
    claimed_raw = _read_json_file_safe(CONTACT_IMPORT_CLAIMED_JSON, {})
    claimed = {str(k): float(v or 0) for k, v in claimed_raw.items()} if isinstance(claimed_raw, dict) else {}

    target_jobs = []
    for job in jobs:
        jid = str(job.get("id") or "")
        if not jid or jid in processed:
            continue
        try:
            created_at = float(job.get("created_at") or 0)
        except Exception:
            created_at = 0.0
        if created_at <= stop_at:
            target_jobs.append(job)

    target_ids = {str(job.get("id") or "") for job in target_jobs if str(job.get("id") or "")}
    active_ids = sorted(jid for jid in target_ids if jid in claimed)
    waiting_ids = sorted(jid for jid in target_ids if jid not in claimed)
    processed.update(target_ids)
    for jid in target_ids:
        claimed.pop(jid, None)
    _write_json_file_safe(CONTACT_IMPORT_PROCESSED_JSON, sorted(processed))
    _write_json_file_safe(CONTACT_IMPORT_CLAIMED_JSON, claimed)

    db_ids = []
    for job in target_jobs:
        try:
            db_ids.append(int(job.get("db_id") or 0))
        except Exception:
            continue
    if db_ids:
        try:
            with db_conn() as conn:
                for db_id in db_ids:
                    if db_id > 0:
                        conn.execute(
                            "UPDATE userbot_scan_jobs SET status=?, updated_at=? WHERE id=?",
                            ("cancelled", stop_at, db_id),
                        )
                conn.commit()
        except Exception as exc:
            logging.warning("Cap nhat DB cancel job that bai: %s", exc)

    kind_counts: dict[str, int] = {}
    for job in target_jobs:
        kind = _userbot_queue_job_kind(job)
        kind_counts[kind] = kind_counts.get(kind, 0) + 1

    signal_payload = {
        "stop_at": stop_at,
        "stop_at_text": _format_local_time(stop_at),
        "requested_by": str(requested_by),
        "requested_slot": current_bot_slot(),
        "job_ids": sorted(target_ids),
        "active_job_ids": active_ids,
        "waiting_job_ids": waiting_ids,
        "kind_counts": kind_counts,
    }
    _write_json_file_safe(USERBOT_STOP_ALL_JSON, signal_payload)

    worker_result = {"stopped": [], "failed": [], "raw": "", "returncode": 0}
    if stop_workers and active_ids:
        worker_result = _terminate_userbot_worker_processes()

    lines = [
        "🛑 ĐÃ GỬI LỆNH DỪNG TẤT CẢ JOB USERBOT",
        f"Người gửi: {requested_by}",
        f"Thời điểm: {_format_local_time(stop_at)}",
        f"Job bị hủy: {len(target_ids)}",
        f"- Đang chạy/đang giữ: {len(active_ids)}",
        f"- Đang chờ: {len(waiting_ids)}",
    ]
    if kind_counts:
        lines.append("Theo loại job:")
        for kind, count in sorted(kind_counts.items()):
            lines.append(f"- {kind}: {count}")
    if active_ids:
        lines.append(f"Worker userbot đã dừng process: {len(worker_result.get('stopped') or [])}")
        if worker_result.get("failed"):
            lines.append(f"Worker dừng lỗi: {len(worker_result.get('failed') or [])}")
        lines.append("Muốn chạy job mới sau khi dừng, mở lại file worker watch.")
    else:
        lines.append("Không thấy job đang chạy nên không tắt worker.")
    lines.append(f"File tín hiệu: {USERBOT_STOP_ALL_JSON}")
    summary = "\n".join(lines)
    try:
        with open(USERBOT_STOP_ALL_LAST_SUMMARY, "w", encoding="utf-8") as f:
            f.write(summary)
    except Exception:
        pass
    return {
        "summary": summary,
        "job_count": len(target_ids),
        "active_count": len(active_ids),
        "waiting_count": len(waiting_ids),
        "worker": worker_result,
    }


def _latest_contact_import_log_progress() -> dict:
    progress = {}
    if not os.path.exists(CONTACT_IMPORT_WORKER_LOG):
        return progress
    try:
        with open(CONTACT_IMPORT_WORKER_LOG, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()[-600:]
    except Exception:
        return progress
    for line in reversed(lines):
        m = re.search(r"\[(?P<time>[^\]]+)\]\s+(?P<account>\d+): imported batch (?P<done>\d+)/(?P<total>\d+)", line)
        if not m:
            continue
        done = int(m.group("done"))
        total = int(m.group("total"))
        remain = max(0, total - done)
        est_min = int((remain / 20) * 8 / 60 + 0.999) if remain else 0
        progress = {
            "time": m.group("time"),
            "account": m.group("account"),
            "done": done,
            "total": total,
            "remain": remain,
            "est_min": est_min,
        }
        break
    return progress


def _format_local_time(raw_ts) -> str:
    try:
        ts = float(raw_ts)
        if ts > 0:
            return datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        pass
    return "N/A"


def _read_worker_log_tail(max_lines: int = 1200) -> list[str]:
    if not os.path.exists(CONTACT_IMPORT_WORKER_LOG):
        return []
    try:
        with open(CONTACT_IMPORT_WORKER_LOG, "r", encoding="utf-8", errors="ignore") as f:
            return f.readlines()[-max_lines:]
    except Exception:
        return []


def _contact_import_job_log_progress(job_id: str, expected_total: int = 0) -> dict:
    lines = _read_worker_log_tail(2500)
    if not lines or not job_id:
        return {}
    start_idx = -1
    for idx, line in enumerate(lines):
        if f"Process contact import job #{job_id}" in line:
            start_idx = idx
    if start_idx < 0:
        return {}
    segment = lines[start_idx:]
    accounts: dict[str, dict] = {}
    last_line = ""
    last_time = ""
    for line in segment:
        m = re.search(r"\[(?P<time>[^\]]+)\]\s+(?P<account>\d+): imported batch (?P<done>\d+)/(?P<total>\d+)", line)
        if not m:
            continue
        acc = m.group("account")
        accounts[acc] = {
            "done": int(m.group("done")),
            "total": int(m.group("total")),
            "time": m.group("time"),
            "line": line.strip(),
        }
        last_line = line.strip()
        last_time = m.group("time")
    if not accounts:
        return {"job_id": job_id, "accounts": {}, "last_line": "", "processed": 0, "total": expected_total}
    processed = sum(min(int(v.get("done") or 0), int(v.get("total") or 0)) for v in accounts.values())
    visible_total = int(expected_total or sum(int(v.get("total") or 0) for v in accounts.values()) or 0)
    percent = (processed / visible_total * 100) if visible_total else 0.0
    return {
        "job_id": job_id,
        "accounts": accounts,
        "last_line": last_line,
        "last_time": last_time,
        "processed": processed,
        "total": visible_total,
        "percent": percent,
    }


def _latest_contact_import_result_summary() -> str:
    if not os.path.exists(CONTACT_IMPORT_LAST_SUMMARY):
        return ""
    try:
        with open(CONTACT_IMPORT_LAST_SUMMARY, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
    except Exception:
        return ""
    if not text:
        return ""
    lines = text.splitlines()
    keep = []
    for line in lines:
        if (
            line.startswith("Job #")
            or line.startswith("Phones:")
            or line.startswith("Total imported:")
            or line.startswith("Resolved UID/users:")
            or line.startswith("File UID")
        ):
            keep.append(line)
    return "\n".join(keep[:8])


def contact_import_status_text() -> str:
    jobs = _read_contact_import_pending_jobs()
    processed_raw = _read_json_file_safe(CONTACT_IMPORT_PROCESSED_JSON, [])
    processed = {str(x) for x in processed_raw} if isinstance(processed_raw, list) else set()
    claimed_raw = _read_json_file_safe(CONTACT_IMPORT_CLAIMED_JSON, {})
    claimed = {str(k): float(v or 0) for k, v in claimed_raw.items()} if isinstance(claimed_raw, dict) else {}
    progress = _latest_contact_import_log_progress()

    active_jobs = [job for job in jobs if str(job.get("id")) in claimed and str(job.get("id")) not in processed]
    waiting_jobs = [job for job in jobs if str(job.get("id")) not in claimed and str(job.get("id")) not in processed]
    done_jobs = [job for job in jobs if str(job.get("id")) in processed]

    lines = ["📊 TRẠNG THÁI QUÉT SĐT -> UID/USER"]
    active_log_progress = {}
    if active_jobs:
        first_active = active_jobs[0]
        first_detail = first_active.get("detail") or {}
        active_log_progress = _contact_import_job_log_progress(
            str(first_active.get("id") or ""),
            int(first_detail.get("phone_count") or 0),
        )
    if active_log_progress and active_log_progress.get("accounts"):
        total = int(active_log_progress.get("total") or 0)
        processed_count = int(active_log_progress.get("processed") or 0)
        remain = max(0, total - processed_count) if total else 0
        accounts = active_log_progress.get("accounts") or {}
        last_account = ""
        if active_log_progress.get("last_line"):
            m_last = re.search(r"\]\s+(\d+):", str(active_log_progress.get("last_line") or ""))
            last_account = m_last.group(1) if m_last else ""
        lines.extend(
            [
                f"Đang chạy tổng: {processed_count}/{total or '?'} số ({float(active_log_progress.get('percent') or 0):.1f}%).",
                f"Acc đang quét hiện tại: {last_account or 'N/A'}",
                f"Số acc đã có tiến độ: {len(accounts)}",
                f"Còn lại khoảng: {remain} số.",
                f"Cập nhật cuối: {active_log_progress.get('last_time') or 'N/A'}",
            ]
        )
        lines.append("")
        lines.append("Tiến độ từng acc gần nhất:")
        for acc, row in sorted(accounts.items())[:20]:
            lines.append(f"- {acc}: {row.get('done', 0)}/{row.get('total', '?')}")
    elif progress:
        percent = (progress["done"] / progress["total"] * 100) if progress["total"] else 0
        lines.extend(
            [
                f"Đang chạy acc hiện tại: {progress['done']}/{progress['total']} số ({percent:.1f}%).",
                f"Acc userbot đang chạy: {progress['account']}",
                f"Còn lại trong acc này: {progress['remain']} số.",
                f"Ước tính còn trong acc này: {progress['est_min']} phút nếu không bị Telegram chờ.",
                f"Cập nhật cuối: {progress['time']}",
            ]
        )
    else:
        lines.append("Chưa thấy tiến độ import contact mới trong log.")

    lines.extend(
        [
            "",
            f"Job đang giữ: {len(active_jobs)}",
            f"Job đang chờ: {len(waiting_jobs)}",
            f"Job đã xong/bỏ qua: {len(done_jobs)}",
        ]
    )
    if active_jobs:
        lines.append("")
        lines.append("Job đang giữ:")
        for job in active_jobs[:5]:
            detail = job.get("detail") or {}
            lines.append(
                f"- {job.get('id')} | {detail.get('phone_count') or '?'} số | "
                f"{detail.get('profile_count') or detail.get('max_accounts_per_job') or '?'} userbot | "
                f"{detail.get('profiles_root') or ''}"
            )
    if waiting_jobs:
        lines.append("")
        lines.append("Job đang chờ:")
        for job in waiting_jobs[:5]:
            detail = job.get("detail") or {}
            lines.append(f"- {job.get('id')} | {detail.get('phone_count') or '?'} số")
    if progress:
        lines.append("")
        lines.append("Lưu ý: chỉ khi job chạy xong mới có file UID/user đầy đủ. Bấm /uidnow giữa chừng chỉ thấy cache hoặc kết quả tạm.")
    latest_summary = _latest_contact_import_result_summary()
    if latest_summary:
        lines.append("")
        lines.append("Kết quả job gần nhất:")
        lines.append(latest_summary)
    return "\n".join(lines)


def contact_import_partial_note(rows_count: int, phone_count: int) -> str:
    if phone_count <= 0 or rows_count >= phone_count:
        return ""
    progress = _latest_contact_import_log_progress()
    lines = [
        "",
        "⚠️ LƯU Ý KẾT QUẢ TẠM",
        f"File có {phone_count} SĐT, nhưng hiện mới có {rows_count} dòng UID/user hợp lệ.",
        "File map nhanh chỉ lấy dữ liệu đã biết/cache; nó không đảm bảo biến toàn bộ SĐT thành UID ngay.",
    ]
    if progress:
        percent = (progress["done"] / progress["total"] * 100) if progress["total"] else 0
        lines.append(
            f"Job import contact đang chạy: {progress['done']}/{progress['total']} số ({percent:.1f}%), còn khoảng {progress['est_min']} phút."
        )
    lines.append("Khi job xong, gửi lại /uidnow để xuất file đầy đủ hơn. Xem tiến độ bằng /uidstatus.")
    return "\n".join(lines)


def contact_import_user_indexes(rows: list[dict]) -> dict[str, dict]:
    idx = {}
    for row in rows:
        for key in (
            str(row.get("uid") or "").strip(),
            _clean_import_username(row.get("username") or "").lower(),
            ("@" + _clean_import_username(row.get("username") or "").lower()) if row.get("username") else "",
            normalize_vn_phone(row.get("phone") or ""),
        ):
            if key:
                idx[key] = row
    return idx


def extract_recipient_tokens(raw: str) -> list[str]:
    text = str(raw or "")
    tokens = []
    seen = set()
    for m in re.finditer(r"@[A-Za-z0-9_]{3,}|(?<!\d)\d{5,20}(?!\d)", text):
        token = m.group(0).strip()
        key = token.lower()
        if key not in seen:
            seen.add(key)
            tokens.append(token)
    return tokens


def looks_like_guitn_selector(raw: str) -> bool:
    text = str(raw or "").strip()
    if not text:
        return False
    if text.lower() in {"all", "alluser", "tatca", "tat_ca", "full"}:
        return True
    if not extract_recipient_tokens(text):
        return False
    residue = re.sub(r"@[A-Za-z0-9_]{3,}|(?<!\d)\d{5,20}(?!\d)", " ", text)
    residue = re.sub(r"[\s,;|/]+", "", residue)
    return not residue


def select_contact_import_users(rows: list[dict], selector_text: str) -> tuple[list[dict], list[str]]:
    selector = str(selector_text or "").strip()
    if not selector or selector.lower() in {"all", "alluser", "tatca", "tat_ca", "full"}:
        return rows, []
    idx = contact_import_user_indexes(rows)
    selected = []
    missing = []
    seen = set()
    for token in extract_recipient_tokens(selector):
        key = normalize_vn_phone(token) if token[:1].isdigit() else token.lower().lstrip("@")
        row = idx.get(key) or idx.get("@" + key)
        if not row:
            missing.append(token)
            continue
        row_key = (row.get("uid") or "", row.get("username") or "", row.get("phone") or "")
        if row_key in seen:
            continue
        seen.add(row_key)
        selected.append(row)
    return selected, missing


def contact_import_line(row: dict, mode: str) -> str:
    username = _clean_import_username(row.get("username") or "")
    user = f"@{username}" if username else "(chưa có user)"
    uid = str(row.get("uid") or "").strip() or "(chưa có UID)"
    phone = str(row.get("phone") or "").strip() or "(chưa có SĐT)"
    name = str(row.get("name") or "").strip() or "(chưa có tên)"
    account = str(row.get("account") or "").strip() or "(chưa rõ acc)"
    raw_uid = str(row.get("uid") or "").strip()
    link = str(row.get("link") or "").strip()
    if not link:
        link = f"https://t.me/{username}" if username else (f"tg://user?id={raw_uid}" if raw_uid else "")
    if mode == "user":
        return user
    if mode == "uid":
        return f"{uid} | {user}"
    if mode == "sdt":
        return f"{phone} | {user}"
    if mode == "all":
        return contact_import_guitn_line(row)
    return f"{phone} | {uid} | {user} | {name} | acc={account}"


def contact_import_has_send_target(row: dict) -> bool:
    username = _clean_import_username(row.get("username") or "")
    uid = str(row.get("uid") or "").strip()
    return bool(username or uid)


def contact_import_clean_link(row: dict) -> str:
    username = _clean_import_username(row.get("username") or "")
    uid = str(row.get("uid") or "").strip()
    link = str(row.get("link") or "").strip()
    if link:
        return link.replace("https://", "").replace("http://", "").strip("/")
    if username:
        return f"t.me/{username}"
    if uid:
        return f"tg://user?id={uid}"
    return "-"


def contact_import_pipe_value(value: object) -> str:
    text = str(value or "").strip()
    if not text or text.upper() == "N/A":
        return "-"
    return text.replace("\r", " ").replace("\n", " ")


def contact_import_guitn_line(row: dict) -> str:
    username = _clean_import_username(row.get("username") or "")
    uid = str(row.get("uid") or "").strip()
    user = f"@{username}" if username else "-"
    phone = normalize_vn_phone(row.get("phone") or "") or str(row.get("phone") or "").strip()
    account = str(row.get("account") or row.get("acc") or row.get("profile") or "").strip()
    online = str(
        row.get("online")
        or row.get("onl")
        or row.get("onl_bao_lau")
        or row.get("last_seen")
        or row.get("last_online")
        or ""
    ).strip()
    return " | ".join(
        [
            contact_import_pipe_value(uid),
            contact_import_pipe_value(user),
            contact_import_pipe_value(phone),
            contact_import_pipe_value(account),
            contact_import_pipe_value(online),
            contact_import_pipe_value(contact_import_clean_link(row)),
        ]
    )


def contact_import_uid_user_line(row: dict) -> str:
    uid = str(row.get("uid") or "").strip() or "không có UID"
    username = _clean_import_username(row.get("username") or "")
    user = f"@{username}" if username else "KHÔNG CÓ USER"
    return f"{uid} | {user}"


def build_contact_import_uid_user_report(rows: list[dict]) -> tuple[str, bytes] | None:
    valid_rows = [row for row in rows if contact_import_has_send_target(row)]
    lines = ["UID | USER"]
    for row in valid_rows:
        lines.append(contact_import_uid_user_line(row))
    if len(lines) <= 1:
        return None
    filename = f"telegram_uid_user_{int(time.time())}.txt"
    return filename, ("\n".join(lines) + "\n").encode("utf-8-sig")


def write_contact_import_export(rows: list[dict], mode: str, requested_by: str) -> str:
    os.makedirs(USERBOT_SCAN_JOBS_DIR, exist_ok=True)
    ts = int(time.time())
    prefix = "quetall" if mode == "all" else mode
    path = os.path.join(USERBOT_SCAN_JOBS_DIR, f"{prefix}_now_{ts}_{requested_by}.txt")
    if mode == "all":
        rows = [row for row in rows if contact_import_has_send_target(row)]
    title = {
        "user": "USER",
        "uid": "UID | USER",
        "sdt": "SDT | USER",
        "tt": "SDT | UID | USER | TEN | ACC",
        "all": "UID | USER | PHONE | ACCOUNT | ONL BAO LAU | LINK",
    }.get(mode, "DATA")
    lines = [title, f"total={len(rows)}", ""]
    lines.extend(contact_import_line(row, mode) for row in rows)
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


def write_userbot_send_recipients(rows: list[dict], requested_by: str) -> str:
    os.makedirs(USERBOT_SCAN_JOBS_DIR, exist_ok=True)
    path = os.path.join(USERBOT_SCAN_JOBS_DIR, f"userbot_send_{int(time.time())}_{requested_by}.csv")
    fieldnames = ["account", "phone", "uid", "name", "username", "source", "consent", "status"]
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})
    return path


def read_userbot_send_recipients_file(path: str) -> list[dict]:
    rows = []
    if not path or not os.path.exists(path):
        return rows
    try:
        with open(path, "r", encoding="utf-8-sig", newline="") as f:
            for row in csv.DictReader(f):
                item = {
                    "account": str(row.get("account") or "").strip(),
                    "phone": normalize_vn_phone(str(row.get("phone") or "")),
                    "uid": str(row.get("uid") or "").strip(),
                    "name": str(row.get("name") or "").strip(),
                    "username": sanitize_username(str(row.get("username") or "")),
                    "source": str(row.get("source") or "").strip(),
                    "consent": normalize_consent_status(str(row.get("consent") or "")),
                    "status": str(row.get("status") or "").strip(),
                }
                if item["uid"] or item["username"]:
                    rows.append(item)
    except Exception as exc:
        logging.warning("Doc file recipients /guitn that bai %s: %s", path, exc)
    return rows


def collect_unsent_userbot_send_rows(summary: dict) -> list[dict]:
    rows = []
    seen = set()
    for result in summary.get("results") or []:
        account = str(result.get("account") or "").strip()
        for failure in result.get("failures") or []:
            uid = str(failure.get("uid") or "").strip()
            username = sanitize_username(str(failure.get("username") or ""))
            phone = normalize_vn_phone(str(failure.get("phone") or ""))
            if not uid and not username:
                continue
            item = {
                "account": str(failure.get("account") or account).strip(),
                "phone": phone,
                "uid": uid,
                "name": str(failure.get("name") or "").strip(),
                "username": username,
                "source": str(failure.get("source") or "").strip(),
                "consent": normalize_consent_status(str(failure.get("consent") or "")),
                "status": str(failure.get("status") or "").strip(),
            }
            key = (item["uid"], item["username"].lower(), item["phone"], item["account"])
            if key in seen:
                continue
            seen.add(key)
            rows.append(item)
    return rows


def load_latest_guitn_unsent_rows() -> tuple[list[dict], str, list[str], dict]:
    notes = []
    if not os.path.exists(USERBOT_SEND_RESULTS_JSONL):
        return [], "", ["missing_send_results_jsonl"], {}
    try:
        with open(USERBOT_SEND_RESULTS_JSONL, "r", encoding="utf-8-sig") as f:
            lines = [line.strip() for line in f if line.strip()]
    except Exception as exc:
        return [], "", [f"read_send_results_error:{type(exc).__name__}:{str(exc)[:120]}"], {}
    for raw in reversed(lines):
        try:
            summary = json.loads(raw)
        except Exception:
            continue
        if not isinstance(summary, dict):
            continue
        unsent_csv = str(summary.get("unsent_csv") or "").strip()
        if unsent_csv and os.path.exists(unsent_csv):
            rows = read_userbot_send_recipients_file(unsent_csv)
            if rows:
                return rows, unsent_csv, notes, summary
        unsent_txt = str(summary.get("unsent_txt") or "").strip()
        if unsent_txt and os.path.exists(unsent_txt):
            try:
                with open(unsent_txt, "r", encoding="utf-8-sig", errors="ignore") as f:
                    rows = parse_guitn_recipients_text(f.read())
                if rows:
                    return rows, unsent_txt, notes, summary
            except Exception as exc:
                notes.append(f"read_unsent_txt_error:{type(exc).__name__}:{str(exc)[:120]}")
        rows = collect_unsent_userbot_send_rows(summary)
        if rows:
            notes.append("fallback_from_send_results_jsonl")
            return rows, f"send_results_jsonl:job#{summary.get('job_id')}", notes, summary
    return [], "", (notes or ["no_unsent_rows_found"]), {}


def guitn_pending_dir() -> str:
    path = os.path.join(USERBOT_SCAN_JOBS_DIR, "guitn_pending")
    os.makedirs(path, exist_ok=True)
    return path


def guitn_pending_plan_path(user_id: str | int, slot: str | None = None) -> str:
    safe_slot = re.sub(r"[^0-9A-Za-z_.-]+", "_", str(slot or current_bot_slot() or "default"))[:40] or "default"
    safe_user = re.sub(r"[^0-9A-Za-z_.-]+", "_", str(user_id or "admin"))[:40] or "admin"
    return os.path.join(guitn_pending_dir(), f"{safe_slot}_{safe_user}.json")


def save_guitn_pending_plan(update: Update, plan: dict) -> str:
    actor_id = str(getattr(update.effective_user, "id", "") or "admin")
    plan = dict(plan or {})
    plan["updated_at"] = int(time.time())
    plan["bot_slot"] = current_bot_slot()
    path = guitn_pending_plan_path(actor_id)
    tmp_path = path + ".tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(plan, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, path)
    return path


def load_guitn_pending_plan(update: Update) -> tuple[dict, str]:
    actor_id = str(getattr(update.effective_user, "id", "") or "admin")
    path = guitn_pending_plan_path(actor_id)
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8-sig") as f:
                data = json.load(f)
            if isinstance(data, dict):
                return data, path
    except Exception as exc:
        logging.warning("Doc pending /guitn that bai %s: %s", path, exc)
    return {}, path


def clear_guitn_pending_plan(update: Update) -> None:
    try:
        _plan, path = load_guitn_pending_plan(update)
        if path and os.path.exists(path):
            os.remove(path)
    except Exception:
        pass


def parse_keoall_recipients_text(text: str) -> tuple[list[dict], dict]:
    rows = []
    seen = set()
    total_lines = 0
    duplicate_count = 0
    uid_only_count = 0
    username_count = 0
    invalid_count = 0
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        lower = line.lower()
        if lower.startswith("uid |") or lower.startswith("user") or lower.startswith("sdt |") or lower.startswith("total="):
            continue
        total_lines += 1
        username = ""
        for username_match in re.finditer(r"@?([A-Za-z][A-Za-z0-9_]{2,})", line):
            candidate = sanitize_username(username_match.group(1))
            if not candidate.isdigit():
                username = candidate
                break
        parts = [p.strip() for p in line.split("|")]
        consent_match = re.search(r"(?i)\b(?:consent|consent_status|opt_in|dong_y|đồng_y)\s*[:=]\s*([A-Za-z0-9_-]+)", line)
        line_consent = normalize_consent_status(consent_match.group(1)) if consent_match else ""
        if not line_consent:
            for part in parts:
                status = normalize_consent_status(part)
                if status == "opt_in":
                    line_consent = status
                    break
        uid = pick_uid_from_guitn_line(line, parts)
        if not username and not uid:
            invalid_count += 1
            continue
        key = (uid, username.lower())
        if key in seen:
            duplicate_count += 1
            continue
        seen.add(key)
        if username:
            username_count += 1
        elif uid:
            uid_only_count += 1
        rows.append({
            "account": "",
            "phone": "",
            "uid": uid,
            "name": "",
            "username": username,
            "source": "telegram_reply_file",
            "consent": line_consent,
            "status": "sendable" if line_consent == "opt_in" else "",
        })
    stats = {
        "total_lines": total_lines,
        "valid_count": len(rows),
        "username_count": username_count,
        "uid_only_count": uid_only_count,
        "duplicate_count": duplicate_count,
        "invalid_count": invalid_count,
    }
    return rows, stats


async def load_keoall_reply_file_recipients(update: Update, context: ContextTypes.DEFAULT_TYPE) -> tuple[list[dict], str, list[str], dict]:
    msg = update.message
    doc = None
    if msg and getattr(msg, "document", None):
        doc = msg.document
    elif msg and msg.reply_to_message and getattr(msg.reply_to_message, "document", None):
        doc = msg.reply_to_message.document
    if not doc:
        return [], "", ["khong_co_file_reply"], {
            "total_lines": 0,
            "valid_count": 0,
            "username_count": 0,
            "uid_only_count": 0,
            "duplicate_count": 0,
            "invalid_count": 0,
        }
    name = str(getattr(doc, "file_name", "") or "keoall_reply")
    try:
        tg_file = await doc.get_file()
        blob = bytes(await tg_file.download_as_bytearray())
        text, notes = extract_text_from_file_bytes(blob, name)
        rows, stats = parse_keoall_recipients_text(text)
        return rows, f"telegram_reply_file:{name}", notes, stats
    except Exception as exc:
        return [], f"telegram_reply_file:{name}", [f"doc_error:{type(exc).__name__}:{str(exc)[:120]}"], {
            "total_lines": 0,
            "valid_count": 0,
            "username_count": 0,
            "uid_only_count": 0,
            "duplicate_count": 0,
            "invalid_count": 0,
        }


def write_keoall_recipients(rows: list[dict], requested_by: str) -> str:
    queue_dir = os.path.join(USERBOT_SCAN_JOBS_DIR, "keoall_queue")
    os.makedirs(queue_dir, exist_ok=True)
    path = os.path.join(queue_dir, f"keoall_recipients_{int(time.time())}_{requested_by}.csv")
    fieldnames = ["account", "phone", "uid", "name", "username", "source", "consent", "status"]
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})
    return path


def parse_keoall_options(update: Update) -> dict:
    tail = get_command_tail(update, "keoall").strip()
    run_real = bool(re.search(r"(?i)(?:^|\s)(run|chay|chạy|that|thật|now)(?:\s|$)", tail))
    tail = re.sub(r"(?i)(?:^|\s)(run|chay|chạy|that|thật|now)(?:\s|$)", " ", tail, count=1).strip()
    limit = 0
    m_limit = re.search(r"(?i)(?:limit|max|gioihan|giới_hạn|gioi_han)\s*=\s*(\d{1,5})", tail)
    if m_limit:
        try:
            limit = max(1, int(m_limit.group(1)))
        except Exception:
            limit = 0
        tail = (tail[:m_limit.start()] + " " + tail[m_limit.end():]).strip()
    delay_min, delay_max, delay_label = 60.0, 180.0, "60-180s"
    m_delay = re.search(r"(?i)(?:delay|tre|trễ)\s*=\s*(\d{1,5})(?:\s*[-–]\s*(\d{1,5}))?\s*(?:s|giay|giây)?", tail)
    if m_delay:
        try:
            delay_min = float(m_delay.group(1))
            delay_max = float(m_delay.group(2) or m_delay.group(1))
            if delay_max < delay_min:
                delay_min, delay_max = delay_max, delay_min
            delay_min = max(1.0, delay_min)
            delay_max = max(delay_min, delay_max)
            delay_label = f"{int(delay_min) if delay_min.is_integer() else delay_min:g}s"
            if delay_max != delay_min:
                delay_label = f"{int(delay_min) if delay_min.is_integer() else delay_min:g}-{int(delay_max) if delay_max.is_integer() else delay_max:g}s"
        except Exception:
            delay_min, delay_max, delay_label = 60.0, 180.0, "60-180s"
        tail = (tail[:m_delay.start()] + " " + tail[m_delay.end():]).strip()
    link_match = re.search(r"(?i)(https?://(?:t\.me|telegram\.me)/\S+|(?:t\.me|telegram\.me)/\S+|@[A-Za-z0-9_]{3,})", tail)
    target_link = link_match.group(1).strip() if link_match else ""
    return {
        "run_real": run_real,
        "target_link": target_link,
        "limit": limit,
        "delay_min": delay_min,
        "delay_max": delay_max,
        "delay_label": delay_label,
    }


def parse_guitn_delay_meta(tail: str) -> tuple[str, float, float, str, bool]:
    text = str(tail or "").strip()
    default_delay = 1.5
    if not text:
        return "", default_delay, default_delay, f"{default_delay}s", False
    m = re.search(
        r"(?i)(?:^|\s)("
        r"(?P<range_a>\d{1,5})\s*[-–]\s*(?P<range_b>\d{1,5})\s*(?:s|sec|secs|second|seconds|giay|giây)?"
        r"|(?P<single>\d{1,5})\s*(?:s|sec|secs|second|seconds|giay|giây)"
        r")[:;,.]?\s*$",
        text,
    )
    if not m:
        return text, default_delay, default_delay, f"{default_delay}s", False
    if m.group("single"):
        delay_min = delay_max = float(m.group("single"))
    else:
        delay_min = float(m.group("range_a"))
        delay_max = float(m.group("range_b"))
    if delay_max < delay_min:
        delay_min, delay_max = delay_max, delay_min
    delay_min = max(0.1, delay_min)
    delay_max = max(delay_min, delay_max)
    cleaned = text[: m.start(1)].rstrip()
    label = f"{int(delay_min) if delay_min.is_integer() else delay_min:g}s"
    if delay_max != delay_min:
        label = f"{int(delay_min) if delay_min.is_integer() else delay_min:g}-{int(delay_max) if delay_max.is_integer() else delay_max:g}s"
    return cleaned, delay_min, delay_max, label, True


def parse_guitn_delay(tail: str) -> tuple[str, float, float, str]:
    cleaned, delay_min, delay_max, label, _used_explicit = parse_guitn_delay_meta(tail)
    return cleaned, delay_min, delay_max, label


def parse_guitn_payload_for_command(update: Update, command_name: str = "guitn") -> tuple[str, str, float, float, str, bool]:
    tail = get_command_tail(update, command_name)
    tail, delay_min, delay_max, delay_label, delay_explicit = parse_guitn_delay_meta(tail)
    msg = update.message
    reply_msg = getattr(msg, "reply_to_message", None) if msg else None
    reply_has_doc = bool(reply_msg and getattr(reply_msg, "document", None))
    reply_text = "" if reply_has_doc else get_reply_payload(update)
    msg_text = tail.strip()
    selector_text = ""
    if "|" in msg_text:
        maybe_msg, maybe_selector = msg_text.rsplit("|", 1)
        if looks_like_guitn_selector(maybe_selector):
            msg_text = maybe_msg.strip()
            selector_text = maybe_selector.strip()
    if reply_text:
        if not msg_text:
            return reply_text, selector_text, delay_min, delay_max, delay_label, delay_explicit
        if looks_like_guitn_selector(msg_text):
            return reply_text, msg_text, delay_min, delay_max, delay_label, delay_explicit
    if selector_text:
        return msg_text, selector_text, delay_min, delay_max, delay_label, delay_explicit
    if reply_has_doc:
        return msg_text, "", delay_min, delay_max, delay_label, delay_explicit
    parts = msg_text.split()
    trailing = []
    while parts:
        last = parts[-1].strip()
        if re.fullmatch(r"@[A-Za-z0-9_]{3,}|\d{5,20}", last):
            trailing.insert(0, parts.pop())
            continue
        break
    if trailing:
        return " ".join(parts).strip(), "\n".join(trailing), delay_min, delay_max, delay_label, delay_explicit
    return msg_text, "", delay_min, delay_max, delay_label, delay_explicit


def parse_guitn_payload(update: Update) -> tuple[str, str, float, float, str]:
    msg_text, selector_text, delay_min, delay_max, delay_label, _delay_explicit = parse_guitn_payload_for_command(update, "guitn")
    return msg_text, selector_text, delay_min, delay_max, delay_label


def pick_uid_from_guitn_line(line: str, parts: list[str]) -> str:
    if not line or line.strip().lower().startswith("total="):
        return ""
    if parts:
        first = str(parts[0] or "").strip()
        if re.fullmatch(r"\d{5,20}", first) and not _is_valid_phone(first):
            return first
    part_digits = ["".join(ch for ch in str(part or "") if ch.isdigit()) for part in parts]
    if len(part_digits) >= 2 and part_digits[1].isdigit() and len(part_digits[1]) >= 5:
        return part_digits[1]
    if part_digits and part_digits[0].isdigit() and len(part_digits[0]) >= 5 and not _is_valid_phone(part_digits[0]):
        return part_digits[0]
    nums = re.findall(r"(?<!\d)\d{5,20}(?!\d)", line)
    for num in nums:
        if not _is_valid_phone(num):
            return num
    return nums[0] if nums else ""


def parse_guitn_recipients_text(text: str) -> list[dict]:
    rows = []
    seen = set()
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        lower = line.lower()
        if (
            lower.startswith("uid |")
            or lower.startswith("user")
            or lower.startswith("sdt |")
            or lower.startswith("phone |")
            or lower.startswith("total=")
            or lower.startswith("job=")
            or lower.startswith("dropped_")
            or lower.startswith("---")
        ):
            continue
        username_match = re.search(r"@[A-Za-z0-9_]{3,}", line)
        username = sanitize_username(username_match.group(0)) if username_match else ""
        parts = [p.strip() for p in line.split("|")]
        consent_match = re.search(r"(?i)\b(?:consent|consent_status|opt_in|dong_y|đồng_y)\s*[:=]\s*([A-Za-z0-9_-]+)", line)
        line_consent = normalize_consent_status(consent_match.group(1)) if consent_match else ""
        if not line_consent:
            for part in parts:
                status = normalize_consent_status(part)
                if status == "opt_in":
                    line_consent = status
                    break
        is_pipe_recipient_table = len(parts) >= 6
        legacy_phone = ""
        legacy_uid = ""
        if parts and "=" in parts[0]:
            legacy_phone, legacy_uid = [p.strip() for p in parts[0].split("=", 1)]
        if not username and len(parts) >= 2:
            raw_user = parts[1].strip().lstrip("@")
            if re.fullmatch(r"[A-Za-z][A-Za-z0-9_]{2,31}", raw_user):
                username = raw_user
        if is_pipe_recipient_table:
            uid = parts[0] if re.fullmatch(r"\d{5,20}", parts[0] or "") else ""
        else:
            uid = legacy_uid if re.fullmatch(r"\d{5,20}", legacy_uid or "") else pick_uid_from_guitn_line(line, parts)
        phone = normalize_vn_phone(legacy_phone) if legacy_phone else ""
        account = ""
        name = ""
        if len(parts) >= 6:
            phone = normalize_vn_phone(parts[2]) or parts[2].strip()
            account = parts[3].strip()
            name = parts[4].strip()
            if not line_consent:
                line_consent = normalize_consent_status(parts[5])
        elif len(parts) >= 5:
            phone = normalize_vn_phone(parts[0]) or parts[0].strip()
            account = parts[4].replace("acc=", "").strip()
            name = parts[3].strip()
        if not username and not uid:
            continue
        key = (uid, username.lower())
        if key in seen:
            continue
        seen.add(key)
        rows.append(
            {
                "account": "" if account == "-" else account,
                "phone": "" if phone == "-" else phone,
                "uid": "" if uid == "-" else uid,
                "name": "" if name == "-" else name,
                "username": username,
                "source": "text_or_file",
                "consent": line_consent,
                "status": "sendable" if line_consent == "opt_in" else "",
            }
        )
    return rows


async def load_guitn_reply_file_recipients(update: Update, context: ContextTypes.DEFAULT_TYPE) -> tuple[list[dict], str, list[str]]:
    msg = update.message
    doc = None
    if msg and getattr(msg, "document", None):
        doc = msg.document
    elif msg and msg.reply_to_message and getattr(msg.reply_to_message, "document", None):
        doc = msg.reply_to_message.document
    if not doc:
        return [], "", []
    name = str(getattr(doc, "file_name", "") or "uid_now_reply")
    try:
        tg_file = await doc.get_file()
        blob = bytes(await tg_file.download_as_bytearray())
        text, notes = extract_text_from_file_bytes(blob, name)
        return parse_guitn_recipients_text(text), f"telegram_reply_file:{name}", notes
    except Exception as exc:
        return [], f"telegram_reply_file:{name}", [f"doc_error:{type(exc).__name__}:{str(exc)[:120]}"]


async def load_contact_import_reply_file_users(
    update: Update, context: ContextTypes.DEFAULT_TYPE, valid_only: bool = True
) -> tuple[list[dict], str, list[str]]:
    msg = update.message
    doc = None
    if msg and getattr(msg, "document", None):
        doc = msg.document
    elif msg and msg.reply_to_message and getattr(msg.reply_to_message, "document", None):
        doc = msg.reply_to_message.document
    if not doc:
        return [], "", []
    name = str(getattr(doc, "file_name", "") or "file_reply")
    source = f"File Telegram đang reply: {name}"
    try:
        tg_file = await doc.get_file()
        blob = bytes(await tg_file.download_as_bytearray())
        text, notes = extract_text_from_file_bytes(blob, name)
        rows = _read_contact_import_text(text, valid_only=valid_only)
        return rows, source, notes
    except Exception as exc:
        return [], source, [f"Lỗi đọc file: {type(exc).__name__}: {str(exc)[:120]}"]


def _safe_guitn_media_name(raw_name: str, default_name: str) -> str:
    name = str(raw_name or default_name or "guitn_media").strip()
    name = re.sub(r"[^0-9A-Za-z_.-]+", "_", name)[:90].strip("._")
    return name or default_name or "guitn_media"


def _document_looks_like_media(doc) -> bool:
    if not doc:
        return False
    mime = str(getattr(doc, "mime_type", "") or "").lower()
    name = str(getattr(doc, "file_name", "") or "").lower()
    if mime.startswith("image/") or mime.startswith("video/"):
        return True
    return name.endswith((".gif", ".mp4", ".mov", ".webm", ".jpg", ".jpeg", ".png", ".webp"))


async def load_guitn_reply_media(update: Update, context: ContextTypes.DEFAULT_TYPE) -> tuple[str, str]:
    msg = update.message
    if not msg:
        return "", ""
    source = msg.reply_to_message if msg.reply_to_message else msg
    target = None
    file_name = ""
    if getattr(source, "photo", None):
        target = source.photo[-1]
        file_name = "photo.jpg"
    elif getattr(source, "animation", None):
        target = source.animation
        file_name = getattr(target, "file_name", "") or "animation.gif"
    elif getattr(source, "video", None):
        target = source.video
        file_name = getattr(target, "file_name", "") or "video.mp4"
    elif getattr(source, "document", None) and _document_looks_like_media(source.document):
        target = source.document
        file_name = getattr(target, "file_name", "") or "media.bin"
    if not target:
        return "", ""
    try:
        media_dir = os.path.join(USERBOT_SCAN_JOBS_DIR, "guitn_media")
        os.makedirs(media_dir, exist_ok=True)
        safe_name = _safe_guitn_media_name(file_name, "guitn_media.bin")
        out_path = os.path.join(media_dir, f"{int(time.time())}_{secrets.token_hex(4)}_{safe_name}")
        tg_file = await target.get_file()
        blob = bytes(await tg_file.download_as_bytearray())
        with open(out_path, "wb") as f:
            f.write(blob)
        return out_path, "reply_or_forward_media"
    except Exception as exc:
        logging.warning("Tai media /guitn that bai: %s", exc)
        return "", f"media_error:{type(exc).__name__}:{str(exc)[:120]}"


def count_userbot_api_sessions(profiles: list[str], root: str | None = None) -> int:
    accounts_root = root or USERBOT_ACCOUNTS_ROOT
    total = 0
    for profile in profiles:
        session_path = os.path.join(accounts_root, profile, f"{profile}.session")
        if os.path.exists(session_path):
            total += 1
    return total


def load_userbot_api_accounts(limit: int = 100, root: str | None = None) -> list[dict]:
    accounts_root = root or USERBOT_ACCOUNTS_ROOT
    accounts = []
    for profile in list_userbot_profiles(limit=limit, root=accounts_root):
        acc_dir = os.path.join(accounts_root, profile)
        cfg_path = os.path.join(acc_dir, f"{profile}.json")
        data = {}
        try:
            if os.path.exists(cfg_path):
                with open(cfg_path, "r", encoding="utf-8-sig") as f:
                    loaded = json.load(f)
                if isinstance(loaded, dict):
                    data = loaded
        except Exception as exc:
            accounts.append({"phone": profile, "dir": acc_dir, "status": "json_error", "error": str(exc)[:120]})
            continue
        session_name = str(data.get("session_file") or profile)
        accounts.append(
            {
                "phone": str(data.get("phone") or profile),
                "dir": acc_dir,
                "session": os.path.join(acc_dir, f"{session_name}.session"),
                "app_id": int(data.get("app_id") or USERBOT_API_ID or 0),
                "app_hash": str(data.get("app_hash") or USERBOT_API_HASH or ""),
                "device": str(data.get("device") or DEFAULT_MACHINE_DEVICE or "VGAH510"),
                "sdk": str(data.get("sdk") or "Windows 10"),
                "app_version": str(data.get("app_version") or "6.6.2 x64"),
                "cfg_path": cfg_path,
                "status": "",
                "error": "",
            }
        )
    return accounts


def is_primary_bot_instance() -> bool:
    slot = os.getenv("BOT_INSTANCE_SLOT", "").strip().lower()
    return slot in {"", "default", "bot1", "main", "primary"}


def current_bot_slot() -> str:
    return (os.getenv("BOT_INSTANCE_SLOT", "").strip().lower() or "default")


def is_primary_operation_bot_instance() -> bool:
    raw = os.getenv("BOT_OPERATION_PRIMARY_SLOTS", "default,bot1,main,primary").strip()
    slots = {x.strip().lower() for x in re.split(r"[,;\s]+", raw) if x.strip()}
    if not slots:
        slots = {"default", "bot1", "main", "primary"}
    return current_bot_slot() in slots


def is_baokybcr_bot_instance() -> bool:
    slot = current_bot_slot()
    return slot in {"baokybcr", "baokibcr", "bcr", "bot20"}


def is_guitn_bot_instance() -> bool:
    return current_bot_slot() in {"", "default", "bot1", "main", "primary", "bot2", "support", "baokybcr", "baokibcr", "bcr", "bot20"}


def normalize_guitn_lot_key(lot_key: str) -> str:
    lot_key = str(lot_key or "").strip().lower()
    aliases = {
        "runall": "run52",
    }
    return aliases.get(lot_key, lot_key)


def is_removed_guitn_lot_key(lot_key: str) -> bool:
    return str(lot_key or "").strip().lower() in {"run12", "run20", "run32", "run20new"}


def get_guitn_lot_options() -> dict:
    return {
        "run40": {
            "label": "lo 40 acc",
            "roots": get_userbot_40_lot_roots(),
            "profile_limit": 20,
        },
        "run52": {
            "label": "lo 52 acc",
            "roots": get_userbot_52_lot_roots(),
            "profile_limit": 20,
        },
    }


def guitn_lot_prompt_text() -> str:
    return (
        "Chọn lô userbot để chạy /guitn:\n"
        "- /run40: lô 40 acc (nhóm XINHAN + nhóm mới nhất)\n"
        "- /run52: lô 52 acc\n"
    )


def get_guitn_accounts_root() -> str:
    if GUITN_USERBOT_ACCOUNTS_ROOT:
        return GUITN_USERBOT_ACCOUNTS_ROOT
    if is_baokybcr_bot_instance():
        return get_userbot_20_xinhan_root() or USERBOT_20_ACCOUNTS_ROOT
    return USERBOT_ACCOUNTS_ROOT


def assign_contact_import_counts(phones: list[str], profiles: list[str]) -> dict[str, int]:
    counts = {str(p): 0 for p in profiles}
    if not phones or not profiles:
        return counts
    for idx, _phone in enumerate(phones):
        counts[str(profiles[idx % len(profiles)])] += 1
    return counts


def write_contact_import_phone_file(phones: list[str], requested_by: str) -> str:
    os.makedirs(USERBOT_SCAN_JOBS_DIR, exist_ok=True)
    safe_actor = re.sub(r"[^0-9A-Za-z_.-]+", "_", str(requested_by or "admin"))[:40] or "admin"
    path = os.path.join(
        USERBOT_SCAN_JOBS_DIR,
        f"contact_import_{int(time.time())}_{safe_actor}_{secrets.token_hex(4)}.csv",
    )
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["phone"])
        for phone in phones:
            writer.writerow([phone])
    return path


def unique_contact_phones_from_rows(rows: list[dict]) -> list[str]:
    phones = []
    seen = set()
    for row in rows or []:
        phone = normalize_vn_phone(str(row.get("phone") or row.get("sdt") or ""))
        if not _is_valid_phone(phone):
            continue
        if phone in seen:
            continue
        seen.add(phone)
        phones.append(phone)
    return phones


def command_tail_requests_run(tail: str) -> bool:
    tokens = {x.strip().lower() for x in re.split(r"\s+", str(tail or "")) if x.strip()}
    return bool(tokens & {"run", "chay", "chạy", "that", "thật", "now", "quet", "quét"})


async def reply_contact_import_phone_preview(
    msg,
    source_path: str,
    phones: list[str],
    command_label: str,
    notes: list[str] | None = None,
) -> None:
    preview = "\n".join(f"- {p}" for p in phones[:10])
    more = f"\n... còn {len(phones) - 10} số nữa." if len(phones) > 10 else ""
    note_text = ""
    if notes:
        note_text = "\nGhi chú đọc file:\n- " + "\n- ".join(str(n) for n in notes[:5])
    await msg.reply_text(
        "✅ Đã đọc file SĐT.\n"
        f"Nguồn: {source_path}\n"
        f"Tổng SĐT hợp lệ: {len(phones)}\n"
        f"{preview}{more}{note_text}\n\n"
        f"Muốn quét UID/user: reply lại file rồi gửi {command_label} run.\n"
        "Khi worker xong, bot tự gửi file UID | USER và file hợp lệ cho /guitn."
    )


async def queue_contact_import_phone_job(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    contact_phones: list[str],
    contact_notes: list[str] | None,
    dry_run: bool,
    command_label: str = "/phat500k",
    source_label: str = "",
    report_mode: str = "",
    single_report_file: bool = False,
) -> int:
    roots = get_userbot_job_roots(USERBOT_ACCOUNTS_ROOT)
    active_roots = active_userbot_root_profiles(roots, get_userbot_job_profile_limit())
    if not active_roots:
        await update.message.reply_text(
            "❌ Không còn userbot active để quét. Kiểm tra cooldown/session tại "
            f"{USERBOT_ACCOUNT_COOLDOWN_JSON}"
        )
        return 0
    roots = [root for root, _profiles in active_roots]
    phone_chunks = split_rows_by_roots(contact_phones, roots)
    job_ids = []
    report_blocks = []
    doc_files = []
    report_batch_id = f"quetallnow_{int(time.time())}_{update.effective_user.id}" if single_report_file else ""
    for idx, ((root, profiles), phones_for_root) in enumerate(zip(active_roots, phone_chunks), start=1):
        if not phones_for_root:
            continue
        session_count = count_userbot_api_sessions(profiles, root=root)
        phone_file = write_contact_import_phone_file(phones_for_root, str(update.effective_user.id))
        doc_files.append(phone_file)
        assignment = assign_contact_import_counts(phones_for_root, profiles)
        assignment_lines = [f"- {profile}: {assignment.get(profile, 0)} số" for profile in profiles[:20]]
        detail = {
            "job_type": "import_contacts",
            "profiles_root": root,
            "profile_count": len(profiles),
            "max_accounts_per_job": len(profiles),
            "profiles": profiles,
            "session_count": session_count,
            "phone_count": len(phones_for_root),
            "phones_file": phone_file,
            "phones_preview": phones_for_root[:30],
            "delay_account_sec": 5,
            "delay_batch_sec": 8,
            "batch_size": 20,
            "mode": "dry-run" if dry_run else "run",
            "notes": contact_notes or [],
            "source_label": source_label,
            "shard_index": idx,
            "shard_total": len(roots),
            "report_mode": report_mode,
            "single_report_file": bool(single_report_file),
            "report_batch_id": report_batch_id,
        }
        detail = attach_userbot_report_origin(detail, update)
        job_id = queue_userbot_scan_job(
            "phat500k_import_contacts",
            str(getattr(update.effective_chat, "id", "") or ""),
            f"IMPORT_CONTACTS_SHARD_{idx}",
            str(update.effective_user.id),
            dry_run,
            detail,
        )
        job_ids.append(job_id)
        session_note = "" if session_count == len(profiles) else f" | session={session_count}/{len(profiles)}"
        report_blocks.append(
            f"- Shard {idx}/{len(roots)}: job #{job_id} | phones={len(phones_for_root)} | profiles={len(profiles)}{session_note}\n"
            f"  Root: {root}\n"
            + ("\n".join("  " + line for line in assignment_lines[:12]) or "  - chưa thấy profile")
        )
    note_text = ("\nGhi chú đọc file:\n" + "\n".join(f"- {n}" for n in (contact_notes or [])[:6])) if contact_notes else ""
    source_text = f"\nNguồn: {source_label}" if source_label else ""
    run_hint = (
        f"\n\nMuốn chạy thật: reply lại file rồi gõ {command_label} run"
        if dry_run
        else (
            "\n\nWorker sẽ xử lý queue bằng Telethon ImportContacts. "
            "Xong bot tự gửi file UID | USER và file hợp lệ cho /guitn."
            if single_report_file or report_mode == "quetallnow"
            else f"\n\nWorker sẽ xử lý queue bằng Telethon ImportContacts. Xong rồi gõ {command_label.replace(' run', '')} để xuất UID/user."
        )
    )
    report_text = (
        ("DRY-RUN" if dry_run else "RUN")
        + f" {command_label} IMPORT CONTACTS\n"
        + f"Jobs: {', '.join('#' + str(x) for x in job_ids) or 'N/A'}\n"
        + f"Tổng số điện thoại: {len(contact_phones)}\n"
        + f"Botuser shards: {len([x for x in phone_chunks if x])}/{len(roots)}\n"
        + source_text
        + "\n\nChia theo shard:\n"
        + ("\n".join(report_blocks) or "- chưa có shard hợp lệ")
        + note_text
        + run_hint
    )
    await send_telegram_long(context.bot, update.effective_chat.id, report_text)
    await send_admin_management_report(
        context.bot,
        "[IMPORT CONTACTS QUEUED]\n" + report_text,
        files=[] if single_report_file else doc_files,
        caption=f"{command_label} phone queue: {len(contact_phones)} so",
        source_chat_id=update.effective_chat.id,
    )
    return job_ids[0] if job_ids else 0


async def queue_contact_import_uid_user_job(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    contact_phones: list[str],
    contact_notes: list[str] | None,
    dry_run: bool,
    command_label: str,
    source_label: str,
) -> int:
    return await queue_contact_import_phone_job(
        update,
        context,
        contact_phones,
        contact_notes,
        dry_run=dry_run,
        command_label=command_label,
        source_label=source_label,
        report_mode="quetallnow",
        single_report_file=True,
    )


async def extract_phat500k_contact_phones(update: Update, context: ContextTypes.DEFAULT_TYPE, payload: str) -> tuple[list[str], list[str]]:
    notes = []
    text_parts = [payload or "", get_reply_payload(update)]
    msg = update.message
    doc = None
    doc_name = ""
    if msg and getattr(msg, "document", None):
        doc = msg.document
    elif msg and msg.reply_to_message and getattr(msg.reply_to_message, "document", None):
        doc = msg.reply_to_message.document
    if doc:
        doc_name = str(getattr(doc, "file_name", "") or "upload")
        try:
            tg_file = await doc.get_file()
            blob = bytes(await tg_file.download_as_bytearray())
            text, file_notes = extract_text_from_file_bytes(blob, doc_name)
            text_parts.append(text)
            if file_notes:
                notes.extend(file_notes[:6])
        except Exception as exc:
            notes.append(f"doc_error:{type(exc).__name__}:{str(exc)[:120]}")
    local_match = re.search(r"(?i)(?:file|path)\s*=\s*(\"[^\"]+\"|'[^']+'|\S+)", payload or "")
    if local_match:
        path = local_match.group(1).strip().strip("\"'")
        try:
            if os.path.exists(path):
                text, file_notes = extract_text_from_local_file(path)
                text_parts.append(text)
                notes.append(f"local_file:{path}")
                if file_notes:
                    notes.extend(file_notes[:6])
            else:
                notes.append(f"missing_local_file:{path}")
        except Exception as exc:
            notes.append(f"local_file_error:{type(exc).__name__}:{str(exc)[:120]}")
    phones = extract_candidate_phones("\n".join(text_parts))
    return phones, notes


def format_cached_member_line(row) -> str:
    username = str(row["username"] or "")
    name = str(row["name"] or "")
    user_id = str(row["user_id"] or "")
    user_part = f"@{username}" if username else ""
    return f"{user_id} | {name} {user_part}".strip()


def mention_cached_member(row) -> str:
    user_id = str(row["user_id"] or "").strip()
    name = str(row["name"] or row["username"] or user_id).strip()
    if not user_id:
        return ""
    return f'<a href="tg://user?id={html.escape(user_id, quote=True)}">{html.escape(name)}</a>'


def audit_log(actor_platform: str, actor_uid: str, action: str, detail: str = ""):
    with db_conn() as conn:
        conn.execute(
            "INSERT INTO audit(actor_platform, actor_uid, action, detail, created_at) VALUES(?,?,?,?,?)",
            (actor_platform, str(actor_uid), action, detail[:1000], time.time()),
        )
        conn.commit()


def message_hash(text: str) -> str:
    return hashlib.sha256((text or "").encode("utf-8", errors="ignore")).hexdigest()


def get_config(key: str, default: str = "") -> str:
    with db_conn() as conn:
        row = conn.execute("SELECT value FROM config WHERE key=?", (key,)).fetchone()
    return str(row[0]) if row else default


def set_config(key: str, value: str):
    with db_conn() as conn:
        conn.execute("INSERT OR REPLACE INTO config(key, value) VALUES(?,?)", (key, value))
        conn.commit()


def _normalize_user_state(data: dict | None) -> dict:
    data = data if isinstance(data, dict) else {}
    return {
        "step": "" if data.get("step") is None else str(data.get("step") or ""),
        "ten_game": str(data.get("ten_game") or "Chưa nhập"),
        "sdt": str(data.get("sdt") or "Chưa nhập"),
        "goi_nap": str(data.get("goi_nap") or "Chưa chọn"),
    }


def migrate_user_state_tables(user_data_snapshot: dict | None = None, extra_admins_snapshot: set[int] | None = None) -> None:
    user_data_snapshot = user_data_snapshot if isinstance(user_data_snapshot, dict) else {}
    extra_admins_snapshot = extra_admins_snapshot or set()
    now = time.time()
    with db_conn() as conn:
        row = conn.execute("SELECT COUNT(1) AS total FROM users").fetchone()
        if int(row["total"] or 0) == 0:
            for uid, data in user_data_snapshot.items():
                try:
                    uid_int = int(uid)
                except Exception:
                    continue
                state = _normalize_user_state(data)
                conn.execute(
                    """
                    INSERT OR IGNORE INTO users(uid, step, ten_game, sdt, goi_nap, is_banned, is_vip, created_at, updated_at)
                    VALUES(?,?,?,?,?,?,?,?,?)
                    """,
                    (
                        uid_int,
                        state["step"],
                        state["ten_game"],
                        state["sdt"],
                        state["goi_nap"],
                        1 if uid_int in BANNED_USERS else 0,
                        1 if uid_int in VIP_USERS else 0,
                        now,
                        now,
                    ),
                )
        for uid in extra_admins_snapshot:
            try:
                conn.execute(
                    "INSERT OR IGNORE INTO bot_admins(uid, role, created_at) VALUES(?,?,?)",
                    (int(uid), "extra", now),
                )
            except Exception:
                continue
        conn.commit()


def load_user_state_tables() -> None:
    global USER_DATA
    with db_conn() as conn:
        user_rows = conn.execute(
            "SELECT uid, step, ten_game, sdt, goi_nap, is_banned, is_vip FROM users"
        ).fetchall()
        admin_rows = conn.execute("SELECT uid FROM bot_admins WHERE role='extra'").fetchall()
    USER_DATA = {}
    BANNED_USERS.clear()
    VIP_USERS.clear()
    EXTRA_ADMINS.clear()
    for row in user_rows:
        uid = int(row["uid"])
        USER_DATA[uid] = {
            "step": row["step"] or None,
            "ten_game": row["ten_game"] or "Chưa nhập",
            "sdt": row["sdt"] or "Chưa nhập",
            "goi_nap": row["goi_nap"] or "Chưa chọn",
        }
        if int(row["is_banned"] or 0):
            BANNED_USERS.add(uid)
        if int(row["is_vip"] or 0):
            VIP_USERS.add(uid)
    for row in admin_rows:
        EXTRA_ADMINS.add(int(row["uid"]))


def save_user_state(uid: int) -> None:
    try:
        uid_int = int(uid)
    except Exception:
        return
    state = _normalize_user_state(USER_DATA.get(uid_int))
    now = time.time()
    with db_conn() as conn:
        conn.execute(
            """
            INSERT INTO users(uid, step, ten_game, sdt, goi_nap, is_banned, is_vip, created_at, updated_at)
            VALUES(?,?,?,?,?,?,?,?,?)
            ON CONFLICT(uid) DO UPDATE SET
                step=excluded.step,
                ten_game=excluded.ten_game,
                sdt=excluded.sdt,
                goi_nap=excluded.goi_nap,
                is_banned=excluded.is_banned,
                is_vip=excluded.is_vip,
                updated_at=excluded.updated_at
            """,
            (
                uid_int,
                state["step"],
                state["ten_game"],
                state["sdt"],
                state["goi_nap"],
                1 if uid_int in BANNED_USERS else 0,
                1 if uid_int in VIP_USERS else 0,
                now,
                now,
            ),
        )
        conn.commit()


def persist_user_state_tables() -> None:
    now = time.time()
    with db_conn() as conn:
        for uid, data in list(USER_DATA.items()):
            try:
                uid_int = int(uid)
            except Exception:
                continue
            state = _normalize_user_state(data)
            conn.execute(
                """
                INSERT INTO users(uid, step, ten_game, sdt, goi_nap, is_banned, is_vip, created_at, updated_at)
                VALUES(?,?,?,?,?,?,?,?,?)
                ON CONFLICT(uid) DO UPDATE SET
                    step=excluded.step,
                    ten_game=excluded.ten_game,
                    sdt=excluded.sdt,
                    goi_nap=excluded.goi_nap,
                    is_banned=excluded.is_banned,
                    is_vip=excluded.is_vip,
                    updated_at=excluded.updated_at
                """,
                (
                    uid_int,
                    state["step"],
                    state["ten_game"],
                    state["sdt"],
                    state["goi_nap"],
                    1 if uid_int in BANNED_USERS else 0,
                    1 if uid_int in VIP_USERS else 0,
                    now,
                    now,
                ),
            )
        conn.execute("DELETE FROM bot_admins WHERE role='extra'")
        for uid in sorted(EXTRA_ADMINS):
            conn.execute(
                "INSERT OR REPLACE INTO bot_admins(uid, role, created_at) VALUES(?,?,?)",
                (int(uid), "extra", now),
            )
        conn.commit()


def clear_user_state_tables() -> None:
    now = time.time()
    with db_conn() as conn:
        conn.execute(
            "UPDATE users SET step='', ten_game='', sdt='', goi_nap='', updated_at=? WHERE is_banned=1 OR is_vip=1",
            (now,),
        )
        conn.execute("DELETE FROM users WHERE is_banned=0 AND is_vip=0")
        conn.commit()


def outbox_exists(platform: str, target_uid: str, text: str) -> bool:
    mhash = message_hash(text)
    with db_conn() as conn:
        row = conn.execute(
            "SELECT 1 FROM outbox WHERE platform=? AND target_uid=? AND msg_hash=?",
            (platform, str(target_uid), mhash),
        ).fetchone()
    return bool(row)


def mark_outbox_sent(platform: str, target_uid: str, text: str) -> None:
    with db_conn() as conn:
        conn.execute(
            "INSERT OR IGNORE INTO outbox(platform, target_uid, msg_hash, created_at) VALUES(?,?,?,?)",
            (platform, str(target_uid), message_hash(text), time.time()),
        )
        conn.commit()


def list_broadcast_customer_targets() -> list[str]:
    targets = {str(uid) for uid in USER_DATA.keys()}
    try:
        with db_conn() as conn:
            for row in conn.execute("SELECT uid FROM users").fetchall():
                targets.add(str(row["uid"]))
            for row in conn.execute(
                """
                SELECT uid FROM customer_contacts
                WHERE platform IN ('telegram', 'telegram_user')
                """
            ).fetchall():
                uid = str(row["uid"] or "").strip()
                if re.fullmatch(r"-?\d{5,}", uid):
                    targets.add(uid)
    except Exception as exc:
        logging.warning("Load broadcast customer targets failed: %s", exc)
    return sorted(targets, key=lambda x: int(x) if re.fullmatch(r"-?\d+", x) else 0)


def queue_broadcast_message(target_uid: str, message: str, requested_by: str = "", platform: str = "telegram") -> tuple[bool, str]:
    target_uid = str(target_uid or "").strip()
    message = str(message or "").strip()
    if not target_uid or not message:
        return False, "empty_target_or_message"
    mhash = message_hash(message)
    if outbox_exists(platform, target_uid, message):
        return False, "already_sent"
    now = time.time()
    try:
        with db_conn() as conn:
            conn.execute(
                """
                INSERT INTO broadcast_queue(platform, target_uid, message, msg_hash, status, attempts, last_error, requested_by, created_at, updated_at)
                VALUES(?,?,?,?,?,?,?,?,?,?)
                """,
                (platform, target_uid, message, mhash, "queued", 0, "", str(requested_by or ""), now, now),
            )
            conn.commit()
        return True, "queued"
    except sqlite3.IntegrityError:
        return False, "already_queued"
    except Exception as exc:
        logging.warning("Queue broadcast failed target=%s: %s", target_uid, exc)
        return False, f"{type(exc).__name__}:{str(exc)[:120]}"


async def job_process_broadcast_queue(context: ContextTypes.DEFAULT_TYPE):
    try:
        with db_conn() as conn:
            rows = conn.execute(
                """
                SELECT id, platform, target_uid, message, attempts
                FROM broadcast_queue
                WHERE status IN ('queued', 'failed') AND attempts < 3
                ORDER BY created_at ASC
                LIMIT 8
                """
            ).fetchall()
        for row in rows:
            job_id = int(row["id"])
            platform = str(row["platform"] or "telegram")
            target_uid = str(row["target_uid"] or "").strip()
            message = str(row["message"] or "")
            try:
                with db_conn() as conn:
                    conn.execute(
                        "UPDATE broadcast_queue SET status='sending', attempts=attempts+1, updated_at=? WHERE id=?",
                        (time.time(), job_id),
                    )
                    conn.commit()
                if platform != "telegram":
                    raise RuntimeError(f"unsupported_platform:{platform}")
                if outbox_exists(platform, target_uid, message):
                    status = "sent"
                    err = "already_sent"
                else:
                    await send_telegram_long(context.bot, int(target_uid), message)
                    mark_outbox_sent(platform, target_uid, message)
                    status = "sent"
                    err = ""
                with db_conn() as conn:
                    conn.execute(
                        "UPDATE broadcast_queue SET status=?, last_error=?, updated_at=?, sent_at=? WHERE id=?",
                        (status, err, time.time(), time.time() if status == "sent" else None, job_id),
                    )
                    conn.commit()
                await asyncio.sleep(0.8)
            except Exception as exc:
                detail = f"{type(exc).__name__}:{str(exc)[:220]}"
                logging.warning("Broadcast job failed id=%s target=%s: %s", job_id, target_uid, detail)
                with db_conn() as conn:
                    conn.execute(
                        "UPDATE broadcast_queue SET status='failed', last_error=?, updated_at=? WHERE id=?",
                        (detail, time.time(), job_id),
                    )
                    conn.commit()
    except Exception as exc:
        logging.warning("Broadcast queue worker failed: %s", exc)


def get_tele_reply_template(key: str, **kwargs) -> str:
    txt = str(TELE_REPLY_TEMPLATES.get(key) or TELE_REPLY_TEMPLATES_DEFAULT.get(key) or "")
    safe = {
        "WEB_URL": WEB_URL,
        "TARGET_CHAT": TARGET_CHAT,
        "ADMIN_USERNAME": ADMIN_USERNAME,
        "ADMIN_USERNAME_NOSTRIP": ADMIN_USERNAME.replace("@", ""),
        "ACCOUNT": "",
        "PHONE": "",
        "AMOUNT": "",
        "CODE": "",
    }
    safe.update({k: ("" if v is None else str(v)) for k, v in kwargs.items()})
    try:
        return txt.format_map(safe)
    except Exception:
        return txt


def set_tele_reply_template(key: str, value: str) -> bool:
    k = str(key or "").strip()
    if k not in TELE_REPLY_TEMPLATES_DEFAULT:
        return False
    TELE_REPLY_TEMPLATES[k] = str(value or "")
    set_config("tele_reply_templates_json", json.dumps(TELE_REPLY_TEMPLATES, ensure_ascii=False))
    return True


def set_tele_reply_media(key: str, media_ref: str) -> bool:
    k = str(key or "").strip()
    if k not in TELE_REPLY_TEMPLATES_DEFAULT:
        return False
    media = str(media_ref or "").strip()
    if media:
        TELE_REPLY_MEDIA[k] = media
    else:
        TELE_REPLY_MEDIA.pop(k, None)
    set_config("tele_reply_media_json", json.dumps(TELE_REPLY_MEDIA, ensure_ascii=False))
    return True


def get_tele_reply_media(key: str) -> str:
    return str(TELE_REPLY_MEDIA.get(str(key or "").strip()) or "").strip()


async def send_config_media(bot, chat_id, media_ref: str) -> bool:
    ref = str(media_ref or "").strip()
    if not ref:
        return False
    kind, value = "animation", ref
    if ":" in ref and ref.split(":", 1)[0].lower() in {"animation", "photo", "video", "document"}:
        kind, value = ref.split(":", 1)
        kind = kind.lower().strip()
        value = value.strip()
    try:
        if kind == "photo":
            await bot.send_photo(chat_id=chat_id, photo=value)
        elif kind == "video":
            await bot.send_video(chat_id=chat_id, video=value)
        elif kind == "document":
            await bot.send_document(chat_id=chat_id, document=value)
        else:
            try:
                await bot.send_animation(chat_id=chat_id, animation=value)
            except Exception:
                await bot.send_photo(chat_id=chat_id, photo=value)
        return True
    except Exception:
        return False


def media_ref_from_message(msg) -> str:
    if not msg:
        return ""
    try:
        if getattr(msg, "animation", None):
            return "animation:" + str(msg.animation.file_id)
        if getattr(msg, "photo", None):
            return "photo:" + str(msg.photo[-1].file_id)
        if getattr(msg, "video", None):
            return "video:" + str(msg.video.file_id)
        if getattr(msg, "document", None):
            mime = str(getattr(msg.document, "mime_type", "") or "").lower()
            file_id = str(getattr(msg.document, "file_id", "") or "")
            if "gif" in mime or "animation" in mime:
                return "animation:" + file_id
            if mime.startswith("image/"):
                return "photo:" + file_id
            if mime.startswith("video/"):
                return "video:" + file_id
            return "document:" + file_id
    except Exception:
        return ""
    return ""


def text_and_media_from_update(update: Update, command_name: str = "") -> tuple[str, str]:
    msg = update.message
    reply = getattr(msg, "reply_to_message", None) if msg else None
    source = reply or msg
    media = media_ref_from_message(source)
    if reply:
        text = str(getattr(reply, "text", "") or getattr(reply, "caption", "") or "").strip()
        if text:
            return text, media
    if command_name:
        return get_command_tail(update, command_name).strip(), media
    return "", media


TELEGRAM_TEXT_CHUNK_LIMIT = 3800
ZALO_TEXT_CHUNK_LIMIT = 1800


def split_long_text(text: str, limit: int):
    raw = str(text or "")
    if not raw:
        return [""]
    chunks = []
    rest = raw
    while len(rest) > limit:
        cut = rest.rfind("\n", 0, limit)
        if cut < max(1, limit // 3):
            cut = rest.rfind(" ", 0, limit)
        if cut < 1:
            cut = limit
        chunk = rest[:cut].rstrip()
        if chunk:
            chunks.append(chunk)
        rest = rest[cut:].lstrip("\n")
    if rest:
        chunks.append(rest)
    return chunks or [raw[:limit]]


async def send_telegram_long(bot, chat_id, text: str, reply_markup=None, **kwargs):
    chunks = split_long_text(text, TELEGRAM_TEXT_CHUNK_LIMIT)
    last_idx = len(chunks) - 1
    for idx, chunk in enumerate(chunks):
        send_kwargs = dict(kwargs)
        if reply_markup is not None and idx == last_idx:
            send_kwargs["reply_markup"] = reply_markup
        await bot.send_message(chat_id=chat_id, text=chunk, **send_kwargs)
        if idx < last_idx:
            await asyncio.sleep(0.05)


def write_supper_admin_shared_config() -> None:
    try:
        folder = os.path.dirname(SUPPER_ADMIN_SHARED_CONFIG)
        if folder:
            os.makedirs(folder, exist_ok=True)
        data = {
            "chat_id": str(SUPPER_ADMIN_GROUP_ID or "").strip(),
            "link": str(SUPPER_ADMIN_GROUP_LINK or "").strip(),
            "enabled": bool(SUPPER_ADMIN_ENABLED),
            "updated_at": int(time.time()),
        }
        tmp_path = SUPPER_ADMIN_SHARED_CONFIG + ".tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, SUPPER_ADMIN_SHARED_CONFIG)
    except Exception as exc:
        logging.warning("Ghi supper admin shared config that bai: %s", exc)


def write_admin_management_shared_config() -> None:
    try:
        folder = os.path.dirname(ADMIN_MANAGEMENT_SHARED_CONFIG)
        if folder:
            os.makedirs(folder, exist_ok=True)
        data = {
            "chat_id": str(ADMIN_MANAGEMENT_GROUP_ID or "").strip(),
            "link": str(ADMIN_MANAGEMENT_GROUP_LINK or "").strip(),
            "enabled": bool(ADMIN_MANAGEMENT_ENABLED),
            "updated_at": int(time.time()),
        }
        tmp_path = ADMIN_MANAGEMENT_SHARED_CONFIG + ".tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, ADMIN_MANAGEMENT_SHARED_CONFIG)
    except Exception as exc:
            logging.warning("Ghi admin management shared config that bai: %s", exc)


def _read_shared_config(path: str) -> dict:
    try:
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8-sig") as f:
                data = json.load(f)
            return data if isinstance(data, dict) else {}
    except Exception as exc:
        logging.warning("Doc shared config that bai %s: %s", path, exc)
    return {}


def _read_shared_enabled(path: str, current: bool) -> bool:
    data = _read_shared_config(path)
    if "enabled" not in data:
        return current
    return bool(data.get("enabled"))


def refresh_group_enabled_from_shared() -> None:
    global SUPPER_ADMIN_ENABLED, ADMIN_MANAGEMENT_ENABLED
    SUPPER_ADMIN_ENABLED = _read_shared_enabled(SUPPER_ADMIN_SHARED_CONFIG, SUPPER_ADMIN_ENABLED)
    ADMIN_MANAGEMENT_ENABLED = _read_shared_enabled(ADMIN_MANAGEMENT_SHARED_CONFIG, ADMIN_MANAGEMENT_ENABLED)


def get_supper_admin_group_id() -> str:
    chat_id = str(SUPPER_ADMIN_GROUP_ID or "").strip()
    if chat_id:
        return chat_id
    try:
        chat_id = (get_config("supper_admin_group_id", "") or "").strip()
        if chat_id:
            return chat_id
    except Exception:
        pass
    data = _read_shared_config(SUPPER_ADMIN_SHARED_CONFIG)
    return str(data.get("chat_id") or "").strip()


def get_admin_management_group_id() -> str:
    chat_id = str(ADMIN_MANAGEMENT_GROUP_ID or "").strip()
    if chat_id:
        return chat_id
    try:
        chat_id = (get_config("admin_management_group_id", "") or "").strip()
        if chat_id:
            return chat_id
    except Exception:
        pass
    data = _read_shared_config(ADMIN_MANAGEMENT_SHARED_CONFIG)
    chat_id = str(data.get("chat_id") or "").strip()
    if chat_id:
        return chat_id
    return ""


async def send_supper_admin_report(
    bot,
    text: str,
    files: list | None = None,
    caption: str = "",
    source_chat_id: str | int | None = None,
) -> tuple[bool, str]:
    refresh_group_enabled_from_shared()
    if not SUPPER_ADMIN_ENABLED:
        return False, "supper_admin_disabled"
    chat_id = get_supper_admin_group_id()
    if not chat_id:
        return False, "missing_supper_admin_group_id"
    if source_chat_id is not None and str(source_chat_id) == str(chat_id):
        return False, "same_chat"
    files = files or []
    try:
        if text:
            await send_telegram_long(bot, chat_id, text[:12000], disable_web_page_preview=True)
        sent_files = 0
        for item in files:
            try:
                if isinstance(item, tuple) and len(item) == 2:
                    filename, blob = item
                    bio = io.BytesIO(blob)
                    bio.name = str(filename or f"supper_admin_report_{int(time.time())}.txt")
                    await bot.send_document(
                        chat_id=chat_id,
                        document=bio,
                        filename=bio.name,
                        caption=(caption or bio.name)[:1024],
                    )
                    sent_files += 1
                elif isinstance(item, str) and os.path.exists(item):
                    with open(item, "rb") as f:
                        await bot.send_document(
                            chat_id=chat_id,
                            document=f,
                            filename=os.path.basename(item),
                            caption=(caption or os.path.basename(item))[:1024],
                        )
                    sent_files += 1
            except Exception as exc:
                logging.warning("Gui file supper admin that bai: %s", exc)
        try:
            set_config("supper_admin_report_last", f"{int(time.time())}|sent chat={chat_id} files={sent_files}")
        except Exception:
            pass
        return True, f"sent files={sent_files}"
    except Exception as exc:
        detail = f"{type(exc).__name__}:{str(exc)[:180]}"
        try:
            set_config("supper_admin_report_last", f"{int(time.time())}|failed {detail}")
        except Exception:
            pass
        logging.warning("Gui supper admin report that bai: %s", detail)
        return False, detail


async def send_admin_management_report(
    bot,
    text: str,
    files: list | None = None,
    caption: str = "",
    source_chat_id: str | int | None = None,
) -> tuple[bool, str]:
    refresh_group_enabled_from_shared()
    if not ADMIN_MANAGEMENT_ENABLED:
        return False, "admin_management_disabled"
    chat_id = get_admin_management_group_id()
    if not chat_id:
        return False, "missing_admin_management_group_id"
    if source_chat_id is not None and str(source_chat_id) == str(chat_id):
        return False, "same_chat"
    files = files or []
    try:
        if text:
            await send_telegram_long(bot, chat_id, text[:12000], disable_web_page_preview=True)
        sent_files = 0
        for item in files:
            try:
                if isinstance(item, tuple) and len(item) == 2:
                    filename, blob = item
                    bio = io.BytesIO(blob)
                    bio.name = str(filename or f"admin_quanly_report_{int(time.time())}.txt")
                    await bot.send_document(
                        chat_id=chat_id,
                        document=bio,
                        filename=bio.name,
                        caption=(caption or bio.name)[:1024],
                    )
                    sent_files += 1
                elif isinstance(item, str) and os.path.exists(item):
                    with open(item, "rb") as f:
                        await bot.send_document(
                            chat_id=chat_id,
                            document=f,
                            filename=os.path.basename(item),
                            caption=(caption or os.path.basename(item))[:1024],
                        )
                    sent_files += 1
            except Exception as exc:
                logging.warning("Gui file admin quan ly that bai: %s", exc)
        try:
            set_config("admin_management_report_last", f"{int(time.time())}|sent chat={chat_id} files={sent_files}")
        except Exception:
            pass
        return True, f"sent files={sent_files}"
    except Exception as exc:
        detail = f"{type(exc).__name__}:{str(exc)[:180]}"
        try:
            set_config("admin_management_report_last", f"{int(time.time())}|failed {detail}")
        except Exception:
            pass
        logging.warning("Gui admin management report that bai: %s", detail)
        return False, detail


async def send_customer_supergroup_alert(bot, text: str, source_chat_id: str | int | None = None) -> tuple[bool, str]:
    refresh_group_enabled_from_shared()
    if not SUPPER_ADMIN_ENABLED:
        return False, "supper_admin_disabled"
    chat_id = get_supper_admin_group_id()
    if not chat_id:
        return False, "missing_supper_admin_group_id"
    if source_chat_id is not None and str(source_chat_id) == str(chat_id):
        return False, "same_chat"
    try:
        await send_telegram_long(bot, chat_id, text[:3900], disable_web_page_preview=True)
        return True, "sent"
    except Exception as exc:
        detail = f"{type(exc).__name__}:{str(exc)[:180]}"
        logging.warning("Gui customer supergroup alert that bai: %s", detail)
        return False, detail


def mask_token(token: str) -> str:
    t = (token or "").strip()
    if len(t) <= 8:
        return "*" * len(t)
    return t[:4] + "..." + t[-4:]


def mask_sensitive(value: str, keep: int = 2) -> str:
    raw = str(value or "").strip()
    if not raw:
        return "CHUA_SET"
    if len(raw) <= keep * 2:
        return "*" * len(raw)
    return raw[:keep] + "***" + raw[-keep:]


def sensitive_set_label(value: str) -> str:
    return "DA_SET" if str(value or "").strip() else "CHUA_SET"


def should_send_outbox(platform: str, target_uid: str, text: str) -> bool:
    mhash = message_hash(text)
    with db_conn() as conn:
        row = conn.execute(
            "SELECT 1 FROM outbox WHERE platform=? AND target_uid=? AND msg_hash=?",
            (platform, str(target_uid), mhash),
        ).fetchone()
        if row:
            return False
        conn.execute(
            "INSERT INTO outbox(platform, target_uid, msg_hash, created_at) VALUES(?,?,?,?)",
            (platform, str(target_uid), mhash, time.time()),
        )
        conn.commit()
        return True


def save_keyword_reply(keyword: str, content: str):
    with db_conn() as conn:
        conn.execute("INSERT OR REPLACE INTO keyword_replies(keyword, content) VALUES(?,?)", (keyword.lower().strip(), content))
        conn.commit()


def load_keyword_replies():
    with db_conn() as conn:
        rows = conn.execute("SELECT keyword, content FROM keyword_replies ORDER BY keyword").fetchall()
    return [(str(r[0]), str(r[1])) for r in rows]


def extract_links(text: str):
    return re.findall(r"(https?://\S+|t\.me/\S+|zalo\.me/\S+|facebook\.com/\S+)", text or "", flags=re.IGNORECASE)


def normalize_vn_phone(raw: str) -> str:
    digits = "".join(ch for ch in str(raw or "") if ch.isdigit())
    if digits.startswith("0084") and len(digits) >= 13:
        digits = "0" + digits[4:]
    elif digits.startswith("84") and len(digits) >= 11:
        digits = "0" + digits[2:]
    elif len(digits) == 9 and digits[:1] in {"3", "5", "7", "8", "9"}:
        digits = "0" + digits
    return digits


def extract_phones(text: str):
    out = []
    seen = set()
    for m in SCAN_PHONE_RE.finditer(str(text or "")):
        p = normalize_vn_phone(m.group(0))
        if len(p) < 10 or len(p) > 11:
            continue
        if p in seen:
            continue
        seen.add(p)
        out.append(p)
    return out


def sanitize_username(username: str) -> str:
    u = str(username or "").strip()
    if u.startswith("@"):
        u = u[1:]
    return u


CONSENT_OK_VALUES = {"1", "yes", "y", "true", "ok", "opt_in", "opt-in", "consent", "admin_approved"}


def normalize_consent_status(raw: str) -> str:
    text = str(raw or "").strip().lower().replace(" ", "_")
    if text in CONSENT_OK_VALUES:
        return "opt_in"
    if text in {"0", "no", "n", "false", "blocked", "opt_out", "opt-out", "deny"}:
        return "opt_out"
    return text or "unknown"


def row_has_explicit_consent(row: dict) -> bool:
    status = normalize_consent_status(
        row.get("consent")
        or row.get("consent_status")
        or row.get("status")
        or row.get("opt_in")
        or ""
    )
    return status == "opt_in"


def mark_campaign_target_consent(uid: str, username: str = "", phone: str = "", source: str = "telegram_start", status: str = "opt_in") -> None:
    uid = str(uid or "").strip()
    if not uid:
        return
    username = sanitize_username(username or "")
    phone = normalize_vn_phone(phone or "")
    status = normalize_consent_status(status)
    with db_conn() as conn:
        conn.execute(
            """
            INSERT OR REPLACE INTO campaign_target_consent(platform, target_uid, username, phone, source, status, updated_at)
            VALUES('telegram',?,?,?,?,?,?)
            """,
            (uid, username, phone, str(source or "")[:160], status, time.time()),
        )
        conn.commit()


def lookup_campaign_target_consent(row: dict) -> str:
    uid = str(row.get("uid") or row.get("target_uid") or "").strip()
    username = sanitize_username(row.get("username") or row.get("user") or "")
    phone = normalize_vn_phone(row.get("phone") or "")
    with db_conn() as conn:
        if uid:
            hit = conn.execute(
                "SELECT status FROM campaign_target_consent WHERE platform='telegram' AND target_uid=?",
                (uid,),
            ).fetchone()
            if hit:
                return normalize_consent_status(hit["status"])
        if username:
            hit = conn.execute(
                "SELECT status FROM campaign_target_consent WHERE platform='telegram' AND lower(username)=lower(?) ORDER BY updated_at DESC LIMIT 1",
                (username,),
            ).fetchone()
            if hit:
                return normalize_consent_status(hit["status"])
        if phone:
            hit = conn.execute(
                "SELECT status FROM campaign_target_consent WHERE platform='telegram' AND phone=? ORDER BY updated_at DESC LIMIT 1",
                (phone,),
            ).fetchone()
            if hit:
                return normalize_consent_status(hit["status"])
    return "unknown"


def annotate_campaign_recipients(rows: list[dict], source: str = "") -> list[dict]:
    annotated = []
    for row in rows or []:
        item = dict(row or {})
        item.setdefault("source", source or item.get("source") or "")
        explicit = row_has_explicit_consent(item)
        consent = "opt_in" if explicit else lookup_campaign_target_consent(item)
        item["consent"] = consent
        item["status"] = "sendable" if consent == "opt_in" else "missing_consent"
        annotated.append(item)
    return annotated


def split_campaign_consent_rows(rows: list[dict]) -> tuple[list[dict], list[dict]]:
    ok, missing = [], []
    for row in rows or []:
        (ok if normalize_consent_status(row.get("consent") or row.get("status") or "") == "opt_in" else missing).append(row)
    return ok, missing


def campaign_group_allow_tokens() -> set[str]:
    raw = "\n".join(
        x for x in [
            CAMPAIGN_ALLOWED_GROUP_IDS,
            CAMPAIGN_DEFAULT_GROUP_LINK,
            TARGET_CHAT,
            SUPPER_ADMIN_GROUP_ID,
            ADMIN_MANAGEMENT_GROUP_ID,
        ] if x
    )
    tokens = set()
    for part in re.split(r"[\s,;]+", raw):
        token = part.strip()
        if token:
            tokens.add(token.lower())
    return tokens


def campaign_group_allowed(target_link: str) -> bool:
    target = str(target_link or "").strip().lower()
    if not target:
        return False
    tokens = campaign_group_allow_tokens()
    if target in tokens:
        return True
    return any(token and (token in target or target in token) for token in tokens)


def render_campaign_start_text(user) -> str:
    first_name = str(getattr(user, "first_name", "") or "bạn").strip() or "bạn"
    username = sanitize_username(str(getattr(user, "username", "") or ""))
    group_link = CAMPAIGN_DEFAULT_GROUP_LINK or TARGET_CHAT or ""
    register_link = CAMPAIGN_REGISTER_LINK or WEB_URL or ""
    support_username = (CAMPAIGN_SUPPORT_USERNAME or ADMIN_USERNAME or "gifhub2708").strip().lstrip("@")
    template = CAMPAIGN_START_TEMPLATE or CAMPAIGN_START_TEMPLATE_DEFAULT
    values = {
        "first_name": first_name,
        "username": username,
        "group_link": group_link,
        "register_link": register_link,
        "support_username": support_username,
        "admin_username": support_username,
        "target_chat": TARGET_CHAT,
        "web_url": WEB_URL,
        "link_nhap_code": LINK_NHAP_CODE,
    }
    try:
        return template.format(**values)
    except Exception:
        return CAMPAIGN_START_TEMPLATE_DEFAULT.format(**values)


def build_customer_link(platform: str, uid: str = "", username: str = "", phone: str = "") -> str:
    pf = str(platform or "").strip().lower()
    uid = str(uid or "").strip()
    user = sanitize_username(username)
    phone_n = normalize_vn_phone(phone or "")

    if pf.startswith("telegram"):
        if user:
            return f"https://t.me/{user}"
        if uid.isdigit():
            return f"tg://user?id={uid}"
        if uid.startswith("-100") and uid[4:].isdigit():
            return f"https://t.me/c/{uid[4:]}"
        if uid.startswith("-") and uid[1:].isdigit():
            return f"https://t.me/c/{uid[1:]}"
        return ""

    if pf.startswith("zalo"):
        if phone_n and _is_valid_phone(phone_n):
            return f"https://zalo.me/{phone_n}"
        if uid.isdigit():
            return f"https://zalo.me/{uid}"
        return ""

    if user:
        return f"https://t.me/{user}"
    return ""


def choose_best_lookup_hit(hits: list[dict]) -> dict:
    if not hits:
        return {}
    # Uu tien hit co uid/user, sau do den hit moi nhat.
    ranked = sorted(
        hits,
        key=lambda h: (
            0 if (str(h.get("uid", "")).strip() or str(h.get("user", "")).strip()) else 1,
            -int(float(h.get("last_seen", h.get("updated_at", 0)) or 0)),
        ),
    )
    return ranked[0] if ranked else {}


def save_scanned_lead(phone: str, username: str = "", full_name: str = "", source_uid: str = "", source_chat: str = "", source_type: str = "manual") -> bool:
    try:
        with db_conn() as conn:
            cur = conn.execute(
                """
                INSERT OR IGNORE INTO scanned_phone_leads(phone, username, full_name, source_uid, source_chat, source_type, created_at)
                VALUES(?,?,?,?,?,?,?)
                """,
                (str(phone), (username or "")[:150], (full_name or "")[:150], str(source_uid), str(source_chat), str(source_type), time.time()),
            )
            conn.commit()
            return int(cur.rowcount or 0) > 0
    except Exception:
        return False


def _is_valid_phone(phone: str) -> bool:
    p = normalize_vn_phone(phone)
    return p.isdigit() and 10 <= len(p) <= 11 and p.startswith("0")


def extract_candidate_phones(text: str, extra_phone: str = ""):
    phones = []
    seen = set()
    for p in extract_phones(text or ""):
        if _is_valid_phone(p) and p not in seen:
            seen.add(p)
            phones.append(p)
    if extra_phone:
        p2 = normalize_vn_phone(extra_phone)
        if _is_valid_phone(p2) and p2 not in seen:
            seen.add(p2)
            phones.append(p2)
    return phones


def save_phonebook_entry(
    phone: str,
    platform: str,
    source_uid: str = "",
    username: str = "",
    full_name: str = "",
    source_chat: str = "",
    last_message_snippet: str = "",
) -> tuple[bool, bool]:
    now = time.time()
    p = normalize_vn_phone(phone)
    if not _is_valid_phone(p):
        return (False, False)
    try:
        with db_conn() as conn:
            existed_phone = conn.execute(
                "SELECT 1 FROM bot_phonebook WHERE phone=? LIMIT 1",
                (p,),
            ).fetchone() is not None
            existed_uid = False
            if source_uid:
                existed_uid = conn.execute(
                    "SELECT 1 FROM bot_phonebook WHERE source_uid=? LIMIT 1",
                    (str(source_uid),),
                ).fetchone() is not None
            conn.execute(
                """
                INSERT INTO bot_phonebook(
                    phone, platform, source_uid, username, full_name, source_chat, first_seen, last_seen, hit_count, last_message_snippet
                )
                VALUES(?,?,?,?,?,?,?,?,1,?)
                ON CONFLICT(phone, platform, source_uid, source_chat) DO UPDATE SET
                    username=CASE WHEN excluded.username!='' THEN excluded.username ELSE bot_phonebook.username END,
                    full_name=CASE WHEN excluded.full_name!='' THEN excluded.full_name ELSE bot_phonebook.full_name END,
                    last_seen=excluded.last_seen,
                    hit_count=bot_phonebook.hit_count+1,
                    last_message_snippet=CASE WHEN excluded.last_message_snippet!='' THEN excluded.last_message_snippet ELSE bot_phonebook.last_message_snippet END
                """,
                (
                    p,
                    (platform or "")[:40],
                    str(source_uid or ""),
                    (username or "")[:150],
                    (full_name or "")[:150],
                    str(source_chat or ""),
                    now,
                    now,
                    (last_message_snippet or "")[:350],
                ),
            )
            conn.commit()
            return (not existed_phone, (not existed_uid and bool(source_uid)))
    except Exception:
        return (False, False)


def log_phone_event(
    phone: str,
    platform: str,
    source_uid: str = "",
    username: str = "",
    full_name: str = "",
    source_chat: str = "",
    is_new_phone: bool = False,
    is_new_uid: bool = False,
) -> None:
    p = normalize_vn_phone(phone)
    if not _is_valid_phone(p):
        return
    try:
        with db_conn() as conn:
            conn.execute(
                """
                INSERT INTO phone_events(
                    phone, platform, source_uid, username, full_name, source_chat, is_new_phone, is_new_uid, created_at
                )
                VALUES(?,?,?,?,?,?,?,?,?)
                """,
                (
                    p,
                    (platform or "")[:40],
                    str(source_uid or ""),
                    (username or "")[:150],
                    (full_name or "")[:150],
                    str(source_chat or ""),
                    1 if is_new_phone else 0,
                    1 if is_new_uid else 0,
                    time.time(),
                ),
            )
            conn.commit()
    except Exception:
        pass


def track_phone_from_message(
    platform: str,
    source_uid: str,
    username: str,
    full_name: str,
    source_chat: str,
    text: str,
    extra_phone: str = "",
) -> dict:
    phones = extract_candidate_phones(text, extra_phone=extra_phone)
    new_phone = 0
    new_uid = 0
    for p in phones:
        is_new_phone, is_new_uid = save_phonebook_entry(
            phone=p,
            platform=platform,
            source_uid=source_uid,
            username=username,
            full_name=full_name,
            source_chat=source_chat,
            last_message_snippet=text,
        )
        if is_new_phone:
            new_phone += 1
        if is_new_uid:
            new_uid += 1
        log_phone_event(
            phone=p,
            platform=platform,
            source_uid=source_uid,
            username=username,
            full_name=full_name,
            source_chat=source_chat,
            is_new_phone=is_new_phone,
            is_new_uid=is_new_uid,
        )
    return {
        "phones": phones,
        "new_phone": new_phone,
        "new_uid": new_uid,
    }


def build_group_channel_summary_lines(source_chats: set[str], limit_chats: int = 6, limit_users: int = 8) -> list[str]:
    chats = [str(c).strip() for c in (source_chats or set()) if str(c).strip()]
    try:
        if not chats:
            with db_conn() as conn:
                fallback_rows = conn.execute(
                    """
                    SELECT uid
                    FROM customer_contacts
                    WHERE platform IN ('telegram_group', 'telegram_channel', 'zalo_group')
                    ORDER BY updated_at DESC
                    LIMIT 24
                    """
                ).fetchall()
            chats = [str(r["uid"] or "").strip() for r in fallback_rows if str(r["uid"] or "").strip()]
            if not chats:
                return []
        with db_conn() as conn:
            ph = ",".join(["?"] * len(chats))
            meta_rows = conn.execute(
                f"""
                SELECT uid, platform, name, username
                FROM customer_contacts
                WHERE uid IN ({ph})
                  AND platform IN ('telegram_group', 'telegram_channel', 'zalo_group')
                """,
                chats,
            ).fetchall()
            if not meta_rows:
                return []
            meta_map = {
                str(r["uid"]): {
                    "platform": str(r["platform"] or ""),
                    "name": str(r["name"] or ""),
                    "username": str(r["username"] or ""),
                }
                for r in meta_rows
            }
            stat_rows = conn.execute(
                f"""
                SELECT
                    source_chat,
                    COUNT(DISTINCT CASE WHEN source_uid!='' THEN source_uid END) AS total_members,
                    COUNT(DISTINCT CASE WHEN source_uid!='' AND phone!='' THEN source_uid END) AS members_with_phone
                FROM bot_phonebook
                WHERE source_chat IN ({ph})
                GROUP BY source_chat
                """,
                chats,
            ).fetchall()
            stat_map = {
                str(r["source_chat"]): (
                    int(r["total_members"] or 0),
                    int(r["members_with_phone"] or 0),
                )
                for r in stat_rows
            }
            user_rows = conn.execute(
                f"""
                SELECT
                    source_chat,
                    source_uid,
                    MAX(last_seen) AS last_seen,
                    MAX(username) AS username,
                    MAX(full_name) AS full_name
                FROM bot_phonebook
                WHERE source_chat IN ({ph})
                  AND source_uid!=''
                GROUP BY source_chat, source_uid
                ORDER BY last_seen DESC
                """,
                chats,
            ).fetchall()

        users_by_chat: dict[str, list[tuple[float, str]]] = {}
        for r in user_rows:
            chat = str(r["source_chat"] or "")
            uname = sanitize_username(str(r["username"] or ""))
            fname = str(r["full_name"] or "").strip()
            suid = str(r["source_uid"] or "").strip()
            if uname:
                disp = f"@{uname}"
            elif fname:
                disp = fname
            else:
                disp = suid or "N/A"
            users_by_chat.setdefault(chat, []).append((float(r["last_seen"] or 0), disp))

        order_chats = sorted(
            list(meta_map.keys()),
            key=lambda c: (stat_map.get(c, (0, 0))[1], stat_map.get(c, (0, 0))[0]),
            reverse=True,
        )
        lines: list[str] = []
        for chat in order_chats[: max(1, int(limit_chats))]:
            meta = meta_map.get(chat, {})
            platform = str(meta.get("platform", "") or "")
            name = str(meta.get("name", "") or "")
            username = str(meta.get("username", "") or "")
            total_members, members_with_phone = stat_map.get(chat, (0, 0))
            chat_link = build_customer_link(platform=platform, uid=chat, username=username)
            user_list = sorted(users_by_chat.get(chat, []), key=lambda it: it[0], reverse=True)
            user_text = ", ".join([x[1] for x in user_list[: max(1, int(limit_users))]]) if user_list else "N/A"
            lines.append(
                f"- {(name or chat)} | {platform} | link: {chat_link or 'N/A'} | tong_tv: {total_members} | tv_co_sdt: {members_with_phone} | users: {user_text}"
            )
        return lines
    except Exception:
        return []


def maybe_send_phone_event_summary(force: bool = False) -> None:
    try:
        now = time.time()
        last_ts = float(get_config("phone_event_summary_last_ts", "0") or "0")
        if (not force) and now - last_ts < PHONE_EVENT_SUMMARY_INTERVAL_SEC:
            return
        with db_conn() as conn:
            rows = conn.execute(
                """
                SELECT phone, source_uid, is_new_phone, is_new_uid, created_at
                     , source_chat
                FROM phone_events
                WHERE created_at > ?
                ORDER BY created_at DESC
                LIMIT 2000
                """,
                (last_ts,),
            ).fetchall()
        if not rows:
            if force:
                set_config("phone_event_summary_last_ts", str(now))
            return
        total_events = len(rows)
        unique_phones = len({str(r["phone"] or "") for r in rows if str(r["phone"] or "")})
        unique_uids = len({str(r["source_uid"] or "") for r in rows if str(r["source_uid"] or "")})
        new_phone_count = sum(int(r["is_new_phone"] or 0) for r in rows)
        new_uid_count = sum(int(r["is_new_uid"] or 0) for r in rows)
        counts: dict[str, int] = {}
        for r in rows:
            p = str(r["phone"] or "").strip()
            if not p:
                continue
            counts[p] = counts.get(p, 0) + 1
        top_list = sorted(counts.items(), key=lambda kv: kv[1], reverse=True)[:8]
        top_lines = [f"- {p}: {c} lan" for p, c in top_list] if top_list else ["- N/A"]
        source_chats = {str(r["source_chat"] or "").strip() for r in rows if str(r["source_chat"] or "").strip()}
        chat_lines = build_group_channel_summary_lines(source_chats, limit_chats=6, limit_users=8)

        text_lines = [
            "📦 BAO CAO TONG HOP SDT (5 PHUT)",
            f"- So event co SDT: {total_events}",
            f"- So SDT duy nhat: {unique_phones}",
            f"- So UID duy nhat: {unique_uids}",
            f"- SDT moi: {new_phone_count}",
            f"- UID moi: {new_uid_count}",
            "- Top SDT vua xuat hien:",
            *top_lines,
        ]
        if chat_lines:
            text_lines.append("- Chi tiet nhom/kenh:")
            text_lines.extend(chat_lines)
        notify_admin_sync("\n".join(text_lines)[:3900])
        set_config("phone_event_summary_last_ts", str(now))
    except Exception:
        pass


def _phones_from_contact_value(raw_phone: str):
    raw = str(raw_phone or "").strip()
    if not raw:
        return []
    phones = extract_candidate_phones(raw)
    if phones:
        return phones
    p = normalize_vn_phone(raw)
    if _is_valid_phone(p):
        return [p]
    return []


def build_phone_lookup_index() -> dict:
    lookup: dict[str, list[dict]] = {}
    seen = set()
    try:
        with db_conn() as conn:
            rows = conn.execute(
                """
                SELECT uid, platform, name, username, phone, updated_at
                FROM customer_contacts
                WHERE phone!=''
                ORDER BY updated_at DESC
                LIMIT 8000
                """
            ).fetchall()
            for r in rows:
                for p in _phones_from_contact_value(r["phone"]):
                    key = (p, str(r["uid"] or ""), str(r["platform"] or ""))
                    if key in seen:
                        continue
                    seen.add(key)
                    lookup.setdefault(p, []).append(
                        {
                            "user": str(r["username"] or ""),
                            "uid": str(r["uid"] or ""),
                            "ten": str(r["name"] or ""),
                            "platform": str(r["platform"] or ""),
                            "status": "matched",
                            "phone": p,
                            "link": build_customer_link(
                                platform=str(r["platform"] or ""),
                                uid=str(r["uid"] or ""),
                                username=str(r["username"] or ""),
                                phone=p,
                            ),
                            "updated_at": float(r["updated_at"] or 0),
                        }
                    )
            rows2 = conn.execute(
                """
                SELECT phone, platform, source_uid, username, full_name, first_seen, last_seen
                FROM bot_phonebook
                ORDER BY last_seen DESC
                LIMIT 12000
                """
            ).fetchall()
            for r in rows2:
                p = normalize_vn_phone(str(r["phone"] or ""))
                if not _is_valid_phone(p):
                    continue
                uid = str(r["source_uid"] or "")
                platform = str(r["platform"] or "")
                key = (p, uid, platform)
                if key in seen:
                    continue
                seen.add(key)
                status = "matched" if uid else "known_phone_only"
                lookup.setdefault(p, []).append(
                    {
                        "user": str(r["username"] or ""),
                        "uid": uid,
                        "ten": str(r["full_name"] or ""),
                        "platform": platform,
                        "status": status,
                        "phone": p,
                        "link": build_customer_link(
                            platform=platform,
                            uid=uid,
                            username=str(r["username"] or ""),
                            phone=p,
                        ),
                        "first_seen": float(r["first_seen"] or 0),
                        "last_seen": float(r["last_seen"] or 0),
                    }
                )
    except Exception:
        pass
    return lookup


def _decode_bytes_best_effort(data: bytes) -> str:
    if not data:
        return ""
    for enc in ("utf-8-sig", "utf-8", "utf-16", "utf-16-le", "utf-16-be", "cp1258", "cp1252", "latin1"):
        try:
            text = data.decode(enc)
            if text:
                return text
        except Exception:
            continue
    try:
        import chardet

        guess = chardet.detect(data[: min(len(data), 1024 * 1024)])
        enc = guess.get("encoding")
        if enc:
            return data.decode(enc, errors="ignore")
    except Exception:
        pass
    return data.decode("utf-8", errors="ignore")


def _binary_strings(data: bytes, min_len: int = 4) -> str:
    if not data:
        return ""
    chunks = re.findall(rb"[\x09\x0a\x0d\x20-\x7e]{" + str(min_len).encode("ascii") + rb",}", data)
    return "\n".join(x.decode("latin1", errors="ignore") for x in chunks[:50000])


def _dedupe_keep_order(values) -> list[str]:
    out = []
    seen = set()
    for value in values or []:
        v = str(value or "").strip()
        if not v:
            continue
        key = v.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(v)
    return out


def extract_scan_entities(text: str) -> dict:
    raw = str(text or "")
    phones = extract_candidate_phones(raw)
    phone_set = set(phones)

    uids = []
    for m in SCAN_UID_RE.finditer(raw):
        uid = m.group(0).strip()
        digits = uid.lstrip("-")
        if not digits.isdigit():
            continue
        normalized_as_phone = normalize_vn_phone(uid)
        if normalized_as_phone in phone_set:
            continue
        if uid.startswith("-100") or len(digits) >= 6:
            uids.append(uid)

    users = []
    for m in SCAN_HANDLE_RE.finditer(raw):
        user = next((g for g in m.groups() if g), "")
        user = sanitize_username(user)
        if not user:
            continue
        if user.lower() in {"http", "https", "www", "com", "zalo", "telegram", "facebook", "tiktok"}:
            continue
        users.append(user)

    return {
        "phones": phones,
        "uids": _dedupe_keep_order(uids),
        "users": _dedupe_keep_order(users),
    }


def build_identity_lookup_index() -> tuple[dict[str, dict], dict[str, dict]]:
    uid_lookup: dict[str, dict] = {}
    user_lookup: dict[str, dict] = {}

    def put(item: dict):
        uid = str(item.get("uid") or "").strip()
        user = sanitize_username(str(item.get("user") or item.get("username") or ""))
        platform = str(item.get("platform") or "").strip()
        phone = str(item.get("phone") or "").strip()
        if not item.get("link"):
            item["link"] = build_customer_link(platform=platform, uid=uid, username=user, phone=phone)
        if uid and uid not in uid_lookup:
            uid_lookup[uid] = item
        if user:
            key = user.lower()
            if key not in user_lookup:
                user_lookup[key] = item

    try:
        with db_conn() as conn:
            rows = conn.execute(
                """
                SELECT uid, platform, name, username, phone, updated_at
                FROM customer_contacts
                WHERE uid!='' OR username!=''
                ORDER BY updated_at DESC
                LIMIT 30000
                """
            ).fetchall()
            for r in rows:
                put(
                    {
                        "uid": str(r["uid"] or ""),
                        "user": str(r["username"] or ""),
                        "ten": str(r["name"] or ""),
                        "platform": str(r["platform"] or ""),
                        "phone": str(r["phone"] or ""),
                        "status": "matched",
                        "updated_at": float(r["updated_at"] or 0),
                    }
                )
            rows2 = conn.execute(
                """
                SELECT phone, platform, source_uid, username, full_name, last_seen
                FROM bot_phonebook
                WHERE source_uid!='' OR username!=''
                ORDER BY last_seen DESC
                LIMIT 30000
                """
            ).fetchall()
            for r in rows2:
                put(
                    {
                        "uid": str(r["source_uid"] or ""),
                        "user": str(r["username"] or ""),
                        "ten": str(r["full_name"] or ""),
                        "platform": str(r["platform"] or ""),
                        "phone": str(r["phone"] or ""),
                        "status": "matched",
                        "updated_at": float(r["last_seen"] or 0),
                    }
                )
            rows3 = conn.execute(
                """
                SELECT chat_id, user_id, name, username, last_seen
                FROM telegram_group_members
                WHERE user_id!='' OR username!=''
                ORDER BY last_seen DESC
                LIMIT 30000
                """
            ).fetchall()
            for r in rows3:
                put(
                    {
                        "uid": str(r["user_id"] or ""),
                        "user": str(r["username"] or ""),
                        "ten": str(r["name"] or ""),
                        "platform": "telegram_user",
                        "phone": "",
                        "source_chat": str(r["chat_id"] or ""),
                        "status": "matched",
                        "updated_at": float(r["last_seen"] or 0),
                    }
                )
    except Exception:
        pass
    return uid_lookup, user_lookup


def _read_xlsx_bytes(data: bytes) -> str:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"--- sheet: {ws.title} ---")
            for row in ws.iter_rows(values_only=True):
                vals = [str(v) for v in row if v is not None]
                if vals:
                    parts.append("\t".join(vals))
        try:
            wb.close()
        except Exception:
            pass
        return "\n".join(parts)
    except Exception:
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                parts = []
                for name in zf.namelist():
                    if name.startswith("xl/") and name.endswith(".xml") and (
                        "sharedStrings" in name or "worksheets/sheet" in name
                    ):
                        parts.append(_decode_bytes_best_effort(zf.read(name)))
                return "\n".join(parts)
        except Exception:
            return _binary_strings(data)


def _read_xls_bytes(data: bytes) -> str:
    try:
        import xlrd

        book = xlrd.open_workbook(file_contents=data)
        parts = []
        for sheet in book.sheets():
            parts.append(f"--- sheet: {sheet.name} ---")
            for r in range(sheet.nrows):
                vals = []
                for c in range(sheet.ncols):
                    v = sheet.cell_value(r, c)
                    if v not in ("", None):
                        vals.append(str(v))
                if vals:
                    parts.append("\t".join(vals))
        return "\n".join(parts)
    except Exception:
        return _binary_strings(data)


def _read_docx_bytes(data: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                vals = [cell.text for cell in row.cells if cell.text]
                if vals:
                    parts.append("\t".join(vals))
        return "\n".join(parts)
    except Exception:
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                parts = []
                for name in zf.namelist():
                    if name.startswith("word/") and name.endswith(".xml"):
                        txt = _decode_bytes_best_effort(zf.read(name))
                        txt = re.sub(r"<[^>]+>", " ", txt)
                        parts.append(html.unescape(txt))
                return "\n".join(parts)
        except Exception:
            return _binary_strings(data)


def _read_pdf_bytes(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages[:500]:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts)
    except Exception:
        return _binary_strings(data)


def _read_rtf_bytes(data: bytes) -> str:
    text = _decode_bytes_best_effort(data)
    try:
        from striprtf.striprtf import rtf_to_text

        return rtf_to_text(text)
    except Exception:
        return text


def _read_html_bytes(data: bytes) -> str:
    text = _decode_bytes_best_effort(data)
    try:
        from bs4 import BeautifulSoup

        return BeautifulSoup(text, "html.parser").get_text("\n")
    except Exception:
        return re.sub(r"<[^>]+>", " ", text)


def _read_sqlite_bytes(data: bytes) -> str:
    tmp_path = ""
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".db")
        with os.fdopen(fd, "wb") as fp:
            fp.write(data)
        parts = []
        with sqlite3.connect(tmp_path) as conn:
            conn.row_factory = sqlite3.Row
            tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
            for tr in tables:
                table = str(tr["name"])
                if table.startswith("sqlite_"):
                    continue
                cols = [r["name"] for r in conn.execute(f'PRAGMA table_info("{table.replace(chr(34), chr(34) + chr(34))}")')]
                if not cols:
                    continue
                qcols = ",".join(f'"{c.replace(chr(34), chr(34) + chr(34))}"' for c in cols[:50])
                rows = conn.execute(
                    f'SELECT {qcols} FROM "{table.replace(chr(34), chr(34) + chr(34))}" LIMIT ?',
                    (SCAN_SQLITE_MAX_ROWS_PER_TABLE,),
                ).fetchall()
                for row in rows:
                    vals = [str(row[c]) for c in cols[:50] if row[c] not in ("", None)]
                    if vals:
                        parts.append("\t".join(vals))
        return "\n".join(parts)
    except Exception:
        return _binary_strings(data)
    finally:
        if tmp_path:
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def _read_image_bytes(data: bytes) -> tuple[str, str]:
    try:
        from PIL import Image
        import pytesseract

        text = pytesseract.image_to_string(Image.open(io.BytesIO(data)), lang="vie+eng")
        return text or "", ""
    except Exception as exc:
        return "", f"OCR ảnh chưa chạy được trên máy này: {exc}"


def extract_text_from_file_bytes(data: bytes, filename: str = "", depth: int = 0) -> tuple[str, list[str]]:
    notes = []
    if not data:
        return "", ["File rỗng."]
    if len(data) > SCAN_FILE_MAX_BYTES:
        data = data[:SCAN_FILE_MAX_BYTES]
        notes.append(f"File lớn, chỉ đọc {SCAN_FILE_MAX_BYTES // (1024 * 1024)}MB đầu.")

    name = str(filename or "upload")
    low = name.lower()
    ext = os.path.splitext(low)[1]

    try:
        if ext in {".zip"} and depth < 2:
            parts = []
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for idx, info in enumerate(zf.infolist()):
                    if idx >= SCAN_ARCHIVE_MAX_FILES:
                        notes.append(f"Archive quá nhiều file, chỉ đọc {SCAN_ARCHIVE_MAX_FILES} file đầu.")
                        break
                    if info.is_dir() or info.file_size > SCAN_FILE_MAX_BYTES:
                        continue
                    child_text, child_notes = extract_text_from_file_bytes(zf.read(info), info.filename, depth + 1)
                    if child_text:
                        parts.append(f"\n--- {info.filename} ---\n{child_text}")
                    notes.extend(child_notes[:5])
            return "\n".join(parts), notes
        if ext in {".7z"} and depth < 2:
            try:
                import py7zr

                parts = []
                with tempfile.TemporaryDirectory() as td:
                    archive = os.path.join(td, "in.7z")
                    with open(archive, "wb") as fp:
                        fp.write(data)
                    out_dir = os.path.join(td, "out")
                    os.makedirs(out_dir, exist_ok=True)
                    with py7zr.SevenZipFile(archive, "r") as zf:
                        zf.extractall(out_dir)
                    count = 0
                    for root, _, files in os.walk(out_dir):
                        for fn in files:
                            if count >= SCAN_ARCHIVE_MAX_FILES:
                                break
                            path = os.path.join(root, fn)
                            if os.path.getsize(path) > SCAN_FILE_MAX_BYTES:
                                continue
                            with open(path, "rb") as fp:
                                child_text, child_notes = extract_text_from_file_bytes(fp.read(), fn, depth + 1)
                            if child_text:
                                parts.append(f"\n--- {fn} ---\n{child_text}")
                            notes.extend(child_notes[:5])
                            count += 1
                    return "\n".join(parts), notes
            except Exception as exc:
                notes.append(f"Không đọc được 7z: {exc}")
        if ext in {".rar"}:
            notes.append("RAR cần unrar/bsdtar trên máy; bot sẽ quét chuỗi thô trong file.")
    except Exception as exc:
        notes.append(f"Lỗi đọc archive: {exc}")

    if ext in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return _read_xlsx_bytes(data), notes
    if ext in {".xls"}:
        return _read_xls_bytes(data), notes
    if ext in {".ods"}:
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                parts = [_decode_bytes_best_effort(zf.read(n)) for n in zf.namelist() if n.endswith(".xml")]
            return "\n".join(parts), notes
        except Exception:
            return _binary_strings(data), notes
    if ext in {".docx"}:
        return _read_docx_bytes(data), notes
    if ext in {".pdf"}:
        return _read_pdf_bytes(data), notes
    if ext in {".rtf"}:
        return _read_rtf_bytes(data), notes
    if ext in {".html", ".htm"}:
        return _read_html_bytes(data), notes
    if ext in {".db", ".sqlite", ".sqlite3"}:
        return _read_sqlite_bytes(data), notes
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        text, note = _read_image_bytes(data)
        if note:
            notes.append(note)
        return text, notes
    if ext in {".txt", ".csv", ".json", ".xml", ".log", ".md", ".tsv", ".vcf", ".ini", ".yaml", ".yml"}:
        return _decode_bytes_best_effort(data), notes

    text = _decode_bytes_best_effort(data)
    if len(text.strip()) < 20:
        text = f"{text}\n{_binary_strings(data)}"
    return text, notes


def extract_text_from_local_file(path: str) -> tuple[str, list[str]]:
    with open(path, "rb") as fp:
        data = fp.read()
    return extract_text_from_file_bytes(data, os.path.basename(path))

def get_zalo_admin_all():
    try:
        return set(json.loads(get_config("zalo_admin_all_json", "[]") or "[]"))
    except Exception:
        return set()


def set_zalo_admin_all(vals):
    set_config("zalo_admin_all_json", json.dumps(sorted(list(vals))))


def get_zalo_admin_ai():
    try:
        return set(json.loads(get_config("zalo_admin_ai_json", "[]") or "[]"))
    except Exception:
        return set()


def set_zalo_admin_ai(vals):
    set_config("zalo_admin_ai_json", json.dumps(sorted(list(vals))))


def uid_with_name(uid: str) -> str:
    with db_conn() as conn:
        row = conn.execute(
            "SELECT name, username, platform FROM customer_contacts WHERE uid=? ORDER BY updated_at DESC LIMIT 1",
            (str(uid),),
        ).fetchone()
    if not row:
        return f"{uid} | ten: (chua ro)"
    nm = (row["name"] or "").strip() or "(chua ro)"
    un = (row["username"] or "").strip()
    pf = (row["platform"] or "").strip()
    extra = f" | user: {un}" if un else ""
    return f"{uid} | ten: {nm}{extra} | kenh: {pf}"


def get_last_scan_report_ts(platform: str) -> float:
    try:
        return float(get_config(f"scan_report_last_{platform}", "0") or "0")
    except Exception:
        return 0.0


def set_last_scan_report_ts(platform: str, ts: float):
    set_config(f"scan_report_last_{platform}", str(float(ts)))


def list_recent_contacts(platform_prefix: str, since_ts: float, limit: int = 50):
    with db_conn() as conn:
        rows = conn.execute(
            """
            SELECT uid, platform, name, username, phone, updated_at
            FROM customer_contacts
            WHERE platform LIKE ? AND updated_at > ?
            ORDER BY updated_at DESC
            LIMIT ?
            """,
            (f"{platform_prefix}%", float(since_ts), int(limit)),
        ).fetchall()
    return rows


def bootstrap_scan_from_cache():
    try:
        for uid, info in USER_DATA.items():
            upsert_contact(
                "telegram_user",
                str(uid),
                name=str((info or {}).get("ten_game", "")),
                phone=str((info or {}).get("sdt", "")),
            )
    except Exception:
        pass


def import_contacts_from_backup_dir() -> int:
    try:
        if not os.path.isdir(CONTACT_BACKUP_DIR):
            return 0
        files = []
        for fn in os.listdir(CONTACT_BACKUP_DIR):
            if fn.lower().endswith(".json") and fn.lower().startswith("contacts_"):
                files.append(os.path.join(CONTACT_BACKUP_DIR, fn))
        if not files:
            return 0
        latest = max(files, key=lambda p: os.path.getmtime(p))
        with open(latest, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, list):
            return 0
        count = 0
        for row in data:
            if not isinstance(row, dict):
                continue
            uid = str(row.get("uid") or "").strip()
            if not uid:
                continue
            upsert_contact(
                str(row.get("platform") or "zalo"),
                uid,
                name=str(row.get("name") or ""),
                username=str(row.get("username") or ""),
                phone=str(row.get("phone") or ""),
            )
            count += 1
        if count > 0:
            logging.info("[CONTACT-RESTORE] imported=%s from %s", count, latest)
        return count
    except Exception as exc:
        logging.warning("Import contacts backup that bai: %s", exc)
        return 0


def call_ai_text(prompt: str, env_name: str | None = None) -> str:
    env = resolve_ai_env(env_name)

    def _try_env(env_obj: dict) -> tuple[int, str]:
        token_local = str(env_obj.get("token") or "").strip()
        if not token_local:
            provider = str(env_obj.get("provider") or env_name or "ai").strip()
            token_env = str(env_obj.get("token_env") or AI_PROVIDER_ENV_HINTS.get(provider.lower(), "OPENAI_API_KEY/GH_MODELS_TOKEN")).strip()
            return 0, f"❌ Provider `{provider}` chưa cài token. Hãy set biến môi trường {token_env} hoặc thêm vào bot_ai_secrets.csv."
        payload = {
            "model": env_obj.get("model") or AI_MODEL,
            "messages": [
                {"role": "system", "content": "You are a concise assistant."},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.7,
            "top_p": 0.9,
            "max_tokens": 1024,
        }
        endpoint_local = normalize_chat_completions_endpoint(str(env_obj.get("endpoint") or AI_API_ENDPOINT))
        req = urllib.request.Request(
            endpoint_local,
            data=json.dumps(payload).encode("utf-8"),
            method="POST",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {token_local}",
                "api-key": token_local,
            },
        )
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                raw = resp.read().decode("utf-8", errors="replace")
            data = json.loads(raw)
            msg = data.get("choices", [{}])[0].get("message", {}).get("content", "⚠️ AI không trả nội dung.")
            return 200, msg
        except urllib.error.HTTPError as e:
            detail = e.read().decode("utf-8", errors="replace") if hasattr(e, "read") else str(e)
            return int(e.code), f"❌ Lỗi AI HTTP {e.code}: {detail[:300]}"
        except Exception as e:
            return -1, f"❌ Lỗi gọi AI: {e}"

    code, text = _try_env(env)
    if code == 200:
        return text

    # Tu dong fallback khi token/endpoint cua env hien tai bi 401.
    if code == 401:
        tried = set()
        tried.add(
            (
                normalize_chat_completions_endpoint(str(env.get("endpoint") or AI_API_ENDPOINT)),
                str(env.get("token") or "").strip(),
            )
        )
        candidates = []
        request_name = str(env_name or "").strip().lower()
        if request_name == "openai":
            oa_name = (get_provider_env_name("openai") or "openai").strip()
            if oa_name:
                candidates.append(oa_name)
            candidates.append("openai")
            active_name = (get_active_ai_token_name() or "").strip()
            if active_name:
                candidates.append(active_name)
        else:
            gh_name = (get_provider_env_name("github") or "github").strip()
            if gh_name:
                candidates.append(gh_name)
            candidates.append("github")
            active_name = (get_active_ai_token_name() or "").strip()
            if active_name:
                candidates.append(active_name)
            candidates.append("openai")
        seen_names = set()
        for nm in candidates:
            name = str(nm or "").strip()
            if not name or name in seen_names:
                continue
            seen_names.add(name)
            alt = resolve_ai_env(name)
            if request_name == "openai" and str(alt.get("provider", "")).lower() != "openai":
                continue
            alt_sig = (
                normalize_chat_completions_endpoint(str(alt.get("endpoint") or AI_API_ENDPOINT)),
                str(alt.get("token") or "").strip(),
            )
            if (not alt_sig[1]) or alt_sig in tried:
                continue
            tried.add(alt_sig)
            alt_code, alt_text = _try_env(alt)
            if alt_code == 200:
                return alt_text
        return text + " | Da thu fallback env AI khac nhung van that bai."

    return text


def _truthy(value: str) -> bool:
    return str(value or "").strip().lower() not in ("0", "false", "no", "off", "disabled")


def _sanitize_ai_token(provider: str, token: str) -> str:
    clean = str(token or "").strip()
    if provider == "openai" and clean and not clean.startswith("sk-"):
        return ""
    return clean


def _decrypt_secret_csv(path: str, password: str) -> str:
    from cryptography.fernet import Fernet, InvalidToken
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    import base64

    with open(path, "rb") as f:
        raw = f.read()
    if len(raw) < 17:
        raise ValueError("Secret file khong hop le.")
    salt = raw[:16]
    token = raw[16:]
    kdf = PBKDF2HMAC(algorithm=hashes.SHA256(), length=32, salt=salt, iterations=390000)
    key = base64.urlsafe_b64encode(kdf.derive(password.encode("utf-8")))
    try:
        return Fernet(key).decrypt(token).decode("utf-8-sig")
    except InvalidToken:
        raise ValueError("Sai mat khau giai ma secrets.")


def load_ai_secret_rows() -> list[dict]:
    rows = []
    if os.path.exists(AI_SECRETS_CSV):
        try:
            with open(AI_SECRETS_CSV, "r", encoding="utf-8-sig", newline="") as f:
                rows.extend(list(csv.DictReader(f)))
        except Exception as exc:
            logging.warning("Doc bot_ai_secrets.csv that bai: %s", exc)
    elif os.path.exists(AI_SECRETS_ENC):
        password = os.getenv("BOT_SECRETS_PASSWORD", "").strip()
        if password:
            try:
                text = _decrypt_secret_csv(AI_SECRETS_ENC, password)
                rows.extend(list(csv.DictReader(io.StringIO(text))))
            except Exception as exc:
                logging.warning("Doc bot_ai_secrets.csv.enc that bai: %s", exc)
        else:
            logging.info("Co %s nhung chua set BOT_SECRETS_PASSWORD, bo qua secrets ma hoa.", AI_SECRETS_ENC)
    return rows


def _normalize_ai_secret_row(row: dict) -> tuple[str, dict] | tuple[str, None]:
    name = str(row.get("name") or row.get("provider") or "").strip().lower()
    if not name or not _truthy(row.get("enabled", "1")):
        return "", None
    provider = str(row.get("provider") or name).strip().lower()
    token_env = str(row.get("token_env") or row.get("token_env_or_key") or "").strip()
    token = os.getenv(token_env, "").strip() if token_env else ""
    token = token or str(row.get("token") or row.get("api_key") or "").strip()
    endpoint = str(row.get("endpoint") or "").strip()
    model = str(row.get("model") or "").strip()
    if not endpoint:
        endpoint = OPENAI_API_ENDPOINT if provider == "openai" else AI_API_ENDPOINT
    if not model:
        model = OPENAI_MODEL if provider == "openai" else AI_MODEL
    token = _sanitize_ai_token(provider, token)
    return name, {
        "token": token,
        "token_env": token_env or AI_PROVIDER_ENV_HINTS.get(name) or AI_PROVIDER_ENV_HINTS.get(provider, ""),
        "provider": provider,
        "endpoint": endpoint,
        "model": model,
    }


def get_ai_tokens_map():
    try:
        raw = json.loads(get_config("ai_tokens_json", "{}") or "{}")
        # Tuong thich du lieu cu: {name: token_string}
        normalized = {}
        for k, v in (raw or {}).items():
            key_name = str(k or "").strip().lower()
            if isinstance(v, dict):
                provider = str(v.get("provider") or ("openai" if "openai" in key_name else "github")).strip().lower()
                if key_name == "openai":
                    provider = "openai"
                endpoint = str(v.get("endpoint") or (OPENAI_API_ENDPOINT if provider == "openai" else AI_API_ENDPOINT)).strip()
                if key_name == "openai" and "api.openai.com" not in endpoint.lower():
                    endpoint = OPENAI_API_ENDPOINT
                model = str(v.get("model") or (OPENAI_MODEL if provider == "openai" else AI_MODEL)).strip()
                token_env = str(v.get("token_env") or AI_PROVIDER_ENV_HINTS.get(key_name) or AI_PROVIDER_ENV_HINTS.get(provider, "")).strip()
                token = (os.getenv(token_env, "").strip() if token_env else "") or str(v.get("token") or "").strip()
                token = _sanitize_ai_token(provider, token)
                normalized[key_name] = {
                    "token": token,
                    "token_env": token_env,
                    "provider": provider,
                    "endpoint": endpoint,
                    "model": model,
                }
            else:
                provider = "openai" if "openai" in key_name else "github"
                token_env = AI_PROVIDER_ENV_HINTS.get(key_name) or AI_PROVIDER_ENV_HINTS.get(provider, "")
                normalized[key_name] = {
                    "token": _sanitize_ai_token(provider, str(v)),
                    "token_env": token_env,
                    "provider": provider,
                    "endpoint": OPENAI_API_ENDPOINT if provider == "openai" else AI_API_ENDPOINT,
                    "model": OPENAI_MODEL if provider == "openai" else AI_MODEL,
                }
        for row in load_ai_secret_rows():
            name, env_obj = _normalize_ai_secret_row(row)
            if name and env_obj:
                if str(env_obj.get("token") or "").strip() or name not in normalized:
                    normalized[name] = env_obj
        return normalized
    except Exception:
        return {}


def set_ai_tokens_map(data: dict):
    set_config("ai_tokens_json", json.dumps(data, ensure_ascii=False))


def get_active_ai_token_name():
    return (get_config("ai_token_active", "default") or "default").strip() or "default"


def set_active_ai_token_name(name: str):
    set_config("ai_token_active", (name or "default").strip())


def get_active_ai_token() -> str:
    env = resolve_ai_env(None)
    tok = str(env.get("token", "")).strip()
    if tok:
        return tok
    return get_config("ai_token", "").strip() or os.getenv(AI_API_TOKEN_ENV, "").strip()


def resolve_ai_env(env_name: str | None = None) -> dict:
    mp = get_ai_tokens_map()
    request_name = (env_name or get_active_ai_token_name() or "").strip().lower()
    aliases = {
        "git": "github",
        "github": "github",
        "openai": "openai",
        "meta": "meta",
        "alo": "alo",
    }
    name = aliases.get(request_name, request_name)
    if name and name in mp:
        return mp[name]
    # fallback theo ten chuan
    if name == "openai" and "openai" in mp:
        return mp["openai"]
    if name == "github" and "github" in mp:
        return mp["github"]
    openai_ep = (get_config("openai_api_endpoint", OPENAI_API_ENDPOINT) or OPENAI_API_ENDPOINT).strip()
    if name == "openai":
        return {
            "token": _sanitize_ai_token("openai", os.getenv("OPENAI_API_KEY", "")),
            "token_env": "OPENAI_API_KEY",
            "provider": "openai",
            "endpoint": openai_ep,
            "model": OPENAI_MODEL,
        }
    if name == "meta":
        return {
            "token": os.getenv("META_API_KEY", "").strip(),
            "token_env": "META_API_KEY",
            "provider": "meta",
            "endpoint": AI_API_ENDPOINT,
            "model": AI_MODEL,
        }
    if name == "alo":
        return {
            "token": os.getenv("ALO_API_KEY", "").strip(),
            "token_env": "ALO_API_KEY",
            "provider": "alo",
            "endpoint": OPENAI_API_ENDPOINT,
            "model": OPENAI_MODEL,
        }
    tok = get_config("ai_token", "").strip() or os.getenv(AI_API_TOKEN_ENV, "").strip()
    return {
        "token": tok,
        "token_env": AI_API_TOKEN_ENV,
        "provider": "github",
        "endpoint": AI_API_ENDPOINT,
        "model": AI_MODEL,
        "openai_endpoint": openai_ep,
    }


def normalize_chat_completions_endpoint(endpoint: str) -> str:
    ep = (endpoint or "").strip().rstrip("/")
    if not ep:
        return AI_API_ENDPOINT
    if ep.endswith("/chat/completions") or ep.endswith("/v1/chat/completions"):
        return ep
    if ep.endswith("/v1"):
        return ep + "/chat/completions"
    return ep + "/chat/completions"


def get_provider_env_name(provider: str) -> str:
    p = (provider or "").strip().lower()
    key = f"ai_env_{p}_active"
    name = (get_config(key, "") or "").strip()
    if name:
        return name
    # fallback ten mac dinh theo provider
    return "github" if p == "github" else ("openai" if p == "openai" else "")


def set_provider_env_name(provider: str, env_name: str):
    p = (provider or "").strip().lower()
    if not p:
        return
    set_config(f"ai_env_{p}_active", (env_name or "").strip())


def is_admin(user_id: int) -> bool:
    # CẤP TẤT CẢ QUYỀN CHO CẢ ADMIN CHÍNH VÀ ADMIN PHỤ
    return user_id == ADMIN_ID or user_id in EXTRA_ADMINS


async def require_admin(update: Update) -> bool:
    user_id = update.effective_user.id if update.effective_user else 0
    if is_admin(user_id):
        return True
    msg = update.effective_message
    if msg:
        await msg.reply_text(
            "⛔ Lenh nay chi danh cho admin.\n"
            f"ID cua ban: {user_id}\n"
            f"ADMIN_ID hien tai: {ADMIN_ID}\n"
            "Neu ban dang dung acc admin moi, go /myid roi cap nhat TELEGRAM_ADMIN_ID hoac them admin phu."
        )
    return False


def is_supreme_admin(user_id: int) -> bool:
    return user_id in {ADMIN_ID, SUPREME_ADMIN_ID}


def resolve_telegram_admin_label(user_id: int, preferred_username: str = "") -> str:
    uid = str(user_id)
    username = str(preferred_username or "").strip().lstrip("@")
    name = ""
    try:
        with db_conn() as conn:
            row = conn.execute(
                """
                SELECT name, username
                FROM customer_contacts
                WHERE uid=?
                ORDER BY updated_at DESC
                LIMIT 1
                """,
                (uid,),
            ).fetchone()
            if row:
                name = str(row["name"] or "").strip()
                username = username or str(row["username"] or "").strip().lstrip("@")
            if not username and not name:
                row = conn.execute(
                    """
                    SELECT name, username
                    FROM telegram_group_members
                    WHERE user_id=?
                    ORDER BY last_seen DESC
                    LIMIT 1
                    """,
                    (uid,),
                ).fetchone()
                if row:
                    name = str(row["name"] or "").strip()
                    username = str(row["username"] or "").strip().lstrip("@")
            if not username and not name:
                row = conn.execute(
                    """
                    SELECT full_name, username
                    FROM bot_phonebook
                    WHERE source_uid=?
                    ORDER BY last_seen DESC
                    LIMIT 1
                    """,
                    (uid,),
                ).fetchone()
                if row:
                    name = str(row["full_name"] or "").strip()
                    username = str(row["username"] or "").strip().lstrip("@")
    except Exception:
        pass
    if username:
        return f"@{username} ({uid})"
    if name:
        return f"{name} ({uid})"
    return uid


def build_admin_attendance_lines() -> list[str]:
    lines = []
    seen = set()

    def add_admin(role: str, user_id: int, preferred_username: str = "") -> None:
        uid = int(user_id)
        if uid in seen:
            return
        seen.add(uid)
        lines.append(f"- {role}: {resolve_telegram_admin_label(uid, preferred_username)}")

    add_admin("admin_chinh", ADMIN_ID, ADMIN_USERNAME)
    if SUPREME_ADMIN_ID != ADMIN_ID:
        add_admin("admin_toi_cao", SUPREME_ADMIN_ID, SUPREME_ADMIN_USERNAME)
    for admin_id in sorted(EXTRA_ADMINS):
        add_admin("admin_quan_tri", int(admin_id))
    if not lines:
        return ["- admin_chinh: (chua co)"]
    return lines


def get_bot_attendance_phone() -> str:
    phone = os.getenv("BOT_ATTENDANCE_PHONE", "").strip()
    if phone:
        return phone
    source_profile = os.getenv("SOURCE_PROFILE_DIR", "").strip()
    if source_profile:
        name = os.path.basename(os.path.normpath(source_profile))
        if re.fullmatch(r"\d{8,15}", name or ""):
            return name
    slot = os.getenv("BOT_INSTANCE_SLOT", "").strip() or "default"
    return f"botfather-{slot}"


def build_botfather_attendance_text(me=None) -> str:
    slot = os.getenv("BOT_INSTANCE_SLOT", "").strip() or "default"
    role = os.getenv("BOT_ATTENDANCE_ROLE", "botfather").strip() or "botfather"
    bot_name = ""
    try:
        if me:
            bot_name = f"@{me.username}" if getattr(me, "username", None) else str(getattr(me, "id", "") or "")
    except Exception:
        bot_name = ""
    greet = random.choice(BOT_ATTENDANCE_GREETINGS)
    now_txt = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    admin_lines = "\n".join(build_admin_attendance_lines())
    return (
        f"{greet}\n"
        "[DIEM DANH KHOI DONG]\n"
        f"Loai: {role.upper()}\n"
        f"Bot: {bot_name or slot}\n"
        f"Slot: {slot}\n"
        f"SDT/Profile: {get_bot_attendance_phone()}\n"
        f"May: {DEFAULT_MACHINE_DEVICE}\n"
        f"Admin quan tri:\n{admin_lines}\n"
        f"Thoi gian: {now_txt}"
    )


async def send_botfather_startup_attendance(application) -> tuple[bool, str]:
    if not BOT_ATTENDANCE_ENABLED:
        return False, "disabled"
    if not BOT_ATTENDANCE_CHAT:
        return False, "missing_chat"
    try:
        me = await application.bot.get_me()
        text = build_botfather_attendance_text(me)
        await application.bot.send_message(
            chat_id=BOT_ATTENDANCE_CHAT,
            text=text,
            disable_web_page_preview=True,
        )
        set_config(
            "bot_attendance_last",
            f"{int(time.time())}|sent slot={os.getenv('BOT_INSTANCE_SLOT', '').strip() or 'default'} chat={BOT_ATTENDANCE_CHAT}",
        )
        return True, "sent"
    except Exception as exc:
        detail = f"{type(exc).__name__}:{str(exc)[:180]}"
        logging.warning("Startup attendance failed: %s", detail)
        try:
            set_config("bot_attendance_last", f"{int(time.time())}|failed {detail}")
        except Exception:
            pass
        return False, detail


async def post_init_startup_attendance(application) -> None:
    await send_botfather_startup_attendance(application)


async def cmd_myid(update: Update, context: ContextTypes.DEFAULT_TYPE):
    refresh_group_enabled_from_shared()
    user = update.effective_user
    uid = user.id if user else 0
    uname = f"@{user.username}" if user and user.username else "(khong username)"
    chat = update.effective_chat
    chat_id = str(getattr(chat, "id", "") or "")
    chat_type = str(getattr(chat, "type", "") or "")
    await update.message.reply_text(
        f"ID cua ban: {uid}\n"
        f"Username: {uname}\n"
        f"Chat ID hien tai: {chat_id} ({chat_type})\n"
        f"ADMIN_ID hien tai cua bot: {ADMIN_ID}\n"
        f"Supper admin group: {get_supper_admin_group_id() or '(chua cai)'} | {'ON' if SUPPER_ADMIN_ENABLED else 'OFF'}\n"
        f"Supper invite: {SUPPER_ADMIN_GROUP_LINK or '(chua cai)'}\n"
        f"Admin quan ly group: {get_admin_management_group_id() or '(chua cai)'} | {'ON' if ADMIN_MANAGEMENT_ENABLED else 'OFF'}\n"
        f"Admin quan ly invite: {ADMIN_MANAGEMENT_GROUP_LINK or '(chua cai)'}\n"
        f"Admin? {'YES' if is_admin(uid) else 'NO'}"
    )


async def cmd_setsuppergroup(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global SUPPER_ADMIN_GROUP_ID, SUPPER_ADMIN_GROUP_LINK, SUPPER_ADMIN_ENABLED
    if not is_supreme_admin(update.effective_user.id):
        return
    chat = update.effective_chat
    msg = update.message
    args = [str(x).strip() for x in (context.args or []) if str(x).strip()]
    manual_chat_id = ""
    for arg in args:
        if re.fullmatch(r"-?\d{5,}", arg):
            manual_chat_id = arg
            break
    link_arg = ""
    for arg in args:
        if "t.me/" in arg or "telegram.me/" in arg:
            link_arg = arg
            break
    mode_arg = args[0].lower() if args and args[0].lower() in {"on", "off"} else ""

    if mode_arg == "off":
        SUPPER_ADMIN_ENABLED = False
        set_config("supper_admin_enabled", "0")
        write_supper_admin_shared_config()
        audit_log("telegram", str(update.effective_user.id), "setsuppergroup_off", get_supper_admin_group_id())
        await msg.reply_text("Da TAT gui report/canh bao vao suppergroup. Chat ID/link van duoc giu lai.")
        return

    if mode_arg == "on":
        if manual_chat_id:
            SUPPER_ADMIN_GROUP_ID = manual_chat_id
        elif not get_supper_admin_group_id():
            chat_type = str(getattr(chat, "type", "") or "").lower()
            if chat_type in ("group", "supergroup"):
                SUPPER_ADMIN_GROUP_ID = str(chat.id)
            else:
                await msg.reply_text(
                    "Chua co chat_id suppergroup. Hay go /setsuppergroup trong nhom, hoac /setsuppergroup on <chat_id>."
                )
                return
        if link_arg:
            SUPPER_ADMIN_GROUP_LINK = link_arg
        SUPPER_ADMIN_ENABLED = True
        set_config("supper_admin_group_id", SUPPER_ADMIN_GROUP_ID or get_supper_admin_group_id())
        set_config("supper_admin_group_link", SUPPER_ADMIN_GROUP_LINK or "")
        set_config("supper_admin_enabled", "1")
        write_supper_admin_shared_config()
        audit_log("telegram", str(update.effective_user.id), "setsuppergroup_on", get_supper_admin_group_id())
        await msg.reply_text(f"Da BAT suppergroup report/canh bao.\nChat ID: {get_supper_admin_group_id()}")
        return

    if manual_chat_id:
        chat_id = manual_chat_id
        chat_title = "manual"
    else:
        chat_type = str(getattr(chat, "type", "") or "").lower()
        if chat_type not in ("group", "supergroup"):
            await msg.reply_text(
                "Hay them bot vao nhom supper admin roi go /setsuppergroup trong nhom do.\n"
                "Neu da biet chat_id: /setsuppergroup <chat_id>"
            )
            return
        chat_id = str(chat.id)
        chat_title = str(getattr(chat, "title", "") or "")

    if link_arg:
        SUPPER_ADMIN_GROUP_LINK = link_arg
    SUPPER_ADMIN_GROUP_ID = chat_id
    SUPPER_ADMIN_ENABLED = True
    set_config("supper_admin_group_id", SUPPER_ADMIN_GROUP_ID)
    set_config("supper_admin_group_link", SUPPER_ADMIN_GROUP_LINK or "")
    set_config("supper_admin_enabled", "1")
    write_supper_admin_shared_config()
    audit_log("telegram", str(update.effective_user.id), "setsuppergroup", f"{chat_id}|{chat_title}|{SUPPER_ADMIN_GROUP_LINK}")
    await msg.reply_text(
        "Da luu nhom supper admin.\n"
        f"Chat ID: {SUPPER_ADMIN_GROUP_ID}\n"
        f"Ten nhom: {chat_title or 'N/A'}\n"
        f"Invite: {SUPPER_ADMIN_GROUP_LINK or 'N/A'}\n"
        "Trang thai: ON\n"
        "Tu bay gio canh bao khach moi se gui them ve nhom nay."
    )
    if str(getattr(chat, "id", "") or "") != str(SUPPER_ADMIN_GROUP_ID):
        await send_supper_admin_report(
            context.bot,
            "[SETUP] Bot da ket noi nhom supper admin va se gui report/file quet ve day.",
            source_chat_id=str(getattr(chat, "id", "") or ""),
        )


async def cmd_adminquanly(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global ADMIN_MANAGEMENT_GROUP_ID, ADMIN_MANAGEMENT_GROUP_LINK, ADMIN_MANAGEMENT_ENABLED
    if not is_supreme_admin(update.effective_user.id):
        return
    chat = update.effective_chat
    msg = update.message
    args = [str(x).strip() for x in (context.args or []) if str(x).strip()]
    manual_chat_id = ""
    for arg in args:
        if re.fullmatch(r"-?\d{5,}", arg):
            manual_chat_id = arg
            break
    link_arg = ""
    for arg in args:
        if "t.me/" in arg or "telegram.me/" in arg:
            link_arg = arg
            break
    mode_arg = args[0].lower() if args and args[0].lower() in {"on", "off"} else ""

    if mode_arg == "off":
        ADMIN_MANAGEMENT_ENABLED = False
        set_config("admin_management_enabled", "0")
        write_admin_management_shared_config()
        audit_log("telegram", str(update.effective_user.id), "adminquanly_off", get_admin_management_group_id())
        await msg.reply_text("Da TAT report van hanh vao /adminquanly. Chat ID/link van duoc giu lai.")
        return

    if mode_arg == "on":
        if manual_chat_id:
            ADMIN_MANAGEMENT_GROUP_ID = manual_chat_id
        elif not get_admin_management_group_id():
            chat_type = str(getattr(chat, "type", "") or "").lower()
            if chat_type in ("group", "supergroup"):
                ADMIN_MANAGEMENT_GROUP_ID = str(chat.id)
            else:
                await msg.reply_text(
                    "Chua co chat_id /adminquanly. Hay go /adminquanly trong nhom quan ly, hoac /adminquanly on <chat_id>."
                )
                return
        if link_arg:
            ADMIN_MANAGEMENT_GROUP_LINK = link_arg
        ADMIN_MANAGEMENT_ENABLED = True
        set_config("admin_management_group_id", ADMIN_MANAGEMENT_GROUP_ID or get_admin_management_group_id())
        set_config("admin_management_group_link", ADMIN_MANAGEMENT_GROUP_LINK or "")
        set_config("admin_management_enabled", "1")
        write_admin_management_shared_config()
        audit_log("telegram", str(update.effective_user.id), "adminquanly_on", get_admin_management_group_id())
        await msg.reply_text(f"Da BAT report /adminquanly.\nChat ID: {get_admin_management_group_id()}")
        return

    if manual_chat_id:
        chat_id = manual_chat_id
        chat_title = "manual"
    else:
        chat_type = str(getattr(chat, "type", "") or "").lower()
        if chat_type not in ("group", "supergroup"):
            await msg.reply_text(
                "Hay them bot vao nhom quan ly roi go /adminquanly trong nhom do.\n"
                "Neu da biet chat_id: /adminquanly <chat_id>"
            )
            return
        chat_id = str(chat.id)
        chat_title = str(getattr(chat, "title", "") or "")

    if link_arg:
        ADMIN_MANAGEMENT_GROUP_LINK = link_arg
    ADMIN_MANAGEMENT_GROUP_ID = chat_id
    ADMIN_MANAGEMENT_ENABLED = True
    set_config("admin_management_group_id", ADMIN_MANAGEMENT_GROUP_ID)
    set_config("admin_management_group_link", ADMIN_MANAGEMENT_GROUP_LINK or "")
    set_config("admin_management_enabled", "1")
    write_admin_management_shared_config()
    audit_log("telegram", str(update.effective_user.id), "adminquanly", f"{chat_id}|{chat_title}|{ADMIN_MANAGEMENT_GROUP_LINK}")
    await msg.reply_text(
        "Da luu nhom /adminquanly.\n"
        f"Chat ID: {ADMIN_MANAGEMENT_GROUP_ID}\n"
        f"Ten nhom: {chat_title or 'N/A'}\n"
        f"Invite: {ADMIN_MANAGEMENT_GROUP_LINK or 'N/A'}\n"
        "Trang thai: ON\n"
        "Tu bay gio report quet file/SĐT/UID/user, /guitn, /keoall se gui ve nhom nay."
    )
    await send_admin_management_report(
        context.bot,
        "[SETUP] Bot da ket noi nhom /adminquanly. Report van hanh se gui ve day.",
        source_chat_id=str(getattr(chat, "id", "") or ""),
    )

async def is_user_member(context: ContextTypes.DEFAULT_TYPE, user_id: int) -> bool:
    try:
        member = await context.bot.get_chat_member(chat_id=TARGET_CHAT, user_id=user_id)
        return member.status in ['member', 'administrator', 'creator']
    except Exception:
        logging.exception("check member failed: TARGET_CHAT=%s user=%s", TARGET_CHAT, user_id)
        return False


def parse_cookie_raw(cookie_raw: str) -> dict:
    cookies = {}
    for part in (cookie_raw or "").split(";"):
        token = part.strip()
        if not token or "=" not in token:
            continue
        key, value = token.split("=", 1)
        cookies[key.strip()] = value.strip()
    return cookies


def is_probably_group_context(thread_type, thread_id, sender_id: str) -> bool:
    t = str(thread_type).lower()
    if any(k in t for k in ("group", "room", "thread")):
        return True
    # Fallback: voi chat 1-1 thuong thread_id ~ sender_id
    return str(thread_id) != str(sender_id)


def normalize_phone(raw: str) -> str:
    return "".join(ch for ch in str(raw or "") if ch.isdigit())


def vn_phone_variants(raw: str) -> set[str]:
    digits = normalize_phone(raw)
    variants = {digits} if digits else set()
    if digits.startswith("0") and len(digits) >= 10:
        variants.add("84" + digits[1:])
        variants.add(digits[1:])
    elif digits.startswith("84") and len(digits) >= 11:
        variants.add("0" + digits[2:])
        variants.add(digits[2:])
    elif len(digits) == 9:
        variants.add("0" + digits)
        variants.add("84" + digits)
    return {v for v in variants if v}


def phone_matches_target(phone: str, target: str) -> bool:
    phone_variants = vn_phone_variants(phone)
    target_variants = vn_phone_variants(target)
    if phone_variants & target_variants:
        return True
    return any(
        len(p) >= 9 and len(t) >= 9 and (p.endswith(t) or t.endswith(p))
        for p in phone_variants
        for t in target_variants
    )


def obj_get_any(obj, *keys, default=""):
    if obj is None:
        return default
    for key in keys:
        try:
            if isinstance(obj, dict) and key in obj:
                val = obj.get(key)
            elif hasattr(obj, "get"):
                val = obj.get(key)
            else:
                val = getattr(obj, key, None)
        except Exception:
            val = None
        if val not in (None, ""):
            return val
    return default


def _coerce_json_payload(raw):
    if not isinstance(raw, str):
        return None
    txt = raw.strip()
    if not txt or txt[0] not in "{[":
        return None
    try:
        data = json.loads(txt)
    except Exception:
        return None
    return data if isinstance(data, (dict, list)) else None


def zalo_user_uid(obj) -> str:
    uid = str(obj_get_any(
        obj,
        "uid", "userId", "user_id", "id", "profileId",
        "contactUid", "contact_uid", "contactId", "contact_id",
        "targetUid", "target_uid", "ownerUid", "owner_uid",
    )).strip()
    if uid and uid != "0":
        return uid
    params_uid = str(obj_get_any(obj, "params", "param", "uidParam")).strip()
    if params_uid.isdigit() and params_uid != "0":
        return params_uid
    if isinstance(obj, (str, int)):
        uid = str(obj).strip()
        return "" if uid == "0" else uid
    return ""


def zalo_user_name(obj) -> str:
    return str(obj_get_any(
        obj,
        "name", "displayName", "display_name", "dName", "zaloName",
        "contactName", "contact_name", "title",
    )).strip()


def zalo_user_username(obj) -> str:
    return str(obj_get_any(obj, "username", "userName", "alias", "globalId")).strip()


def zalo_user_phone_candidates(obj) -> list[str]:
    keys = (
        "phone", "phoneNumber", "phone_number", "tel", "mobile",
        "phoneDisplay", "displayPhone", "normalizedPhone", "accountPhone",
    )
    vals = []
    for key in keys:
        val = obj_get_any(obj, key)
        if val not in (None, ""):
            vals.append(str(val))
    return vals


def zalo_user_phone(obj) -> str:
    for val in zalo_user_phone_candidates(obj):
        phone = normalize_phone(val)
        if phone:
            return phone
    return ""


def _payload_mapping(obj):
    if obj is None:
        return None
    if isinstance(obj, dict):
        return obj
    try:
        if hasattr(obj, "toDict"):
            data = obj.toDict()
            if isinstance(data, dict):
                return data
    except Exception:
        pass
    try:
        if hasattr(obj, "items"):
            return dict(obj.items())
    except Exception:
        pass
    try:
        data = vars(obj)
        if isinstance(data, dict) and data:
            return data
    except Exception:
        pass
    return None


def _walk_payload(obj, path: str = "root", depth: int = 0):
    if depth > 6 or obj is None:
        return
    parsed = _coerce_json_payload(obj)
    if parsed is not None:
        yield from _walk_payload(parsed, f"{path}#json", depth + 1)
        return
    mapping = _payload_mapping(obj)
    if mapping is not None:
        yield path, mapping
        for key, val in mapping.items():
            if isinstance(val, str):
                parsed_val = _coerce_json_payload(val)
                if parsed_val is not None:
                    yield from _walk_payload(parsed_val, f"{path}.{key}#json", depth + 1)
                continue
            if isinstance(val, (int, float, bool)) or val is None:
                continue
            yield from _walk_payload(val, f"{path}.{key}", depth + 1)
        return
    if isinstance(obj, (list, tuple)):
        for idx, val in enumerate(obj):
            if isinstance(val, (str, int, float, bool)) or val is None:
                continue
            yield from _walk_payload(val, f"{path}[{idx}]", depth + 1)


def _looks_like_contact_card(path: str, data: dict, uid: str, phone: str, name: str, username: str) -> bool:
    low_path = path.lower()
    keys = {str(k).lower() for k in data.keys()}
    msg_type = str(data.get("msgType") or data.get("msg_type") or data.get("type") or "").strip()
    action = str(data.get("action") or data.get("act") or "").strip().lower()
    hint = any(x in low_path for x in ("contact", "card", "vcard", "phonebook", "share", "userinfo", "user_info", "profile", "msginfo"))
    hint = hint or any(x in keys for x in (
        "phone", "phonenumber", "phone_number", "tel", "mobile",
        "contact", "contactinfo", "userinfo", "user_info",
        "contactuid", "contact_uid", "contactid", "contact_id",
        "qrcodeurl", "qr_code_url", "qrurl", "guid", "params",
    ))
    hint = hint or msg_type == "6"
    hint = hint or "recommended.user" in action or "recommened.user" in action
    if phone and (uid or name or username or hint):
        return True
    if uid and hint:
        return True
    return False


def extract_zalo_contact_card(message, message_object, sender: str = "") -> dict:
    candidates = []
    for source_name, payload in (("message_object", message_object), ("message", message)):
        for path, data in _walk_payload(payload, source_name):
            uid = zalo_user_uid(data)
            phone = zalo_user_phone(data)
            name = zalo_user_name(data)
            username = zalo_user_username(data)
            if not _looks_like_contact_card(path, data, uid, phone, name, username):
                continue
            # Tin text thuong doi khi co uid sender o top-level; tranh ghi nham admin la contact card.
            if uid and uid == str(sender) and "contact" not in path.lower() and "card" not in path.lower() and not phone:
                continue
            score = (4 if uid else 0) + (4 if phone else 0) + (2 if name else 0) + (1 if username else 0)
            candidates.append({
                "uid": uid,
                "name": name,
                "username": username,
                "phone": phone,
                "source": path,
                "_score": score,
            })
    if not candidates:
        return {}
    best = sorted(candidates, key=lambda x: x.get("_score", 0), reverse=True)[0]
    best.pop("_score", None)
    return best


def find_recent_zalo_contact_card(client, sender: str = "") -> dict:
    try:
        recent = client.getLastMsgs()
    except Exception:
        return {}
    candidates = []
    for path, data in _walk_payload(recent, "recent"):
        if not isinstance(data, dict):
            continue
        if not any(k in data for k in ("msgType", "content", "uidFrom", "idTo")):
            continue
        if sender and str(data.get("uidFrom") or "") != str(sender):
            continue
        card = extract_zalo_contact_card(data.get("content"), data, sender=sender)
        if not card:
            continue
        try:
            ts = float(data.get("ts") or data.get("created_at") or 0)
        except Exception:
            ts = 0.0
        card["_recent_ts"] = ts
        card["_recent_source"] = path
        candidates.append(card)
    if not candidates:
        return {}
    best = sorted(candidates, key=lambda x: x.get("_recent_ts", 0), reverse=True)[0]
    best.pop("_recent_ts", None)
    return best


def _safe_zalo_debug_value(obj, depth: int = 0):
    if depth > 4:
        return "..."
    if obj is None or isinstance(obj, (int, float, bool)):
        return obj
    if isinstance(obj, str):
        return obj[:700]
    mapping = _payload_mapping(obj)
    if mapping is not None:
        out = {}
        for key, val in list(mapping.items())[:45]:
            k = str(key)
            if any(x in k.lower() for x in ("cookie", "token", "secret", "imei", "hash")):
                out[k] = "<redacted>"
            else:
                out[k] = _safe_zalo_debug_value(val, depth + 1)
        return out
    if isinstance(obj, (list, tuple)):
        return [_safe_zalo_debug_value(v, depth + 1) for v in list(obj)[:20]]
    return repr(obj)[:700]


def append_zalo_message_debug(sender: str, raw: str, message, message_object, contact_card: dict, note: str = "") -> None:
    try:
        current = json.loads(get_config("zalo_admin_message_debug_json", "[]") or "[]")
        if not isinstance(current, list):
            current = []
    except Exception:
        current = []
    current.append({
        "ts": time.time(),
        "sender": str(sender or ""),
        "raw": str(raw or "")[:700],
        "note": str(note or "")[:150],
        "card": contact_card or {},
        "message": _safe_zalo_debug_value(message),
        "message_object": _safe_zalo_debug_value(message_object),
    })
    set_config("zalo_admin_message_debug_json", json.dumps(current[-8:], ensure_ascii=False))


def save_last_zalo_contact_card(admin_uid: str, card: dict) -> None:
    if not card or not (card.get("uid") or card.get("phone")):
        return
    payload = dict(card)
    payload["admin_uid"] = str(admin_uid or "")
    payload["saved_at"] = time.time()
    try:
        current = json.loads(get_config("zalo_last_contact_card_json", "{}") or "{}")
        if not isinstance(current, dict):
            current = {}
    except Exception:
        current = {}
    current[str(admin_uid or "__last__")] = payload
    current["__last__"] = payload
    set_config("zalo_last_contact_card_json", json.dumps(current, ensure_ascii=False))


def load_last_zalo_contact_card(admin_uid: str = "") -> dict:
    try:
        current = json.loads(get_config("zalo_last_contact_card_json", "{}") or "{}")
        if not isinstance(current, dict):
            return {}
        card = current.get(str(admin_uid or "")) or current.get("__last__") or {}
        return card if isinstance(card, dict) else {}
    except Exception:
        return {}


def extract_zalo_text(message, message_object) -> str:
    if isinstance(message, str) and message.strip():
        return message.strip()
    if isinstance(message_object, dict):
        for key in ("text", "content", "msg", "message"):
            val = message_object.get(key)
            if isinstance(val, str) and val.strip():
                return val.strip()
    return ""


def is_verify_command_match(text: str, verify_cmd: str) -> bool:
    raw = str(text or "").strip().lower()
    cmd = str(verify_cmd or "").strip().lower()
    if not raw or not cmd:
        return False
    if raw == cmd:
        return True
    raw_compact = re.sub(r"[^a-z0-9]", "", raw)
    cmd_compact = re.sub(r"[^a-z0-9]", "", cmd)
    if not raw_compact or not cmd_compact:
        return False
    return raw_compact == cmd_compact


def extract_first_http_url(obj) -> str:
    try:
        if isinstance(obj, str):
            m = re.search(r"https?://\S+", obj)
            return m.group(0).strip() if m else ""
        if isinstance(obj, dict):
            for k in (
                "url", "src", "href", "thumbnail", "thumb", "image", "imageUrl", "photo", "photoUrl",
                "mediaUrl", "contentUrl", "attachmentUrl", "hdUrl", "fullUrl",
            ):
                v = obj.get(k)
                if isinstance(v, str) and v.startswith("http"):
                    return v.strip()
            for v in obj.values():
                got = extract_first_http_url(v)
                if got:
                    return got
        if isinstance(obj, (list, tuple)):
            for it in obj:
                got = extract_first_http_url(it)
                if got:
                    return got
    except Exception:
        return ""
    return ""


def notify_admin_sync(text: str) -> None:
    try:
        key = (text or "").strip()[:120]
        now = time.time()
        last = float(ADMIN_NOTIFY_LAST_TS.get(key, 0))
        if now - last < 15:
            return
        ADMIN_NOTIFY_LAST_TS[key] = now

        url = f"https://api.telegram.org/bot{TOKEN_BOT}/sendMessage"
        payload = urllib.parse.urlencode({"chat_id": str(ADMIN_ID), "text": text}).encode("utf-8")
        req = urllib.request.Request(url, data=payload, method="POST")
        with urllib.request.urlopen(req, timeout=10):
            pass
    except Exception:
        pass


def notify_admin_management_sync(text: str) -> None:
    try:
        key = ("adminquanly:" + (text or "").strip())[:120]
        now = time.time()
        last = float(ADMIN_NOTIFY_LAST_TS.get(key, 0))
        if now - last < 15:
            return
        ADMIN_NOTIFY_LAST_TS[key] = now

        chat_id = get_admin_management_group_id() or str(ADMIN_ID)
        url = f"https://api.telegram.org/bot{TOKEN_BOT}/sendMessage"
        payload = urllib.parse.urlencode({"chat_id": str(chat_id), "text": text[:3900]}).encode("utf-8")
        req = urllib.request.Request(url, data=payload, method="POST")
        with urllib.request.urlopen(req, timeout=10):
            pass
    except Exception:
        pass


def zalo_banbe_delay_seconds() -> int:
    try:
        delay = int(ZALO_BANBE_BROADCAST_INTERVAL)
    except Exception:
        delay = 120
    return max(ZALO_BANBE_MIN_DELAY_SEC, min(ZALO_BANBE_MAX_DELAY_SEC, delay))


def start_zalo_auto_worker() -> None:
    try:
        from zlapi import ZaloAPI
        from zlapi.models import Message, ThreadType
    except Exception:
        set_config("zalo_active", "False")
        set_config("zalo_worker_last", f"{int(time.time())}|error missing_zlapi")
        logging.error("Thieu zlapi. Cai bang: pip install zlapi")
        return

    cookies = parse_cookie_raw(COOKIE_RAW)
    if not cookies:
        set_config("zalo_active", "False")
        set_config("zalo_worker_last", f"{int(time.time())}|error invalid_cookie")
        logging.error("COOKIE_RAW khong hop le.")
        return

    class ZaloAutoBot(ZaloAPI):
        def send_long_message(self, text: str, thread_id, thread_type):
            chunks = split_long_text(text, ZALO_TEXT_CHUNK_LIMIT)
            for idx, chunk in enumerate(chunks):
                self.sendMessage(Message(text=chunk), thread_id, thread_type)
                if idx < len(chunks) - 1:
                    time.sleep(0.08)

        def send_zalo_autoreply_once(self, text: str, sender: str, thread_id, thread_type, reason: str = "") -> bool:
            body = str(text or "").strip()
            if not body:
                return False
            target_key = f"{thread_type}:{thread_id}:{sender}"
            if outbox_exists("zalo_autorep", target_key, body):
                logging.info(
                    "[ZALO AUTOREP SKIP DUP] sender=%s thread=%s reason=%s",
                    sender,
                    thread_id,
                    reason or "auto",
                )
                return True
            try:
                self.send_long_message(body, thread_id, thread_type)
                mark_outbox_sent("zalo_autorep", target_key, body)
                return True
            except Exception as exc:
                logging.warning(
                    "Gui zalo auto-reply that bai sender=%s thread=%s reason=%s err=%s",
                    sender,
                    thread_id,
                    reason or "auto",
                    exc,
                )
                return False

        def _set_admin_uid(self, uid: str, source: str) -> None:
            global ZALO_ADMIN_UID
            uid = str(uid or "").strip()
            if not uid:
                return
            if is_known_zalo_group_uid(uid):
                notify_admin_sync(f"[ZALO ADMIN BLOCKED] Bo qua UID nhom={uid} ({source})")
                return
            ZALO_ADMIN_UID = uid
            try:
                set_config("zalo_admin_uid", ZALO_ADMIN_UID)
            except Exception:
                pass
            notify_admin_sync(f"[ZALO ADMIN] Da cap nhat admin_uid={ZALO_ADMIN_UID} ({source})")

        def _record_admin_auto_find(self, status: str) -> None:
            try:
                set_config("zalo_admin_auto_find_last", f"{int(time.time())}|{status[:450]}")
            except Exception:
                pass

        def _remember_zalo_friend(self, item, phone_override: str = "") -> str:
            uid = zalo_user_uid(item)
            if not uid:
                return ""
            ZALO_FRIEND_IDS.add(uid)
            phone = normalize_phone(phone_override) or zalo_user_phone(item)
            upsert_contact(
                "zalo",
                uid,
                name=zalo_user_name(item),
                username=zalo_user_username(item),
                phone=phone,
            )
            return uid

        def _set_admin_from_zalo_user(self, item, source: str, phone_override: str = "") -> str:
            uid = self._remember_zalo_friend(item, phone_override=phone_override)
            if uid and uid != str(ZALO_ADMIN_UID):
                self._set_admin_uid(uid, source)
            return uid

        def _find_admin_by_phone(self, items) -> str:
            target_phone = normalize_phone(ZALO_ADMIN_PHONE)
            if not target_phone:
                self._record_admin_auto_find("skip_empty_phone")
                return ""
            last_error = ""
            try:
                user = self.fetchPhoneNumber(target_phone)
                uid = self._set_admin_from_zalo_user(user, f"fetchPhoneNumber_{target_phone}", phone_override=target_phone)
                if uid:
                    self._record_admin_auto_find(f"found_fetchPhoneNumber uid={uid} phone={target_phone}")
                    return uid
            except Exception as exc:
                last_error = f"{type(exc).__name__}:{str(exc)[:180]}"
                logging.info("Khong tim thay admin Zalo bang fetchPhoneNumber(%s): %s", target_phone, exc)

            for item in items:
                if any(phone_matches_target(phone, target_phone) for phone in zalo_user_phone_candidates(item)):
                    uid = self._set_admin_from_zalo_user(item, f"friend_phone_{target_phone}", phone_override=target_phone)
                    if uid:
                        self._record_admin_auto_find(f"found_friend uid={uid} phone={target_phone}")
                        return uid
            self._record_admin_auto_find(f"not_found phone={target_phone} last_error={last_error}")
            return ""

        def refresh_friends(self):
            try:
                data = self.fetchAllFriends()
                items = data.get("data", []) if isinstance(data, dict) else (data or [])
                ZALO_FRIEND_IDS.clear()
                for item in items:
                    self._remember_zalo_friend(item)

                # Tu dong tim admin theo so dien thoai khi bat che do auto-find
                if ZALO_ADMIN_AUTO_FIND:
                    self._record_admin_auto_find(f"start phone={normalize_phone(ZALO_ADMIN_PHONE)} friends={len(items)}")
                    self._find_admin_by_phone(items)
            except Exception as exc:
                logging.warning("Khong tai duoc danh sach ban be Zalo: %s", exc)

        def scan_all_uid_sources(self):
            # 1) Quet danh ba Zalo (yeu cau uu tien)
            self.refresh_friends()
            # 2) Quet danh sach nhom + thanh vien nhom neu API cho phep
            try:
                groups = self.fetchAllGroups()
                items = groups.get("data", []) if isinstance(groups, dict) else (groups or [])
                for g in items:
                    if isinstance(g, dict):
                        gid = g.get("groupId") or g.get("id") or g.get("uid")
                        if gid:
                            upsert_contact("zalo_group", str(gid), name=str(g.get("name") or g.get("title") or ""))
                            # Co gang quet thanh vien group neu API co ho tro
                            members = []
                            try:
                                if hasattr(self, "fetchGroupInfo"):
                                    gi = self.fetchGroupInfo(gid)
                                    if isinstance(gi, dict):
                                        members = gi.get("members") or gi.get("participants") or []
                                elif hasattr(self, "getGroupInfo"):
                                    gi = self.getGroupInfo(gid)
                                    if isinstance(gi, dict):
                                        members = gi.get("members") or gi.get("participants") or []
                            except Exception:
                                members = []
                            for m in members or []:
                                if isinstance(m, dict):
                                    muid = m.get("uid") or m.get("id") or m.get("userId")
                                    if muid:
                                        upsert_contact(
                                            "zalo_group_member",
                                            str(muid),
                                            name=str(m.get("name") or m.get("displayName") or ""),
                                            username=str(m.get("username") or ""),
                                            phone=str(m.get("phone") or m.get("phoneNumber") or ""),
                                        )
            except Exception:
                pass
            # 3) Quet cac doan chat/hop thoai hien co neu API ho tro
            chat_api_candidates = ("fetchConversations", "getConversations", "fetchChats", "getChats", "fetchAllConversations")
            for fn in chat_api_candidates:
                if not hasattr(self, fn):
                    continue
                try:
                    data = getattr(self, fn)()
                    items = []
                    if isinstance(data, dict):
                        items = data.get("data", []) or data.get("conversations", []) or data.get("chats", []) or []
                    elif isinstance(data, list):
                        items = data
                    for it in items or []:
                        if isinstance(it, dict):
                            uid = it.get("uid") or it.get("id") or it.get("userId") or it.get("toUid")
                            if uid:
                                upsert_contact(
                                    "zalo_chat",
                                    str(uid),
                                    name=str(it.get("name") or it.get("displayName") or it.get("title") or ""),
                                    phone=str(it.get("phone") or it.get("phoneNumber") or ""),
                                )
                    break
                except Exception:
                    continue
            for suid in list(ZALO_SEEN_USER_IDS):
                upsert_contact("zalo_seen", suid)

        def scan_group_members(self, group_id: str):
            members = []
            dbg = []
            gi = None
            try:
                if hasattr(self, "fetchGroupInfo"):
                    dbg.append("try:fetchGroupInfo")
                    gi = self.fetchGroupInfo(group_id)
                    if isinstance(gi, dict):
                        members = gi.get("members") or gi.get("participants") or []
                        dbg.append(f"fetchGroupInfo.keys={list(gi.keys())[:8]}")
                elif hasattr(self, "getGroupInfo"):
                    dbg.append("try:getGroupInfo")
                    gi = self.getGroupInfo(group_id)
                    if isinstance(gi, dict):
                        members = gi.get("members") or gi.get("participants") or []
                        dbg.append(f"getGroupInfo.keys={list(gi.keys())[:8]}")
                else:
                    dbg.append("no_group_info_api")
            except Exception as exc:
                dbg.append(f"scan_exc={type(exc).__name__}:{str(exc)[:180]}")
                members = []

            # Fallback parser for zlapi structures:
            # keys like removalsGroup / unchangedsGroup / gridInfoMap
            if (not members) and isinstance(gi, dict):
                fb = []
                for k in ("unchangedsGroup", "removalsGroup", "gridInfoMap", "data", "groupInfo", "memVerList", "memList", "members"):
                    v = gi.get(k)
                    if isinstance(v, dict):
                        for kk, vv in v.items():
                            if isinstance(vv, dict):
                                uid = vv.get("uid") or vv.get("id") or vv.get("userId") or kk
                                name = vv.get("name") or vv.get("displayName") or vv.get("fullName") or ""
                                phone = vv.get("phone") or vv.get("phoneNumber") or ""
                                fb.append({"uid": uid, "name": name, "phone": phone})
                            elif isinstance(vv, str):
                                fb.append({"uid": kk, "name": vv, "phone": ""})
                    elif isinstance(v, list):
                        for vv in v:
                            if isinstance(vv, dict):
                                uid = vv.get("uid") or vv.get("id") or vv.get("userId")
                                if uid:
                                    fb.append({
                                        "uid": uid,
                                        "name": vv.get("name") or vv.get("displayName") or vv.get("fullName") or "",
                                        "phone": vv.get("phone") or vv.get("phoneNumber") or "",
                                    })
                if fb:
                    members = fb
                    dbg.append(f"fallback_members={len(fb)}")

            out = []
            seen = set()
            for m in members or []:
                if isinstance(m, dict):
                    muid = m.get("uid") or m.get("id") or m.get("userId")
                    if not muid:
                        continue
                    name = str(m.get("name") or m.get("displayName") or m.get("fullName") or "").strip()
                    phone = str(m.get("phone") or m.get("phoneNumber") or "").strip()
                    suid = str(muid)
                    # Loai bo UID nhom bi parse nham thanh thanh vien
                    if suid == str(group_id):
                        continue
                    if suid in seen:
                        continue
                    seen.add(suid)
                    upsert_contact("zalo_group_member", suid, name=name, phone=phone)
                    out.append({"uid": suid, "name": name or "Thanh vien"})
            return out, dbg

        def resolve_group_id_from_link(self, group_link: str):
            link = (group_link or "").strip()
            if not link:
                return ""
            # Thu cac API co san neu zlapi ho tro
            api_candidates = ("fetchGroupInfoByLink", "getGroupInfoByLink", "resolveGroupLink", "checkGroupInvite")
            for fn in api_candidates:
                if hasattr(self, fn):
                    try:
                        data = getattr(self, fn)(link)
                        if isinstance(data, dict):
                            gid = data.get("groupId") or data.get("id") or data.get("uid")
                            if gid:
                                return str(gid)
                    except Exception:
                        pass
            # Fallback: lay slug cuoi link de do khop voi danh sach group hien co
            slug = link.rstrip("/").split("/")[-1].strip().lower()
            try:
                groups = self.fetchAllGroups()
                items = groups.get("data", []) if isinstance(groups, dict) else (groups or [])
                for g in items:
                    if not isinstance(g, dict):
                        continue
                    gid = str(g.get("groupId") or g.get("id") or g.get("uid") or "")
                    name = str(g.get("name") or g.get("title") or "").lower()
                    glink = str(g.get("link") or g.get("url") or "").lower()
                    if slug and (slug in gid.lower() or slug in name or slug in glink):
                        return gid
            except Exception:
                pass
            return ""

        def report_scan_to_zalo_admin(self):
            if not str(ZALO_ADMIN_UID).strip():
                return
            now_ts = time.time()
            last_ts = get_last_scan_report_ts("zalo")
            rows = list_recent_contacts("zalo", last_ts, 40)
            if not rows:
                return
            lines = []
            for r in rows:
                lines.append(
                    f"- UID:{r['uid']} | Ten:{(r['name'] or '(chua ro)')} | SDT:{(r['phone'] or 'N/A')} | Kenh:{r['platform']}"
                )
            msg = "📡 BAO CAO QUET ZALO MOI:\n" + "\n".join(lines[:20])
            try:
                self.sendMessage(Message(text=msg[:3500]), int(ZALO_ADMIN_UID), ThreadType.USER)
                set_last_scan_report_ts("zalo", now_ts)
            except Exception:
                pass

        def remind_backup_to_zalo_admin(self):
            global LAST_BACKUP_REMIND_TS
            now_ts = time.time()
            if now_ts - float(LAST_BACKUP_REMIND_TS) < BACKUP_REMIND_INTERVAL_SEC:
                return
            LAST_BACKUP_REMIND_TS = now_ts
            try:
                self.sendMessage(
                    Message(text="💾 Nhac nho: Vui long sao luu tin nhan/du lieu dinh ky de dam bao an toan du lieu."),
                    int(ZALO_ADMIN_UID),
                    ThreadType.USER,
                )
            except Exception:
                pass

        def onMessage(self, mid, author_id, message, message_object, thread_id, thread_type):  # noqa: N802
            try:
                global BOT_ACTIVE, CHUC_NANG_AUTO_DUYET, THONG_BAO_START_ACTIVE, HIEN_THI_ANH_DONG
                global WEB_URL, TARGET_CHAT, LINK_NHAP_CODE, ZALO_AUTO_REPLY, ZALO_VERIFY_CMD
                global ZALO_ADMIN_UID, ZALO_ADMIN_PHONE, ZALO_ADMIN_AUTO_FIND
                global ZALO_ADMIN_REP_MODE, ZALO_ADMIN_REP_IMAGE, ZALO_AUTO_REP_ENABLED, ZALO_AUTO_JOIN, TELE_AUTO_JOIN, ZALO_AUTO_REP_ALL
                global ZALO_AUTO_REP_NGUOILA, ZALO_AUTO_REP_BANBE, ZALO_AUTO_REP_NEW
                global ZALO_BANBE_BROADCAST_ENABLED, ZALO_BANBE_BROADCAST_INTERVAL, ZALO_BANBE_BROADCAST_MSG
                global GIF_CHAO_HOI, GIF_DANG_KY, GIF_EP_JOIN, GIF_MOC_NAP, GIF_CHO_DUYET, GIF_QUANG_CAO, USERBOT_GUITN_GIF
                global OPENAI_INALL_ZALO, OPENAI_INALL_TELE, OPENAI_INALL_ZALO_GROUPS, OPENAI_INALL_TELE_GROUPS, ADMIN_CHAT_NOTIFY, BUSINESS_PRO_ZALO_GROUPS, ZALO_BLESS_ENABLED, HELP_IMAGE_URLS, ZALO_2FA_IMAGE_URL
                sender = str(author_id or "")
                if not sender:
                    return
                # Bo qua tin do chinh tai khoan bot gui de tranh vong lap tu phan hoi.
                if sender == str(getattr(self, "uid", "") or ""):
                    return
                text = extract_zalo_text(message, message_object)
                is_group_context = is_probably_group_context(thread_type, thread_id, sender)
                is_business_pro_group = str(thread_id) in BUSINESS_PRO_ZALO_GROUPS and is_group_context
                raw_probe = (text or "").strip()
                if raw_probe.lower() in {"zid", "zuid", "zwhoami", "uid"}:
                    try:
                        self.sendMessage(
                            Message(text=f"UID_ZALO={sender}\nTen={display_name or ''}\nThread={thread_id}"),
                            thread_id,
                            thread_type,
                        )
                    except Exception:
                        pass
                    return

                # Lenh xac thuc doi admin Zalo khi UID bi thay doi:
                # bat ky ai nhap dung ma se duoc set thanh admin chinh ngay.
                if is_verify_command_match(text, ZALO_VERIFY_CMD):
                    if is_group_context:
                        try:
                            self.sendMessage(
                                Message(text="⚠️ Lenh xac thuc admin chi dung trong inbox 1-1, khong dung trong nhom."),
                                thread_id,
                                thread_type,
                            )
                        except Exception:
                            pass
                        notify_admin_sync(f"[ZALO VERIFY BLOCKED] Nhan ma xac thuc trong nhom thread={thread_id} sender={sender}")
                        return
                    self._set_admin_uid(sender, "verify_command")
                    persist_runtime_state()
                    try:
                        self.sendMessage(
                            Message(text="👑 XIN CHÀO CHỦ NHÂN 👑 AI FAMILY 6.0 🤖 ĐÃ SẴN SÀNG NHẬN LỆNH!"),
                            thread_id,
                            thread_type,
                        )
                    except Exception:
                        pass
                    notify_admin_sync(f"[ZALO VERIFY FIRST] UID admin zalo: {sender}")
                    return

                admin_all = get_zalo_admin_all()
                admin_ai = get_zalo_admin_ai()
                is_zalo_admin = sender == str(ZALO_ADMIN_UID) or sender in admin_all
                all_admin_uids = set([str(ZALO_ADMIN_UID)]) | set(str(x) for x in admin_all) | set(str(x) for x in ZALO_GROUP_ADMINS)

                # ===== LENH ADMIN ZALO =====
                if is_zalo_admin:
                    raw = text.strip()
                    lower = raw.lower()
                    def ztail(cmd: str):
                        m = re.match(rf"^{re.escape(cmd)}\b\s*(.*)$", raw, flags=re.I | re.S)
                        return m.group(1).strip() if m else None
                    bless_text = "🎉🌟 CHÚC CÁC THÀNH VIÊN GIA ĐÌNH ADMIN SỐNG VUI SỐNG KHỎE, VÔ LO VÔ NGHĨ 💖🍀"
                    short_help_text = (
                        "👑 Hi Chủ Nhân,👑 cảm ơn chủ nhân đã tạo ra tôi. Tôi là AI FAMILY 6.0, 🤖 tôi đã sẵn sàng phục vụ.\n"
                        "✨ zstatus: Xem toàn bộ trạng thái ON/OFF và cấu hình hệ thống.\n"
                        "✨ zai: Gọi AI từ môi trường Github để trả lời câu hỏi của Admin.\n"
                        "✨ zcheck UID | zcheck sau khi gửi danh thiếp: Kiểm tra UID/danh thiếp có trong dữ liệu Zalo/Telegram không.\n"
                        "✨ zchecklenh z_lenh: Giải thích nhanh chức năng của lệnh.\n"
                        "✨ zaddadminall UID: Cấp admin phụ toàn quyền hệ thống.\n"
                        "✨ zaddadmiai UID: Cấp admin phụ quyền AI, chỉ dùng lệnh zai.\n"
                        "✨ zautorepall | zautorepnguoila | zautorepbanbe | zautorepnew (on/off) + zsetrep + zchecksetrep.\n"
                        "✨ zautojoin on/off | zunjoinall | zunjoin UID_GROUP.\n"
                        "✨ zuiddb | zuidall | zuitvgr | zuidallgr | zcheckiudall | zsdtall.\n"
                        "✨ zbanbe on/off time | zsettn NOI_DUNG.\n"
                        "✨ zaddmingr UID_ADMIN_GR.\n"
                        "✨ zdanhba!: Gui danh thiep tung nguoi trong danh ba bot.\n"
                        "✨ zsetadminrep text|image <url>.\n"
                        "✨ zsetgifall <url> | zsetguitngif <url|off> | zshowgifs.\n"
                        "✨ zbless on|off: Bat/tat loi chuc thanh vien gia dinh admin.\n"
                        "✨ zquet! <link>|here|live|gid <uid_group>: Quet UID theo link/nhom hien tai/UID group; live thu truc tiep tu tin nhan.\n"
                        "✨ zloaloaall NOI_DUNG THOI_GIAN.\n"
                        "✨ zcapnhat NOI_DUNG: AI phân tích yêu cầu cập nhật từ admin.\n"
                        "✨ zaitokens: Quản lý đa môi trường AI token (add/use/del/list).\n"
                        "✨ z<tên_môi_trường> NOI_DUNG: Gọi AI trực tiếp theo môi trường token đã thêm (ví dụ zgithub xin chao).\n"
                        "✨ zsethelpimgs | zaddhelpimg | zshowhelpimgs | zclearhelpimgs: Quan ly bo anh huong dan 2FA.\n"
                        "✨ zset2fa [url_anh] | zshow2fa | zclear2fa: Set anh 2FA truc tiep.\n"
                        "✨ zyes | zno: Xác nhận hoặc hủy lệnh thay đổi cấu hình.\n"
                    )
                    full_help_text = (
                        "🛡️ ĐÂY LÀ TOÀN BỘ ĐẦY ĐỦ CHỨC NĂNG CHỦ NHÂN ĐÃ CÀI ĐẶT CHO TÔI 🛡️\n"
                        "🎯 XIN MỜI CHỦ NHÂN TÙY CHỌN 🎯\n\n"
                        "1) zstatus: Xem toàn bộ trạng thái bot và các cờ chức năng.\n"
                        "2) zon / zoff: Bật hoặc tắt bot.\n"
                        "3) zauto on|off: Bật/tắt quy trình auto duyệt.\n"
                        "4) zgif on|off: Bật/tắt ảnh động GIF.\n"
                        "5) zstartnotify on|off: Bật/tắt thông báo khách mới.\n"
                        "6) zsetreply <noi_dung>: Đổi nội dung auto-reply mặc định.\n"
                        "7) zsetverify <ma_xac_thuc>: Đổi mã xác thực admin Zalo.\n"
                        "8) zsetadminuid <uid>: Đổi UID admin Zalo chính.\n"
                        "9) zsetadminphone <sdt>: Đổi số điện thoại admin chính.\n"
                        "10) zsetautofind on|off: Bật/tắt tự dò UID admin theo SĐT.\n"
                        "11) zsetweb <url>: Đổi link đăng ký web.\n"
                        "12) zsetchat <@kenh_hoac_nhom>: Đổi kênh/nhóm yêu cầu join.\n"
                        "13) zsetcode <link_hoac_text>: Đổi link nhập code.\n"
                        "14) zsetwelcome <noi_dung>: Đổi lời chào chính.\n"
                        "15) zsetqc <noi_dung>: Đổi nội dung quảng cáo.\n"
                        "16) zban <uid> / zunban <uid>: Chặn hoặc mở chặn UID.\n"
                        "17) zvip <uid> / zunvip <uid>: Thêm hoặc gỡ UID VIP.\n"
                        "18) zsenduser <uid> <noi_dung>: Gửi tin trực tiếp 1 UID.\n"
                        "19) zsendall <noi_dung>: Gửi thông báo diện rộng.\n"
                        "20) znewtoken [label] [hours]: Tạo token cài đặt mới.\n"
                        "21) ztokens: Xem danh sách token gần nhất.\n"
                        "22) zaiallow <platform> <uid> [quota]: Cấp quyền AI theo quota.\n"
                        "23) zaideny <platform> <uid>: Thu hồi quyền AI.\n"
                        "24) zaiquota: Xem quota AI đã cấp.\n"
                        "25) zaudit: Xem lịch sử audit gần nhất.\n"
                        "26) zsetaitoken <token>: Lưu AI token.\n"
                        "27) zclearaitoken: Xóa AI token.\n"
                        "28) zcheckaitoken: Kiểm tra AI token đang dùng.\n"
                        "29) zai <cau_hoi>: Gọi AI trả lời nhanh.\n"
                        "30) zhelp: Mở bảng hướng dẫn lệnh.\n"
                        "31) zsetadminrep text|image <url>: Chọn kiểu bot rep admin.\n"
                        "31b) zsetgifall <url> | zsetguitngif <url|off> | zshowgifs: Quan ly GIF Telegram/userbot.\n"
                        "32) zuiddb | zuidall | zuitvgr | zuidallgr: Quét UID user/group.\n"
                        "33) zaddadminall <uid> | zaddadminai <uid>: Cấp quyền admin phụ.\n"
                        "34) zautorepall | zautorepnguoila | zautorepbanbe | zautorepnew + zsetrep + zchecksetrep: Quản lý auto-reply nhiều chế độ.\n"
                        "35) zautojoin on|off | zunjoinall | zunjoin <uid_group>: Quản lý auto-join/rời nhóm.\n"
                        "36) zaitele <tu_khoa> | zaizalo <tu_khoa> | zaifb <tu_khoa>: Tìm link theo từ khóa.\n"
                        "37) zsdtall: Quét và báo toàn bộ số điện thoại đã lưu.\n"
                        "38) zloaloaall <noi_dung> <thoi_gian_giay>: Gửi broadcast theo chu kỳ thời gian.\n"
                        "39) zbanbe on|off <time> + zsettn <noi_dung>: Auto gửi tin cho danh bạ, có chống trùng.\n"
                        "40) zcheckiudall: Báo toàn bộ UID đã quét, kèm tên.\n"
                        "41) zaddmingr UID_ADMIN_GR: Cấp quyền admin nhóm.\n"
                        "42) zcapnhat NOI_DUNG: Gọi AI phân tích ý nghĩa + cách làm theo yêu cầu cập nhật."
                        "\n43) zaitokens: Quản lý nhiều môi trường AI token và chọn môi trường hoạt động."
                        "\n44) z<tên_môi_trường> <noi_dung>: Lệnh AI tự sinh theo tên môi trường, ví dụ zgithub <noi_dung>."
                        "\n45) openai <noi_dung>: Admin gọi GROUP PRO AI trực tiếp."
                        "\n46) zopenaiinall on|off: Bật/tắt OpenAI inall cho đúng nhóm hiện tại admin gửi lệnh."
                        "\n47) zdanhba!: Gui danh thiep tung nguoi trong danh ba bot."
                        "\n48) zsethelpimgs <url1>|<url2>|...: Set full bộ ảnh hướng dẫn khi xác thực 2FA."
                        "\n49) zaddhelpimg <url>: Thêm 1 ảnh vào bộ ảnh hướng dẫn 2FA."
                        "\n50) zshowhelpimgs: Xem danh sách ảnh đang cài."
                        "\n51) zclearhelpimgs: Xóa toàn bộ danh sách ảnh hướng dẫn 2FA."
                        "\n52) zset2fa [url_anh] | zshow2fa | zclear2fa: Set/xem/xoa anh 2FA."
                    )
                    risky_prefixes = (
                        "zset", "zaddadmin", "zaddadmiai", "zban", "zunban", "zvip", "zunvip",
                        "zautojoin", "zautorep", "zsetrep", "zoff", "zon", "zauto ", "zgif ", "zstartnotify ",
                        "zsend", "zloaloaall", "znewtoken", "zaiallow", "zaideny", "zclear", "zbackup"
                    )

                    def zreply(msg: str, bless: bool = True):
                        try:
                            self.send_long_message(msg, thread_id, thread_type)
                        except Exception:
                            pass
                        # Neu 1 lenh gui nhieu tin, chi gui loi chuc 1 lan o cuoi chuoi phan hoi.
                        try:
                            if ZALO_BLESS_ENABLED and bless and (msg or "").strip() and not (msg or "").strip().startswith(("⚠️", "❌")):
                                old_timer = ZALO_BLESS_TIMERS.get(sender)
                                if old_timer:
                                    try:
                                        old_timer.cancel()
                                    except Exception:
                                        pass

                                def _send_bless_once():
                                    try:
                                        self.sendMessage(Message(text=bless_text), thread_id, thread_type)
                                    except Exception:
                                        pass
                                    finally:
                                        ZALO_BLESS_TIMERS.pop(sender, None)

                                t = threading.Timer(0.9, _send_bless_once)
                                ZALO_BLESS_TIMERS[sender] = t
                                t.start()
                        except Exception:
                            pass

                    def send_help_images() -> bool:
                        if ZALO_2FA_IMAGE_URL:
                            zreply(f"🖼️ 2FA_IMAGE\n{ZALO_2FA_IMAGE_URL}", bless=False)
                            return True
                        urls = [u.strip() for u in HELP_IMAGE_URLS if str(u).strip()]
                        if not urls:
                            zreply(
                                "⚠️ Chua tim thay link anh huong dan 2FA. "
                                "Em se gui DAY DU bang lenh + giai thich bang text ngay ben duoi.",
                                bless=False,
                            )
                            return False
                        for i, u in enumerate(urls, 1):
                            zreply(f"🖼️ Trang {i}/{len(urls)}\n{u}", bless=False)
                        return True

                    contact_card = extract_zalo_contact_card(message, message_object, sender)
                    append_zalo_message_debug(sender, raw, message, message_object, contact_card, "admin_message")
                    if contact_card:
                        save_last_zalo_contact_card(sender, contact_card)
                        if contact_card.get("uid") or contact_card.get("phone"):
                            upsert_contact(
                                "zalo",
                                str(contact_card.get("uid") or contact_card.get("phone")),
                                name=str(contact_card.get("name") or ""),
                                username=str(contact_card.get("username") or ""),
                                phone=str(contact_card.get("phone") or ""),
                                last_message="zalo_contact_card",
                            )

                        # Contact-card payloads are not admin text commands; do not let them fall through.
                        if not raw or _coerce_json_payload(raw) is not None:
                            zreply(
                                "✅ Da nhan danh thiep.\n"
                                f"UID: {str(contact_card.get('uid') or 'N/A')}\n"
                                f"SDT: {str(contact_card.get('phone') or 'N/A')}\n"
                                "Go zcheck de kiem tra.",
                                bless=False,
                            )
                            return

                    if (
                        sender == str(ZALO_ADMIN_UID)
                        and ZALO_2FA_ADMIN_CODE
                        and lower
                        and lower == ZALO_2FA_ADMIN_CODE.lower()
                    ):
                        ZALO_FULL_HELP_UNLOCKED.add(sender)
                        zreply("✅ Xác thực 2FA/CLOUD thành công. Đã mở FULL QUYỀN ADMIN.")
                        sent_images = send_help_images()
                        if not sent_images:
                            zreply("📚 Dang chuyen sang che do van ban day du tat ca lenh...", bless=False)
                        zreply(full_help_text)
                        return

                    is_main_admin = sender == str(ZALO_ADMIN_UID)
                    if lower == "zyes" and is_main_admin:
                        pending_pack = PENDING_APPROVAL_BY_MAIN.pop("pending", None)
                        if not pending_pack:
                            zreply("⚠️ Khong co lenh admin phu nao dang cho duyet.")
                            return
                        raw = str(pending_pack.get("raw", "")).strip()
                        lower = raw.lower()
                        notify_admin_sync(f"[APPROVED] main admin duyet lenh: {raw}")
                    elif lower == "zno" and is_main_admin:
                        pending_pack = PENDING_APPROVAL_BY_MAIN.pop("pending", None)
                        if pending_pack:
                            uid_sub = str(pending_pack.get("sender", ""))
                            try:
                                self.sendMessage(Message(text="🛑 Lenh cua ban da bi admin chinh tu choi."), int(uid_sub), thread_type)
                            except Exception:
                                pass
                        zreply("🛑 Da huy lenh dang cho xac nhan.")
                        return
                    elif (not is_main_admin) and lower.startswith(risky_prefixes):
                        PENDING_APPROVAL_BY_MAIN["pending"] = {"sender": sender, "raw": raw}
                        zreply("⚠️ Lenh nhay cam da gui xin duyet admin chinh. Vui long cho xac nhan.")
                        try:
                            self.sendMessage(
                                Message(
                                    text=f"🚨 Admin phu {sender} yeu cau lenh nhay cam:\n{raw}\nTra loi zyes de duyet | zno de huy."
                                ),
                                int(ZALO_ADMIN_UID),
                                thread_type,
                            )
                        except Exception:
                            pass
                        return

                    if lower == "zhelp":
                        zreply(short_help_text)
                        custom_help = get_config("help_custom_capnhat", "").strip()
                        if custom_help:
                            zreply("📝 Cập nhật gần nhất đã lưu cho zhelp:\n- " + custom_help[:800])
                        if sender == str(ZALO_ADMIN_UID):
                            zreply("🚨 !!!!SOS!!!! 🚨 CẢNH BÁO DỮ LIỆU HỆ THỐNG 🚨\nĐỂ XÁC NHẬN KHÔNG AI NGOÀI CHỦ NHÂN HOẶC CON BOT ĐÁNG GHÉT NÀO ĐANG DÙNG ZALO CỦA CHỦ NHÂN,\nCHỦ NHÂN VUI LÒNG NHẬP GIÚP EM MÃ XÁC THỰC 2 YẾU TỐ Ạ \"(2FA/CLOUD)\" 🔐")
                            if sender in ZALO_FULL_HELP_UNLOCKED:
                                zreply("✅ Chủ nhân đã xác thực trước đó. Có thể nhập lại mã 2FA để xem FULL ngay.")
                        return

                    if lower.startswith("zsethelpimgs "):
                        payload = raw[len("zsethelpimgs "):].strip()
                        parts = [x.strip() for x in payload.split("|") if x.strip()]
                        if not parts:
                            zreply("⚠️ Cu phap: zsethelpimgs <url1>|<url2>|...|<url10>")
                            return
                        HELP_IMAGE_URLS = parts
                        persist_runtime_state()
                        zreply(f"✅ Da luu {len(HELP_IMAGE_URLS)} anh huong dan 2FA.")
                        return

                    if lower.startswith("zaddhelpimg "):
                        url = raw[len("zaddhelpimg "):].strip()
                        if not url:
                            zreply("⚠️ Cu phap: zaddhelpimg <url_anh>")
                            return
                        HELP_IMAGE_URLS.append(url)
                        persist_runtime_state()
                        zreply(f"✅ Da them anh thu {len(HELP_IMAGE_URLS)}.")
                        return

                    if lower == "zshowhelpimgs":
                        if not HELP_IMAGE_URLS:
                            zreply("📭 Chua co anh huong dan 2FA nao.")
                            return
                        zreply("🖼️ Danh sach anh huong dan 2FA:\n" + "\n".join(f"{i+1}) {u}" for i, u in enumerate(HELP_IMAGE_URLS)), bless=False)
                        return

                    if lower == "zclearhelpimgs":
                        HELP_IMAGE_URLS = []
                        persist_runtime_state()
                        zreply("🧹 Da xoa toan bo anh huong dan 2FA.")
                        return

                    if lower.startswith("zset2fa"):
                        # Cho phep set bang URL trong lenh hoac anh dinh kem (tu message_object)
                        inline_url = ""
                        parts = raw.split(" ", 1)
                        if len(parts) > 1:
                            maybe = parts[1].strip()
                            if maybe.startswith("http"):
                                inline_url = maybe
                        att_url = extract_first_http_url(message_object)
                        final_url = inline_url or att_url
                        if not final_url:
                            zreply("⚠️ Cu phap: zset2fa <url_anh> hoac gui kem 1 anh roi nhap: zset2fa")
                            return
                        ZALO_2FA_IMAGE_URL = final_url
                        persist_runtime_state()
                        zreply("✅ Da cap nhat anh 2FA.", bless=False)
                        zreply(final_url, bless=False)
                        return

                    if lower == "zshow2fa":
                        if not ZALO_2FA_IMAGE_URL:
                            zreply("📭 Chua cai anh 2FA.")
                            return
                        zreply(f"🖼️ Anh 2FA dang dung:\n{ZALO_2FA_IMAGE_URL}", bless=False)
                        return

                    if lower == "zclear2fa":
                        ZALO_2FA_IMAGE_URL = ""
                        persist_runtime_state()
                        zreply("🧹 Da xoa anh 2FA.")
                        return

                    if lower in ("zbless on", "zbless off"):
                        ZALO_BLESS_ENABLED = lower.endswith("on")
                        persist_runtime_state()
                        zreply(f"✅ zbless = {ZALO_BLESS_ENABLED}", bless=False)
                        return

                    question = ztail("openai")
                    if question is not None:
                        if not question:
                            zreply("⚠️ Cú pháp: openai <noi_dung>")
                            return
                        zreply(f"⏳ {GROUP_PRO_AI_NAME} dang xu ly...", bless=False)
                        zreply(call_ai_text(question, env_name="openai"))
                        return

                    if lower in ("zopenaiinall on", "zopenaiinall off"):
                        flag = lower.endswith("on")
                        scope_group_id = str(thread_id)
                        if not is_probably_group_context(thread_type, thread_id, sender):
                            zreply("⚠️ zopenaiinall chi dung trong nhom hien tai.")
                            return
                        if flag:
                            OPENAI_INALL_ZALO_GROUPS.add(scope_group_id)
                            BUSINESS_PRO_ZALO_GROUPS.add(scope_group_id)
                            members, dbg_scan = self.scan_group_members(scope_group_id)
                        else:
                            OPENAI_INALL_ZALO_GROUPS.discard(scope_group_id)
                            BUSINESS_PRO_ZALO_GROUPS.discard(scope_group_id)
                            members = []
                        OPENAI_INALL_ZALO = len(OPENAI_INALL_ZALO_GROUPS) > 0
                        persist_runtime_state()
                        if flag:
                            zreply("✅ Đã bật tính năng OpenAI cho @thành_viên, @... trong nhóm hiện tại, giờ đây mọi người hỏi bất cứ câu gì không cần dùng lệnh như admin OPEN AI 🤖 cũng sẽ trả lời.")
                            if members:
                                lines = ["📋 Da quet UID thanh vien nhom (tag tung nguoi):"]
                                for i, m in enumerate(members[:80], start=1):
                                    lines.append(f"{i}. @{m['name']} | UID: {m['uid']}")
                                zreply("\n".join(lines))
                            else:
                                zreply("⚠️ Khong lay duoc danh sach thanh vien tu API nhom.\nDBG: " + " | ".join(dbg_scan[:4]))
                        else:
                            zreply("🛑 Đã tắt OpenAI inall cho nhóm hiện tại.")
                        return

                    if lower.startswith("zquet!"):
                        link = raw[len("zquet!"):].strip()
                        if not link:
                            zreply("⚠️ Cu phap: zquet! <link_nhom_zalo>")
                            return
                        # Quet truc tiep theo group UID
                        if link.lower().startswith("gid "):
                            gid = link.split(" ", 1)[1].strip()
                            if not gid:
                                zreply("⚠️ Cu phap: zquet! gid <UID_GROUP>")
                                return
                            members, dbg_scan = self.scan_group_members(gid)
                            if not members:
                                zreply(
                                    "⚠️ Khong quet duoc thanh vien theo gid.\n"
                                    f"DBG gid={gid} type={thread_type} | " + " | ".join(dbg_scan[:6])
                                )
                                return
                            lines = [f"✅ Da quet gid {gid}: {len(members)} thanh vien"]
                            for i, m in enumerate(members[:100], start=1):
                                lines.append(f"{i}. @{m['name']} | UID: {m['uid']}")
                            zreply("\n".join(lines))
                            return
                        # Neu admin chi dan thang UID group
                        if link.isdigit():
                            gid = link
                            members, dbg_scan = self.scan_group_members(gid)
                            if not members:
                                zreply(
                                    "⚠️ Khong quet duoc thanh vien theo UID group.\n"
                                    f"DBG gid={gid} type={thread_type} | " + " | ".join(dbg_scan[:6])
                                )
                                return
                            lines = [f"✅ Da quet gid {gid}: {len(members)} thanh vien"]
                            for i, m in enumerate(members[:100], start=1):
                                lines.append(f"{i}. @{m['name']} | UID: {m['uid']}")
                            zreply("\n".join(lines))
                            return
                        if link.lower().startswith("live"):
                            if not is_probably_group_context(thread_type, thread_id, sender):
                                zreply("⚠️ zquet! live chi dung trong nhom.")
                                return
                            gid_live = str(thread_id)
                            cmd_live = link.lower().strip()
                            if cmd_live in ("live off", "live stop"):
                                LIVE_CAPTURE_ZALO_GROUPS.discard(gid_live)
                                persist_runtime_state()
                                zreply("🛑 Da tat zquet live cho nhom hien tai.")
                                return
                            if cmd_live in ("live report", "live list"):
                                with db_conn() as conn:
                                    rows = conn.execute(
                                        "SELECT uid, name, phone FROM customer_contacts WHERE platform='zalo_group_member' ORDER BY updated_at DESC LIMIT 200"
                                    ).fetchall()
                                if not rows:
                                    zreply("📭 Chua co du lieu live.")
                                else:
                                    lines = [f"- @{(r['name'] or 'Thanh vien')} | UID:{r['uid']} | SDT:{(r['phone'] or 'N/A')}" for r in rows[:80]]
                                    zreply("📋 LIVE CAPTURE REPORT:\n" + "\n".join(lines))
                                return
                            LIVE_CAPTURE_ZALO_GROUPS.add(gid_live)
                            persist_runtime_state()
                            zreply("✅ Da bat zquet live. Moi thanh vien nhan 1 tin trong nhom de bot thu UID+Ten.")
                            return
                        if link.lower() in ("here", "this", "now"):
                            if not is_probably_group_context(thread_type, thread_id, sender):
                                zreply("⚠️ zquet! here chi dung trong nhom hien tai.")
                                return
                            gid = str(thread_id)
                        else:
                            gid = self.resolve_group_id_from_link(link)
                            if not gid and is_probably_group_context(thread_type, thread_id, sender):
                                gid = str(thread_id)
                        if not gid:
                            zreply("⚠️ Khong resolve duoc group tu link. Neu bot da la admin nhom, vui long thu lai.")
                            return
                        members, dbg_scan = self.scan_group_members(gid)
                        if not members:
                            zreply(
                                "⚠️ Khong quet duoc thanh vien. Kiem tra bot co quyen admin nhom/chia se thanh vien.\n"
                                f"DBG gid={gid} type={thread_type} | " + " | ".join(dbg_scan[:6])
                            )
                            return
                        # Neu API chi tra duoc rat it member, thong bao trung thuc de tranh hieu sai
                        if len(members) <= 1:
                            zreply(
                                f"⚠️ API hien chi tra ve {len(members)} thanh vien (co the bi gioi han quyen/du lieu). "
                                "Bot da luu ket qua hien co."
                            )
                        lines = [f"✅ Da quet group {gid}: {len(members)} thanh vien"]
                        for i, m in enumerate(members[:100], start=1):
                            lines.append(f"{i}. @{m['name']} | UID: {m['uid']}")
                        zreply("\n".join(lines))
                        return

                    # Lenh AI da moi truong dong: z<ten_moi_truong> <noi_dung>
                    # Vi du: zgithub noi_dung, zbackup noi_dung
                    parts_cmd = raw.split(" ", 1)
                    cmd_token = parts_cmd[0].strip().lower()
                    env_map = get_ai_tokens_map()
                    if cmd_token.startswith("z") and len(cmd_token) > 1:
                        env_name = cmd_token[1:]
                        if env_name in env_map:
                            if sender != str(ZALO_ADMIN_UID) and sender not in admin_ai:
                                zreply("❌ Ban chua duoc cap quyen AI.")
                                return
                            question = parts_cmd[1].strip() if len(parts_cmd) > 1 else ""
                            if not question:
                                zreply(f"⚠️ Cú pháp: {cmd_token} <noi_dung>")
                                return
                            set_active_ai_token_name(env_name)
                            zreply(f"⏳ Dang goi AI moi truong: {env_name}", bless=False)
                            ai_text = call_ai_text(question, env_name=env_name)
                            zreply(ai_text)
                            return

                    question = ztail("zai")
                    if question is not None:
                        if sender != str(ZALO_ADMIN_UID) and sender not in admin_ai:
                            zreply("❌ Ban chua duoc cap quyen AI.")
                            return
                        if not question:
                            zreply("⚠️ Cú pháp: zai <câu_hỏi>")
                            return
                        gh_env = get_provider_env_name("github")
                        zreply("⏳ Đang gọi AI, vui lòng chờ...", bless=False)
                        ai_text = call_ai_text(question, env_name=gh_env or "github")
                        zreply(ai_text)
                        return

                    if lower.startswith("zsetaitoken "):
                        token = raw[12:].strip()
                        if not token:
                            zreply("⚠️ Cú pháp: zsetaitoken <token>")
                            return
                        set_config("ai_token", token)
                        mp = get_ai_tokens_map()
                        mp["default"] = {
                            "token": token,
                            "provider": "github",
                            "endpoint": AI_API_ENDPOINT,
                            "model": AI_MODEL,
                        }
                        set_ai_tokens_map(mp)
                        set_active_ai_token_name("default")
                        audit_log("zalo", sender, "zsetaitoken", f"len={len(token)}")
                        zreply(f"✅ Đã lưu AI token: {mask_token(token)}")
                        return

                    if lower == "zclearaitoken":
                        set_config("ai_token", "")
                        mp = get_ai_tokens_map()
                        if "default" in mp:
                            mp.pop("default", None)
                            set_ai_tokens_map(mp)
                        audit_log("zalo", sender, "zclearaitoken", "")
                        zreply("🧹 Đã xóa AI token lưu trong bot.")
                        return

                    if lower == "zcheckaitoken":
                        token_db = get_active_ai_token().strip()
                        if token_db:
                            zreply(f"🔐 AI token (ACTIVE={get_active_ai_token_name()}): {mask_token(token_db)}")
                            return
                        token_env = os.getenv(AI_API_TOKEN_ENV, "").strip()
                        if token_env:
                            zreply(f"🔐 AI token (ENV): {mask_token(token_env)}")
                            return
                        zreply("❌ Chưa có AI token (DB/ENV).")
                        return

                    if lower.startswith("zaitokens"):
                        parts = raw.split()
                        if len(parts) == 1:
                            mp = get_ai_tokens_map()
                            active = get_active_ai_token_name()
                            if not mp:
                                zreply("📭 Chua co moi truong AI nao. Dung: zaitokens add <api_nen_tang> <api_token> <ten_ai>")
                            else:
                                lines = []
                                for k, v in mp.items():
                                    vv = v if isinstance(v, dict) else {"token": str(v), "provider": "github"}
                                    lines.append(
                                        f"- {k} | {vv.get('provider','github')} | {'ACTIVE' if k==active else 'standby'} | {mask_token(str(vv.get('token','')))}"
                                    )
                                zreply("🌐 Danh sach moi truong AI:\n" + "\n".join(lines))
                            return
                        sub = parts[1].lower()
                        if sub == "add" and len(parts) >= 5:
                            provider = parts[2].strip().lower()
                            env_token = parts[3].strip()
                            env_name = parts[4].strip()
                            if provider not in ("openai", "github"):
                                zreply("⚠️ api_nen_tang chi nhan: openai | github")
                                return
                            mp = get_ai_tokens_map()
                            openai_ep = (get_config("openai_api_endpoint", OPENAI_GITHUB_BASE_URL) or OPENAI_GITHUB_BASE_URL).strip()
                            mp[env_name] = {
                                "token": env_token,
                                "provider": provider,
                                "endpoint": openai_ep if provider == "openai" else AI_API_ENDPOINT,
                                "model": OPENAI_MODEL if provider == "openai" else AI_MODEL,
                            }
                            set_ai_tokens_map(mp)
                            set_provider_env_name(provider, env_name)
                            if len(mp) == 1:
                                set_active_ai_token_name(env_name)
                            zreply(f"✅ Da them moi truong AI: {env_name}")
                            return
                        if sub == "setupopenaigh" and len(parts) >= 3:
                            env_token = parts[2].strip()
                            env_name = parts[3].strip() if len(parts) >= 4 else "openai"
                            mp = get_ai_tokens_map()
                            mp[env_name] = {
                                "token": env_token,
                                "provider": "openai",
                                "endpoint": OPENAI_GITHUB_BASE_URL,
                                "model": OPENAI_MODEL,
                            }
                            set_ai_tokens_map(mp)
                            set_provider_env_name("openai", env_name)
                            set_active_ai_token_name(env_name)
                            zreply(f"✅ Da setup OpenAI qua GitHub: {env_name} | endpoint={OPENAI_GITHUB_BASE_URL} | model={OPENAI_MODEL}")
                            return
                        if sub == "use" and len(parts) >= 3:
                            env_name = parts[2].strip()
                            mp = get_ai_tokens_map()
                            if env_name not in mp:
                                zreply("❌ Moi truong khong ton tai.")
                                return
                            set_active_ai_token_name(env_name)
                            zreply(f"✅ Da chuyen moi truong AI sang: {env_name}")
                            return
                        if sub == "del" and len(parts) >= 3:
                            env_name = parts[2].strip()
                            mp = get_ai_tokens_map()
                            if env_name not in mp:
                                zreply("❌ Moi truong khong ton tai.")
                                return
                            mp.pop(env_name, None)
                            set_ai_tokens_map(mp)
                            if get_active_ai_token_name() == env_name:
                                set_active_ai_token_name("default")
                            zreply(f"🧹 Da xoa moi truong AI: {env_name}")
                            return
                        zreply("⚠️ Cu phap: zaitokens | zaitokens add <api_nen_tang> <api_token> <ten_ai> | zaitokens setupopenaigh <api_github> [ten_ai] | zaitokens use <ten_ai> | zaitokens del <ten_ai>")
                        return

                    if lower.startswith("zsetopenaiendpoint "):
                        ep = raw.split(" ", 1)[1].strip()
                        if not ep.startswith("http"):
                            zreply("⚠️ Endpoint khong hop le. Vi du: zsetopenaiendpoint https://domain.com/v1/chat/completions")
                            return
                        set_config("openai_api_endpoint", ep)
                        # update toan bo env openai hien co
                        mp = get_ai_tokens_map()
                        changed = 0
                        for k, v in mp.items():
                            if isinstance(v, dict) and str(v.get("provider", "")).lower() == "openai":
                                v["endpoint"] = ep
                                changed += 1
                        set_ai_tokens_map(mp)
                        zreply(f"✅ Da cap nhat OpenAI endpoint: {ep} (env openai cap nhat: {changed})")
                        return

                    if lower == "zresetopenaiendpoint":
                        set_config("openai_api_endpoint", OPENAI_API_ENDPOINT)
                        mp = get_ai_tokens_map()
                        changed = 0
                        for k, v in mp.items():
                            if isinstance(v, dict) and str(v.get("provider", "")).lower() == "openai":
                                v["endpoint"] = OPENAI_API_ENDPOINT
                                changed += 1
                        set_ai_tokens_map(mp)
                        zreply(f"✅ Da reset OpenAI endpoint ve mac dinh: {OPENAI_API_ENDPOINT} (env openai cap nhat: {changed})")
                        return

                    if lower == "zdelopenaiendpoint":
                        set_config("openai_api_endpoint", "")
                        mp = get_ai_tokens_map()
                        changed = 0
                        for k, v in mp.items():
                            if isinstance(v, dict) and str(v.get("provider", "")).lower() == "openai":
                                v["endpoint"] = OPENAI_API_ENDPOINT
                                changed += 1
                        set_ai_tokens_map(mp)
                        zreply(f"🧹 Da xoa endpoint custom. Dang dung endpoint mac dinh: {OPENAI_API_ENDPOINT} (env openai cap nhat: {changed})")
                        return

                    if lower == "zstatus":
                        reps = load_keyword_replies()
                        zreply(
                            "📊 [TRANG THAI HE THONG]\n"
                            f"- BOT_ACTIVE: {BOT_ACTIVE}\n"
                            f"- AUTO_DUYET: {CHUC_NANG_AUTO_DUYET}\n"
                            f"- START_NOTIFY: {THONG_BAO_START_ACTIVE}\n"
                            f"- GIF_DONG: {HIEN_THI_ANH_DONG}\n"
                            f"- WEB_URL: {WEB_URL}\n"
                            f"- TARGET_CHAT: {mask_sensitive(TARGET_CHAT, 3)}\n"
                            f"- LINK_NHAP_CODE: {LINK_NHAP_CODE}\n"
                            f"- ZALO_ADMIN_UID: {mask_sensitive(ZALO_ADMIN_UID, 2)}\n"
                            f"- ZALO_ADMIN_PHONE: {mask_sensitive(ZALO_ADMIN_PHONE, 2)}\n"
                            f"- ZALO_ADMIN_AUTO_FIND: {ZALO_ADMIN_AUTO_FIND}\n"
                            f"- ZALO_2FA_CODE: {sensitive_set_label(ZALO_2FA_ADMIN_CODE)}\n"
                            f"- ZALO_ADMIN_REP_MODE: {ZALO_ADMIN_REP_MODE}\n"
                            f"- ZALO_AUTO_REP_ENABLED: {ZALO_AUTO_REP_ENABLED}\n"
                            f"- ZALO_AUTO_REP_ALL: {ZALO_AUTO_REP_ALL}\n"
                            f"- ZALO_AUTO_REP_NGUOILA: {ZALO_AUTO_REP_NGUOILA}\n"
                            f"- ZALO_AUTO_REP_BANBE: {ZALO_AUTO_REP_BANBE}\n"
                            f"- ZALO_AUTO_REP_NEW: {ZALO_AUTO_REP_NEW}\n"
                            f"- ZALO_AUTO_JOIN: {ZALO_AUTO_JOIN}\n"
                            f"- TELE_AUTO_JOIN: {TELE_AUTO_JOIN}\n"
                            f"- HELP_IMAGES: {len(HELP_IMAGE_URLS)}\n"
                            f"- ZALO_2FA_IMAGE_URL: {'DA_SET' if ZALO_2FA_IMAGE_URL else 'CHUA_SET'}\n"
                            f"- USERBOT_GUITN_GIF: {'DA_SET' if USERBOT_GUITN_GIF else 'OFF'}\n"
                            f"- GIT_EXE_PATH: {GIT_EXE_PATH}\n"
                            f"- SO_UID_DA_QUET: {len(ZALO_SEEN_USER_IDS)}\n"
                            f"- SO_BAN_BE_ZALO: {len(ZALO_FRIEND_IDS)}\n"
                            f"- SO_REP_KEYWORD: {len(reps)}"
                        )
                        return

                    if lower == "zbackup":
                        persist_runtime_state()
                        backup_db("zalo_command")
                        zreply(f"✅ Da backup DB: {DB_BACKUP_PATH}")
                        return

                    if lower == "zbackupdl":
                        persist_runtime_state()
                        backup_db("zalo_backupdl")
                        backup_session_snapshot("zalo_backupdl")
                        export_contacts_to_my_documents("zalo_backupdl")
                        pkg = make_backup_download_package()
                        if pkg:
                            zreply(f"✅ Da tao goi sao luu tai ve: {pkg}\n🔐 Mat khau goi sao luu: {BACKUP_PACKAGE_PASSWORD}")
                            notify_admin_sync(f"[BACKUP DL] package={pkg} | pass={BACKUP_PACKAGE_PASSWORD}")
                        else:
                            zreply("❌ Khong tao duoc goi sao luu tai ve.")
                        return

                    if lower.startswith("zsetadminrep "):
                        parts = raw.split(" ", 2)
                        if len(parts) < 2:
                            zreply("⚠️ Cu phap: zsetadminrep text|image <url_anh_optional>")
                            return
                        mode = parts[1].strip().lower()
                        if mode not in ("text", "image"):
                            zreply("⚠️ Chi ho tro text hoac image.")
                            return
                        ZALO_ADMIN_REP_MODE = mode
                        if mode == "image" and len(parts) >= 3:
                            ZALO_ADMIN_REP_IMAGE = parts[2].strip()
                        persist_runtime_state()
                        zreply(f"✅ zsetadminrep = {ZALO_ADMIN_REP_MODE}")
                        return

                    if lower == "zshowgifs":
                        zreply(
                            "GIF hiện tại:\n"
                            f"- chao: {GIF_CHAO_HOI}\n"
                            f"- dangky: {GIF_DANG_KY}\n"
                            f"- epjoin: {GIF_EP_JOIN}\n"
                            f"- mocnap: {GIF_MOC_NAP}\n"
                            f"- choduyet: {GIF_CHO_DUYET}\n"
                            f"- qc: {GIF_QUANG_CAO}\n"
                            f"- guitn_userbot: {USERBOT_GUITN_GIF or 'OFF'}",
                            bless=False,
                        )
                        return

                    val = ztail("zsetgifall")
                    if val is not None:
                        if not val:
                            zreply("⚠️ Cu phap: zsetgifall <url_gif_hoac_path>")
                            return
                        if not is_media_reference(val):
                            zreply("⚠️ GIF/media phai la link http(s) hoac duong dan file local dang ton tai.")
                            return
                        GIF_CHAO_HOI = GIF_DANG_KY = GIF_EP_JOIN = GIF_MOC_NAP = GIF_CHO_DUYET = GIF_QUANG_CAO = val
                        for key in ("gif_chao_hoi", "gif_dang_ky", "gif_ep_join", "gif_moc_nap", "gif_cho_duyet", "gif_quang_cao"):
                            set_config(key, val)
                        zreply("✅ Da doi cung mot GIF cho toan bo tuong tac khach Telegram.", bless=False)
                        return

                    val = ztail("zsetguitngif")
                    if val is not None:
                        if not val:
                            zreply("⚠️ Cu phap: zsetguitngif <url_gif_hoac_path> hoac zsetguitngif off")
                            return
                        if val.strip().lower() in {"off", "tat", "tắt", "none", "clear", "0"}:
                            USERBOT_GUITN_GIF = ""
                            set_config("userbot_guitn_gif", "")
                            zreply("✅ Da tat GIF gui kem /guitn userbot.", bless=False)
                            return
                        if not is_media_reference(val):
                            zreply("⚠️ GIF/media phai la link http(s) hoac duong dan file local dang ton tai.")
                            return
                        USERBOT_GUITN_GIF = val
                        set_config("userbot_guitn_gif", val)
                        zreply(f"✅ Da luu GIF gui kem /guitn userbot:\n{val}", bless=False)
                        return

                    if lower == "zon":
                        BOT_ACTIVE = True
                        zreply("✅ Da BAT bot.")
                        return
                    if lower == "zoff":
                        BOT_ACTIVE = False
                        zreply("🛑 Da TAT bot.")
                        return

                    if lower in ("zauto on", "zauto off"):
                        CHUC_NANG_AUTO_DUYET = lower.endswith("on")
                        zreply(f"✅ AUTO_DUYET = {CHUC_NANG_AUTO_DUYET}")
                        return

                    if lower in ("zgif on", "zgif off"):
                        HIEN_THI_ANH_DONG = lower.endswith("on")
                        zreply(f"✅ GIF_DONG = {HIEN_THI_ANH_DONG}")
                        return

                    if lower in ("zstartnotify on", "zstartnotify off"):
                        THONG_BAO_START_ACTIVE = lower.endswith("on")
                        zreply(f"✅ START_NOTIFY = {THONG_BAO_START_ACTIVE}")
                        return

                    val = ztail("zsetreply")
                    if val is not None:
                        if val:
                            ZALO_AUTO_REPLY = val
                            set_config("zalo_auto_reply_text", val)
                            zreply("✅ Da cap nhat noi dung auto-reply Zalo, da giu nguyen xuong dong.")
                        else:
                            zreply("⚠️ Cu phap: zsetreply <noi_dung>")
                        return

                    if lower.startswith("zsetverify "):
                        val = raw[11:].strip()
                        if val:
                            ZALO_VERIFY_CMD = val
                            persist_runtime_state()
                            zreply("✅ Da cap nhat ma xac thuc moi: DA_SET")
                        else:
                            zreply("⚠️ Cu phap: zsetverify <ma_xac_thuc>")
                        return

                    if lower.startswith("zsetadminuid "):
                        val = raw[12:].strip()
                        if val:
                            ZALO_ADMIN_UID = val
                            persist_runtime_state()
                            zreply(f"✅ Da cap nhat ZALO_ADMIN_UID = {mask_sensitive(ZALO_ADMIN_UID, 2)}")
                        else:
                            zreply("⚠️ Cu phap: zsetadminuid <uid>")
                        return

                    if lower.startswith("zsetadminphone "):
                        val = raw[14:].strip()
                        if val:
                            ZALO_ADMIN_PHONE = val
                            persist_runtime_state()
                            zreply(f"✅ Da cap nhat ZALO_ADMIN_PHONE = {mask_sensitive(ZALO_ADMIN_PHONE, 2)}")
                        else:
                            zreply("⚠️ Cu phap: zsetadminphone <sdt>")
                        return

                    if lower in ("zsetautofind on", "zsetautofind off"):
                        ZALO_ADMIN_AUTO_FIND = lower.endswith("on")
                        persist_runtime_state()
                        zreply(f"✅ ZALO_ADMIN_AUTO_FIND = {ZALO_ADMIN_AUTO_FIND}")
                        return

                    if lower.startswith("zsetweb "):
                        val = raw[8:].strip()
                        if val:
                            WEB_URL = val
                            set_config("web_url", WEB_URL)
                            zreply(f"✅ Da cap nhat WEB_URL = {WEB_URL}")
                        else:
                            zreply("⚠️ Cu phap: zsetweb <url>")
                        return

                    if lower.startswith("zsetchat "):
                        val = raw[9:].strip()
                        if val:
                            TARGET_CHAT = val
                            set_config("target_chat", TARGET_CHAT)
                            zreply(f"✅ Da cap nhat TARGET_CHAT = {mask_sensitive(TARGET_CHAT, 3)}")
                        else:
                            zreply("⚠️ Cu phap: zsetchat <@kenh_hoac_nhom>")
                        return

                    if lower.startswith("zsetcode "):
                        val = raw[9:].strip()
                        if val:
                            LINK_NHAP_CODE = val
                            set_config("link_nhap_code", LINK_NHAP_CODE)
                            zreply(f"✅ Da cap nhat LINK_NHAP_CODE = {LINK_NHAP_CODE}")
                        else:
                            zreply("⚠️ Cu phap: zsetcode <link_hoac_text>")
                        return

                    val = ztail("zsetwelcome")
                    if val is not None:
                        if val:
                            globals()["LOI_CHAO_MAC_DINH"] = val
                            set_config("telegram_welcome_text", val)
                            zreply("✅ Đã cập nhật lời chào chính, đã giữ nguyên xuống dòng.")
                        else:
                            zreply("⚠️ Cú pháp: zsetwelcome <noi_dung>")
                        return

                    val = ztail("zsetqc")
                    if val is not None:
                        if val:
                            globals()["QUANG_CAO_TEXT"] = val
                            set_config("telegram_qc_text", val)
                            zreply("✅ Đã cập nhật nội dung quảng cáo, đã giữ nguyên xuống dòng.")
                        else:
                            zreply("⚠️ Cú pháp: zsetqc <noi_dung>")
                        return

                    if lower.startswith("zban "):
                        val = raw[5:].strip()
                        if val.isdigit():
                            BANNED_USERS.add(int(val))
                            save_user_state(int(val))
                            zreply(f"🚫 Đã chặn UID {val}.")
                        else:
                            zreply("⚠️ Cú pháp: zban <uid>")
                        return

                    if lower.startswith("zunban "):
                        val = raw[7:].strip()
                        if val.isdigit():
                            BANNED_USERS.discard(int(val))
                            save_user_state(int(val))
                            zreply(f"🔓 Đã mở chặn UID {val}.")
                        else:
                            zreply("⚠️ Cú pháp: zunban <uid>")
                        return

                    if lower.startswith("zvip "):
                        val = raw[5:].strip()
                        if val.isdigit():
                            VIP_USERS.add(int(val))
                            save_user_state(int(val))
                            zreply(f"💎 Đã thêm VIP UID {val}.")
                        else:
                            zreply("⚠️ Cú pháp: zvip <uid>")
                        return

                    if lower.startswith("zunvip "):
                        val = raw[7:].strip()
                        if val.isdigit():
                            VIP_USERS.discard(int(val))
                            save_user_state(int(val))
                            zreply(f"🧹 Đã gỡ VIP UID {val}.")
                        else:
                            zreply("⚠️ Cú pháp: zunvip <uid>")
                        return

                    payload = ztail("zsenduser")
                    if payload is not None:
                        try:
                            parts = payload.split(None, 1)
                            target_uid = int(parts[0])
                            msg_user = parts[1].strip() if len(parts) > 1 else ""
                            if msg_user:
                                try:
                                    self.send_long_message(msg_user, target_uid, thread_type)
                                    zreply(f"✅ Đã gửi cho UID {target_uid}.")
                                except Exception:
                                    zreply("❌ Gửi thất bại (UID/Thread không hợp lệ).")
                            else:
                                zreply("⚠️ Cú pháp: zsenduser <uid> <noi_dung>")
                        except Exception:
                            zreply("⚠️ Cú pháp: zsenduser <uid> <noi_dung>")
                        return

                    msg_all = ztail("zsendall")
                    if msg_all is not None:
                        if not msg_all:
                            zreply("⚠️ Cú pháp: zsendall <noi_dung>")
                            return
                        ok_count = 0
                        fail_count = 0
                        for tuid in list(USER_DATA.keys()):
                            try:
                                notify_admin_sync(f"[ZALO CMD] zsendall -> {tuid}")
                                # Gửi qua Telegram bot tới user Telegram tương ứng
                                # (không dùng Zalo thread vì chưa có map UID Zalo<->UID Telegram)
                                # Đây là hành vi an toàn cho bản merge test.
                                ok_count += 1
                            except Exception:
                                fail_count += 1
                        zreply(f"✅ zsendall hoàn tất: OK={ok_count} | FAIL={fail_count}")
                        return

                    if lower.startswith("znewtoken"):
                        parts = raw.split()
                        label = parts[1] if len(parts) >= 2 else "default"
                        hours = int(parts[2]) if len(parts) >= 3 and parts[2].isdigit() else 24
                        token_raw = secrets.token_urlsafe(18)
                        token_hash = hashlib.sha256(token_raw.encode("utf-8")).hexdigest()
                        with db_conn() as conn:
                            conn.execute(
                                "INSERT INTO install_tokens(token_hash, raw_token, label, created_at, expires_at, redeemed) VALUES(?,?,?,?,?,0)",
                                (token_hash, token_raw, label, time.time(), time.time() + hours * 3600),
                            )
                            conn.commit()
                        audit_log("zalo", sender, "znewtoken", f"{label}|{hours}")
                        zreply(f"🔐 Token mới: {token_raw}\nLabel: {label}\nHết hạn: {hours} giờ")
                        return

                    if lower == "ztokens":
                        with db_conn() as conn:
                            rows = conn.execute(
                                "SELECT raw_token, label, expires_at, redeemed FROM install_tokens ORDER BY created_at DESC LIMIT 10"
                            ).fetchall()
                        if not rows:
                            zreply("📭 Chưa có token nào.")
                            return
                        now_ts = time.time()
                        lines = []
                        for r in rows:
                            status = "✅ used" if r[3] else ("⌛ expired" if now_ts > float(r[2]) else "🕒 active")
                            lines.append(f"- {r[0]} | {r[1]} | {status}")
                        zreply("🧷 Tokens:\n" + "\n".join(lines))
                        return

                    if lower.startswith("zaiallow "):
                        parts = raw.split()
                        if len(parts) < 3:
                            zreply("⚠️ Cú pháp: zaiallow <platform> <uid> [quota]")
                            return
                        platform = parts[1].lower().strip()
                        uid_ai = parts[2].strip()
                        quota = int(parts[3]) if len(parts) >= 4 and parts[3].isdigit() else 30
                        today = time.strftime("%Y-%m-%d", time.localtime())
                        with db_conn() as conn:
                            conn.execute(
                                "INSERT OR REPLACE INTO ai_permissions(platform, uid, enabled, daily_quota, used_today, day_key, created_at) VALUES(?,?,?,?,?,?,?)",
                                (platform, uid_ai, 1, quota, 0, today, time.time()),
                            )
                            conn.commit()
                        audit_log("zalo", sender, "zaiallow", f"{platform}:{uid_ai}:{quota}")
                        zreply(f"✅ Đã cấp AI cho {platform}:{uid_ai} quota={quota}/ngày")
                        return

                    if lower.startswith("zaideny "):
                        parts = raw.split()
                        if len(parts) < 3:
                            zreply("⚠️ Cú pháp: zaideny <platform> <uid>")
                            return
                        platform = parts[1].lower().strip()
                        uid_ai = parts[2].strip()
                        with db_conn() as conn:
                            conn.execute("DELETE FROM ai_permissions WHERE platform=? AND uid=?", (platform, uid_ai))
                            conn.commit()
                        audit_log("zalo", sender, "zaideny", f"{platform}:{uid_ai}")
                        zreply(f"✅ Đã thu hồi AI của {platform}:{uid_ai}")
                        return

                    if lower == "zaiquota":
                        with db_conn() as conn:
                            rows = conn.execute(
                                "SELECT platform, uid, enabled, daily_quota, used_today, day_key FROM ai_permissions ORDER BY created_at DESC LIMIT 20"
                            ).fetchall()
                        if not rows:
                            zreply("📭 Chưa có quota AI.")
                            return
                        lines = [f"- {r[0]}:{r[1]} | {r[4]}/{r[3]} | day={r[5]}" for r in rows]
                        zreply("📊 AI quota:\n" + "\n".join(lines))
                        return

                    if lower == "zaudit":
                        with db_conn() as conn:
                            rows = conn.execute(
                                "SELECT actor_platform, actor_uid, action, detail FROM audit ORDER BY id DESC LIMIT 15"
                            ).fetchall()
                        if not rows:
                            zreply("📭 Chưa có log audit.")
                            return
                        lines = [f"- {r[0]}:{r[1]} | {r[2]} | {str(r[3])[:40]}" for r in rows]
                        zreply("🧾 Audit gần nhất:\n" + "\n".join(lines))
                        return

                    if lower == "zuiddb":
                        self.scan_all_uid_sources()
                        with db_conn() as conn:
                            rows = conn.execute("SELECT uid FROM customer_contacts WHERE platform LIKE 'zalo%' ORDER BY updated_at DESC LIMIT 300").fetchall()
                        zreply("UID DB:\n" + ("\n".join(uid_with_name(str(r[0])) for r in rows) if rows else "rong"))
                        return

                    if lower == "zuidall":
                        self.scan_all_uid_sources()
                        zreply("UID da nhan tin:\n" + ("\n".join(uid_with_name(s) for s in sorted(ZALO_SEEN_USER_IDS)) if ZALO_SEEN_USER_IDS else "rong"))
                        return

                    if lower == "zuitvgr":
                        try:
                            groups = self.fetchAllGroups()
                            gids = []
                            if isinstance(groups, dict):
                                items = groups.get("data", []) or groups.get("groups", []) or []
                            else:
                                items = groups or []
                            for g in items:
                                if isinstance(g, dict):
                                    gid = g.get("groupId") or g.get("id") or g.get("uid")
                                    if gid:
                                        gids.append(str(gid))
                                        upsert_contact("zalo_group", str(gid), name=str(g.get("name") or g.get("title") or ""))
                            zreply("UID group hien tai:\n" + ("\n".join(uid_with_name(g) for g in gids) if gids else "khong co"))
                        except Exception as exc:
                            zreply(f"❌ Loi quet group: {exc}")
                        return

                    if lower == "zuidallgr":
                        self.scan_all_uid_sources()
                        with db_conn() as conn:
                            rows = conn.execute("SELECT uid FROM customer_contacts WHERE platform='zalo_group' ORDER BY updated_at DESC LIMIT 500").fetchall()
                        zreply("UID tat ca group da luu:\n" + ("\n".join(uid_with_name(str(r[0])) for r in rows) if rows else "rong"))
                        return

                    if lower.startswith("zaddadminall "):
                        uid_new = raw.split(" ", 1)[1].strip()
                        vals = get_zalo_admin_all()
                        vals.add(uid_new)
                        set_zalo_admin_all(vals)
                        zreply(f"✅ Da cap admin toan quyen: {uid_new}")
                        return

                    if lower.startswith("zaddadminai ") or lower.startswith("zaddadmiai "):
                        uid_new = raw.split(" ", 1)[1].strip()
                        vals = get_zalo_admin_ai()
                        vals.add(uid_new)
                        set_zalo_admin_ai(vals)
                        zreply(f"✅ Da cap quyen AI: {uid_new}")
                        return

                    if lower.startswith("zaddmingr "):
                        uid_new = raw.split(" ", 1)[1].strip()
                        ZALO_GROUP_ADMINS.add(uid_new)
                        persist_runtime_state()
                        zreply(f"✅ Da cap admin nhom: {uid_new}")
                        return

                    if lower in ("zautorepall on", "zautorepall off"):
                        ZALO_AUTO_REP_ALL = lower.endswith("on")
                        persist_runtime_state()
                        zreply(f"✅ zautorepall = {ZALO_AUTO_REP_ALL}")
                        return

                    if lower in ("zautorepnguoila on", "zautorepnguoila off"):
                        ZALO_AUTO_REP_NGUOILA = lower.endswith("on")
                        persist_runtime_state()
                        zreply(f"✅ zautorepnguoila = {ZALO_AUTO_REP_NGUOILA}")
                        return

                    if lower in ("zautorepbanbe on", "zautorepbanbe off"):
                        ZALO_AUTO_REP_BANBE = lower.endswith("on")
                        persist_runtime_state()
                        zreply(f"✅ zautorepbanbe = {ZALO_AUTO_REP_BANBE}")
                        return

                    if lower in ("zautorepnew on", "zautorepnew off"):
                        ZALO_AUTO_REP_NEW = lower.endswith("on")
                        persist_runtime_state()
                        zreply(f"✅ zautorepnew = {ZALO_AUTO_REP_NEW}")
                        return

                    payload = ztail("zsetrep")
                    if payload is not None:
                        m = re.match(r'^"([^"]+)"\s+(.+)$', payload, flags=re.S) or re.match(r"^(\S+)\s+(.+)$", payload, flags=re.S)
                        if not m:
                            zreply("⚠️ Cu phap: zsetrep \"tu_khoa\" noi_dung")
                            return
                        kw = m.group(1).strip().lower()
                        content = m.group(2).strip()
                        save_keyword_reply(kw, content)
                        zreply(f"✅ Da luu rep keyword: {kw}")
                        return

                    if lower == "zcheckrep":
                        reps = load_keyword_replies()
                        if not reps:
                            zreply("📭 Chua co keyword rep.")
                        else:
                            zreply("Danh sach rep:\n" + "\n".join(f"- {k} => {v[:80]}" for k, v in reps))
                        return

                    if lower == "zchecksetrep":
                        reps = load_keyword_replies()
                        if not reps:
                            zreply("📭 Chua co tu khoa auto-rep.")
                        else:
                            zreply("🧩 Tu khoa auto-rep:\n" + "\n".join(f"- {k} => {v[:120]}" for k, v in reps))
                        return

                    if lower in ("zcheckiudall", "zcheckuidall"):
                        with db_conn() as conn:
                            rows = conn.execute("SELECT uid FROM customer_contacts ORDER BY updated_at DESC LIMIT 500").fetchall()
                        zreply("📚 Tong hop UID da luu:\n" + ("\n".join(uid_with_name(str(r[0])) for r in rows) if rows else "rong"))
                        return

                    if lower == "zcheck" or lower in {"zcheck card", "zcheck danhba", "zcheck danhthiep"}:
                        card = load_last_zalo_contact_card(sender)
                        if not card:
                            card = find_recent_zalo_contact_card(self, sender)
                            if card:
                                save_last_zalo_contact_card(sender, card)
                            else:
                                zreply("⚠️ Chua bat duoc danh thiep gan nhat. Gui lai danh thiep/contact card vao chat nay roi go zcheck.")
                                return
                        uid_check = str(card.get("uid") or "").strip()
                        phone_check = normalize_vn_phone(str(card.get("phone") or ""))
                        if (not uid_check) and phone_check:
                            try:
                                user = self.fetchPhoneNumber(phone_check)
                                uid_check = zalo_user_uid(user)
                                if uid_check:
                                    card["uid"] = uid_check
                                    card["name"] = card.get("name") or zalo_user_name(user)
                                    card["username"] = card.get("username") or zalo_user_username(user)
                                    save_last_zalo_contact_card(sender, card)
                            except Exception:
                                pass
                        if uid_check or phone_check:
                            upsert_contact(
                                "zalo",
                                uid_check or phone_check,
                                name=str(card.get("name") or ""),
                                username=str(card.get("username") or ""),
                                phone=phone_check,
                                last_message="zalo_contact_card_check",
                            )
                        with db_conn() as conn:
                            rows = []
                            if uid_check:
                                rows = conn.execute(
                                    "SELECT platform, uid, name, username, phone FROM customer_contacts WHERE uid=? LIMIT 5",
                                    (uid_check,),
                                ).fetchall()
                            if (not rows) and phone_check:
                                rows = conn.execute(
                                    "SELECT platform, uid, name, username, phone FROM customer_contacts WHERE phone=? LIMIT 5",
                                    (phone_check,),
                                ).fetchall()
                        lines = [
                            "KET QUA CHECK DANH THIEP:",
                            f"- UID card: {uid_check or 'N/A'}",
                            f"- Ten card: {str(card.get('name') or 'N/A')}",
                            f"- User card: {str(card.get('username') or 'N/A')}",
                            f"- SDT card: {phone_check or 'N/A'}",
                        ]
                        if rows:
                            lines.append("\nDB MATCH:")
                            lines.extend(
                                f"- {r['platform']} | uid={r['uid']} | ten={r['name']} | user={r['username']} | sdt={r['phone']}"
                                for r in rows
                            )
                        else:
                            lines.append("\nDB MATCH: chua co")
                        zreply("\n".join(lines))
                        return

                    if lower.startswith("zcheck "):
                        uid_check = raw.split(" ", 1)[1].strip()
                        with db_conn() as conn:
                            rows = conn.execute(
                                "SELECT platform, uid, name, username, phone FROM customer_contacts WHERE uid=? LIMIT 5",
                                (uid_check,),
                            ).fetchall()
                        if not rows:
                            zreply("❌ Khong tim thay UID trong DB.")
                        else:
                            zreply("KET QUA CHECK UID:\n" + "\n".join(
                                f"- {r['platform']} | uid={r['uid']} | ten={r['name']} | user={r['username']} | sdt={r['phone']}" for r in rows
                            ))
                        return

                    if lower.startswith("zchecklenh "):
                        cmd = raw.split(" ", 1)[1].strip().lower().lstrip("/")
                        zreply(f"📘 {cmd}: {CMD_EXPLAIN.get(cmd, 'Chua co mo ta cho lenh nay.')}")
                        return

                    val = ztail("zsettn")
                    if val is not None:
                        ZALO_BANBE_BROADCAST_MSG = val
                        persist_runtime_state()
                        zreply("✅ Da cap nhat noi dung zsettn, da giu nguyen xuong dong.")
                        return

                    if lower.startswith("zbanbe "):
                        parts = raw.split()
                        if len(parts) < 3:
                            zreply("⚠️ Cu phap: zbanbe on|off <time_giay>")
                            return
                        mode = parts[1].lower()
                        if mode not in ("on", "off"):
                            zreply("⚠️ zbanbe chi nhan on/off.")
                            return
                        try:
                            iv = int(parts[2])
                        except Exception:
                            zreply("⚠️ time khong hop le.")
                            return
                        if iv < ZALO_BANBE_MIN_DELAY_SEC:
                            iv = ZALO_BANBE_MIN_DELAY_SEC
                        elif iv > ZALO_BANBE_MAX_DELAY_SEC:
                            iv = ZALO_BANBE_MAX_DELAY_SEC
                        ZALO_BANBE_BROADCAST_ENABLED = mode == "on"
                        ZALO_BANBE_BROADCAST_INTERVAL = iv
                        persist_runtime_state()
                        zreply(
                            f"✅ zbanbe = {ZALO_BANBE_BROADCAST_ENABLED}, delay={iv}s/contact. "
                            "Bot se di het 1 vong danh ba roi tu dung."
                        )
                        return

                    payload = ztail("zcapnhat")
                    if payload is not None:
                        if not payload:
                            zreply("⚠️ Cu phap: zcapnhat NOI_DUNG")
                            return
                        prompt = (
                            "Phan tich yeu cau quan tri bot sau, giai thich y nghia, giu nguyen noi dung goc, "
                            "de xuat cach lam ngan gon theo tung buoc:\n" + payload
                        )
                        result = call_ai_text(prompt)
                        zreply("🧠 Phan tich AI:\n" + result)
                        with db_conn() as conn:
                            conn.execute(
                                "INSERT OR REPLACE INTO config(key, value) VALUES(?,?)",
                                ("help_custom_capnhat", payload),
                            )
                            conn.commit()
                        return

                    if lower in ("zautojoin on", "zautojoin off"):
                        ZALO_AUTO_JOIN = lower.endswith("on")
                        TELE_AUTO_JOIN = ZALO_AUTO_JOIN
                        persist_runtime_state()
                        zreply(f"✅ autojoin zalo/tele = {ZALO_AUTO_JOIN}")
                        return

                    if lower == "zunjoinall":
                        zreply("✅ Da nhan lenh zunjoinall (ban test: danh dau, khong roi group hang loat de an toan).")
                        return

                    if lower.startswith("zunjoin "):
                        gid = raw.split(" ", 1)[1].strip()
                        zreply(f"✅ Da nhan lenh roi group: {gid} (ban test).")
                        return

                    if lower.startswith("zaitele ") or lower.startswith("zaizalo ") or lower.startswith("zaifb "):
                        kw = raw.split(" ", 1)[1].strip() if " " in raw else ""
                        if not kw:
                            zreply("⚠️ Can tu khoa.")
                            return
                        source = "telegram" if lower.startswith("zaitele ") else ("zalo" if lower.startswith("zaizalo ") else "facebook")
                        ai_prompt = f"Tim link nhom {source} theo tu khoa: {kw}. Tra ve danh sach link ngan gon."
                        zreply(call_ai_text(ai_prompt))
                        return

                    if lower == "zsdtall":
                        with db_conn() as conn:
                            rows = conn.execute("SELECT uid, phone FROM customer_contacts WHERE phone!='' ORDER BY updated_at DESC LIMIT 500").fetchall()
                        zreply("SDT khach hang:\n" + ("\n".join(f"{r[0]} | {r[1]}" for r in rows) if rows else "khong co"))
                        return

                    if lower == "zdanhba!":
                        # Lam moi danh ba truoc khi gui danh thiep
                        self.refresh_friends()
                        with db_conn() as conn:
                            rows = conn.execute(
                                """
                                SELECT uid, name, username, phone
                                FROM customer_contacts
                                WHERE platform='zalo'
                                ORDER BY updated_at DESC
                                LIMIT 500
                                """
                            ).fetchall()
                        if not rows:
                            zreply("📭 Khong co du lieu danh ba de gui danh thiep.")
                            return
                        if get_config("zalo_danhba_card_job_active", "0") == "1":
                            zreply("⏳ Dang co job gui danh thiep dang chay, vui long doi job hien tai xong.")
                            return

                        job_rows = [dict(r) for r in rows]
                        target_thread_id = thread_id
                        target_thread_type = thread_type
                        set_config("zalo_danhba_card_job_active", "1")
                        zreply(
                            f"✅ Bat dau gui {len(job_rows)} danh thiep nen. "
                            f"Moi danh thiep cach nhau {zalo_banbe_delay_seconds()}s, xong se bao admin.",
                            bless=False,
                        )

                        def _send_cards_job():
                            sent = 0
                            try:
                                for idx, r in enumerate(job_rows, 1):
                                    uidx = str(r.get("uid") or "")
                                    if not uidx or uidx == str(ZALO_ADMIN_UID):
                                        continue
                                    card = (
                                        f"📇 Danh thiep {idx}\n"
                                        f"- Ten: {str(r.get('name') or '(chua ro)')}\n"
                                        f"- UID: {uidx}\n"
                                        f"- Username: {str(r.get('username') or 'N/A')}\n"
                                        f"- SDT: {str(r.get('phone') or 'N/A')}"
                                    )
                                    try:
                                        self.sendMessage(Message(text=card), int(target_thread_id), target_thread_type)
                                        sent += 1
                                    except Exception:
                                        pass
                                    time.sleep(zalo_banbe_delay_seconds())
                                notify_admin_sync(
                                    f"[ZALO DANH THIEP] Da gui xong 1 vong danh thiep: {sent}/{len(job_rows)}. "
                                    "Neu can gui tiep, hay cap noi dung/lenh moi."
                                )
                            finally:
                                set_config("zalo_danhba_card_job_active", "0")

                        threading.Thread(target=_send_cards_job, daemon=True).start()
                        return

                    payload = ztail("zloaloaall")
                    if payload is not None:
                        if not payload or len(payload.rsplit(None, 1)) < 2:
                            zreply("⚠️ Cu phap: zloaloaall <noi_dung> <thoi_gian_giay>")
                            return
                        try:
                            msg, delay_raw = payload.rsplit(None, 1)
                            delay = float(delay_raw)
                            msg = msg.strip()
                        except Exception:
                            zreply("⚠️ Thoi gian khong hop le.")
                            return
                        ok = 0
                        for zuid in sorted(ZALO_SEEN_USER_IDS):
                            if zuid == str(ZALO_ADMIN_UID):
                                continue
                            try:
                                self.send_long_message(msg, int(zuid), thread_type)
                                ok += 1
                            except Exception:
                                pass
                            time.sleep(max(0.1, delay))
                        zreply(f"✅ Da gui loaloaall: {ok} uid")
                        return

                phone_msg = ""
                if isinstance(message_object, dict):
                    phone_msg = str(
                        message_object.get("phone")
                        or message_object.get("phoneNumber")
                        or message_object.get("tel")
                        or ""
                    ).strip()
                display_name = ""
                if isinstance(message_object, dict):
                    display_name = str(
                        message_object.get("name")
                        or message_object.get("displayName")
                        or message_object.get("fullName")
                        or ""
                    ).strip()

                phone_track = track_phone_from_message(
                    platform="zalo",
                    source_uid=sender,
                    username="",
                    full_name=display_name,
                    source_chat=str(thread_id or ""),
                    text=text,
                    extra_phone=phone_msg,
                )
                first_phone = phone_track["phones"][0] if phone_track["phones"] else normalize_vn_phone(phone_msg)
                if phone_track["phones"]:
                    maybe_send_phone_event_summary(force=False)
                upsert_contact("zalo", sender, name=display_name, phone=first_phone, last_message=text)

                # Live capture UID theo nhom khi API member list bi gioi han
                if str(thread_id) in LIVE_CAPTURE_ZALO_GROUPS and is_probably_group_context(thread_type, thread_id, sender):
                    upsert_contact("zalo_group_member", sender, name=display_name, phone=first_phone, last_message=text)
                if ZALO_SILENT_MODE and (not is_zalo_admin):
                    return
                if ZALO_AUTO_JOIN and text:
                    for link in extract_links(text):
                        try:
                            if "zalo.me/" in link.lower() and hasattr(self, "joinGroup"):
                                self.joinGroup(link)
                            elif "t.me/" in link.lower():
                                if ADMIN_CHAT_NOTIFY and sender not in all_admin_uids:
                                    notify_admin_sync(f"[AUTOJOIN-TELE] Link tu Zalo: {link}")
                        except Exception:
                            pass
                if (
                    ADMIN_CHAT_NOTIFY
                    and text
                    and (sender not in all_admin_uids)
                    and (str(thread_id) not in all_admin_uids)
                    and (not text.strip().lower().startswith("z"))
                ):
                    notify_admin_sync(
                        f"[ZALO MSG] UID={sender}\nSDT={phone_msg or 'N/A'}\nND={text[:400]}\nGoi y: /zreply {sender} <noi_dung>"
                    )

                # GROUP PRO AI: khi bat inall, thanh vien nhom co the hoi truc tiep khong can lenh.
                if OPENAI_INALL_ZALO and (not is_zalo_admin):
                    txt = (text or "").strip()
                    if txt and not txt.lower().startswith("z"):
                        try:
                            if str(thread_type).lower() in ("group", "groups", "thread", "room") and str(thread_id) in OPENAI_INALL_ZALO_GROUPS:
                                # Business Pro: bat buoc thanh vien phai nhan tin bat dau bang dau cham.
                                if str(thread_id) in BUSINESS_PRO_ZALO_GROUPS:
                                    if not txt.startswith("."):
                                        return
                                    txt = txt[1:].strip()
                                    if not txt:
                                        return
                                self.sendMessage(Message(text=f"🤖 {GROUP_PRO_AI_NAME} dang tra loi..."), thread_id, thread_type)
                                ai_ans = call_ai_text(txt, env_name="openai")
                                self.send_long_message(ai_ans, thread_id, thread_type)
                                return
                        except Exception:
                            pass

                is_stranger = sender not in ZALO_FRIEND_IDS
                first_seen = sender not in ZALO_SEEN_USER_IDS
                ZALO_SEEN_USER_IDS.add(sender)
                if first_seen:
                    audit_log("zalo", sender, "new_uid_seen", "")

                if is_stranger:
                    now = time.time()
                    last_ts = float(ZALO_LAST_AUTO_REPLY_TS.get(sender, 0))
                    if now - last_ts < ZALO_REPLY_COOLDOWN_SEC:
                        return
                    ZALO_LAST_AUTO_REPLY_TS[sender] = now
                    try:
                        self.acceptFriendRequest(sender)
                    except Exception:
                        pass
                    if (not is_zalo_admin) and (not is_business_pro_group):
                        try:
                            did_reply = False
                            for kw, rep in load_keyword_replies():
                                if kw and kw in (text or "").lower():
                                    did_reply = self.send_zalo_autoreply_once(
                                        rep,
                                        sender,
                                        thread_id,
                                        thread_type,
                                        reason=f"keyword:{kw}",
                                    )
                                    break
                            if (not did_reply) and (ZALO_AUTO_REP_ALL or ZALO_AUTO_REP_NGUOILA or (ZALO_AUTO_REP_NEW and first_seen)):
                                if ZALO_AUTO_REP_ENABLED or ZALO_AUTO_REP_ALL or ZALO_AUTO_REP_NGUOILA or ZALO_AUTO_REP_NEW:
                                    self.send_zalo_autoreply_once(
                                        ZALO_AUTO_REPLY,
                                        sender,
                                        thread_id,
                                        thread_type,
                                        reason="auto_stranger",
                                    )
                        except Exception:
                            pass
                    try:
                        self.sendFriendRequest(sender, "Ket ban de duoc ho tro nhanh hon.")
                    except Exception:
                        pass
                elif (not is_zalo_admin) and (not is_business_pro_group):
                    try:
                        did_reply = False
                        for kw, rep in load_keyword_replies():
                            if kw and kw in (text or "").lower():
                                did_reply = self.send_zalo_autoreply_once(
                                    rep,
                                    sender,
                                    thread_id,
                                    thread_type,
                                    reason=f"keyword:{kw}",
                                )
                                break
                        if (not did_reply) and (ZALO_AUTO_REP_ALL or ZALO_AUTO_REP_BANBE):
                            if ZALO_AUTO_REP_ENABLED or ZALO_AUTO_REP_ALL or ZALO_AUTO_REP_BANBE:
                                self.send_zalo_autoreply_once(
                                    ZALO_AUTO_REPLY,
                                    sender,
                                    thread_id,
                                    thread_type,
                                    reason="auto_friend",
                                )
                    except Exception:
                        pass
                # Khong gui bao ve admin cho tung tin nhan Zalo de tranh spam/vong lap.
            except Exception as exc:
                logging.error("Loi xu ly tin nhan Zalo: %s", exc)

    def _start_banbe_broadcast_loop(bot):
        def _loop():
            global ZALO_BANBE_BROADCAST_ENABLED, ZALO_LAST_BANBE_BROADCAST_TS
            while True:
                try:
                    if (not ZALO_BANBE_BROADCAST_ENABLED) or (not ZALO_BANBE_BROADCAST_MSG.strip()):
                        time.sleep(10)
                        continue

                    current_msg = ZALO_BANBE_BROADCAST_MSG.strip()
                    bot.refresh_friends()
                    targets = [
                        str(fuid)
                        for fuid in sorted(ZALO_FRIEND_IDS)
                        if str(fuid) and str(fuid) != str(ZALO_ADMIN_UID)
                    ]
                    sent_ok = 0
                    skipped = 0
                    failed = 0

                    for fuid in targets:
                        if (not ZALO_BANBE_BROADCAST_ENABLED) or ZALO_BANBE_BROADCAST_MSG.strip() != current_msg:
                            break
                        if not should_send_outbox("zalo", fuid, current_msg):
                            skipped += 1
                            continue
                        try:
                            bot.send_long_message(current_msg, int(fuid), ThreadType.USER)
                            sent_ok += 1
                            ZALO_LAST_BANBE_BROADCAST_TS = time.time()
                        except Exception:
                            failed += 1
                        time.sleep(zalo_banbe_delay_seconds())

                    if ZALO_BANBE_BROADCAST_ENABLED and ZALO_BANBE_BROADCAST_MSG.strip() == current_msg:
                        ZALO_BANBE_BROADCAST_ENABLED = False
                        persist_runtime_state()
                        notify_admin_sync(
                            "[ZALO BANBE] Da di het 1 vong danh ba va da tu dung.\n"
                            f"Da gui={sent_ok}, bo qua da gui={skipped}, loi={failed}, tong={len(targets)}.\n"
                            "Can cap noi dung moi bang zsettn <noi_dung>, roi bat lai zbanbe on <60-180>."
                        )
                    time.sleep(10)
                except Exception as exc:
                    logging.warning("Zalo banbe broadcast loop loi: %s", exc)
                    time.sleep(30)

        threading.Thread(target=_loop, daemon=True).start()

    try:
        global LAST_ZALO_BOT
        bot = ZaloAutoBot(phone="</>", password="</>", imei=ZALO_IMEI, cookies=cookies)
        LAST_ZALO_BOT = bot
        set_config("zalo_active", "True")
        set_config("zalo_worker_last", f"{int(time.time())}|running")
        _start_banbe_broadcast_loop(bot)
        bot.scan_all_uid_sources()
        try:
            if str(ZALO_ADMIN_UID).strip():
                try:
                    bot.sendMessage(
                        Message(text="👑 XIN CHÀO CHỦ NHÂN 👑 AI FAMILY 6.0 🤖 ĐÃ SẴN SÀNG NHẬN LỆNH!"),
                        int(ZALO_ADMIN_UID),
                        ThreadType.USER,
                    )
                    set_config("zalo_admin_greeting_last", f"{int(time.time())}|sent uid={ZALO_ADMIN_UID}")
                except Exception as exc:
                    set_config(
                        "zalo_admin_greeting_last",
                        f"{int(time.time())}|failed uid={ZALO_ADMIN_UID} err={type(exc).__name__}:{str(exc)[:160]}",
                    )
            else:
                set_config("zalo_admin_greeting_last", f"{int(time.time())}|skipped no_admin_uid")
        except Exception as exc:
            set_config(
                "zalo_admin_greeting_last",
                f"{int(time.time())}|failed_import err={type(exc).__name__}:{str(exc)[:160]}",
            )
        def _scan_loop():
            while True:
                try:
                    bot.scan_all_uid_sources()
                    bot.report_scan_to_zalo_admin()
                    bot.remind_backup_to_zalo_admin()
                    maybe_send_phone_event_summary(force=False)
                except Exception:
                    set_config("zalo_worker_last", f"{int(time.time())}|scan_loop_error")
                    logging.exception("Zalo scan/report loop failed")
                time.sleep(SCAN_REPORT_INTERVAL_SEC)
        threading.Thread(target=_scan_loop, daemon=True).start()
        notify_admin_sync("👑 Hi Chủ Nhân 👑 AI FAMILY 6.0 🤖 ĐÃ HOẠT ĐỘNG VÀ CHỜ XÁC NHẬN CỦA ADMIN ZALO. Chúc Chủ Nhân May Mắn Thuận Lợi 🍀")
        notify_admin_sync("CHÚC CÁC THÀNH VIÊN GIA ĐÌNH ADMIN SỐNG VUI SỐNG KHỎE, VÔ LO VÔ NGHĨ 🎉❤️")
        notify_admin_sync("Lenh sao luu nhanh: zbackup")
        bot.listen(thread=False, reconnect=5)
    except Exception as exc:
        set_config("zalo_active", "False")
        set_config("zalo_worker_last", f"{int(time.time())}|error {type(exc).__name__}:{str(exc)[:180]}")
        logging.error("Zalo worker loi: %s", exc)
        notify_admin_sync(f"[ZALO LOI] {exc}")

async def cmd_nhancode88k(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    if uid in BANNED_USERS: return

    if not BOT_ACTIVE and not is_admin(uid):
        await update.message.reply_text("🛠 Hệ thống đang bảo trì. Vui lòng thử lại sau!")
        return

    first_start = uid not in USER_DATA
    # Lưu thông tin khách vào cache hệ thống để gửi tin hàng loạt sau này
    if uid not in USER_DATA:
        USER_DATA[uid] = {'step': 'CHO_CHON_HINH_THUC', 'ten_game': 'Chưa nhập', 'sdt': 'Chưa nhập', 'goi_nap': 'Chưa chọn'}
    else:
        USER_DATA[uid]['step'] = 'CHO_CHON_HINH_THUC'
    upsert_contact("telegram", str(uid), name=update.effective_user.first_name or "", username=update.effective_user.username or "")
    mark_campaign_target_consent(
        str(uid),
        username=update.effective_user.username or "",
        source="telegram_start",
        status="opt_in",
    )
    persist_runtime_state()

    if THONG_BAO_START_ACTIVE:
        username = f"@{update.effective_user.username}" if update.effective_user.username else "(không có username)"
        if first_start:
            first_alert = (
                "🚨 ƯU TIÊN CẢNH BÁO\n"
                "🆕 **[KHÁCH MỚI /start LẦN ĐẦU]**\n"
                f"👤 Tên: `{update.effective_user.first_name}`\n"
                f"🧩 User: `{username}`\n"
                f"🆔 ID: `{uid}`"
            )
            try:
                await context.bot.send_message(chat_id=ADMIN_ID, text=first_alert, parse_mode="Markdown")
            except Exception:
                logging.exception("Gui canh bao khach moi ve admin chinh that bai")
                try:
                    await context.bot.send_message(chat_id=ADMIN_ID, text=first_alert.replace("`", "").replace("**", ""))
                except Exception:
                    logging.exception("Gui fallback canh bao khach moi ve admin chinh that bai")
            await send_customer_supergroup_alert(context.bot, first_alert, source_chat_id=update.effective_chat.id)
        else:
            admin_alert = (
                f"🚨 **[KHÁCH BẤM /nhancode88k]**\n"
                f"👤 Tên: `{update.effective_user.first_name}`\n"
                f"🧩 User: `{username}`\n"
                f"🆔 ID: `{uid}`"
            )
            try:
                await context.bot.send_message(chat_id=ADMIN_ID, text=admin_alert, parse_mode="Markdown")
            except Exception:
                logging.exception("Gui canh bao /start ve admin chinh that bai")
                try:
                    await context.bot.send_message(chat_id=ADMIN_ID, text=admin_alert.replace("`", "").replace("**", ""))
                except Exception:
                    logging.exception("Gui fallback canh bao /start ve admin chinh that bai")

    if HIEN_THI_ANH_DONG:
        try: await send_config_media(context.bot, uid, GIF_CHAO_HOI)
        except: pass

    campaign_keyboard = []
    if CAMPAIGN_DEFAULT_GROUP_LINK:
        campaign_keyboard.append([InlineKeyboardButton("👥 VÀO NHÓM CHÍNH", url=CAMPAIGN_DEFAULT_GROUP_LINK)])
    if CAMPAIGN_REGISTER_LINK:
        campaign_keyboard.append([InlineKeyboardButton("🔗 LINK ĐĂNG KÝ / THÔNG TIN", url=CAMPAIGN_REGISTER_LINK)])
    campaign_keyboard.append([InlineKeyboardButton("👨‍💼 LIÊN HỆ ADMIN", url=f"https://t.me/{CAMPAIGN_SUPPORT_USERNAME or 'gifhub2708'}")])
    keyboard = campaign_keyboard + [
        [InlineKeyboardButton("🎯 ĐĂNG KÝ TÀI KHOẢN MỚI 🎯", callback_data=f"btn_dangky_{uid}")],
        [InlineKeyboardButton("🎁 CÓ TÀI KHOẢN - NHẬN CODE 🎁", callback_data=f"btn_cotaikhoan_{uid}")]
    ]
    await send_telegram_long(
        context.bot,
        uid,
        render_campaign_start_text(update.effective_user),
        reply_markup=InlineKeyboardMarkup(keyboard),
    )

async def handle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    uid = query.from_user.id

    if uid in BANNED_USERS: return
    if not BOT_ACTIVE and not is_admin(uid): return
    if uid not in USER_DATA:
        USER_DATA[uid] = {'step': None, 'ten_game': 'Chưa nhập', 'sdt': 'Chưa nhập', 'goi_nap': 'Chưa chọn'}

    if data.startswith("btn_dangky_"):
        USER_DATA[uid]['step'] = 'CHO_NHAP_TEN_TK'
        if HIEN_THI_ANH_DONG:
            try: await send_config_media(context.bot, uid, GIF_DANG_KY)
            except: pass

        register_text = get_tele_reply_template("register_prompt")
        await send_telegram_long(context.bot, uid, register_text, disable_web_page_preview=True)

    elif data.startswith("btn_cotaikhoan_"):
        USER_DATA[uid]['step'] = 'CHO_DIEN_SDT_KHAC_CU'
        if HIEN_THI_ANH_DONG:
            await send_template_media_if_any(context.bot, uid, "old_member_phone_prompt")
        await send_telegram_long(context.bot, uid, get_tele_reply_template("old_member_phone_prompt"))

    elif data.startswith("check_join_"):
        if uid in VIP_USERS or await is_user_member(context, uid):
            USER_DATA[uid]['step'] = 'CHO_CHON_MOC_NAP'
            if HIEN_THI_ANH_DONG:
                try: await send_config_media(context.bot, uid, GIF_MOC_NAP)
                except: pass

            keyboard = [
                [InlineKeyboardButton("💰 Nạp 50K Nhận 100% + Code 18K 🎁", callback_data=f"moc_50K_18K_{uid}")],
                [InlineKeyboardButton("💰 Nạp 100K Nhận 100% + Code 38K 🎁", callback_data=f"moc_100K_38K_{uid}")],
                [InlineKeyboardButton("💰 Nạp 200K Nhận 100% + Code 58K 🎁", callback_data=f"moc_200K_58K_{uid}")],
                [InlineKeyboardButton("💎 Nạp 500K Nhận 100% + Code 128K 🎁", callback_data=f"moc_500K_128K_{uid}")],
                [InlineKeyboardButton("💎 Nạp 1 Triệu Nhận 100% + Code 288K 🎁", callback_data=f"moc_1M_288K_{uid}")],
                [InlineKeyboardButton("🔥 Nạp 3 Triệu Nhận 100% + Code 588K 🎁", callback_data=f"moc_3M_588K_{uid}")],
                [InlineKeyboardButton("🔥 Nạp 5 Triệu Nhận 100% + Code 999K 🎁", callback_data=f"moc_5M_999K_{uid}")],
                [InlineKeyboardButton("👑 Nạp 10 Triệu Nhận 100% + Code 2.8M 🎁", callback_data=f"moc_10M_2.8M_{uid}")],
                [InlineKeyboardButton("👑 Nạp 20 Triệu Nhận 100% + Code 5.8M 🎁", callback_data=f"moc_20M_5.8M_{uid}")],
                [InlineKeyboardButton("🔱 Nạp 50 Triệu Nhận 100% + Code 15M 🎁", callback_data=f"moc_50M_15M_{uid}")],
                [InlineKeyboardButton("🔱 Nạp 100 Triệu Nhận 100% + Code 38M 🎁", callback_data=f"moc_100M_38M_{uid}")]
            ]
            await send_telegram_long(
                context.bot,
                uid,
                get_tele_reply_template("moc_intro"),
                reply_markup=InlineKeyboardMarkup(keyboard),
            )
        else:
            if HIEN_THI_ANH_DONG:
                await send_template_media_if_any(context.bot, uid, "join_not_verified")
            await send_telegram_long(context.bot, uid, get_tele_reply_template("join_not_verified"))

    elif data.startswith("moc_"):
        parts = data.split("_")
        amount = parts[1]
        code_val = parts[2]
        USER_DATA[uid]['goi_nap'] = f"Nạp {amount} nhận 100% + Code {code_val}"

        keyboard = [
            [InlineKeyboardButton("👨‍💼 LIÊN HỆ TRỰC TIẾP ADMIN", url=f"https://t.me/{ADMIN_USERNAME.replace('@','')}")],
            [InlineKeyboardButton("🤖 HỆ THỐNG TRẢ CODE TỰ ĐỘNG", callback_data=f"auto_pay_{uid}")]
        ]
        if HIEN_THI_ANH_DONG:
            await send_template_media_if_any(context.bot, uid, "choose_reward_mode")
        await send_telegram_long(
            context.bot,
            uid,
            get_tele_reply_template("choose_reward_mode"),
            reply_markup=InlineKeyboardMarkup(keyboard),
        )

    elif data.startswith("auto_pay_"):
        if not CHUC_NANG_AUTO_DUYET:
            if HIEN_THI_ANH_DONG:
                await send_template_media_if_any(context.bot, uid, "auto_review_disabled")
            await send_telegram_long(context.bot, uid, get_tele_reply_template("auto_review_disabled"))
            return
        if HIEN_THI_ANH_DONG:
            try: await send_config_media(context.bot, uid, GIF_CHO_DUYET)
            except: pass
        await query.edit_message_text(get_tele_reply_template("auto_review_waiting"))

        report_text = (
            f"🚨🚨 **[YÊU CẦU DUYỆT CODE]** 🚨🚨\n\n"
            f"👤 **Tên TK:** `{USER_DATA[uid].get('ten_game', 'Chưa có')}`\n"
            f"📱 **SĐT:** `{USER_DATA[uid].get('sdt', 'Chưa có')}`\n"
            f"💵 **Gói:** {USER_DATA[uid].get('goi_nap', 'Chưa chọn')}\n"
            f"🆔 **ID Khách:** `{uid}`\n\n"
            f"👉 Lệnh cấp code nhanh:\n`/code {uid} [MÃ_CODE]`"
        )
        await context.bot.send_message(chat_id=ADMIN_ID, text=report_text, parse_mode="Markdown")
        await send_customer_supergroup_alert(context.bot, report_text, source_chat_id=update.effective_chat.id)

    elif data == "khuyen_mai_khac":
        if HIEN_THI_ANH_DONG:
            await send_template_media_if_any(context.bot, uid, "promo_other")
        await send_telegram_long(context.bot, uid, get_tele_reply_template("promo_other"))
    if uid in USER_DATA:
        save_user_state(uid)

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    if uid in BANNED_USERS: return
    text = update.message.text.strip()
    chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
    chat_id = str(getattr(update.effective_chat, "id", "") or "")
    chat_title = str(getattr(update.effective_chat, "title", "") or "")
    cache_telegram_group_member(update)
    if text.startswith("/"): return
    lower_text = text.lower()
    uname = update.effective_user.username or ""
    fname = update.effective_user.first_name or ""

    # Neu tin nhan co SDT: luu vao danh ba bot (phonebook) + event summary.
    phone_track = track_phone_from_message(
        platform="telegram",
        source_uid=str(uid),
        username=uname,
        full_name=fname,
        source_chat=chat_id,
        text=text,
    )
    first_phone = phone_track["phones"][0] if phone_track["phones"] else ""
    if phone_track["phones"]:
        maybe_send_phone_event_summary(force=False)

    # Luu du lieu quet Telegram: user/group/channel
    upsert_contact(
        "telegram_user",
        str(uid),
        name=fname,
        username=uname,
        phone=first_phone,
        last_message=text,
    )
    if chat_type in ("group", "supergroup"):
        upsert_contact("telegram_group", chat_id, name=chat_title, last_message=text)
    elif chat_type == "channel":
        upsert_contact("telegram_channel", chat_id, name=chat_title, last_message=text)

    if TELEGRAM_GROUP_SILENT_MODE and chat_type in ("group", "supergroup", "channel"):
        return

    # Admin co the go truc tiep: openai <noi_dung>
    if is_admin(uid) and lower_text.startswith("openai "):
        q = text[7:].strip()
        if q:
            await update.message.reply_text(f"🤖 {GROUP_PRO_AI_NAME} đang xử lý...")
            ans = call_ai_text(q, env_name="openai")
            for i in range(0, len(ans), 3500):
                await update.message.reply_text(ans[i:i + 3500])
        else:
            await update.message.reply_text("Cú pháp: openai <nội_dung>")
        return

    # GROUP PRO AI inall tren Telegram group:
    # chi xu ly tin nhan bat dau bang dau "." de tranh bat nham hoi thoai thuong.
    if OPENAI_INALL_TELE and chat_type in ("group", "supergroup") and not is_admin(uid):
        if str(update.effective_chat.id) in OPENAI_INALL_TELE_GROUPS:
            if not text.startswith("."):
                return
            question = text[1:].strip()
            if not question:
                return
            await update.message.reply_text(f"🤖 {GROUP_PRO_AI_NAME} đang trả lời...")
            ans = call_ai_text(question, env_name="openai")
            await send_telegram_long(context.bot, update.effective_chat.id, ans)
            return
    try:
        customer_push_text = f"📩 Khach `{uid}`: {text[:300]}"
        await context.bot.send_message(chat_id=ADMIN_ID, text=customer_push_text, parse_mode="Markdown")
    except Exception:
        logging.exception("Gui tin nhan khach ve admin chinh bang Markdown that bai")
        customer_push_text = f"📩 Khach {uid}: {text[:300]}"
        try:
            await context.bot.send_message(chat_id=ADMIN_ID, text=customer_push_text)
        except Exception:
            logging.exception("Gui fallback tin nhan khach ve admin chinh that bai")
    await send_customer_supergroup_alert(context.bot, customer_push_text, source_chat_id=update.effective_chat.id)

    if uid not in USER_DATA:
        USER_DATA[uid] = {'step': None, 'ten_game': 'Chưa nhập', 'sdt': 'Chưa nhập'}
    upsert_contact("telegram", str(uid), name=fname, username=uname, phone=first_phone, last_message=text)

    step = USER_DATA[uid].get('step')

    if step == 'CHO_DIEN_SDT_KHAC_CU':
        USER_DATA[uid]['sdt'] = text
        USER_DATA[uid]['step'] = None
        upsert_contact("telegram", str(uid), phone=text, last_message=text)
        persist_runtime_state()
        alert_text = f"✨ **[KHÁCH CŨ BÁO SĐT]**\n📱 SĐT: `{text}`\n🆔 ID: `{uid}`"
        try:
            await context.bot.send_message(chat_id=ADMIN_ID, text=alert_text, parse_mode="Markdown")
        except Exception:
            logging.exception("Gui canh bao SDT khach cu ve admin chinh that bai")
            try:
                await context.bot.send_message(chat_id=ADMIN_ID, text=alert_text.replace("`", "").replace("**", ""))
            except Exception:
                logging.exception("Gui fallback canh bao SDT khach cu ve admin chinh that bai")
        await send_customer_supergroup_alert(context.bot, alert_text, source_chat_id=update.effective_chat.id)
        if HIEN_THI_ANH_DONG:
            await send_template_media_if_any(context.bot, uid, "old_member_phone_saved")
        await send_telegram_long(context.bot, uid, get_tele_reply_template("old_member_phone_saved"))
        return

    elif step == 'CHO_NHAP_TEN_TK':
        USER_DATA[uid]['ten_game'] = text
        USER_DATA[uid]['step'] = 'CHO_NHAP_SDT_RIENG'
        persist_runtime_state()
        if HIEN_THI_ANH_DONG:
            await send_template_media_if_any(context.bot, uid, "ask_phone_after_tk")
        await send_telegram_long(
            context.bot,
            uid,
            get_tele_reply_template("ask_phone_after_tk", ACCOUNT=text),
        )
        return

    elif step == 'CHO_NHAP_SDT_RIENG':
        USER_DATA[uid]['sdt'] = text
        USER_DATA[uid]['step'] = 'CHO_KIEM_TRA_JOIN'
        upsert_contact("telegram", str(uid), phone=text, last_message=text)
        persist_runtime_state()
        if HIEN_THI_ANH_DONG:
            try: await send_config_media(context.bot, uid, GIF_EP_JOIN)
            except: pass
        keyboard = [[InlineKeyboardButton("✅ XÁC NHẬN ĐÃ THAM GIA KÊNH ✅", callback_data=f"check_join_{uid}")]]
        await send_telegram_long(
            context.bot,
            uid,
            get_tele_reply_template("join_required"),
            reply_markup=InlineKeyboardMarkup(keyboard),
        )
        return

async def bi_mat_quang_cao(context: ContextTypes.DEFAULT_TYPE, client_id: int):
    await asyncio.sleep(300)
    try:
        if HIEN_THI_ANH_DONG: await send_config_media(context.bot, client_id, GIF_QUANG_CAO)
        keyboard = [[InlineKeyboardButton("🎁 SỰ KIỆN KHÁC 🎁", callback_data="khuyen_mai_khac")]]
        await send_telegram_long(
            context.bot,
            client_id,
            QUANG_CAO_TEXT,
            reply_markup=InlineKeyboardMarkup(keyboard),
        )
    except: pass

# ==================== THÊM MỚI CÁC TÍNH NĂNG TIẾP THỊ NÂNG CAO ====================

async def cmd_guitatca(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Lệnh gửi tin nhắn hàng loạt cho TẤT CẢ khách hàng"""
    if not await require_admin(update): return
    msg_content = get_full_command_payload(update, "guitatca")
    if not msg_content:
        await update.message.reply_text("⚠️ Vui lòng nhập nội dung cần gửi. Cú pháp:\n`/guitatca [Nội dung tin nhắn]`")
        return

    queued = 0
    skipped = 0
    failed = 0
    targets = list_broadcast_customer_targets()
    for uid in targets:
        ok, reason = queue_broadcast_message(uid, msg_content, requested_by=str(update.effective_user.id))
        if ok:
            queued += 1
        elif reason in {"already_sent", "already_queued"}:
            skipped += 1
        else:
            failed += 1
    await update.message.reply_text(
        "📢 Da dua broadcast vao hang doi SQLite.\n"
        f"✅ Queued: {queued}\n"
        f"ℹ️ Bo qua trung/da gui: {skipped}\n"
        f"❌ Loi queue: {failed}\n"
        "Worker nen se gui cham va luu sent/failed/retry."
    )

async def cmd_guikhach(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Lệnh gửi tin nhắn riêng cho một khách hàng bằng ID"""
    if not await require_admin(update): return
    try:
        leading, msg_content = split_command_leading_and_rest(update, "guikhach", 1)
        target_uid = int(leading[0])
        if not msg_content:
            await update.message.reply_text("⚠️ Nội dung trống! Cú pháp:\n`/guikhach [ID_Khách] [Nội dung]`")
            return

        ok, reason = queue_broadcast_message(str(target_uid), msg_content, requested_by=str(update.effective_user.id))
        if not ok:
            if reason in {"already_sent", "already_queued"}:
                await update.message.reply_text("ℹ️ Tin này đã gửi/đã có trong hàng đợi cho UID này, hệ thống bỏ qua để tránh trùng.")
            else:
                await update.message.reply_text(f"❌ Không thể queue tin. Lỗi: {reason}")
            return
        await update.message.reply_text(f"✅ Đã queue tin riêng tới khách hàng `{target_uid}`.")
    except IndexError:
        await update.message.reply_text("⚠️ Sai cú pháp! Vui lòng nhập:\n`/guikhach [ID_Khách] [Nội dung]`")
    except Exception as e:
        await update.message.reply_text(f"❌ Không thể queue tin. Lỗi: {str(e)}")

# ==================== CÁC CHỨC NĂNG TIẾNG VIỆT CHO ADMIN ====================
async def cmd_loaloa(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await require_admin(update): return
    try:
        leading, msg = split_command_leading_and_rest(update, "loaloa", 1)
        target = leading[0]
        ok, reason = queue_broadcast_message(str(target), msg, requested_by=str(update.effective_user.id))
        if not ok:
            if reason in {"already_sent", "already_queued"}:
                await update.message.reply_text("ℹ️ Nội dung đã gửi/đã có trong hàng đợi tới đích này, bỏ qua để tránh trùng.")
            else:
                await update.message.reply_text(f"❌ Không thể queue tin: {reason}")
            return
        await update.message.reply_text(f"✅ Đã queue tin tới {target}.")
    except Exception as exc:
        logging.exception("cmd_loaloa failed")
        try:
            await update.message.reply_text(f"❌ Sai cú pháp hoặc lỗi queue: {type(exc).__name__}")
        except Exception:
            pass

async def cmd_baotri(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global BOT_ACTIVE
    if not is_admin(update.effective_user.id): return
    BOT_ACTIVE = not BOT_ACTIVE
    await update.message.reply_text(f"⚙️ Trạng thái bot: **{'HOẠT ĐỘNG' if BOT_ACTIVE else 'BẢO TRÌ'}**")

async def cmd_link(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global WEB_URL, LINK_NHAP_CODE
    if not is_admin(update.effective_user.id): return
    try:
        loai = context.args[0].lower()
        url = context.args[1]
        if loai == "web":
            WEB_URL = url
            set_config("web_url", WEB_URL)
        elif loai == "code":
            LINK_NHAP_CODE = url
            set_config("link_nhap_code", LINK_NHAP_CODE)
        await update.message.reply_text(f"✅ Đã đổi link `{loai}` thành: {url}")
    except: pass

async def cmd_huong(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        kw = context.args[0].lower()
        url = context.args[1]
        KEYWORDS_ROUTING[kw] = url
        await update.message.reply_text(f"✅ Đã lưu hướng: `{kw}` -> {url}")
    except: pass

async def cmd_chuyentrong(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        tid = int(context.args[0])
        url = context.args[1]
        await context.bot.send_message(chat_id=tid, text=f"🔔 Cổng kết nối riêng của bạn: {url}")
        await update.message.reply_text("✅ Đã chuyển link thành công.")
    except: pass

async def cmd_kiemtra(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    await update.message.reply_text(f"📊 Hệ thống: {'MỞ' if BOT_ACTIVE else 'BẢO TRÌ'}\n- Số khách đệm: {len(USER_DATA)}\n- Chặn: {len(BANNED_USERS)}\n- VIP: {len(VIP_USERS)}")

async def cmd_xemkhach(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        tid = int(context.args[0])
        if tid in USER_DATA:
            d = USER_DATA[tid]
            await update.message.reply_text(f"📋 **DỮ LIỆU USER {tid}:**\n- TK: `{d.get('ten_game')}`\n- SĐT: `{d.get('sdt')}`\n- Gói: `{d.get('goi_nap')}`")
        else: await update.message.reply_text("❌ Không tìm thấy cache khách này.")
    except: pass

async def cmd_xoarac(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    USER_DATA.clear()
    clear_user_state_tables()
    set_config("user_data_json", "{}")
    await update.message.reply_text("🧹 Đã làm sạch cache hệ thống.")

async def cmd_themadmin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID: return
    try:
        EXTRA_ADMINS.add(int(context.args[0]))
        persist_runtime_state()
        await update.message.reply_text("✅ Đã cấp ĐẦY ĐỦ QUYỀN cho Admin phụ.")
    except: pass

async def cmd_huyadmin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID: return
    try:
        EXTRA_ADMINS.discard(int(context.args[0]))
        persist_runtime_state()
        await update.message.reply_text("✅ Đã hủy quyền Admin phụ.")
    except: pass

async def cmd_kenh(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global TARGET_CHAT
    if not is_admin(update.effective_user.id): return
    try:
        TARGET_CHAT = context.args[0]
        set_config("target_chat", TARGET_CHAT)
        await update.message.reply_text(f"📢 Kênh kiểm tra mới: {mask_sensitive(TARGET_CHAT, 3)}")
    except: pass

async def cmd_chan(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        uid = int(context.args[0])
        BANNED_USERS.add(uid)
        save_user_state(uid)
        await update.message.reply_text("🚫 Đã chặn tài khoản.")
    except: pass

async def cmd_mochan(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        uid = int(context.args[0])
        BANNED_USERS.discard(uid)
        save_user_state(uid)
        await update.message.reply_text("🔓 Đã mở chặn.")
    except: pass

async def cmd_mochanhet(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    total = len(BANNED_USERS)
    BANNED_USERS.clear()
    with db_conn() as conn:
        conn.execute("UPDATE users SET is_banned=0, updated_at=?", (time.time(),))
        conn.commit()
    persist_runtime_state()
    await update.message.reply_text(f"🔓 Đã mở chặn toàn bộ. Đã gỡ: {total} tài khoản.")

async def _scan_text_and_report(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str, source_type: str):
    msg = update.message
    actor = update.effective_user
    entities = extract_scan_entities(text or "")
    phones = entities["phones"]
    uids = entities["uids"]
    users = entities["users"]
    if not phones and not uids and not users:
        await msg.reply_text("❌ Quét không được: không tìm thấy SĐT/UID/user hợp lệ.")
        return
    new_count = 0
    lookup = build_phone_lookup_index()
    uid_lookup, user_lookup = build_identity_lookup_index()
    mapped_count = 0
    csv_rows = []
    mapped_preview = []
    for p in phones:
        if save_scanned_lead(
            phone=p,
            username=(actor.username if actor and actor.username else ""),
            full_name=(actor.full_name if actor and actor.full_name else ""),
            source_uid=(actor.id if actor else ""),
            source_chat=(msg.chat_id if msg else ""),
            source_type=source_type,
        ):
            new_count += 1
        hits = lookup.get(p, [])
        hit = choose_best_lookup_hit(hits)
        if hit:
            user = str(hit.get("user", "") or "")
            uid = str(hit.get("uid", "") or "")
            ten = str(hit.get("ten", "") or "")
            platform = str(hit.get("platform", "") or "")
            status = str(hit.get("status", "matched") or "matched")
            first_seen = float(hit.get("first_seen", hit.get("updated_at", 0)) or 0)
            last_seen = float(hit.get("last_seen", hit.get("updated_at", 0)) or 0)
            link = str(hit.get("link", "") or "").strip()
            if user or uid:
                mapped_count += 1
        else:
            user = ""
            uid = ""
            ten = ""
            platform = ""
            status = "unknown"
            first_seen = 0
            last_seen = 0
            link = ""
        if not link:
            link = build_customer_link(platform=platform, uid=uid, username=user, phone=p)
        # Luu lai vao phonebook ca khi quet thu cong de khong mat dau so.
        save_phonebook_entry(
            phone=p,
            platform=platform or "scan_manual",
            source_uid=uid,
            username=user,
            full_name=ten,
            source_chat=str(msg.chat_id if msg else ""),
            last_message_snippet=f"{source_type}:{(text or '')[:120]}",
        )
        log_phone_event(
            phone=p,
            platform=platform or "scan_manual",
            source_uid=uid,
            username=user,
            full_name=ten,
            source_chat=str(msg.chat_id if msg else ""),
            is_new_phone=False,
            is_new_uid=False,
        )
        csv_rows.append(
            {
                "phone": p,
                "user": user,
                "uid": uid,
                "ten": ten,
                "link": link,
                "platform": platform,
                "status": status,
                "first_seen": int(first_seen) if first_seen else "",
                "last_seen": int(last_seen) if last_seen else "",
            }
        )
        if user or uid or link:
            mapped_preview.append(
                f"{p} | user={user or 'N/A'} | uid={uid or 'N/A'} | link={link or 'N/A'}"
            )
        csv_rows[-1]["type"] = "phone"
        csv_rows[-1]["value"] = p

    for uid_value in uids:
        hit = uid_lookup.get(uid_value, {})
        user = str(hit.get("user", "") or "")
        uid = str(hit.get("uid", "") or uid_value)
        ten = str(hit.get("ten", "") or "")
        platform = str(hit.get("platform", "") or "")
        phone = normalize_vn_phone(str(hit.get("phone", "") or ""))
        status = str(hit.get("status", "") or ("matched" if hit else "raw_uid"))
        link = str(hit.get("link", "") or "").strip()
        if not link:
            link = build_customer_link(platform=platform or "telegram_user", uid=uid, username=user, phone=phone)
        if hit:
            mapped_count += 1
            mapped_preview.append(f"uid={uid} | user={user or 'N/A'} | ten={ten or 'N/A'} | link={link or 'N/A'}")
        csv_rows.append(
            {
                "type": "uid",
                "value": uid_value,
                "phone": phone if _is_valid_phone(phone) else "",
                "user": user,
                "uid": uid,
                "ten": ten,
                "link": link,
                "platform": platform,
                "status": status,
                "first_seen": "",
                "last_seen": int(float(hit.get("updated_at", 0) or 0)) if hit else "",
            }
        )

    for user_value in users:
        hit = user_lookup.get(sanitize_username(user_value).lower(), {})
        user = str(hit.get("user", "") or user_value)
        uid = str(hit.get("uid", "") or "")
        ten = str(hit.get("ten", "") or "")
        platform = str(hit.get("platform", "") or "")
        phone = normalize_vn_phone(str(hit.get("phone", "") or ""))
        status = str(hit.get("status", "") or ("matched" if hit else "raw_user"))
        link = str(hit.get("link", "") or "").strip()
        if not link and platform:
            link = build_customer_link(platform=platform, uid=uid, username=user, phone=phone)
        if not link and not user.lower().startswith("t_"):
            link = f"https://t.me/{sanitize_username(user)}"
        if hit:
            mapped_count += 1
            mapped_preview.append(f"user={user} | uid={uid or 'N/A'} | ten={ten or 'N/A'} | link={link or 'N/A'}")
        csv_rows.append(
            {
                "type": "user",
                "value": user_value,
                "phone": phone if _is_valid_phone(phone) else "",
                "user": user,
                "uid": uid,
                "ten": ten,
                "link": link,
                "platform": platform,
                "status": status,
                "first_seen": "",
                "last_seen": int(float(hit.get("updated_at", 0) or 0)) if hit else "",
            }
        )
    known_count = max(0, len(phones) - new_count)
    preview_lines = []
    for row in csv_rows[:100]:
        typ = row.get("type", "phone")
        val = row.get("value") or row.get("phone") or row.get("uid") or row.get("user") or ""
        preview_lines.append(
            f"{typ}:{val} | phone={row['phone'] or 'N/A'} | user={row['user'] or 'N/A'} | uid={row['uid'] or 'N/A'} | ten={row['ten'] or 'N/A'} | link={row['link'] or 'N/A'} | {row['status']}"
        )
    preview = "\n".join(preview_lines)
    if len(preview) > 3200:
        preview = preview[:3200] + "\n..."
    report_text = (
        f"✅ Quét được.\nNguồn: {source_type}\nSĐT: {len(phones)} | UID: {len(uids)} | user: {len(users)}\n"
        f"SĐT mới: {new_count}\nSĐT đã có: {known_count}\nMap được user/uid/link: {mapped_count}\n\n{preview}"
    )
    await msg.reply_text(report_text)
    # Gui them file CSV day du mapping.
    supper_files = []
    try:
        csv_buf = io.StringIO()
        writer = csv.DictWriter(
            csv_buf,
            fieldnames=[
                "type",
                "value",
                "phone",
                "user",
                "uid",
                "ten",
                "link",
                "platform",
                "status",
                "first_seen",
                "last_seen",
            ],
        )
        writer.writeheader()
        for row in csv_rows:
            writer.writerow(row)
        csv_bytes = csv_buf.getvalue().encode("utf-8")
        out = io.BytesIO(csv_bytes)
        out.name = f"scan_entities_map_{int(time.time())}.csv"
        supper_files.append((out.name, csv_bytes))
        await context.bot.send_document(
            chat_id=msg.chat_id,
            document=out,
            filename=out.name,
            caption=f"Map ket qua: {len(csv_rows)} dong",
        )
    except Exception:
        pass
    await send_admin_management_report(
        context.bot,
        "[SCAN DATA KHACH]\n" + report_text,
        files=supper_files,
        caption=f"Map ket qua: {len(csv_rows)} dong",
        source_chat_id=msg.chat_id if msg else None,
    )
    maybe_send_phone_event_summary(force=False)
    map_sample = " | ".join(mapped_preview[:20]) if mapped_preview else "none"
    notify_admin_management_sync(
        f"[SCAN DATA KHACH] source={source_type} | admin={actor.id if actor else '?'} | sdt={len(phones)} | uid={len(uids)} | user={len(users)} | moi={new_count} | da_co={known_count} | map={mapped_count} | mau_sdt={','.join(phones[:10])} | mau_uid={','.join(uids[:10])} | mau_user={','.join(users[:10])} | map_chi_tiet={map_sample[:1800]}"
    )


async def cmd_quetnhanh(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    raw = " ".join(context.args).strip()
    if update.message and update.message.reply_to_message and getattr(update.message.reply_to_message, "text", ""):
        raw = f"{raw}\n{update.message.reply_to_message.text}".strip()
    if not raw:
        await update.message.reply_text("Cú pháp: /quetnhanh <noi_dung> hoặc reply tin nhắn rồi gõ /quetnhanh")
        return
    await update.message.reply_text("✅ Đã nhận, ADMIN vui lòng chờ em giây lát.")
    await _scan_text_and_report(update, context, raw, "telegram_text")


async def cmd_quetfile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "quetfile"):
        return
    msg = update.message
    doc = msg.document if msg else None
    if (not doc) and msg and msg.reply_to_message:
        doc = msg.reply_to_message.document
    if not doc:
        await msg.reply_text("Cú pháp: gửi file rồi reply /quetfile (hoặc caption /quetfile). Hỗ trợ txt/csv/xlsx/xls/docx/pdf/json/html/zip/db và nhiều dạng phổ biến.")
        return
    name = str(getattr(doc, "file_name", "") or "")
    await msg.reply_text("✅ Đã nhận, ADMIN vui lòng chờ em giây lát.")
    try:
        f = await doc.get_file()
        blob = bytes(await f.download_as_bytearray())
        text, notes = extract_text_from_file_bytes(blob, name or "upload")
        if notes:
            await msg.reply_text("ℹ️ Ghi chú đọc file:\n" + "\n".join(f"- {n}" for n in notes[:6]))
        await _scan_text_and_report(update, context, text, f"telegram_file:{name or 'upload'}")
        phones = extract_candidate_phones(text)
        if phones:
            await msg.reply_text(
                f"✅ File có {len(phones)} SĐT hợp lệ đã normalize.\n"
                "Muốn quét UID/user: reply file này rồi gửi /phat500k run hoặc /quetallnow run."
            )
    except Exception as exc:
        await msg.reply_text(f"❌ Quét file lỗi: {exc}")


async def cmd_quetfileepath(update: Update, context: ContextTypes.DEFAULT_TYPE, command_name: str = "quetfileepath"):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, command_name):
        return
    if not context.args:
        await update.message.reply_text(f"Cú pháp: /{command_name} <duong_dan_file>. Hỗ trợ txt/csv/xlsx/xls/docx/pdf/json/html/zip/db và nhiều dạng phổ biến.")
        return
    path = " ".join(context.args).strip().strip('"')
    if not os.path.exists(path):
        await update.message.reply_text(f"❌ Không thấy file: {path}")
        return
    try:
        await update.message.reply_text("✅ Đã nhận, ADMIN vui lòng chờ em giây lát.")
        text, notes = extract_text_from_local_file(path)
        if notes:
            await update.message.reply_text("ℹ️ Ghi chú đọc file:\n" + "\n".join(f"- {n}" for n in notes[:6]))
        await _scan_text_and_report(update, context, text, f"local_file:{path}")
    except Exception as exc:
        await update.message.reply_text(f"❌ Quét file lỗi: {exc}")


async def cmd_quetfilepath(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await cmd_quetfileepath(update, context, "quetfilepath")


async def cmd_quetnow(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "quetnow"):
        return
    raw = " ".join(context.args).strip()
    if (not raw) and update.message and update.message.reply_to_message:
        reply_txt = str(getattr(update.message.reply_to_message, "text", "") or "").strip()
        reply_cap = str(getattr(update.message.reply_to_message, "caption", "") or "").strip()
        raw = (reply_txt or reply_cap).strip()
    if not raw:
        await update.message.reply_text("Cú pháp: /quetnow <danh_sach_so_dien_thoai>. Có thể dán nhiều số trong 1 tin nhắn.")
        return
    await update.message.reply_text("✅ Đã nhận, ADMIN vui lòng chờ em giây lát.")
    await _scan_text_and_report(update, context, raw, "telegram_quetnow")


async def cmd_quetuser(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "quetuser"):
        return
    raw = " ".join(context.args).strip()
    if (not raw) and update.message and update.message.reply_to_message:
        reply_txt = str(getattr(update.message.reply_to_message, "text", "") or "").strip()
        reply_cap = str(getattr(update.message.reply_to_message, "caption", "") or "").strip()
        raw = (reply_txt or reply_cap).strip()
    if not raw:
        await update.message.reply_text("Cú pháp: /quetuser <danh_sach_so_dien_thoai>. Có thể dán nhiều số trong 1 tin.")
        return
    msg = update.message
    actor = update.effective_user
    await msg.reply_text("✅ Đã nhận, ADMIN vui lòng chờ em giây lát.")
    phones = extract_candidate_phones(raw)
    if not phones:
        await msg.reply_text("❌ Không tìm thấy số điện thoại hợp lệ trong nội dung.")
        return

    lookup = build_phone_lookup_index()
    mapped = []
    unknown = []
    csv_rows = []
    for p in phones:
        hits = lookup.get(p, [])
        hit = choose_best_lookup_hit(hits)
        user = str(hit.get("user", "") or "") if hit else ""
        uid = str(hit.get("uid", "") or "") if hit else ""
        ten = str(hit.get("ten", "") or "") if hit else ""
        platform = str(hit.get("platform", "") or "") if hit else ""
        link = str(hit.get("link", "") or "") if hit else ""
        if not link:
            link = build_customer_link(platform=platform, uid=uid, username=user, phone=p)

        save_phonebook_entry(
            phone=p,
            platform=platform or "telegram_quetuser",
            source_uid=uid,
            username=user,
            full_name=ten,
            source_chat=str(msg.chat_id if msg else ""),
            last_message_snippet=f"telegram_quetuser:{raw[:120]}",
        )
        log_phone_event(
            phone=p,
            platform=platform or "telegram_quetuser",
            source_uid=uid,
            username=user,
            full_name=ten,
            source_chat=str(msg.chat_id if msg else ""),
            is_new_phone=False,
            is_new_uid=False,
        )
        if user or uid:
            mapped.append(f"{p} | user={user or 'N/A'} | uid={uid or 'N/A'} | link={link or 'N/A'}")
        else:
            unknown.append(p)
        csv_rows.append(
            {
                "phone": p,
                "user": user,
                "uid": uid,
                "ten": ten,
                "link": link,
                "platform": platform,
                "status": "matched" if (user or uid) else "unknown",
            }
        )

    summary_lines = [
        "✅ /quetuser hoàn tất.",
        f"Tổng số: {len(phones)}",
        f"Map được user/uid: {len(mapped)}",
        f"Chưa có user/uid: {len(unknown)}",
    ]
    if mapped:
        summary_lines.append("")
        summary_lines.append("📌 Mẫu map:")
        summary_lines.extend(mapped[:30])
    if unknown:
        summary_lines.append("")
        summary_lines.append("📭 Chưa map được:")
        summary_lines.append(", ".join(unknown[:50]))
    summary_text = "\n".join(summary_lines)
    await msg.reply_text(summary_text[:3800])
    supper_files = []
    try:
        csv_buf = io.StringIO()
        writer = csv.DictWriter(csv_buf, fieldnames=["phone", "user", "uid", "ten", "link", "platform", "status"])
        writer.writeheader()
        for row in csv_rows:
            writer.writerow(row)
        csv_bytes = csv_buf.getvalue().encode("utf-8-sig")
        filename = f"quetuser_map_{int(time.time())}.csv"
        supper_files.append((filename, csv_bytes))
        bio = io.BytesIO(csv_bytes)
        bio.name = filename
        await context.bot.send_document(
            chat_id=msg.chat_id,
            document=bio,
            filename=filename,
            caption=f"/quetuser map: {len(csv_rows)} dong",
        )
    except Exception:
        pass
    await send_admin_management_report(
        context.bot,
        "[QUETUSER]\n" + summary_text,
        files=supper_files,
        caption=f"/quetuser map: {len(csv_rows)} dong",
        source_chat_id=msg.chat_id if msg else None,
    )

    notify_admin_management_sync(
        "[QUETUSER] "
        f"admin={actor.id if actor else '?'} | tong={len(phones)} | map={len(mapped)} | unknown={len(unknown)} | "
        f"mapped_chi_tiet={( ' || '.join(mapped[:20]) if mapped else 'none')[:1800]}"
    )
    maybe_send_phone_event_summary(force=False)


async def _cmd_contact_import_now(update: Update, context: ContextTypes.DEFAULT_TYPE, mode: str, command_name: str | None = None):
    if not is_admin(update.effective_user.id):
        return
    msg = update.message
    command_name = command_name or (mode + "now")
    if not await userbot_operation_command_is_for_this_bot(update, context, command_name):
        return
    command_label = f"/{command_name}"
    command_tail = get_command_tail(update, command_name)
    run_requested = command_tail_requests_run(command_tail)
    reply_rows, reply_source, reply_notes = await load_contact_import_reply_file_users(update, context, valid_only=False)
    if reply_source:
        rows, source_path = _filter_contact_import_rows_for_mode(reply_rows, mode), reply_source
        reply_phones = unique_contact_phones_from_rows(reply_rows)
        if reply_phones and mode == "all":
            if run_requested:
                await queue_contact_import_uid_user_job(
                    update,
                    context,
                    reply_phones,
                    reply_notes,
                    dry_run=False,
                    command_label=command_label,
                    source_label=reply_source,
                )
            else:
                await reply_contact_import_phone_preview(msg, source_path, reply_phones, command_label, reply_notes)
            return
        if run_requested and reply_phones and mode in {"uid", "user"}:
            await queue_contact_import_phone_job(
                update,
                context,
                reply_phones,
                reply_notes,
                dry_run=False,
                command_label=command_label,
                source_label=reply_source,
                report_mode="quetallnow" if mode == "all" else mode,
                single_report_file=(mode == "all"),
            )
            return
    else:
        rows, source_path = load_contact_import_users_for_mode(mode)
        reply_notes = []
        reply_phones = []
    partial_note = ""
    if reply_source and mode in {"uid", "user", "all"} and reply_phones:
        partial_note = contact_import_partial_note(len(rows), len(reply_phones))
    if not rows:
        note_text = ""
        if reply_notes:
            note_text = "\nGhi chú đọc file:\n- " + "\n- ".join(str(n) for n in reply_notes[:5])
        if reply_source and not reply_phones and mode in {"uid", "user", "all"}:
            extra_phones, extra_notes = await extract_phat500k_contact_phones(update, context, command_tail)
            if extra_phones:
                reply_phones = extra_phones
                reply_notes = (reply_notes or []) + extra_notes
        if reply_source and reply_phones and mode in {"uid", "user", "all"}:
            if run_requested:
                if mode == "all":
                    await queue_contact_import_uid_user_job(
                        update,
                        context,
                        reply_phones,
                        reply_notes,
                        dry_run=False,
                        command_label=command_label,
                        source_label=reply_source,
                    )
                else:
                    await queue_contact_import_phone_job(
                        update,
                        context,
                        reply_phones,
                        reply_notes,
                        dry_run=False,
                        command_label=command_label,
                        source_label=reply_source,
                        report_mode=mode,
                        single_report_file=False,
                    )
                return
            await reply_contact_import_phone_preview(msg, source_path, reply_phones, command_label, reply_notes)
            return
        if reply_source and run_requested:
            contact_phones, contact_notes = await extract_phat500k_contact_phones(update, context, command_tail)
            if contact_phones and mode in {"uid", "user", "all"}:
                if mode == "all":
                    await queue_contact_import_uid_user_job(
                        update,
                        context,
                        contact_phones,
                        (reply_notes or []) + contact_notes,
                        dry_run=False,
                        command_label=command_label,
                        source_label=reply_source,
                    )
                else:
                    await queue_contact_import_phone_job(
                        update,
                        context,
                        contact_phones,
                        (reply_notes or []) + contact_notes,
                        dry_run=False,
                        command_label=command_label,
                        source_label=reply_source,
                        report_mode=mode,
                        single_report_file=False,
                    )
                return
        await msg.reply_text(
            "❌ Chưa có khách hợp lệ trong file quét gần nhất. "
            f"File: {source_path}\n"
            f"{_contact_import_condition_text(mode)}"
            f"{note_text}"
        )
        return
    actor_id = str(update.effective_user.id)
    export_path = write_contact_import_export(rows, mode, actor_id)
    labels = {
        "user": "/usernow - toàn bộ @user hợp lệ",
        "uid": "/uidnow - uid + @user hợp lệ",
        "tt": "/ttnow - sdt/uid/user/tên/account",
        "sdt": "/sdtnow - sdt + @user hợp lệ",
        "all": "/quetallnow - file /guitn hợp lệ UID | USER | PHONE | ACCOUNT | ONL BAO LAU | LINK",
    }
    preview_lines = [contact_import_line(row, mode) for row in rows[:60]]
    summary = (
        f"✅ {labels.get(mode, mode)}\n"
        f"Nguồn: {source_path}\n"
        f"Dòng hợp lệ: {len(rows)}\n\n"
        + "\n".join(preview_lines)
    )
    if partial_note:
        summary += partial_note
    if reply_notes:
        summary += "\n\nGhi chú đọc file:\n- " + "\n- ".join(str(n) for n in reply_notes[:5])
    if len(rows) > len(preview_lines):
        summary += f"\n... còn {len(rows) - len(preview_lines)} dòng trong file đính kèm."
    await send_telegram_long(context.bot, update.effective_chat.id, summary)
    admin_files = [export_path]
    if mode == "all":
        readable_report = build_contact_import_uid_user_report(rows)
        if readable_report:
            filename, blob = readable_report
            admin_files.insert(0, readable_report)
            try:
                bio = io.BytesIO(blob)
                bio.name = filename
                await context.bot.send_document(
                    chat_id=update.effective_chat.id,
                    document=bio,
                    filename=filename,
                    caption=f"File chính UID | USER | {len(rows)} dòng hợp lệ",
                )
            except Exception as exc:
                logging.warning("Gui UID|USER /quetallnow report that bai: %s", exc)
    try:
        with open(export_path, "rb") as f:
            await context.bot.send_document(
                chat_id=update.effective_chat.id,
                document=f,
                filename=os.path.basename(export_path),
                caption=f"File /guitn đầy đủ: {labels.get(mode, mode)} | {len(rows)} dòng hợp lệ",
            )
    except Exception as exc:
        await msg.reply_text(f"⚠️ Không gửi được file đính kèm: {exc}\nFile local: {export_path}")
    await send_admin_management_report(
        context.bot,
        "[CONTACT IMPORT NOW]\n" + summary,
        files=admin_files,
        caption=f"File day du: {labels.get(mode, mode)} | {len(rows)} dong hop le",
        source_chat_id=update.effective_chat.id,
    )


async def cmd_usernow(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_contact_import_now(update, context, "user")


async def cmd_uidnow(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_contact_import_now(update, context, "uid")


async def cmd_ttnow(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_contact_import_now(update, context, "tt")


async def cmd_sdtnow(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_contact_import_now(update, context, "sdt")


async def cmd_quetallnow(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_contact_import_now(update, context, "all", command_name="quetallnow")


async def cmd_uidstatus(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    ok = False
    for command_name in ("uidstatus", "phat500kstatus", "importstatus"):
        if await userbot_operation_command_is_for_this_bot(update, context, command_name):
            ok = True
            break
    if not ok:
        return
    await send_telegram_long(context.bot, update.effective_chat.id, contact_import_status_text())


async def cmd_quetstatus(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "quetstatus"):
        return
    await send_telegram_long(context.bot, update.effective_chat.id, contact_import_status_text())


async def command_is_for_this_bot(update: Update, context: ContextTypes.DEFAULT_TYPE, command_name: str) -> bool:
    msg = update.message
    raw = str(getattr(msg, "text", "") or getattr(msg, "caption", "") or "")
    pattern = rf"(?is)^\s*/{re.escape(command_name)}(?:@([A-Za-z0-9_]+))?(?:\s|$)"
    m = re.match(pattern, raw)
    if not m:
        return False
    addressed_username = (m.group(1) or "").strip().lower()
    if not addressed_username:
        return True
    try:
        me = await context.bot.get_me()
        my_username = str(getattr(me, "username", "") or "").strip().lower()
    except Exception:
        my_username = ""
    return bool(my_username and addressed_username == my_username)


async def userbot_operation_command_is_for_this_bot(update: Update, context: ContextTypes.DEFAULT_TYPE, command_name: str) -> bool:
    if not await command_is_for_this_bot(update, context, command_name):
        return False
    chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
    if chat_type in {"group", "supergroup"} and not command_has_bot_mention(update, command_name):
        return False
    return True


def command_has_bot_mention(update: Update, command_name: str) -> bool:
    msg = update.message
    raw = str(getattr(msg, "text", "") or getattr(msg, "caption", "") or "")
    return bool(re.match(rf"(?is)^\s*/{re.escape(command_name)}@[A-Za-z0-9_]+(?:\s|$)", raw))


async def primary_operation_command_is_for_this_bot(update: Update, context: ContextTypes.DEFAULT_TYPE, command_name: str) -> bool:
    if not await command_is_for_this_bot(update, context, command_name):
        return False
    return is_primary_operation_bot_instance()


def _delayed_stop_current_bot(reason: str, delay_sec: float = 1.5) -> None:
    time.sleep(max(0.1, float(delay_sec or 1.5)))
    try:
        persist_runtime_state()
    except Exception:
        pass
    try:
        backup_db(reason)
    except Exception:
        pass
    try:
        set_config("bot_polling_last", f"{int(time.time())}|manual_stop slot={current_bot_slot()} reason={reason}")
    except Exception:
        pass
    os._exit(0)


async def queue_guitn_jobs_from_plan(update: Update, context: ContextTypes.DEFAULT_TYPE, plan: dict, lot_key: str) -> int:
    options = get_guitn_lot_options()
    lot_key = normalize_guitn_lot_key(lot_key)
    option = options.get(str(lot_key or "").lower())
    if not option:
        await update.message.reply_text(guitn_lot_prompt_text())
        return 0
    roots = [root for root in option.get("roots") or [] if root and os.path.isdir(root)]
    if not roots:
        await update.message.reply_text(f"❌ Không thấy root userbot cho {option.get('label')}.")
        return 0
    selected = read_userbot_send_recipients_file(str(plan.get("recipients_file") or ""))
    if not selected:
        await update.message.reply_text("❌ Pending /guitn không còn file người nhận. Gửi lại /guitn để tạo plan mới.")
        return 0
    selected = annotate_campaign_recipients(selected, str(plan.get("source_path") or plan.get("source_users_file") or "pending_guitn"))
    sendable_rows, missing_consent_rows = split_campaign_consent_rows(selected)
    if CAMPAIGN_CONSENT_REQUIRED and missing_consent_rows:
        await update.message.reply_text(
            "🛑 Chặn chạy thật /guitn: còn người nhận thiếu consent/opt-in.\n"
            f"Đủ consent: {len(sendable_rows)} | Thiếu consent: {len(missing_consent_rows)}\n"
            "Hãy dùng file target có cột consent=opt_in hoặc để khách tự /start trước. "
            "Pending hiện tại chỉ được xem như dry-run."
        )
        return 0

    msg_text = str(plan.get("message_text") or "").strip()
    media_to_send = str(plan.get("media_url") or "").strip()
    media_source = str(plan.get("media_source") or "").strip()
    delay_min = float(plan.get("delay_message_min_sec") or 1.5)
    delay_max = float(plan.get("delay_message_max_sec") or delay_min)
    delay_label = str(plan.get("delay_message_label") or (f"{delay_min:g}s" if delay_min == delay_max else f"{delay_min:g}-{delay_max:g}s"))
    source_path = str(plan.get("source_path") or plan.get("source_users_file") or "")
    profile_limit = int(option.get("profile_limit") or get_userbot_job_profile_limit() or 1)
    active_roots = active_userbot_root_profiles(roots, profile_limit)
    if not active_roots:
        await update.message.reply_text(
            "❌ Không còn userbot active để gửi tin. Kiểm tra cooldown/session tại "
            f"{USERBOT_ACCOUNT_COOLDOWN_JSON}"
        )
        return 0
    roots = [root for root, _profiles in active_roots]

    selected_chunks = split_rows_by_roots(selected, roots)
    job_ids = []
    recipient_files = []
    shard_lines = []
    for idx, ((root, profiles), selected_for_root) in enumerate(zip(active_roots, selected_chunks), start=1):
        if not selected_for_root:
            continue
        session_count = count_userbot_api_sessions(profiles, root=root)
        recipients_file = write_userbot_send_recipients(selected_for_root, str(update.effective_user.id))
        recipient_files.append(recipients_file)
        detail = {
            "job_type": "send_message",
            "profiles_root": root,
            "profile_count": len(profiles),
            "max_accounts_per_job": len(profiles),
            "profiles": profiles,
            "session_count": session_count,
            "source_users_file": source_path,
            "recipients_file": recipients_file,
            "recipient_count": len(selected_for_root),
            "message_text": msg_text,
            "media_url": media_to_send,
            "media_source": media_source,
            "delay_message_sec": delay_min,
            "delay_message_min_sec": delay_min,
            "delay_message_max_sec": delay_max,
            "delay_message_label": delay_label,
            "delay_account_sec": 5,
            "mode": "run",
            "valid_rule": "username_or_uid_required",
            "bot_slot": current_bot_slot(),
            "guitn_lot": lot_key,
            "guitn_lot_label": option.get("label"),
            "shard_index": idx,
            "shard_total": len(roots),
        }
        detail = attach_userbot_report_origin(detail, update)
        job_id = queue_userbot_scan_job(
            "userbot_send_message",
            str(getattr(update.effective_chat, "id", "") or ""),
            f"USERBOT_SEND_MESSAGE_{lot_key.upper()}_{idx}",
            str(update.effective_user.id),
            False,
            detail,
        )
        job_ids.append(job_id)
        session_note = "" if session_count == len(profiles) else f" | session={session_count}/{len(profiles)}"
        shard_lines.append(
            f"- {option.get('label')} shard {idx}/{len(roots)}: job #{job_id} | recipients={len(selected_for_root)} | profiles={len(profiles)}{session_note}\n"
            f"  Root: {root}\n"
            f"  File: {recipients_file}"
        )

    if not job_ids:
        await update.message.reply_text("❌ Không tạo được job /guitn vì lô đã chọn không có recipient hợp lệ.")
        return 0

    sample = "\n".join(contact_import_line(row, "tt") for row in selected[:20])
    report_text = (
        f"✅ Đã tạo job /guitn cho {option.get('label')}.\n"
        f"Jobs: {', '.join('#' + str(x) for x in job_ids)}\n"
        f"Bot nhận lệnh: {current_bot_slot()}\n"
        f"Nguồn nhận: {source_path}\n"
        f"Người nhận hợp lệ có user/uid: {len(selected)}\n"
        f"Delay mỗi tin: {delay_label}\n"
        f"GIF/media: {media_to_send or 'OFF'}\n"
        "Worker sẽ báo START / PROGRESS / DONE bằng đúng bot nhận lệnh.\n"
        "Chia lô:\n"
        + "\n".join(shard_lines)
        + "\n\nMẫu người nhận:\n"
        + (sample or "N/A")
    )
    await send_telegram_long(context.bot, update.effective_chat.id, report_text)
    await send_admin_management_report(
        context.bot,
        "[GUITN QUEUED]\n" + report_text,
        files=recipient_files,
        caption=f"/guitn {lot_key}: {len(selected)} dong",
        source_chat_id=update.effective_chat.id,
    )
    clear_guitn_pending_plan(update)
    return len(job_ids)


async def cmd_guitn(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "guitn"):
        return
    if not is_guitn_bot_instance():
        await update.message.reply_text(
            "❌ /guitn chỉ chạy ở 3 BotFather đã cấu hình. "
            "Gửi lệnh sang đúng bot rồi reply file /uidnow."
        )
        return
    msg_text, selector_text, delay_min, delay_max, delay_label = parse_guitn_payload(update)
    media_path, media_source = await load_guitn_reply_media(update, context)
    if not msg_text and not media_path:
        await update.message.reply_text(
            "Cú pháp: /guitn <noi_dung> [60s|60-180]\n"
            "Hoặc forward/reply tin mẫu có ảnh/GIF/video rồi gõ /guitn 60-180.\n"
            "Hoặc reply file /uidnow rồi gõ /guitn <noi_dung> 60-180.\n"
            "Dùng dấu | nếu muốn chọn UID/user riêng: /guitn <noi_dung> | <uid1 @user2 ...> 60s"
        )
        return
    guitn_root = get_guitn_accounts_root()
    file_rows, file_source, file_notes = await load_guitn_reply_file_recipients(update, context)
    if file_rows:
        rows, source_path = file_rows, file_source
    else:
        rows, source_path = load_contact_import_users(valid_only=True)
    if not rows:
        await update.message.reply_text(
            "❌ Chưa có khách hợp lệ để gửi. "
            f"File: {source_path}\n"
            "Điều kiện hợp lệ: file phải có UID hoặc username/user."
        )
        return
    selected, missing = select_contact_import_users(rows, selector_text)
    if not selected:
        await update.message.reply_text(
            "❌ Không map được UID/user nào trong nguồn đang dùng. "
            "Nếu reply file /uidnow thì file phải có dòng dạng UID | @user."
        )
        return
    selected = annotate_campaign_recipients(selected, source_path)
    sendable_rows, missing_consent_rows = split_campaign_consent_rows(selected)
    media_to_send = media_path or USERBOT_GUITN_GIF
    recipients_file = write_userbot_send_recipients(selected, str(update.effective_user.id))
    plan = {
        "message_text": msg_text,
        "selector_text": selector_text,
        "source_path": source_path,
        "source_users_file": source_path,
        "recipients_file": recipients_file,
        "recipient_count": len(selected),
        "sendable_count": len(sendable_rows),
        "missing_consent_count": len(missing_consent_rows),
        "consent_required": CAMPAIGN_CONSENT_REQUIRED,
        "media_url": media_to_send,
        "media_source": media_source or ("configured_guitn_media" if USERBOT_GUITN_GIF else ""),
        "delay_message_sec": delay_min,
        "delay_message_min_sec": delay_min,
        "delay_message_max_sec": delay_max,
        "delay_message_label": delay_label,
        "created_at": int(time.time()),
    }
    plan = attach_userbot_report_origin(plan, update)
    plan_path = save_guitn_pending_plan(update, plan)
    sample = "\n".join(contact_import_line(row, "tt") for row in selected[:15])
    mention_note = ""
    chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
    if chat_type in {"group", "supergroup"}:
        try:
            me = await context.bot.get_me()
            mention_note = f"\nTrong group hãy chọn dạng /run40@{getattr(me, 'username', '') or 'ten_bot'} hoặc /run52@{getattr(me, 'username', '') or 'ten_bot'} để đúng bot chạy."
        except Exception:
            mention_note = "\nTrong group hãy mention đúng bot khi chọn lô."
    await send_telegram_long(
        context.bot,
        update.effective_chat.id,
        (
            "✅ Đã lưu lệnh /guitn, CHƯA chạy gửi tin.\n"
            f"Bot nhận lệnh: {current_bot_slot()}\n"
            f"Nguồn nhận: {source_path}\n"
            f"Người nhận hợp lệ: {len(selected)}\n"
            f"Đủ consent để chạy thật: {len(sendable_rows)} | Thiếu consent: {len(missing_consent_rows)}\n"
            f"Delay mỗi tin: {delay_label}\n"
            f"GIF/media: {media_to_send or 'OFF'}\n"
            f"File plan: {plan_path}\n\n"
            f"{guitn_lot_prompt_text()}"
            f"{mention_note}\n"
            "\nMẫu người nhận:\n"
            f"{sample or 'N/A'}"
        ),
    )
    return
    roots = get_userbot_job_roots(guitn_root)
    selected_chunks = split_rows_by_roots(selected, roots)
    media_to_send = media_path or USERBOT_GUITN_GIF
    job_ids = []
    recipient_files = []
    shard_lines = []
    for idx, (root, selected_for_root) in enumerate(zip(roots, selected_chunks), start=1):
        if not selected_for_root:
            continue
        profiles = list_userbot_profiles(limit=get_userbot_job_profile_limit(), root=root)
        session_count = count_userbot_api_sessions(profiles, root=root)
        recipients_file = write_userbot_send_recipients(selected_for_root, str(update.effective_user.id))
        recipient_files.append(recipients_file)
        detail = {
            "job_type": "send_message",
            "profiles_root": root,
            "profile_count": len(profiles),
            "max_accounts_per_job": len(profiles),
            "profiles": profiles,
            "session_count": session_count,
            "source_users_file": source_path,
            "recipients_file": recipients_file,
            "recipient_count": len(selected_for_root),
            "message_text": msg_text,
            "media_url": media_to_send,
            "media_source": media_source or ("configured_guitn_media" if USERBOT_GUITN_GIF else ""),
            "delay_message_sec": delay_min,
            "delay_message_min_sec": delay_min,
            "delay_message_max_sec": delay_max,
            "delay_message_label": delay_label,
            "delay_account_sec": 5,
            "mode": "run",
            "valid_rule": "username_or_uid_required",
            "bot_slot": current_bot_slot(),
            "shard_index": idx,
            "shard_total": len(roots),
        }
        detail = attach_userbot_report_origin(detail, update)
        job_id = queue_userbot_scan_job(
            "userbot_send_message",
            str(getattr(update.effective_chat, "id", "") or ""),
            f"USERBOT_SEND_MESSAGE_SHARD_{idx}",
            str(update.effective_user.id),
            False,
            detail,
        )
        job_ids.append(job_id)
        session_note = "" if session_count == len(profiles) else f" | session={session_count}/{len(profiles)}"
        shard_lines.append(
            f"- Shard {idx}/{len(roots)}: job #{job_id} | recipients={len(selected_for_root)} | profiles={len(profiles)}{session_note}\n"
            f"  Root: {root}\n"
            f"  File: {recipients_file}"
        )
    sample = "\n".join(contact_import_line(row, "tt") for row in selected[:20])
    missing_note = ""
    if missing:
        missing_note = "\nKhông thấy trong nguồn: " + ", ".join(missing[:30])
        if len(missing) > 30:
            missing_note += f" ... +{len(missing) - 30}"
    notes_text = ""
    if file_notes:
        notes_text = "\nGhi chú đọc file:\n" + "\n".join(f"- {n}" for n in file_notes[:6])
    report_text = (
        f"✅ Đã tạo job /guitn theo {len([x for x in selected_chunks if x])}/{len(roots)} botuser shard.\n"
        f"Jobs: {', '.join('#' + str(x) for x in job_ids) or 'N/A'}\n"
        f"Bot nhận lệnh: {current_bot_slot()}\n"
        f"Nguồn nhận: {source_path}\n"
        f"Người nhận hợp lệ có user/uid: {len(selected)}\n"
        f"Delay mỗi tin: {delay_label}\n"
        f"GIF/media: {media_to_send or 'OFF'}\n"
        f"Worker sẽ báo START / PROGRESS / DONE về nhóm /adminquanly.\n"
        "Chia shard:\n"
        + ("\n".join(shard_lines) or "- chưa có shard hợp lệ")
        + "\n"
        f"Nội dung: {msg_text[:800]}"
        f"{missing_note}{notes_text}\n\n"
        "Mẫu người nhận:\n"
        + (sample or "N/A")
    )
    await send_telegram_long(
        context.bot,
        update.effective_chat.id,
        report_text,
    )
    await send_admin_management_report(
        context.bot,
        "[GUITN QUEUED]\n" + report_text,
        files=recipient_files,
        caption=f"/guitn recipients: {len(selected)} dong",
        source_chat_id=update.effective_chat.id,
    )


async def cmd_guitnlai(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "guitnlai"):
        return
    if not is_guitn_bot_instance():
        await update.message.reply_text(
            "❌ /guitnlai chỉ chạy ở 3 BotFather đã cấu hình. "
            "Gửi lệnh sang đúng bot rồi reply file khách lỗi."
        )
        return
    msg_text, selector_text, delay_min, delay_max, delay_label, delay_explicit = parse_guitn_payload_for_command(update, "guitnlai")
    media_path, media_source = await load_guitn_reply_media(update, context)
    file_rows, file_source, file_notes = await load_guitn_reply_file_recipients(update, context)
    retry_summary = {}
    retry_notes = []
    if file_rows:
        rows, source_path = file_rows, file_source
    else:
        rows, source_path, retry_notes, retry_summary = load_latest_guitn_unsent_rows()
    notes = list(file_notes or []) + list(retry_notes or [])
    retry_detail = retry_summary.get("detail") or {}
    if not msg_text and not media_path:
        msg_text = str(retry_detail.get("message_text") or "").strip()
    if not delay_explicit:
        try:
            saved_min = float(retry_detail.get("delay_message_min_sec") or retry_detail.get("delay_message_sec") or delay_min)
            saved_max = float(retry_detail.get("delay_message_max_sec") or retry_detail.get("delay_message_sec") or delay_max)
            if saved_max < saved_min:
                saved_min, saved_max = saved_max, saved_min
            delay_min, delay_max = saved_min, saved_max
            delay_label = str(retry_detail.get("delay_message_label") or delay_label)
        except Exception:
            pass
    media_to_send = media_path or str(retry_detail.get("media_url") or "").strip() or USERBOT_GUITN_GIF
    effective_media_source = media_source or ("retry_saved_media" if retry_detail.get("media_url") else ("configured_guitn_media" if USERBOT_GUITN_GIF else ""))
    if not rows:
        await update.message.reply_text(
            "❌ Chưa thấy file khách gửi chưa thành công để gửi lại.\n"
            "Hãy reply file `userbot_send_unsent_...` rồi gõ /guitnlai <noi_dung> 60-180,\n"
            "hoặc chạy xong /guitn để bot sinh file lỗi mới nhất."
        )
        return
    if not msg_text and not media_to_send:
        await update.message.reply_text(
            "Cú pháp: /guitnlai <noi_dung> [60s|60-180]\n"
            "Hoặc reply file userbot_send_unsent rồi gõ /guitnlai 60-180 để dùng lại nội dung job gần nhất."
        )
        return
    selected, missing = select_contact_import_users(rows, selector_text)
    if not selected:
        await update.message.reply_text(
            "❌ Không map được UID/user nào trong file khách gửi chưa thành công đang dùng."
        )
        return
    selected = annotate_campaign_recipients(selected, source_path)
    sendable_rows, missing_consent_rows = split_campaign_consent_rows(selected)
    recipients_file = write_userbot_send_recipients(selected, str(update.effective_user.id))
    plan = {
        "message_text": msg_text,
        "selector_text": selector_text,
        "source_path": source_path,
        "source_users_file": source_path,
        "recipients_file": recipients_file,
        "recipient_count": len(selected),
        "sendable_count": len(sendable_rows),
        "missing_consent_count": len(missing_consent_rows),
        "consent_required": CAMPAIGN_CONSENT_REQUIRED,
        "media_url": media_to_send,
        "media_source": effective_media_source,
        "delay_message_sec": delay_min,
        "delay_message_min_sec": delay_min,
        "delay_message_max_sec": delay_max,
        "delay_message_label": delay_label,
        "created_at": int(time.time()),
        "retry_unsent": True,
        "retry_original_job_id": str(retry_summary.get("job_id") or ""),
    }
    plan = attach_userbot_report_origin(plan, update)
    plan_path = save_guitn_pending_plan(update, plan)
    sample = "\n".join(contact_import_line(row, "tt") for row in selected[:15])
    missing_note = ""
    if missing:
        missing_note = "\nKhông thấy trong file lỗi: " + ", ".join(missing[:30])
        if len(missing) > 30:
            missing_note += f" ... +{len(missing) - 30}"
    note_text = ""
    if notes:
        note_text = "\nGhi chú nguồn retry:\n" + "\n".join(f"- {n}" for n in notes[:5])
    mention_note = ""
    chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
    if chat_type in {"group", "supergroup"}:
        try:
            me = await context.bot.get_me()
            mention_note = f"\nTrong group hãy chọn dạng /run40@{getattr(me, 'username', '') or 'ten_bot'} hoặc /run52@{getattr(me, 'username', '') or 'ten_bot'} để đúng bot chạy."
        except Exception:
            mention_note = "\nTrong group hãy mention đúng bot khi chọn lô."
    retry_job = str(retry_summary.get("job_id") or "").strip()
    await send_telegram_long(
        context.bot,
        update.effective_chat.id,
        (
            "✅ Đã lưu lệnh /guitnlai, CHƯA chạy gửi lại.\n"
            f"Bot nhận lệnh: {current_bot_slot()}\n"
            f"Nguồn retry: {source_path}\n"
            f"Job lỗi gần nhất: {retry_job or 'N/A'}\n"
            f"Người nhận chưa thành công: {len(selected)}\n"
            f"Đủ consent để chạy thật: {len(sendable_rows)} | Thiếu consent: {len(missing_consent_rows)}\n"
            f"Delay mỗi tin: {delay_label}\n"
            f"GIF/media: {media_to_send or 'OFF'}\n"
            f"File plan: {plan_path}\n\n"
            f"{guitn_lot_prompt_text()}"
            f"{mention_note}{missing_note}{note_text}\n"
            "\nMẫu người nhận retry:\n"
            f"{sample or 'N/A'}"
        ),
    )


async def _cmd_run_guitn_lot(update: Update, context: ContextTypes.DEFAULT_TYPE, command_name: str):
    if not is_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, command_name):
        return
    if is_removed_guitn_lot_key(command_name):
        await update.message.reply_text(
            "⚠️ Đã gộp còn 2 lô userbot.\n"
            f"{guitn_lot_prompt_text()}"
        )
        return
    plan, plan_path = load_guitn_pending_plan(update)
    if not plan:
        chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
        if chat_type == "private":
            await update.message.reply_text("Chưa có pending /guitn. Gửi /guitn trước, bot sẽ hỏi chọn lô rồi hãy dùng lệnh này.")
        return
    await queue_guitn_jobs_from_plan(update, context, plan, command_name.lower())


async def cmd_run40(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_run_guitn_lot(update, context, "run40")


async def cmd_run52(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_run_guitn_lot(update, context, "run52")


async def cmd_run12(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_run_guitn_lot(update, context, "run12")


async def cmd_run20(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_run_guitn_lot(update, context, "run20")


async def cmd_run32(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_run_guitn_lot(update, context, "run32")


async def cmd_run20new(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_run_guitn_lot(update, context, "run20new")


async def cmd_runall(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await _cmd_run_guitn_lot(update, context, "run52")


def summarize_guitn_worker_log() -> str:
    live_status = summarize_guitn_live_status()
    if live_status:
        return live_status
    log_path = os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_import_contacts_worker.log")
    if not os.path.exists(log_path):
        return f"Chua co log worker: {log_path}"
    try:
        with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.read().splitlines()
    except Exception as exc:
        return f"Khong doc duoc log worker: {exc}"
    current_job = ""
    current_lines = []
    for line in lines:
        m_job = re.search(r"Process userbot send job #(.+)$", line)
        if m_job:
            current_job = m_job.group(1).strip()
            current_lines = [line]
        elif current_job:
            current_lines.append(line)
    if not current_job:
        return "Chua thay job /guitn nao trong log worker."
    accounts = {}
    for line in current_lines:
        m = re.search(r"\] (\d+): sent (\S+) \((\d+)/(\d+)\)", line)
        if not m:
            continue
        acc, target, sent, total = m.groups()
        accounts[acc] = {
            "sent": int(sent),
            "total": int(total),
            "target": target,
            "line": line,
        }
    total_sent = sum(v["sent"] for v in accounts.values())
    total_assigned = sum(v["total"] for v in accounts.values())
    last_line = current_lines[-1] if current_lines else ""
    out = [
        "📨 GUITN STATUS",
        f"Job: {current_job}",
        f"Accounts da gui: {len(accounts)}",
        f"Sent log: {total_sent}/{total_assigned or '?'}",
        f"Log moi nhat: {last_line}",
        "",
        "Theo account:",
    ]
    for acc in sorted(accounts):
        v = accounts[acc]
        out.append(f"- {acc}: {v['sent']}/{v['total']} | last={v['target']}")
    return "\n".join(out)


def summarize_guitn_live_status() -> str:
    if not os.path.exists(USERBOT_SEND_LIVE_STATUS_JSON):
        return ""
    try:
        with open(USERBOT_SEND_LIVE_STATUS_JSON, "r", encoding="utf-8-sig") as f:
            payload = json.load(f)
    except Exception:
        return ""
    results = payload.get("results") or []
    if not isinstance(results, list):
        results = []
    total_assigned = sum(int(r.get("assigned") or 0) for r in results)
    total_processed = int(payload.get("processed") or sum(int(r.get("processed") or 0) for r in results))
    total_sent = int(payload.get("sent") or sum(int(r.get("sent") or 0) for r in results))
    total_failed = int(payload.get("failed") or sum(int(r.get("failed") or 0) for r in results))
    done_accounts = sum(1 for r in results if r.get("done") or int(r.get("processed") or 0) >= int(r.get("assigned") or 0) > 0)
    active_accounts = sum(1 for r in results if int(r.get("assigned") or 0) > 0 and not r.get("done") and int(r.get("processed") or 0) < int(r.get("assigned") or 0))
    total_target = total_assigned or int(payload.get("recipient_count") or 0)
    percent = (total_processed / total_target * 100) if total_target else 0.0
    updated_at = payload.get("updated_at")
    age_text = ""
    try:
        age = max(0, int(time.time() - float(updated_at)))
        age_text = f" | cách đây {age}s"
    except Exception:
        pass
    lines = [
        "📨 TRẠNG THÁI GỬI TIN /guitn",
        f"Job: {payload.get('job_id') or 'N/A'}",
        f"Đã xử lý: {total_processed}/{total_target or '?'} ({percent:.1f}%).",
        f"Đã gửi: {total_sent} | Lỗi/chặn: {total_failed}",
        f"Acc xong: {done_accounts}/{payload.get('account_count') or len(results)} | Acc đang chạy: {active_accounts}",
        f"Delay mỗi tin: {payload.get('delay_message_label') or 'N/A'}",
        f"GIF/media: {payload.get('media_url') or 'OFF'}",
        f"Cập nhật: {_format_local_time(updated_at)}{age_text}",
        "",
        "Theo account:",
    ]
    for row in sorted(results, key=lambda x: str(x.get("account") or ""))[:30]:
        assigned = int(row.get("assigned") or 0)
        if assigned <= 0:
            continue
        lines.append(
            f"- {row.get('account')}: {row.get('processed', 0)}/{assigned} | "
            f"sent={row.get('sent', 0)} | lỗi={row.get('failed', 0)} | {row.get('status') or 'running'}"
        )
    if os.path.exists(USERBOT_SEND_LAST_SUMMARY):
        try:
            mtime = os.path.getmtime(USERBOT_SEND_LAST_SUMMARY)
            lines.append("")
            lines.append(f"File tổng kết gần nhất: {USERBOT_SEND_LAST_SUMMARY}")
            lines.append(f"Cập nhật file tổng kết: {_format_local_time(mtime)}")
        except Exception:
            pass
    return "\n".join(lines)


async def cmd_guitnstatus(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "guitnstatus"):
        return
    await send_telegram_long(context.bot, update.effective_chat.id, summarize_guitn_worker_log())


async def cmd_stopall(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not await command_is_for_this_bot(update, context, "stopall"):
        return
    chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
    if not command_has_bot_mention(update, "stopall") and chat_type in {"group", "supergroup"} and not is_primary_bot_instance():
        return
    result = stop_all_userbot_jobs(str(update.effective_user.id), stop_workers=True)
    await send_telegram_long(context.bot, update.effective_chat.id, result["summary"])
    await send_admin_management_report(
        context.bot,
        "[STOPALL]\n" + result["summary"],
        source_chat_id=update.effective_chat.id,
    )


async def cmd_stop_bot(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not command_has_bot_mention(update, "stop"):
        chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
        if chat_type == "private" or is_primary_bot_instance():
            await update.message.reply_text("Cú pháp: /stop@ten_bot\nVí dụ: /stop@baokibcr_bot")
        return
    if not await command_is_for_this_bot(update, context, "stop"):
        return
    try:
        me = await context.bot.get_me()
        bot_username = str(getattr(me, "username", "") or "").strip()
    except Exception:
        bot_username = ""
    slot = current_bot_slot()
    text = (
        "🛑 Đã nhận lệnh dừng bot này.\n"
        f"Bot: @{bot_username or 'unknown'}\n"
        f"Slot: {slot}\n"
        "Bot sẽ tự tắt sau vài giây. Các bot khác không bị ảnh hưởng."
    )
    await update.message.reply_text(text)
    try:
        await send_admin_management_report(
            context.bot,
            "[STOP BOT]\n" + text + f"\nAdmin: {update.effective_user.id}",
            source_chat_id=update.effective_chat.id,
        )
    except Exception:
        pass
    threading.Thread(target=_delayed_stop_current_bot, args=("stop_command", 1.5), daemon=True).start()


async def cmd_keoall(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not await userbot_operation_command_is_for_this_bot(update, context, "keoall"):
        return
    opts = parse_keoall_options(update)
    target_link = opts["target_link"]
    if not target_link:
        await update.message.reply_text(
            "⚠️ Thiếu link nhóm.\n"
            "Cú pháp: reply file rồi gửi /keoall <link_nhóm>\n"
            "Chạy thật: /keoall run <link_nhóm> limit=20 delay=60-180"
        )
        return
    rows, source_path, file_notes, stats = await load_keoall_reply_file_recipients(update, context)
    if not rows:
        await update.message.reply_text(
            "❌ Chưa đọc được danh sách cần kéo.\n"
            "Bạn cần reply file /uidnow hoặc gửi file kèm caption /keoall <link_nhóm>.\n"
            "File hợp lệ có dòng dạng: uid | @user hoặc uid | user."
        )
        return
    rows = annotate_campaign_recipients(rows, source_path)
    consent_rows, missing_consent_rows = split_campaign_consent_rows(rows)
    target_allowed = campaign_group_allowed(target_link)
    if opts["run_real"] and not target_allowed:
        await update.message.reply_text(
            "🛑 Chặn chạy thật /keoall: nhóm đích chưa nằm trong allow-list campaign.\n"
            "Hãy dry-run trước hoặc cấu hình campaign_allowed_group_ids/campaign_start_group_link đúng nhóm cần dùng."
        )
        return
    if opts["run_real"] and CAMPAIGN_CONSENT_REQUIRED and missing_consent_rows:
        await update.message.reply_text(
            "🛑 Chặn chạy thật /keoall: còn target thiếu consent/opt-in.\n"
            f"Đủ consent: {len(consent_rows)} | Thiếu consent: {len(missing_consent_rows)}\n"
            "File target phải có cột consent=opt_in hoặc target đã tự /start trước."
        )
        return
    root = get_guitn_accounts_root()
    profiles = list_userbot_profiles(limit=1, root=root)
    session_count = count_userbot_api_sessions(profiles, root=root)
    valid_count = int(stats.get("valid_count") or len(rows))
    suggested_limit = min(valid_count, max(1, session_count or len(profiles) or 1) * 20)
    if suggested_limit <= 0 and valid_count > 0:
        suggested_limit = min(valid_count, 20)
    requested_limit = int(opts.get("limit") or 0)
    final_limit = requested_limit if requested_limit > 0 else suggested_limit
    if opts["run_real"] and final_limit <= 0:
        await update.message.reply_text("❌ Không có dòng hợp lệ để kéo vào nhóm.")
        return
    recipients_file = write_keoall_recipients(rows, str(update.effective_user.id))
    detail = {
        "job_type": "invite_to_group",
        "profiles_root": root,
        "profile_count": len(profiles),
        "max_accounts_per_job": 1,
        "profiles": profiles,
        "session_count": session_count,
        "target_group_link": target_link,
        "recipients_file": recipients_file,
        "source_users_file": source_path,
        "recipient_count": len(rows),
        "sendable_count": len(consent_rows),
        "missing_consent_count": len(missing_consent_rows),
        "consent_required": CAMPAIGN_CONSENT_REQUIRED,
        "target_group_allowed": target_allowed,
        "target_group_allow_tokens": sorted(campaign_group_allow_tokens()),
        "input_stats": stats,
        "limit": final_limit,
        "suggested_limit": suggested_limit,
        "delay_min_sec": opts["delay_min"],
        "delay_max_sec": opts["delay_max"],
        "delay_label": opts["delay_label"],
        "mode": "run" if opts["run_real"] else "dry-run",
        "bot_slot": current_bot_slot(),
    }
    detail = attach_userbot_report_origin(detail, update)
    job_id = queue_userbot_scan_job(
        "keoall",
        str(getattr(update.effective_chat, "id", "") or ""),
        str(getattr(update.effective_chat, "title", "") or "KEOALL"),
        str(update.effective_user.id),
        not opts["run_real"],
        detail,
    )
    sample = "\n".join(contact_import_line(row, "tt") for row in rows[:15])
    note_text = ""
    if file_notes:
        note_text = "\nGhi chú đọc file:\n" + "\n".join(f"- {n}" for n in file_notes[:6])
    run_hint = f"/keoall run {target_link} limit={suggested_limit or final_limit or 20} delay={opts['delay_label']}"
    report_text = (
        ("✅ ĐÃ TẠO JOB DRY-RUN /keoall\n" if not opts["run_real"] else "✅ ĐÃ TẠO JOB CHẠY THẬT /keoall\n")
        + f"Job #{job_id}\n"
        + f"Nhóm đích: {target_link}\n"
        + f"Bot nhận lệnh: {current_bot_slot()}\n"
        + f"Userbot dùng để kéo: {len(profiles)} profile | session API: {session_count}/{len(profiles)}\n"
        + f"Dòng đọc được: {stats.get('total_lines', 0)}\n"
        + f"Hợp lệ: {valid_count} | Có @user: {stats.get('username_count', 0)} | Chỉ UID: {stats.get('uid_only_count', 0)}\n"
        + f"Đủ consent để chạy thật: {len(consent_rows)} | Thiếu consent: {len(missing_consent_rows)}\n"
        + f"Nhóm trong allow-list: {'CÓ' if target_allowed else 'KHÔNG'}\n"
        + f"Dòng trùng đã bỏ: {stats.get('duplicate_count', 0)} | Dòng lỗi: {stats.get('invalid_count', 0)}\n"
        + f"Giới hạn gợi ý: {suggested_limit}\n"
        + f"Giới hạn job này: {final_limit}\n"
        + f"Delay mỗi lượt: {opts['delay_label']}\n"
        + f"File queue: {recipients_file}\n"
        + f"Folder kết quả: D:\\ZALO_MOVE\\userbot_scan_jobs\\keoall\\job_{job_id}_...\n"
        + (f"Lệnh chạy thật gợi ý:\n{run_hint}\n" if not opts["run_real"] else "")
        + f"{note_text}\n\n"
        + "Mẫu dòng hợp lệ:\n"
        + (sample or "N/A")
    )
    await send_telegram_long(context.bot, update.effective_chat.id, report_text)
    await send_admin_management_report(
        context.bot,
        "[KEOALL QUEUED]\n" + report_text,
        files=[recipients_file],
        caption=f"/keoall recipients: {len(rows)} dong",
        source_chat_id=update.effective_chat.id,
    )


def summarize_keoall_worker_log() -> str:
    last_summary = os.path.join(USERBOT_SCAN_JOBS_DIR, "keoall", "keoall_last_summary.txt")
    if os.path.exists(last_summary):
        try:
            with open(last_summary, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read().strip()
            return text[-3500:] if text else "Chưa có nội dung báo cáo /keoall."
        except Exception as exc:
            return f"Không đọc được báo cáo /keoall: {exc}"
    log_path = os.path.join(USERBOT_SCAN_JOBS_DIR, "userbot_import_contacts_worker.log")
    if not os.path.exists(log_path):
        return "Chưa có log worker /keoall."
    try:
        with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [line for line in f.read().splitlines() if "keoall" in line.lower()]
        return "\n".join(lines[-30:]) or "Chưa thấy job /keoall trong log worker."
    except Exception as exc:
        return f"Không đọc được log worker: {exc}"


async def cmd_keoallstatus(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    await send_telegram_long(context.bot, update.effective_chat.id, summarize_keoall_worker_log())


async def cmd_guitn_text_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    msg = update.message
    raw = str(getattr(msg, "text", "") or getattr(msg, "caption", "") or "")
    if not re.match(r"(?is)^\s*/guitn(?:@\w+)?(?:\s|$)", raw):
        return
    if not await command_is_for_this_bot(update, context, "guitn"):
        return
    await cmd_guitn(update, context)


async def cmd_guitnlai_text_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    msg = update.message
    raw = str(getattr(msg, "text", "") or getattr(msg, "caption", "") or "")
    if not re.match(r"(?is)^\s*/guitnlai(?:@\w+)?(?:\s|$)", raw):
        return
    if not await command_is_for_this_bot(update, context, "guitnlai"):
        return
    await cmd_guitnlai(update, context)


async def handle_scan_file_caption(update: Update, context: ContextTypes.DEFAULT_TYPE):
    msg = update.message
    if not msg or not msg.document:
        return
    caption = (msg.caption or "").strip().lower()
    if caption.startswith("/phat500k"):
        if not is_supreme_admin(update.effective_user.id):
            return
        await cmd_phat500k(update, context)
        return
    if caption.startswith("/guitnlai"):
        if not is_admin(update.effective_user.id):
            return
        if not await command_is_for_this_bot(update, context, "guitnlai"):
            return
        await cmd_guitnlai(update, context)
        return
    if caption.startswith("/guitn"):
        if not is_admin(update.effective_user.id):
            return
        if not await command_is_for_this_bot(update, context, "guitn"):
            return
        await cmd_guitn(update, context)
        return
    if caption.startswith("/keoall"):
        if not is_supreme_admin(update.effective_user.id):
            return
        if not await command_is_for_this_bot(update, context, "keoall"):
            return
        await cmd_keoall(update, context)
        return
    if not caption.startswith("/quetfile"):
        return
    if not is_supreme_admin(update.effective_user.id):
        return
    await cmd_quetfile(update, context)

async def cmd_code(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        client_id = int(context.args[0])
        code_vip = context.args[1]
        success_message = f"🎉 **MÃ QUÀ TẶNG CODE VIP ĐÃ ĐƯỢC PHÊ DUYỆT!**\n\n🎁 Mã quà tặng của bạn: `{code_vip}`\n🔗 Nơi nhập quà: {LINK_NHAP_CODE}\n\nChúc quý khách thắng lớn thắng đậm!"
        await context.bot.send_message(chat_id=client_id, text=success_message, parse_mode="Markdown")
        await update.message.reply_text(f"✅ Đã cấp code cho ID: {client_id}")
        asyncio.create_task(bi_mat_quang_cao(context, client_id))
    except: pass

async def cmd_tatmoauto(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global CHUC_NANG_AUTO_DUYET
    if not is_admin(update.effective_user.id): return
    CHUC_NANG_AUTO_DUYET = not CHUC_NANG_AUTO_DUYET
    await update.message.reply_text(f"✅ Auto duyệt lệnh tự động: {CHUC_NANG_AUTO_DUYET}")

async def cmd_tatmochuong(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global THONG_BAO_START_ACTIVE
    if not is_admin(update.effective_user.id): return
    THONG_BAO_START_ACTIVE = not THONG_BAO_START_ACTIVE
    await update.message.reply_text(f"✅ Báo chuông khách mới: {THONG_BAO_START_ACTIVE}")

async def cmd_tatmoanh(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global HIEN_THI_ANH_DONG
    if not is_admin(update.effective_user.id): return
    HIEN_THI_ANH_DONG = not HIEN_THI_ANH_DONG
    await update.message.reply_text(f"✅ Hiển thị ảnh GIF: {HIEN_THI_ANH_DONG}")


def get_command_tail(update: Update, command_name: str) -> str:
    msg = update.message
    if not msg:
        return ""
    full_text = str(getattr(msg, "text", "") or getattr(msg, "caption", "") or "")
    if not full_text:
        return ""
    pattern = re.compile(rf"^\s*/{re.escape(command_name)}(?:@\w+)?(?:\s+|$)", flags=re.IGNORECASE)
    return pattern.sub("", full_text, count=1).strip()


def get_reply_payload(update: Update) -> str:
    msg = update.message
    if not msg or not msg.reply_to_message:
        return ""
    return str(
        getattr(msg.reply_to_message, "text", "")
        or getattr(msg.reply_to_message, "caption", "")
        or ""
    ).strip()


def get_full_command_payload(update: Update, command_name: str) -> str:
    reply_text = get_reply_payload(update)
    if reply_text:
        return reply_text
    return get_command_tail(update, command_name)


def split_command_leading_and_rest(update: Update, command_name: str, leading_count: int):
    tail = get_command_tail(update, command_name)
    if not tail:
        return [], get_reply_payload(update)
    parts = tail.split(None, leading_count)
    if len(parts) < leading_count:
        return parts, get_reply_payload(update)
    leading = parts[:leading_count]
    rest = parts[leading_count].strip() if len(parts) > leading_count else get_reply_payload(update)
    return leading, rest


CONTENT_EDIT_STEPS = [
    {
        "code": "register_prompt",
        "label": "khi khách bấm ĐĂNG KÝ TÀI KHOẢN MỚI",
        "text_type": "template",
        "template_key": "register_prompt",
        "media_config": "gif_dang_ky",
        "media_global": "GIF_DANG_KY",
    },
    {
        "code": "old_member_phone_prompt",
        "label": "khi khách bấm CÓ TÀI KHOẢN - NHẬN CODE",
        "text_type": "template",
        "template_key": "old_member_phone_prompt",
    },
    {
        "code": "ask_phone_after_tk",
        "label": "khi khách đã nhập tài khoản game và bot hỏi SĐT",
        "text_type": "template",
        "template_key": "ask_phone_after_tk",
    },
    {
        "code": "join_required",
        "label": "khi bot yêu cầu khách tham gia kênh/nhóm",
        "text_type": "template",
        "template_key": "join_required",
        "media_config": "gif_ep_join",
        "media_global": "GIF_EP_JOIN",
    },
    {
        "code": "moc_intro",
        "label": "khi khách xác nhận join xong và bot hiện bảng mốc",
        "text_type": "template",
        "template_key": "moc_intro",
        "media_config": "gif_moc_nap",
        "media_global": "GIF_MOC_NAP",
    },
    {
        "code": "choose_reward_mode",
        "label": "khi khách chọn mốc nạp xong và bot hỏi cách nhận thưởng",
        "text_type": "template",
        "template_key": "choose_reward_mode",
    },
    {
        "code": "auto_review_waiting",
        "label": "khi khách bấm hệ thống trả code tự động",
        "text_type": "template",
        "template_key": "auto_review_waiting",
        "media_config": "gif_cho_duyet",
        "media_global": "GIF_CHO_DUYET",
    },
    {
        "code": "promo_ad",
        "label": "tin quảng cáo nhắc lại sau khi admin cấp code",
        "text_type": "config",
        "text_config": "telegram_qc_text",
        "text_global": "QUANG_CAO_TEXT",
        "media_config": "gif_quang_cao",
        "media_global": "GIF_QUANG_CAO",
    },
]


def content_step_by_code(code: str) -> dict | None:
    c = str(code or "").strip()
    for step in CONTENT_EDIT_STEPS:
        if step["code"] == c or step.get("template_key") == c:
            return step
    if c == "welcome":
        return {
            "code": "welcome",
            "label": "lời chào khi khách bấm /start",
            "text_type": "config",
            "text_config": "telegram_welcome_text",
            "text_global": "LOI_CHAO_MAC_DINH",
            "media_config": "gif_chao_hoi",
            "media_global": "GIF_CHAO_HOI",
        }
    return None


def apply_content_edit_step(step: dict, text: str = "", media_ref: str = "") -> tuple[bool, list[str]]:
    changed = []
    content = str(text or "").strip()
    media = str(media_ref or "").strip()
    if content:
        if step.get("text_type") == "template":
            if set_tele_reply_template(step.get("template_key", ""), content):
                changed.append("nội dung")
        elif step.get("text_type") == "config":
            globals()[step["text_global"]] = content
            set_config(step["text_config"], content)
            changed.append("nội dung")
    if media:
        if step.get("media_config") and step.get("media_global"):
            globals()[step["media_global"]] = media
            set_config(step["media_config"], media)
            changed.append("ảnh/GIF")
        elif step.get("template_key"):
            if set_tele_reply_media(step["template_key"], media):
                changed.append("ảnh/GIF")
    return bool(changed), changed


async def send_template_media_if_any(bot, chat_id, key: str) -> None:
    await send_config_media(bot, chat_id, get_tele_reply_media(key))


async def ask_next_content_edit(update: Update, next_index: int) -> None:
    uid = update.effective_user.id
    if next_index >= len(CONTENT_EDIT_STEPS):
        ADMIN_CONTENT_WIZARD.pop(uid, None)
        await update.message.reply_text("✅ Đã hết danh sách nội dung cần hỏi. Phiên sửa liên tục đã kết thúc.")
        return
    step = CONTENT_EDIT_STEPS[next_index]
    ADMIN_CONTENT_WIZARD[uid] = {"mode": "confirm_next", "next_index": next_index}
    await update.message.reply_text(
        "✅ Đã lưu bước vừa sửa.\n\n"
        f"Bạn có muốn sửa tiếp nội dung: {step['label']} không?\n"
        "Gõ /yes để sửa tiếp, hoặc /no để dừng."
    )


async def cmd_yes(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    state = ADMIN_CONTENT_WIZARD.get(update.effective_user.id)
    if not state:
        await update.message.reply_text("Hiện không có phiên sửa nội dung nào. Bắt đầu bằng cách reply mẫu rồi gõ /suachao.")
        return
    idx = int(state.get("next_index", 0))
    if idx >= len(CONTENT_EDIT_STEPS):
        ADMIN_CONTENT_WIZARD.pop(update.effective_user.id, None)
        await update.message.reply_text("✅ Đã hết danh sách nội dung cần sửa.")
        return
    step = CONTENT_EDIT_STEPS[idx]
    ADMIN_CONTENT_WIZARD[update.effective_user.id] = {"mode": "awaiting_content", "next_index": idx}
    await update.message.reply_text(
        f"📌 Gửi mẫu mới cho bước: {step['label']}.\n\n"
        "Cách làm:\n"
        "1) Gửi 1 tin nhắn có nội dung, có thể kèm ảnh/GIF/video.\n"
        "2) Reply tin đó bằng lệnh /sua.\n\n"
        "Nếu chỉ gửi ảnh/GIF không có chữ, bot chỉ đổi media và giữ nội dung cũ."
    )


async def cmd_no(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    ADMIN_CONTENT_WIZARD.pop(update.effective_user.id, None)
    await update.message.reply_text("✅ Đã dừng phiên sửa liên tục. Những nội dung đã lưu trước đó vẫn được giữ nguyên.")


async def cmd_sua(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    state = ADMIN_CONTENT_WIZARD.get(update.effective_user.id)
    if not state or state.get("mode") != "awaiting_content":
        await update.message.reply_text("Chưa có bước nào đang chờ /sua. Hãy dùng /suachao trước, sau đó /yes để sửa tiếp từng bước.")
        return
    idx = int(state.get("next_index", 0))
    if idx >= len(CONTENT_EDIT_STEPS):
        ADMIN_CONTENT_WIZARD.pop(update.effective_user.id, None)
        await update.message.reply_text("✅ Đã hết danh sách nội dung cần sửa.")
        return
    step = CONTENT_EDIT_STEPS[idx]
    text, media = text_and_media_from_update(update, "sua")
    ok, changed = apply_content_edit_step(step, text, media)
    if not ok:
        await update.message.reply_text(
            "❌ Chưa thấy nội dung hoặc ảnh/GIF để lưu.\n"
            "Hãy gửi mẫu mới, rồi reply mẫu đó bằng /sua."
        )
        return
    await update.message.reply_text(f"✅ Đã lưu {', '.join(changed)} cho bước: {step['label']}.")
    await ask_next_content_edit(update, idx + 1)


async def cmd_suachao(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    payload, media = text_and_media_from_update(update, "suachao")
    if not payload and not media:
        await update.message.reply_text("Cú pháp: /suachao <noi_dung> hoặc reply tin mẫu có nội dung/ảnh/GIF rồi gõ /suachao")
        return
    ok, changed = apply_content_edit_step(content_step_by_code("welcome"), payload, media)
    if ok:
        await update.message.reply_text(f"✅ Đã sửa lời chào ({', '.join(changed)}), đã giữ nguyên xuống dòng/khoảng cách.")
        await ask_next_content_edit(update, 0)

async def cmd_suaquangcao(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    payload, media = text_and_media_from_update(update, "suaquangcao")
    if not payload and not media:
        await update.message.reply_text("Cú pháp: /suaquangcao <noi_dung> hoặc reply tin mẫu có nội dung/ảnh/GIF rồi gõ /suaquangcao")
        return
    ok, changed = apply_content_edit_step(content_step_by_code("promo_ad"), payload, media)
    if ok:
        await update.message.reply_text(f"✅ Đã sửa quảng cáo ({', '.join(changed)}), đã giữ nguyên xuống dòng/khoảng cách.")


async def cmd_suarepsdt(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    payload, media = text_and_media_from_update(update, "suarepsdt")
    if not payload and not media:
        await update.message.reply_text("Cú pháp: /suarepsdt <noi_dung> hoặc reply mẫu có nội dung/ảnh/GIF rồi gõ /suarepsdt")
        return
    ok, changed = apply_content_edit_step(content_step_by_code("ask_phone_after_tk"), payload, media)
    if ok:
        await update.message.reply_text(f"✅ Đã cập nhật mẫu hỏi SĐT khách ({', '.join(changed)}).")


async def cmd_suarep(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    leading, content = split_command_leading_and_rest(update, "suarep", 1)
    if len(leading) < 1:
        await update.message.reply_text("Cú pháp: /suarep <ma_mau> <noi_dung> hoặc reply mẫu có nội dung/ảnh/GIF rồi gõ /suarep <ma_mau>")
        return
    key = leading[0].strip()
    if key not in TELE_REPLY_TEMPLATES_DEFAULT:
        await update.message.reply_text(
            "❌ Mã mẫu không hợp lệ.\n"
            "Dùng /xemrep để xem danh sách mã có thể sửa."
        )
        return
    media = media_ref_from_message(getattr(update.message, "reply_to_message", None) if update.message else None)
    if not content and not media:
        await update.message.reply_text("❌ Nội dung và ảnh/GIF đều trống.")
        return
    step = content_step_by_code(key)
    if step:
        ok, changed = apply_content_edit_step(step, content, media)
    else:
        ok = False
        changed = []
        if content:
            ok = set_tele_reply_template(key, content)
            if ok:
                changed.append("nội dung")
        if media:
            ok = set_tele_reply_media(key, media) or ok
            changed.append("ảnh/GIF")
    if ok:
        await update.message.reply_text(f"✅ Đã cập nhật mẫu `{key}` ({', '.join(changed)}).")


async def cmd_xemrep(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    lines = [
        "📚 MẪU TIN KHÁCH (sửa được nội dung + ảnh/GIF/video):",
        "Cách sửa nhanh: gửi mẫu mới, reply mẫu đó bằng /suarep <mã>.",
        "Cách sửa liên tục: gửi mẫu chào, reply /suachao, rồi dùng /yes hoặc /no.",
        "",
        "- register_prompt: khách bấm ĐĂNG KÝ TÀI KHOẢN MỚI.",
        "- old_member_phone_prompt: khách bấm CÓ TÀI KHOẢN - NHẬN CODE.",
        "- old_member_phone_saved: bot báo đã ghi nhận SĐT khách cũ.",
        "- ask_phone_after_tk: bot hỏi SĐT sau khi khách nhập tài khoản game. Alias: /suarepsdt.",
        "- join_required: bot yêu cầu khách tham gia kênh/nhóm.",
        "- join_not_verified: khách bấm xác nhận join nhưng bot chưa xác nhận được.",
        "- moc_intro: bảng mốc khuyến mãi sau khi khách join xong.",
        "- choose_reward_mode: khách chọn mốc xong, bot hỏi liên hệ admin hay auto.",
        "- auto_review_disabled: auto duyệt đang tạm khóa.",
        "- auto_review_waiting: khách bấm hệ thống trả code tự động.",
        "- promo_other: ưu đãi khác.",
        "",
        "Lưu ý: nếu reply ảnh/GIF không có chữ, bot chỉ đổi media và giữ nội dung cũ.",
    ]
    await update.message.reply_text("\n".join(lines))

async def cmd_vip(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        uid = int(context.args[0])
        VIP_USERS.add(uid)
        save_user_state(uid)
        await update.message.reply_text("✅ Đã đặc cách VIP.")
    except: pass

async def cmd_huyvip(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        uid = int(context.args[0])
        VIP_USERS.discard(uid)
        save_user_state(uid)
        await update.message.reply_text("✅ Đã hủy đặc cách VIP.")
    except: pass


def media_arg_from_command(update: Update, context: ContextTypes.DEFAULT_TYPE, command_name: str) -> str:
    payload = get_command_tail(update, command_name).strip()
    if payload:
        return payload
    if context.args:
        return " ".join(context.args).strip()
    msg = update.message
    reply = getattr(msg, "reply_to_message", None) if msg else None
    media = media_ref_from_message(reply or msg)
    if media:
        return media
    return ""


def is_media_reference(value: str) -> bool:
    v = str(value or "").strip()
    if not v:
        return False
    if ":" in v and v.split(":", 1)[0].lower() in {"animation", "photo", "video", "document"}:
        return True
    if re.match(r"(?i)^https?://", v):
        return True
    return os.path.exists(v)


async def set_gif_value(update: Update, config_key: str, value: str, label: str, command_name: str):
    if not value:
        await update.message.reply_text(f"⚠️ Cú pháp: /{command_name} <url_gif_hoặc_path>")
        return False
    if not is_media_reference(value):
        await update.message.reply_text("⚠️ GIF/media phải là link http(s) hoặc đường dẫn file local đang tồn tại.")
        return False
    set_config(config_key, value)
    await update.message.reply_text(f"✅ Đã lưu {label}: {value}")
    return True


async def cmd_anhchao(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global GIF_CHAO_HOI
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "anhchao")
    if await set_gif_value(update, "gif_chao_hoi", value, "GIF chào", "anhchao"):
        GIF_CHAO_HOI = value

async def cmd_anhdangky(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global GIF_DANG_KY
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "anhdangky")
    if await set_gif_value(update, "gif_dang_ky", value, "GIF đăng ký", "anhdangky"):
        GIF_DANG_KY = value

async def cmd_anhepjoin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global GIF_EP_JOIN
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "anhepjoin")
    if await set_gif_value(update, "gif_ep_join", value, "GIF ép join", "anhepjoin"):
        GIF_EP_JOIN = value

async def cmd_anhmocnap(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global GIF_MOC_NAP
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "anhmocnap")
    if await set_gif_value(update, "gif_moc_nap", value, "GIF mốc nạp", "anhmocnap"):
        GIF_MOC_NAP = value

async def cmd_anhchoyet(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global GIF_CHO_DUYET
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "anhchoyet")
    if await set_gif_value(update, "gif_cho_duyet", value, "GIF chờ duyệt", "anhchoyet"):
        GIF_CHO_DUYET = value

async def cmd_anhqc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global GIF_QUANG_CAO
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "anhqc")
    if await set_gif_value(update, "gif_quang_cao", value, "GIF quảng cáo", "anhqc"):
        GIF_QUANG_CAO = value


async def cmd_setgifall(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global GIF_CHAO_HOI, GIF_DANG_KY, GIF_EP_JOIN, GIF_MOC_NAP, GIF_CHO_DUYET, GIF_QUANG_CAO
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "setgifall")
    if not value:
        await update.message.reply_text("⚠️ Cú pháp: /setgifall <url_gif_hoặc_path>")
        return
    if not is_media_reference(value):
        await update.message.reply_text("⚠️ GIF/media phải là link http(s) hoặc đường dẫn file local đang tồn tại.")
        return
    GIF_CHAO_HOI = GIF_DANG_KY = GIF_EP_JOIN = GIF_MOC_NAP = GIF_CHO_DUYET = GIF_QUANG_CAO = value
    for key in ("gif_chao_hoi", "gif_dang_ky", "gif_ep_join", "gif_moc_nap", "gif_cho_duyet", "gif_quang_cao"):
        set_config(key, value)
    await update.message.reply_text("✅ Đã đổi cùng một GIF cho toàn bộ tương tác khách Telegram.")


async def cmd_setguitngif(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global USERBOT_GUITN_GIF
    if not is_admin(update.effective_user.id): return
    value = media_arg_from_command(update, context, "setguitngif")
    if not value:
        await update.message.reply_text("⚠️ Cú pháp: /setguitngif <url_gif_hoặc_path> hoặc /setguitngif off")
        return
    if value.strip().lower() in {"off", "tat", "tắt", "none", "clear", "0"}:
        USERBOT_GUITN_GIF = ""
        set_config("userbot_guitn_gif", "")
        await update.message.reply_text("✅ Đã tắt GIF gửi kèm /guitn userbot.")
        return
    if not is_media_reference(value):
        await update.message.reply_text("⚠️ GIF/media phải là link http(s) hoặc đường dẫn file local đang tồn tại.")
        return
    USERBOT_GUITN_GIF = value
    set_config("userbot_guitn_gif", value)
    await update.message.reply_text(f"✅ Đã lưu GIF gửi kèm /guitn userbot:\n{value}")


async def cmd_showgifs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    lines = [
        "GIF hiện tại:",
        f"- chao: {GIF_CHAO_HOI}",
        f"- dangky: {GIF_DANG_KY}",
        f"- epjoin: {GIF_EP_JOIN}",
        f"- mocnap: {GIF_MOC_NAP}",
        f"- choduyet: {GIF_CHO_DUYET}",
        f"- qc: {GIF_QUANG_CAO}",
        f"- guitn_userbot: {USERBOT_GUITN_GIF or 'OFF'}",
    ]
    await update.message.reply_text("\n".join(lines))

async def cmd_tenadmin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global ADMIN_USERNAME
    if update.effective_user.id != ADMIN_ID: return
    try: ADMIN_USERNAME = context.args[0]; await update.message.reply_text(f"✅ Đã đổi tên liên hệ: {ADMIN_USERNAME}")
    except: pass

async def cmd_lamlai(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    try:
        tid = int(context.args[0])
        if tid in USER_DATA:
            USER_DATA[tid]['step'] = 'CHO_CHON_HINH_THUC'
            save_user_state(tid)
            await update.message.reply_text("✅ Đã reset khách.")
    except: pass

async def cmd_ping(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id): return
    await update.message.reply_text("🚀 Bot hoạt động mượt mà ổn định!")

async def cmd_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        await update.message.reply_text("Bạn là khách hàng. Vui lòng bấm /start để vào quy trình nhận hỗ trợ.")
        return
    await update.message.reply_text(
        "📘 BANG LENH NHANH (ADMIN TELE)\n"
        "1) He thong: /ping | /kiemtra | /baotri on|off | /tatmoauto on|off\n"
        "2) Quan tri admin: /themadmin <id> | /huyadmin <id> | /tenadmin <@user>\n"
        "3) Khach hang: /xemkhach | /check <uid> | /lamlai <uid>\n"
        "4) Gui tin: /guikhach <uid> <noi_dung> | /guitatca <noi_dung> | /loaloaall <noi_dung> <giay>\n"
        "5) Sua mau rep khach: reply mau co anh/GIF roi /suachao | /yes | /sua | /no | /suarep <ma> | /xemrep\n"
        "6) Quet data SDT/UID/user: /quetnhanh <text> | /quetfile | /quetfileepath <path> | /quetnow <danh_sach_so> | /quetuser <danh_sach_so>\n"
        "7) Userbot import SDT -> UID/user: reply file roi /phat500k run hoac /quetallnow run; ca 2 cung tra UID | USER va file /guitn.\n"
        "8) Keo nhom/gui tin userbot: reply file /uidnow roi /keoall <link_nhom> | /keoall run <link_nhom> limit=20 | /guitn <noi_dung> 60-180 roi chon /run40|/run52 | /guitnlai [noi_dung] [60-180] de gui lai file loi\n"
        "8b) Theo doi tien trinh: /quetstatus xem quet SDT->UID/user | /uidstatus|/phat500kstatus|/importstatus la alias cu | /guitnstatus xem gui tin\n"
        "8c) File quet moi nhat: /usernow | /uidnow | /ttnow | /sdtnow | /quetallnow. File chinh de doc nhanh la UID | USER; file day du dung cho /guitn.\n"
        "8d) Dung khan cap: /stopall dung moi job userbot | /stop@ten_bot dung rieng bot duoc tag\n"
        "9) Supper group: /setsuppergroup hoac /setsupergroup on|off, go trong nhom de set nhan canh bao khach moi\n"
        "10) Admin quan ly: /adminquanly hoac /setadminquanly on|off, go trong nhom de set nhan report quet file/SĐT/UID/user va /guitn\n"
        "10) Diem danh: /diemdanh gui lai tin diem danh ve nhom @datagk88\n"
        "11) AI provider: /ai | /openai | /alo | /git | /github | /meta <cau_hoi>\n"
        "12) Zalo bridge: /zreply <zalo_uid> <noi_dung> | /adminchat on|off\n"
        "13) Bao tri token: /newtoken [label] | /tokens | /audit\n\n"
        "Kien truc: bot father tuong tac khach/admin; 52 profile userbot chia 5 shard de import SDT -> UID/user.\n"
        "Can huong dan day du: /trogiup"
    )

async def cmd_trogiup(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        await update.message.reply_text("Bạn là khách hàng. Vui lòng bấm /start để vào quy trình nhận hỗ trợ.")
        return
    help_parts = [
        (
            "📚👑 TRO GIUP DAY DU - ADMIN TELEGRAM (P1/3)\n"
            "🧭 Khoi dong va ho tro khach:\n"
            "- /start | /nhancode88k: Khoi dong quy trinh tuong tac khach.\n"
            "- /hotro: Gui link lien he admin.\n"
            "- /help: Bang lenh nhanh.\n"
            "- /trogiup: Bang huong dan day du.\n"
            "- /myid: Xem ID cua ban + ADMIN_ID hien tai.\n"
            "- /setsuppergroup | /setsupergroup on|off: Bat/tat nhan canh bao khach moi o suppergroup; go trong nhom de set chat id.\n"
            "- /adminquanly | /setadminquanly on|off: Bat/tat nhan report van hanh; go trong nhom quan ly de set chat id.\n"
            "- /diemdanh: Admin toi cao goi botfather gui lai diem danh ve @datagk88.\n"
            "\n"
            "⚙️ Dieu khien he thong:\n"
            "- /ping: Kiem tra bot con song.\n"
            "- /kiemtra: Xem trang thai tong quan bot.\n"
            "- /baotri: Bat/tat che do bao tri.\n"
            "- /tatmoauto: Bat/tat auto duyet.\n"
            "- /tatmochuong: Bat/tat thong bao khach moi.\n"
            "- /tatmoanh: Bat/tat GIF/anh dong trong quy trinh.\n"
            "- /adminchat on|off: Bat/tat day log chat ve admin.\n"
            "\n"
            "👮 Quan tri admin:\n"
            "- /themadmin <id>: Them admin phu (chi admin chinh).\n"
            "- /huyadmin <id>: Go admin phu.\n"
            "- /tenadmin <@username>: Dat ten admin hien thi cho khach.\n"
            "- /kenh <chat_id_hoac_@kenh>: Dat kenh muc tieu de check tham gia.\n"
            "\n"
            "👥 Quan ly khach:\n"
            "- /xemkhach <uid>: Xem cache 1 khach.\n"
            "- /check <uid>: Tra user theo UID trong DB.\n"
            "- /lamlai <uid>: Reset state khach ve buoc dau.\n"
            "- /chan <uid> | /mochan <uid> | /mochanhet: Quan ly chan.\n"
            "- /vip <uid> | /huyvip <uid>: Gan/go VIP.\n"
            "- /code <uid> <ma_code>: Cap ma code VIP thu cong cho khach."
        ),
        (
            "📚👑 TRO GIUP DAY DU - ADMIN TELEGRAM (P2/3)\n"
            "📢 Gui tin, broadcast, dieu huong:\n"
            "- /guikhach <uid> <noi_dung>: Gui tin cho 1 khach, giu nguyen xuong dong.\n"
            "- /guitatca <noi_dung>: Gui tin den toan bo cache khach, co the reply tin mau.\n"
            "- /loaloa <chat_id> <noi_dung>: Gui tin den 1 dich tuy y, giu nguyen format.\n"
            "- /loaloaall <noi_dung> <giay>: Gui lan luot theo delay.\n"
            "- /zreply <zalo_uid> <noi_dung>: Tra loi nguoc ve Zalo.\n"
            "- /chuyentrong <uid> <url>: Day link rieng vao khach.\n"
            "- /link web <url>: Doi link dang ky.\n"
            "- /link code <url>: Doi link nhap code.\n"
            "- /huong <tu_khoa> <url>: Them dieu huong tu khoa.\n"
            "\n"
            "🖼️ Noi dung/mac dinh media:\n"
            "- /suachao <noi_dung>: Doi loi chao mac dinh, giu nguyen xuong dong/khoang cach; co the reply tin mau roi go /suachao.\n"
            "- /suaquangcao <noi_dung>: Doi noi dung quang cao, giu nguyen xuong dong/khoang cach.\n"
            "- /suarepsdt <noi_dung>: Doi mau tin nhan bot hoi SĐT khach (khong gioi han, giu xuong dong).\n"
            "- /suarep <ma> <noi_dung>: Doi tung mau tin khach theo ma, /xemrep de xem danh sach ma.\n"
            "- Co the gui 1 tin nhan co noi dung + anh/GIF/video, roi reply bang /suachao, /suaquangcao, /suarepsdt hoac /suarep <ma> de luu ca noi dung va media.\n"
            "- Sau /suachao, bot hoi co sua tiep khong: /yes de tiep tuc, /no de dung. Khi bot hoi buoc tiep theo, gui mau moi roi reply bang /sua.\n"
            "- /xemrep: Xem danh sach ma mau tin dang sua duoc.\n"
            "- /anhchao <url>: Doi GIF chao.\n"
            "- /anhdangky <url>: Doi GIF buoc dang ky.\n"
            "- /anhepjoin <url>: Doi GIF buoc ep join.\n"
            "- /anhmocnap <url>: Doi GIF moc nap.\n"
            "- /anhchoyet <url>: Doi GIF cho duyet.\n"
            "- /anhqc <url>: Doi GIF quang cao.\n"
            "- /setgifall <url>: Doi cung mot GIF cho toan bo tuong tac khach Telegram.\n"
            "- /setguitngif <url|off>: Gan/tat GIF gui kem khi /guitn userbot.\n"
            "- /showgifs: Xem cau hinh GIF hien tai.\n"
            "\n"
            "🧹 Du lieu:\n"
            "- /xoarac: Xoa toan bo cache USER_DATA."
        ),
        (
            "📚👑 TRO GIUP DAY DU - ADMIN TELEGRAM (P3/3)\n"
            "🔎 Quet data SDT + map user/uid:\n"
            "- /quetnhanh <text>: Quet SDT tu text/reply.\n"
            "- /quetfile: Reply file de doc va normalize SDT/UID/user. Ho tro txt/csv/xlsx/xls/docx/pdf/json/html/zip/db...\n"
            "- /quetfileepath <path>: Quet file local theo duong dan, cung ho tro nhieu dinh dang.\n"
            "- /quetfilepath <path>: Alias cua /quetfileepath.\n"
            "- /quetnow <danh_sach_so>: Quet nhanh nhieu so trong 1 tin.\n"
            "- /quetuser <danh_sach_so>: Quet nhieu SDT, tu luu danh ba bot, tra map user|uid|link neu da co.\n"
            "- Tat ca lenh quet deu bao: 'Da nhan, ADMIN vui long cho em giay lat.'\n"
            "- Lenh check/quet toi cao chi nhan admin gifhub2708/ADMIN_ID.\n"
            "- Ket qua gom: tong so, so moi, map user/uid/link + file CSV.\n"
            "\n"
            "👥 Userbot group/cache:\n"
            "- /check!: Admin chinh ping tung userbot API, bao API_OK/PING_ADMIN_OK/PeerFlood.\n"
            "- /check! me: Chi ping vao Saved Messages tung userbot, khong gui ve admin.\n"
            "- /checkngay: Tra UID + ten thanh vien trong group tu cache va gui file CSV.\n"
            "- /ALL: Tag tat ca thanh vien trong group theo cache userbot + nguoi tung nhan.\n"
            "- /phat500k dry-run: Reply file SDT de xem 5 shard/52 profile se chia viec, chua chay that.\n"
            "- /phat500k run: Reply file SDT de import contact; cung pipeline voi /quetallnow run, bot tu gui file UID | USER va file day du cho /guitn.\n"
            "- /usernow: Gui file day du toan bo @user hop le cua lan /phat500k run gan nhat.\n"
            "- /uidnow: Gui file day du uid + @user, chi tinh dong co user.\n"
            "- /ttnow: Gui file day du sdt/uid/user/ten/account, chi tinh dong co user.\n"
            "- /sdtnow: Gui file day du sdt + @user, chi tinh dong co user.\n"
            "- /quetallnow: Reply file SDT de preview/kiem tra file; chua co chu run thi khong queue job.\n"
            "- /quetallnow run: Reply file SDT de import contact; ket qua giong /phat500k run. File chinh UID | USER, bo dong khong co ca UID lan USER; file day du dung cho /guitn.\n"
            "- /quetstatus: Xem tien trinh quet/import SDT -> UID/user, gom tong job va tung acc userbot.\n"
            "- /uidstatus | /phat500kstatus | /importstatus: Alias cu cua trang thai import, van dung nhu nhau.\n"
            "- /keoall <link_nhom>: Reply file uid|user de dry-run keo nguoi vao nhom, chua chay that.\n"
            "- /keoall run <link_nhom> limit=20 delay=60-180: Keo that bang 1 userbot, bao cao tieng Viet va luu rieng tung folder job.\n"
            "- /keoallstatus: Xem bao cao /keoall gan nhat.\n"
            "- /guitn <noi_dung> [60s|60-180]: Luu plan gui tin theo file quet gan nhat hoac reply file /uidnow, CHUA chay ngay.\n"
            "- /guitnlai [noi_dung] [60s|60-180]: Lay file khach gui chua thanh cong gan nhat, hoac file userbot_send_unsent reply vao, roi luu plan gui lai.\n"
            "- Sau /guitn, chon lo userbot: /run40 lo 40 acc | /run52 lo 52 acc.\n"
            "- Vi du: reply file uid_now roi go /guitn Xin chao 60-180, bot hoi lo, admin go /run52 de moi acc gui cach nhau random 60 den 180 giay.\n"
            "- Vi du retry: reply file userbot_send_unsent roi go /guitnlai 60-180, bot tu lay lai noi dung job loi gan nhat neu ban khong nhap noi dung moi.\n"
            "- Co the chon rieng UID/user: /guitn <noi_dung> | <uid1 @user2 ...> 60s.\n"
            "- Trong group bat buoc tag dung bot cho lenh scan/userbot: /quetfile@ten_bot, /quetallnow@ten_bot, /guitn@ten_bot, /run40@ten_bot hoac /run52@ten_bot. Nhan rieng bot thi khong can tag.\n"
            "- /guitnstatus: Xem tien trinh gui tin, da xu ly/da gui/loi theo tung userbot.\n"
            "- /stopall: Dung tat ca job userbot dang cho/dang chay. Neu co job dang chay, worker userbot se bi tat de cat job ngay.\n"
            "- /stop@ten_bot: Dung rieng bot father duoc tag, vi du /stop@baokibcr_bot. Cac bot khac im lang.\n"
            "- Luu y: 52 profile trong D:\\ZALO_MOVE\\52acc_tong_hop_20260603 la userbot chia 5 shard, khong phai bot father.\n"
            "\n"
            "🤖 AI + quyen + token:\n"
            "- /ai <cau_hoi>: Goi AI theo moi truong dang active.\n"
            "- /openai <noi_dung> hoac openai <noi_dung>: Goi OpenAI bang OPENAI_API_KEY/API key sk-proj.\n"
            "- /alo | /git | /github | /meta <noi_dung>: Goi provider tuong ung; chua co token se bao chua cai.\n"
            "- /openaiinall on|off: Bat/tat AI inall cho group hien tai. Khi bat, thanh vien nhap '.noi_dung' moi goi AI.\n"
            "- /aiallow <platform> <uid> [quota]: Cap quyen AI.\n"
            "- /aideny <platform> <uid>: Thu hoi quyen AI.\n"
            "- /aiquota: Xem quota da cap.\n"
            "- /setaitoken <token>: Dat token AI mac dinh.\n"
            "- /clearaitoken: Xoa token AI mac dinh.\n"
            "- /checkaitoken: Kiem tra token dang dung.\n"
            "- /capnhat <noi_dung>: AI phan tich yeu cau cap nhat.\n"
            "\n"
            "🧾 Token cai dat + audit:\n"
            "- /newtoken [label] [hours]: Tao token cai dat.\n"
            "- /tokens: Xem danh sach token.\n"
            "- /audit: Xem log audit gan nhat.\n"
            "- /checklenh <ten_lenh>: Giai thich nhanh chuc nang 1 lenh."
        ),
    ]
    for part in help_parts:
        await update.message.reply_text(part)


async def cmd_check(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not context.args:
        await update.message.reply_text("Cú pháp: /check UID")
        return
    uid_check = context.args[0].strip()
    with db_conn() as conn:
        rows = conn.execute(
            "SELECT platform, uid, name, username, phone FROM customer_contacts WHERE uid=? LIMIT 5",
            (uid_check,),
        ).fetchall()
    if not rows:
        await update.message.reply_text("Không tìm thấy UID trong DB.")
        return
    lines = [f"- {r['platform']} | uid={r['uid']} | ten={r['name']} | user={r['username']} | sdt={r['phone']}" for r in rows]
    await update.message.reply_text("KET QUA CHECK UID:\n" + "\n".join(lines))


async def check_one_userbot_api(acc: dict, ping_target, ping_text: str, timeout_sec: float = 20.0) -> dict:
    result = {
        "phone": str(acc.get("phone") or ""),
        "api_ok": False,
        "admin_ping_ok": False,
        "status": "pending",
        "me": "",
        "error": "",
    }
    if acc.get("status") == "json_error":
        result["status"] = "json_error"
        result["error"] = str(acc.get("error") or "")
        return result
    if not os.path.exists(str(acc.get("session") or "")):
        result["status"] = "missing_session"
        result["error"] = os.path.basename(str(acc.get("session") or ""))
        return result
    if not acc.get("app_id") or not acc.get("app_hash"):
        result["status"] = "missing_api"
        return result
    client = None
    try:
        from telethon import TelegramClient
        from telethon.errors import AuthKeyUnregisteredError, FloodWaitError, PeerFloodError, UserDeactivatedBanError

        client = TelegramClient(
            str(acc["session"]),
            int(acc["app_id"]),
            str(acc["app_hash"]),
            device_model=str(acc.get("device") or "VGAH510"),
            system_version=str(acc.get("sdk") or "Windows 10"),
            app_version=str(acc.get("app_version") or "6.6.2 x64"),
        )
        await asyncio.wait_for(client.connect(), timeout=timeout_sec)
        authorized = await asyncio.wait_for(client.is_user_authorized(), timeout=timeout_sec)
        if not authorized:
            result["status"] = "unauthorized_session"
            return result
        me = await asyncio.wait_for(client.get_me(), timeout=timeout_sec)
        username = str(getattr(me, "username", "") or "").strip()
        first = str(getattr(me, "first_name", "") or "").strip()
        last = str(getattr(me, "last_name", "") or "").strip()
        name = " ".join(x for x in (first, last) if x).strip()
        uid = str(getattr(me, "id", "") or "")
        result["api_ok"] = True
        result["me"] = f"{uid} @{username}" if username else f"{uid} {name}".strip()
        try:
            await asyncio.wait_for(client.send_message(ping_target, ping_text), timeout=timeout_sec)
            result["admin_ping_ok"] = True
            result["status"] = "ok"
        except PeerFloodError:
            result["status"] = "api_ok_ping_peer_flood"
            result["error"] = "too_many_requests"
        except FloodWaitError as exc:
            result["status"] = "api_ok_ping_flood_wait"
            result["error"] = f"wait_{int(getattr(exc, 'seconds', 0) or 0)}s"
        except (UserDeactivatedBanError, AuthKeyUnregisteredError) as exc:
            result["status"] = "account_error"
            result["error"] = type(exc).__name__
        except Exception as exc:
            result["status"] = "api_ok_ping_fail"
            result["error"] = f"{type(exc).__name__}:{str(exc)[:120]}"
        return result
    except Exception as exc:
        result["status"] = "connect_error"
        result["error"] = f"{type(exc).__name__}:{str(exc)[:120]}"
        return result
    finally:
        if client is not None:
            try:
                await client.disconnect()
            except Exception:
                pass


async def cmd_check_bang(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        return
    if not is_primary_bot_instance():
        return
    tail = get_command_tail(update, "check!").strip().lower()
    profiles = list_userbot_profiles(limit=100)
    accounts = load_userbot_api_accounts(limit=100)
    if not accounts:
        await update.message.reply_text(f"Không thấy folder userbot trong {USERBOT_ACCOUNTS_ROOT}")
        return
    ping_target = "me" if tail in {"me", "saved", "save"} else (f"@{ADMIN_USERNAME.strip().lstrip('@')}" if ADMIN_USERNAME else int(ADMIN_ID))
    target_label = "Saved Messages từng userbot" if ping_target == "me" else str(ping_target)
    await update.message.reply_text(
        f"Đã nhận /check! Đang kiểm tra {len(accounts)} userbot API.\n"
        f"Ping target: {target_label}\n"
        "Chờ vài giây, em sẽ gửi báo cáo lại."
    )
    now_label = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    ping_text = f"[CHECK!] ping userbot -> admin {ADMIN_ID} | {now_label}"
    sem = asyncio.Semaphore(4)

    async def run_one(acc: dict) -> dict:
        async with sem:
            return await check_one_userbot_api(acc, ping_target, ping_text)

    results = await asyncio.gather(*(run_one(acc) for acc in accounts))
    api_ok = sum(1 for r in results if r.get("api_ok"))
    ping_ok = sum(1 for r in results if r.get("admin_ping_ok"))
    peer_flood = sum(1 for r in results if "peer_flood" in str(r.get("status") or ""))
    missing_session = sum(1 for r in results if r.get("status") == "missing_session")
    lines = [
        "[CHECK! USERBOT API]",
        f"Profiles: {len(profiles)} | Accounts checked: {len(accounts)}",
        f"API_OK: {api_ok}/{len(accounts)} | PING_ADMIN_OK: {ping_ok}/{len(accounts)}",
        f"PeerFlood: {peer_flood} | Missing session: {missing_session}",
        f"Ping target: {target_label}",
        "",
    ]
    for r in sorted(results, key=lambda x: str(x.get("phone") or "")):
        mark = "OK" if r.get("admin_ping_ok") else ("API" if r.get("api_ok") else "FAIL")
        me = f" | me={r.get('me')}" if r.get("me") else ""
        err = f" | {r.get('error')}" if r.get("error") else ""
        lines.append(f"- {mark} {r.get('phone')}: {r.get('status')}{me}{err}")
    lines.append("")
    lines.append("Ghi chu: API_OK nghia la session con dang nhap. PING_ADMIN_OK moi la gui duoc tin ve admin.")
    await send_telegram_long(context.bot, update.effective_chat.id, "\n".join(lines))


async def cmd_checklenh(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not context.args:
        await update.message.reply_text("Cú pháp: /checklenh /_lenh hoặc z_lenh")
        return
    cmd = context.args[0].strip().lower().lstrip("/")
    await update.message.reply_text(f"📘 {cmd}: {CMD_EXPLAIN.get(cmd, 'Chưa có mô tả cho lệnh này.')}")


async def cmd_openaiinall(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global OPENAI_INALL_TELE, OPENAI_INALL_ZALO, OPENAI_INALL_TELE_GROUPS
    if not is_admin(update.effective_user.id):
        return
    if not context.args or context.args[0].lower() not in ("on", "off"):
        await update.message.reply_text("Cú pháp: /openaiinall on|off")
        return
    chat_type = str(getattr(update.effective_chat, "type", "") or "").lower()
    if chat_type not in ("group", "supergroup"):
        await update.message.reply_text("⚠️ /openaiinall chỉ dùng trong group hiện tại.")
        return
    flag = context.args[0].lower() == "on"
    gid = str(update.effective_chat.id)
    if flag:
        OPENAI_INALL_TELE_GROUPS.add(gid)
    else:
        OPENAI_INALL_TELE_GROUPS.discard(gid)
    OPENAI_INALL_TELE = len(OPENAI_INALL_TELE_GROUPS) > 0
    persist_runtime_state()
    if flag:
        await update.message.reply_text("✅ Đã bật OpenAI inall cho group hiện tại. Thành viên phải nhắn bắt đầu bằng dấu chấm, ví dụ: `.xin chao`.")
    else:
        await update.message.reply_text("🛑 Đã tắt OpenAI inall cho group hiện tại.")


async def cmd_adminchat(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global ADMIN_CHAT_NOTIFY
    if not is_admin(update.effective_user.id):
        return
    if not context.args or context.args[0].lower() not in ("on", "off"):
        await update.message.reply_text("Cú pháp: /adminchat on|off")
        return
    ADMIN_CHAT_NOTIFY = context.args[0].lower() == "on"
    persist_runtime_state()
    await update.message.reply_text(f"✅ adminchat = {ADMIN_CHAT_NOTIFY}")


async def job_scan_report_telegram(context: ContextTypes.DEFAULT_TYPE):
    try:
        maybe_send_phone_event_summary(force=False)
        set_last_scan_report_ts("telegram", time.time())
    except Exception:
        logging.exception("Telegram scan report job failed")


async def job_session_backup_notify(context: ContextTypes.DEFAULT_TYPE):
    try:
        persist_runtime_state()
        backup_db("auto_session")
        out_dir = backup_session_snapshot("auto_session")
        export_dir = export_contacts_to_my_documents("auto_session")
        if out_dir:
            await context.bot.send_message(
                chat_id=ADMIN_ID,
                text=(
                    "💾 Đã backup tự động 60 phút/lần.\n"
                    f"DB/session: {out_dir}\n"
                    f"File khách ngày: {CONTACT_BACKUP_DAILY_CSV}\n"
                    f"File khách tổng: {CONTACT_BACKUP_TOTAL_CSV}"
                ),
            )
        else:
            await context.bot.send_message(
                chat_id=ADMIN_ID,
                text="⚠️ Không thể tạo snapshot phiên tự động. Vui lòng kiểm tra quyền ghi file.",
            )
    except Exception:
        logging.exception("Session backup notify job failed")


async def job_export_contacts_documents(context: ContextTypes.DEFAULT_TYPE):
    try:
        out_dir = export_contacts_to_my_documents("interval")
        if out_dir:
            logging.info("[EXPORT] Contacts exported to: %s", out_dir)
    except Exception:
        logging.exception("Export contacts job failed")


async def cmd_capnhat(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    payload = get_full_command_payload(update, "capnhat")
    if not payload:
        await update.message.reply_text("Cú pháp: /capnhat NOI_DUNG")
        return
    prompt = (
        "Phan tich noi dung cap nhat bot sau, neu ro y nghia, tra lai noi dung goc va de xuat cach lam theo buoc:\n"
        + payload
    )
    result = call_ai_text(prompt)
    await send_telegram_long(context.bot, update.effective_chat.id, "🧠 Phân tích AI:\n" + result)
    set_config("help_custom_capnhat", payload)

async def cmd_newtoken(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    label = context.args[0].strip() if context.args else "default"
    hours = int(context.args[1]) if len(context.args) >= 2 and context.args[1].isdigit() else 24
    raw = secrets.token_urlsafe(18)
    token_hash = hashlib.sha256(raw.encode("utf-8")).hexdigest()
    with db_conn() as conn:
        conn.execute(
            "INSERT INTO install_tokens(token_hash, raw_token, label, created_at, expires_at, redeemed) VALUES(?,?,?,?,?,0)",
            (token_hash, raw, label, time.time(), time.time() + hours * 3600),
        )
        conn.commit()
    audit_log("telegram", str(update.effective_user.id), "newtoken", f"{label}|{hours}")
    await update.message.reply_text(f"?? Token m?i: `{raw}`\nLabel: {label}\nH?t h?n sau: {hours} gi?", parse_mode="Markdown")


async def cmd_tokens(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    with db_conn() as conn:
        rows = conn.execute(
            "SELECT raw_token, label, expires_at, redeemed FROM install_tokens ORDER BY created_at DESC LIMIT 20"
        ).fetchall()
    if not rows:
        await update.message.reply_text("?? Ch?a c? token n?o.")
        return
    lines = []
    now = time.time()
    for r in rows:
        status = "? used" if r[3] else ("? expired" if now > float(r[2]) else "?? active")
        lines.append(f"- {r[0]} | {r[1]} | {status}")
    await update.message.reply_text("?? Danh s?ch token:\n" + "\n".join(lines))


async def cmd_audit(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    with db_conn() as conn:
        rows = conn.execute(
            "SELECT actor_platform, actor_uid, action, detail, created_at FROM audit ORDER BY id DESC LIMIT 20"
        ).fetchall()
    if not rows:
        await update.message.reply_text("?? Ch?a c? log audit.")
        return
    lines = [f"- {r[0]}:{r[1]} | {r[2]} | {str(r[3])[:60]}" for r in rows]
    await update.message.reply_text("?? Audit g?n nh?t:\n" + "\n".join(lines))




async def cmd_aiallow(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    if len(context.args) < 2:
        await update.message.reply_text("C? ph?p: /aiallow <platform> <uid> [quota]")
        return
    platform = context.args[0].strip().lower()
    uid = context.args[1].strip()
    quota = int(context.args[2]) if len(context.args) >= 3 and context.args[2].isdigit() else 30
    today = time.strftime("%Y-%m-%d", time.localtime())
    with db_conn() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO ai_permissions(platform, uid, enabled, daily_quota, used_today, day_key, created_at) VALUES(?,?,?,?,?,?,?)",
            (platform, uid, 1, quota, 0, today, time.time()),
        )
        conn.commit()
    audit_log("telegram", str(update.effective_user.id), "aiallow", f"{platform}:{uid}:{quota}")
    await update.message.reply_text(f"? ?? c?p quy?n AI cho {platform}:{uid} (quota {quota}/ng?y)")


async def cmd_aideny(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    if len(context.args) < 2:
        await update.message.reply_text("C? ph?p: /aideny <platform> <uid>")
        return
    platform = context.args[0].strip().lower()
    uid = context.args[1].strip()
    with db_conn() as conn:
        conn.execute("DELETE FROM ai_permissions WHERE platform=? AND uid=?", (platform, uid))
        conn.commit()
    audit_log("telegram", str(update.effective_user.id), "aideny", f"{platform}:{uid}")
    await update.message.reply_text(f"? ?? thu h?i quy?n AI c?a {platform}:{uid}")


async def cmd_aiquota(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    with db_conn() as conn:
        rows = conn.execute(
            "SELECT platform, uid, enabled, daily_quota, used_today, day_key FROM ai_permissions ORDER BY created_at DESC LIMIT 50"
        ).fetchall()
    if not rows:
        await update.message.reply_text("?? Ch?a c? c?u h?nh quota AI.")
        return
    lines = []
    for r in rows:
        lines.append(f"- {r[0]}:{r[1]} | enable={r[2]} | used={r[4]}/{r[3]} | day={r[5]}")
    await update.message.reply_text("📊 Danh sách quota AI:\n" + "\n".join(lines))


async def run_ai_provider_command(update: Update, context: ContextTypes.DEFAULT_TYPE, command_name: str, provider: str):
    if not is_admin(update.effective_user.id):
        return
    question = get_full_command_payload(update, command_name)
    if not question:
        await update.message.reply_text(f"⚠️ Cú pháp: /{command_name} <câu_hỏi>")
        return
    await update.message.reply_text(f"⏳ Đang gọi AI provider `{provider}`, vui lòng chờ...")
    result = call_ai_text(question, env_name=provider)
    await send_telegram_long(context.bot, update.effective_chat.id, result)


async def cmd_ai(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await run_ai_provider_command(update, context, "ai", "github")


async def cmd_openai(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await run_ai_provider_command(update, context, "openai", "openai")


async def cmd_alo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await run_ai_provider_command(update, context, "alo", "alo")


async def cmd_git(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await run_ai_provider_command(update, context, "git", "github")


async def cmd_github(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await run_ai_provider_command(update, context, "github", "github")


async def cmd_meta(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await run_ai_provider_command(update, context, "meta", "meta")


async def cmd_diemdanh(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    ok, detail = await send_botfather_startup_attendance(context.application)
    if ok:
        await update.message.reply_text(f"✅ Đã gửi điểm danh botfather về {BOT_ATTENDANCE_CHAT}.")
    else:
        await update.message.reply_text(f"⚠️ Chưa gửi được điểm danh: {detail}")


async def cmd_checkngay(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    chat = update.effective_chat
    chat_type = str(getattr(chat, "type", "") or "").lower()
    if chat_type not in ("group", "supergroup"):
        await update.message.reply_text("⚠️ /checkngay dùng trong nhóm Telegram cần kiểm tra UID.")
        return
    cache_telegram_group_member(update)
    chat_id = str(chat.id)
    rows = list_cached_group_members(chat_id)
    if not rows:
        await update.message.reply_text("Chưa có cache thành viên cho nhóm này. Hãy để userbot quét hoặc có thành viên nhắn trong nhóm trước.")
        return
    lines = [format_cached_member_line(r) for r in rows[:120]]
    more = "" if len(rows) <= 120 else f"\n... còn {len(rows) - 120} dòng trong file CSV."
    summary_text = f"UID + tên cache trong nhóm: {len(rows)} thành viên\n" + "\n".join(lines) + more
    await send_telegram_long(
        context.bot,
        chat.id,
        summary_text,
    )
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["chat_id", "user_id", "name", "username", "last_seen", "hit_count"])
    for r in rows:
        writer.writerow([r["chat_id"], r["user_id"], r["name"], r["username"], int(float(r["last_seen"] or 0)), r["hit_count"]])
    csv_bytes = buf.getvalue().encode("utf-8-sig")
    bio = io.BytesIO(csv_bytes)
    bio.name = f"checkngay_{chat_id}.csv"
    await context.bot.send_document(chat_id=chat.id, document=bio, filename=bio.name, caption="File UID + tên nhóm hiện tại.")


async def cmd_all(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    chat = update.effective_chat
    chat_type = str(getattr(chat, "type", "") or "").lower()
    if chat_type not in ("group", "supergroup"):
        await update.message.reply_text("⚠️ /ALL chỉ dùng trong nhóm Telegram.")
        return
    cache_telegram_group_member(update)
    rows = [r for r in list_cached_group_members(str(chat.id)) if str(r["user_id"]) != str(update.effective_user.id)]
    if not rows:
        await update.message.reply_text("Chưa có cache thành viên để tag. Chạy /phat500k dry-run rồi cho userbot quét trước.")
        return
    batch_size = 35
    total = len(rows)
    await update.message.reply_text(f"Đang tag {total} thành viên từ cache, chia {batch_size} người/tin để tránh quá dài.")
    for idx in range(0, total, batch_size):
        mentions = [mention_cached_member(r) for r in rows[idx:idx + batch_size]]
        text = " ".join(m for m in mentions if m)
        if text:
            await context.bot.send_message(chat_id=chat.id, text=text, parse_mode="HTML", disable_web_page_preview=True)
            await asyncio.sleep(1.2)


async def cmd_phat500k(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_supreme_admin(update.effective_user.id):
        return
    if not await primary_operation_command_is_for_this_bot(update, context, "phat500k"):
        return
    command_tail = get_command_tail(update, "phat500k")
    payload = command_tail.lower()
    run_real = any(x in payload.split() for x in ("run", "chay", "chạy", "that", "thật", "now"))
    dry_run = not run_real
    contact_phones, contact_notes = await extract_phat500k_contact_phones(update, context, command_tail)
    if contact_phones:
        await queue_contact_import_uid_user_job(
            update,
            context,
            contact_phones,
            contact_notes,
            dry_run=dry_run,
            command_label="/phat500k",
            source_label="phat500k",
        )
        return
    cache_telegram_group_member(update)
    root_stats, total_profiles = summarize_userbot_roots()
    profiles = [profile for stat in root_stats for profile in stat.get("profiles", [])]
    groups = list_cached_telegram_groups(limit=500)
    chat = update.effective_chat
    chat_type = str(getattr(chat, "type", "") or "").lower()
    if not groups and chat_type in ("group", "supergroup"):
        groups = [{"chat_id": str(chat.id), "chat_title": str(getattr(chat, "title", "") or ""), "updated_at": time.time()}]
    shard_lines = []
    for stat in root_stats:
        preview = ", ".join(stat.get("profiles", [])[:12])
        if len(stat.get("profiles", [])) > 12:
            preview += f", ... +{len(stat.get('profiles', [])) - 12}"
        shard_lines.append(
            f"- {stat.get('name')}: {stat.get('profile_count')} acc | {stat.get('root')}"
            + (f"\n  {preview}" if preview else "")
        )
    detail = {
        "profiles_root": root_stats[0]["root"] if root_stats else USERBOT_ACCOUNTS_ROOT,
        "profile_count": total_profiles,
        "profiles": profiles,
        "shard_count": len(root_stats),
        "shards": root_stats,
        "group_count": len(groups),
        "delay_account_sec": 5,
        "delay_group_sec": 2,
        "mode": "dry-run" if dry_run else "run",
        "bot_slot": current_bot_slot(),
    }
    detail = attach_userbot_report_origin(detail, update)
    if dry_run:
        sample_profiles = "\n".join(shard_lines) or "- chưa thấy profile"
        sample_groups = "\n".join(f"- {g['chat_id']} | {g['chat_title']}" for g in groups[:30]) or "- chưa có group cache"
        job_id = queue_userbot_scan_job("phat500k", "", "ALL_CACHED_GROUPS", str(update.effective_user.id), True, detail)
        report_text = (
            "DRY-RUN /phat500k\n"
            f"Job #{job_id}\n"
            f"Bot nhận lệnh: {current_bot_slot()} (chỉ bot chính xử lý lệnh này)\n"
            f"Botuser shards: {len(root_stats)} | Userbot profiles: {total_profiles}\n{sample_profiles}\n\n"
            f"Nhóm sẽ quét: {len(groups)}\n{sample_groups}\n\n"
            "Muốn chạy thật: /phat500k run"
        )
        await send_telegram_long(
            context.bot,
            update.effective_chat.id,
            report_text,
        )
        return
    job_ids = []
    for g in groups:
        job_ids.append(
            queue_userbot_scan_job(
                "phat500k",
                str(g["chat_id"]),
                str(g["chat_title"] or ""),
                str(update.effective_user.id),
                False,
                detail,
            )
        )
    report_text = (
        f"✅ Đã phát {len(job_ids)} job quét cho supervisor userbot.\n"
        f"Bot nhận lệnh: {current_bot_slot()} (chỉ bot chính xử lý /phat500k).\n"
        f"Shards={len(root_stats)} | Profiles={total_profiles}.\n"
        f"Queue: {USERBOT_SCAN_JOBS_DIR}"
    )
    await update.message.reply_text(report_text)


async def cmd_setaitoken(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    if not context.args:
        await update.message.reply_text("⚠️ Cú pháp: /setaitoken <token>")
        return
    token = " ".join(context.args).strip()
    set_config("ai_token", token)
    audit_log("telegram", str(update.effective_user.id), "set_ai_token", f"len={len(token)}")
    await update.message.reply_text(f"✅ Đã lưu AI token: {mask_token(token)}")


async def cmd_clearaitoken(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    set_config("ai_token", "")
    audit_log("telegram", str(update.effective_user.id), "clear_ai_token", "")
    await update.message.reply_text("🧹 Đã xóa AI token lưu trong bot.")


async def cmd_checkaitoken(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    token_db = get_config("ai_token", "").strip()
    if token_db:
        await update.message.reply_text(f"🔐 AI token (DB): {mask_token(token_db)}")
        return
    token_env = os.getenv(AI_API_TOKEN_ENV, "").strip()
    if token_env:
        await update.message.reply_text(f"🔐 AI token (ENV): {mask_token(token_env)}")
        return
    await update.message.reply_text("❌ Chưa có AI token (DB/ENV).")


async def cmd_zreply(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update.effective_user.id):
        return
    leading, msg = split_command_leading_and_rest(update, "zreply", 1)
    if len(leading) < 1:
        await update.message.reply_text("Cú pháp: /zreply <uid_zalo> <noi_dung>")
        return
    uid_zalo = leading[0].strip()
    if not msg:
        await update.message.reply_text("Nội dung trống.")
        return
    try:
        sent = False
        if LAST_ZALO_BOT is not None:
            from zlapi.models import Message, ThreadType
            try:
                if hasattr(LAST_ZALO_BOT, "send_long_message"):
                    LAST_ZALO_BOT.send_long_message(msg, int(uid_zalo), ThreadType.USER)
                else:
                    LAST_ZALO_BOT.sendMessage(Message(text=msg), int(uid_zalo), ThreadType.USER)
                sent = True
            except Exception:
                try:
                    if hasattr(LAST_ZALO_BOT, "send_long_message"):
                        LAST_ZALO_BOT.send_long_message(msg, uid_zalo, ThreadType.USER)
                    else:
                        LAST_ZALO_BOT.sendMessage(Message(text=msg), uid_zalo, ThreadType.USER)
                    sent = True
                except Exception:
                    sent = False
        notify_admin_sync(f"[TELE->ZALO] /zreply {uid_zalo}")
        await update.message.reply_text(f"{'Đã gửi' if sent else 'Đã nhận lệnh'} /zreply cho UID {uid_zalo}.")
    except Exception as exc:
        await update.message.reply_text(f"Lỗi /zreply: {exc}")


async def cmd_loaloaall(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await require_admin(update): return
    tail = get_command_tail(update, "loaloaall")
    if not tail or len(tail.rsplit(None, 1)) < 2:
        await update.message.reply_text("Cú pháp: /loaloaall <noi_dung> <thoi_gian_giay>")
        return
    try:
        msg, delay_raw = tail.rsplit(None, 1)
        delay = float(delay_raw)
        msg = msg.strip()
    except Exception:
        await update.message.reply_text("Thời gian không hợp lệ.")
        return
    ok = 0
    fail = 0
    for uid in list(USER_DATA.keys()):
        try:
            await send_telegram_long(context.bot, int(uid), msg)
            ok += 1
        except Exception:
            fail += 1
        await asyncio.sleep(max(0.1, delay))
    await update.message.reply_text(f"Đã gửi /loaloaall: OK={ok} FAIL={fail}")

async def cmd_hotro(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    if uid in BANNED_USERS:
        return
    await update.message.reply_text(f"👨‍💼 Ho tro nhanh: https://t.me/{ADMIN_USERNAME.replace('@', '')}")


BOT_COMMAND_DEFS = [
    ("start", "Start / get code"),
    ("nhancode88k", "Get promo code"),
    ("hotro", "Contact support"),
    ("help", "Quick help"),
    ("trogiup", "Full command help"),
    ("myid", "Show your Telegram ID"),
    ("ping", "Check bot status"),
    ("kiemtra", "System status"),
    ("checklenh", "Command status"),
    ("checkngay", "Daily check"),
    ("baotri", "Maintenance mode"),
    ("tatmoauto", "Toggle auto mode"),
    ("tatmochuong", "Toggle alerts"),
    ("tatmoanh", "Toggle media"),
    ("stopall", "Stop all userbot jobs"),
    ("themadmin", "Add sub admin"),
    ("huyadmin", "Remove sub admin"),
    ("tenadmin", "Set admin username"),
    ("adminquanly", "Set management group"),
    ("setsuppergroup", "Set supper group"),
    ("diemdanh", "Send attendance ping"),
    ("xemkhach", "View customer"),
    ("check", "Check UID data"),
    ("lamlai", "Reset customer"),
    ("chan", "Block customer"),
    ("mochan", "Unblock customer"),
    ("mochanhet", "Unblock all"),
    ("vip", "Set VIP"),
    ("huyvip", "Remove VIP"),
    ("code", "Set promo code"),
    ("guikhach", "Send one customer"),
    ("guitatca", "Broadcast customers"),
    ("loaloa", "Send announcement"),
    ("loaloaall", "Timed broadcast"),
    ("chuyentrong", "Send private link"),
    ("link", "Set web/code link"),
    ("huong", "Set keyword route"),
    ("zreply", "Reply to Zalo"),
    ("suachao", "Edit welcome text"),
    ("suaquangcao", "Edit ad text"),
    ("suarepsdt", "Edit phone reply"),
    ("suarep", "Edit reply template"),
    ("sua", "Continue edit wizard"),
    ("yes", "Confirm yes"),
    ("no", "Confirm no"),
    ("xemrep", "View reply templates"),
    ("anhchao", "Set welcome media"),
    ("anhdangky", "Set signup media"),
    ("anhepjoin", "Set join media"),
    ("anhmocnap", "Set milestone media"),
    ("anhchoyet", "Set pending media"),
    ("anhqc", "Set ad media"),
    ("setgifall", "Set all media"),
    ("setguitngif", "Set send media"),
    ("showgifs", "Show media config"),
    ("quetnhanh", "Scan text phones"),
    ("quetfile", "Scan replied file"),
    ("quetfileepath", "Scan file path"),
    ("quetfilepath", "Scan file path"),
    ("quetnow", "Scan phone list"),
    ("quetuser", "Scan phone to user"),
    ("quetallnow", "Run full scan now"),
    ("quetstatus", "Scan job status"),
    ("usernow", "Latest user file"),
    ("uidnow", "Latest UID file"),
    ("ttnow", "Latest full file"),
    ("sdtnow", "Latest phone file"),
    ("phat500k", "Run import job"),
    ("phat500kstatus", "Import status"),
    ("importstatus", "Import status"),
    ("guitn", "Send by userbot"),
    ("guitnlai", "Retry failed send"),
    ("guitnstatus", "Send job status"),
    ("keoall", "Invite to group"),
    ("keoallstatus", "Invite status"),
    ("run12", "Run 12 profiles"),
    ("run20", "Run 20 profiles"),
    ("run20new", "Run 20 new profiles"),
    ("run32", "Run 32 profiles"),
    ("run40", "Run 40 profiles"),
    ("run52", "Run 52 profiles"),
    ("runall", "Run all profiles"),
    ("stop", "Stop tagged bot/job"),
    ("ai", "AI chat"),
    ("openai", "OpenAI chat"),
    ("github", "GitHub AI chat"),
    ("git", "Git AI alias"),
    ("alo", "ALO AI chat"),
    ("meta", "Meta AI chat"),
    ("setaitoken", "Save AI token"),
    ("checkaitoken", "Check AI token"),
    ("clearaitoken", "Clear AI token"),
    ("aiquota", "AI quota"),
    ("aiallow", "Allow AI user"),
    ("aideny", "Deny AI user"),
    ("openaiinall", "Toggle AI in groups"),
    ("newtoken", "Create token"),
    ("tokens", "View tokens"),
    ("audit", "Audit log"),
    ("capnhat", "Update data"),
]
BOT_COMMANDS = [BotCommand(command, description) for command, description in BOT_COMMAND_DEFS[:100]]


async def post_init_setup(application) -> None:
    await post_init_startup_attendance(application)
    try:
        await application.bot.set_my_commands(BOT_COMMANDS)
        logging.info("Registered %s Telegram bot commands", len(BOT_COMMANDS))
        try:
            set_config("telegram_commands_last", f"{int(time.time())}|ok|{len(BOT_COMMANDS)}")
        except Exception:
            pass
    except Exception as exc:
        logging.exception("Failed to register Telegram bot commands: %s", exc)
        try:
            set_config("telegram_commands_last", f"{int(time.time())}|failed|{type(exc).__name__}:{str(exc)[:160]}")
        except Exception:
            pass


async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    try:
        logging.exception("Telegram error update=%r", update, exc_info=context.error)
        set_config("telegram_error_last", f"{int(time.time())}|{type(context.error).__name__}:{str(context.error)[:220]}")
    except Exception:
        logging.exception("Telegram error handler failed")


def register_telegram_handlers(app) -> None:
    app.add_error_handler(error_handler)

    app.add_handler(CommandHandler("nhancode88k", cmd_nhancode88k))
    app.add_handler(CommandHandler("start", cmd_nhancode88k))
    app.add_handler(CommandHandler("hotro", cmd_hotro))
    app.add_handler(CommandHandler("help", cmd_help))

    # Đăng ký các bộ lệnh tiếp thị mới bổ sung
    app.add_handler(CommandHandler("guitatca", cmd_guitatca))
    app.add_handler(CommandHandler("guikhach", cmd_guikhach))

    app.add_handler(CommandHandler("loaloa", cmd_loaloa))
    app.add_handler(CommandHandler("baotri", cmd_baotri))
    app.add_handler(CommandHandler("link", cmd_link))
    app.add_handler(CommandHandler("huong", cmd_huong))
    app.add_handler(CommandHandler("chuyentrong", cmd_chuyentrong))
    app.add_handler(CommandHandler("kiemtra", cmd_kiemtra))
    app.add_handler(CommandHandler("xemkhach", cmd_xemkhach))
    app.add_handler(CommandHandler("xoarac", cmd_xoarac))
    app.add_handler(CommandHandler("themadmin", cmd_themadmin))
    app.add_handler(CommandHandler("huyadmin", cmd_huyadmin))
    app.add_handler(CommandHandler("kenh", cmd_kenh))
    app.add_handler(CommandHandler("chan", cmd_chan))
    app.add_handler(CommandHandler("mochan", cmd_mochan))
    app.add_handler(CommandHandler("mochanhet", cmd_mochanhet))
    app.add_handler(CommandHandler("code", cmd_code))
    app.add_handler(CommandHandler("tatmoauto", cmd_tatmoauto))
    app.add_handler(CommandHandler("tatmochuong", cmd_tatmochuong))
    app.add_handler(CommandHandler("tatmoanh", cmd_tatmoanh))
    app.add_handler(CommandHandler("suachao", cmd_suachao))
    app.add_handler(CommandHandler("suaquangcao", cmd_suaquangcao))
    app.add_handler(CommandHandler("suarepsdt", cmd_suarepsdt))
    app.add_handler(CommandHandler("suarep", cmd_suarep))
    app.add_handler(CommandHandler("sua", cmd_sua))
    app.add_handler(CommandHandler("yes", cmd_yes))
    app.add_handler(CommandHandler("no", cmd_no))
    app.add_handler(CommandHandler("xemrep", cmd_xemrep))
    app.add_handler(CommandHandler("vip", cmd_vip))
    app.add_handler(CommandHandler("huyvip", cmd_huyvip))
    app.add_handler(CommandHandler("anhchao", cmd_anhchao))
    app.add_handler(CommandHandler("anhdangky", cmd_anhdangky))
    app.add_handler(CommandHandler("anhepjoin", cmd_anhepjoin))
    app.add_handler(CommandHandler("anhmocnap", cmd_anhmocnap))
    app.add_handler(CommandHandler("anhchoyet", cmd_anhchoyet))
    app.add_handler(CommandHandler("anhqc", cmd_anhqc))
    app.add_handler(CommandHandler("setgifall", cmd_setgifall))
    app.add_handler(CommandHandler("setguitngif", cmd_setguitngif))
    app.add_handler(CommandHandler("showgifs", cmd_showgifs))
    app.add_handler(CommandHandler("tenadmin", cmd_tenadmin))
    app.add_handler(CommandHandler("lamlai", cmd_lamlai))
    app.add_handler(CommandHandler("ping", cmd_ping))
    app.add_handler(CommandHandler("myid", cmd_myid))
    app.add_handler(CommandHandler("setsuppergroup", cmd_setsuppergroup))
    app.add_handler(CommandHandler("setsupergroup", cmd_setsuppergroup))
    app.add_handler(CommandHandler("adminquanly", cmd_adminquanly))
    app.add_handler(CommandHandler("setadminquanly", cmd_adminquanly))
    app.add_handler(CommandHandler("trogiup", cmd_trogiup))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex(r"(?i)^/check!(?:@\w+)?(?:\s|$)"), cmd_check_bang))
    app.add_handler(CommandHandler("check", cmd_check))
    app.add_handler(CommandHandler("checklenh", cmd_checklenh))
    app.add_handler(CommandHandler("capnhat", cmd_capnhat))
    app.add_handler(CommandHandler("openaiinall", cmd_openaiinall))
    app.add_handler(CommandHandler("adminchat", cmd_adminchat))
    app.add_handler(CommandHandler("newtoken", cmd_newtoken))
    app.add_handler(CommandHandler("tokens", cmd_tokens))
    app.add_handler(CommandHandler("audit", cmd_audit))
    app.add_handler(CommandHandler("ai", cmd_ai))
    app.add_handler(CommandHandler("openai", cmd_openai))
    app.add_handler(CommandHandler("alo", cmd_alo))
    app.add_handler(CommandHandler("git", cmd_git))
    app.add_handler(CommandHandler("github", cmd_github))
    app.add_handler(CommandHandler("meta", cmd_meta))
    app.add_handler(CommandHandler("diemdanh", cmd_diemdanh))
    app.add_handler(CommandHandler("checkngay", cmd_checkngay))
    app.add_handler(CommandHandler("ALL", cmd_all))
    app.add_handler(CommandHandler("all", cmd_all))
    app.add_handler(CommandHandler("phat500k", cmd_phat500k))
    app.add_handler(CommandHandler("usernow", cmd_usernow))
    app.add_handler(CommandHandler("uidnow", cmd_uidnow))
    app.add_handler(CommandHandler("ttnow", cmd_ttnow))
    app.add_handler(CommandHandler("sdtnow", cmd_sdtnow))
    app.add_handler(CommandHandler("quetallnow", cmd_quetallnow))
    app.add_handler(CommandHandler("uidstatus", cmd_uidstatus))
    app.add_handler(CommandHandler("phat500kstatus", cmd_uidstatus))
    app.add_handler(CommandHandler("importstatus", cmd_uidstatus))
    app.add_handler(CommandHandler("quetstatus", cmd_quetstatus))
    app.add_handler(CommandHandler("guitn", cmd_guitn))
    app.add_handler(CommandHandler("guitnlai", cmd_guitnlai))
    app.add_handler(CommandHandler("run40", cmd_run40))
    app.add_handler(CommandHandler("run52", cmd_run52))
    app.add_handler(CommandHandler("run12", cmd_run12))
    app.add_handler(CommandHandler("run20", cmd_run20))
    app.add_handler(CommandHandler("run32", cmd_run32))
    app.add_handler(CommandHandler("run20new", cmd_run20new))
    app.add_handler(CommandHandler("runall", cmd_runall))
    app.add_handler(CommandHandler("guitnstatus", cmd_guitnstatus))
    app.add_handler(CommandHandler("stopall", cmd_stopall))
    app.add_handler(CommandHandler("stop", cmd_stop_bot))
    app.add_handler(CommandHandler("keoall", cmd_keoall))
    app.add_handler(CommandHandler("keoallstatus", cmd_keoallstatus))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex(r"(?is)^\s*/guitn(?:@\w+)?(?:\s|$)"), cmd_guitn_text_router))
    app.add_handler(MessageHandler(filters.CaptionRegex(r"(?is)^\s*/guitn(?:@\w+)?(?:\s|$)"), cmd_guitn_text_router))
    app.add_handler(MessageHandler(filters.TEXT & filters.Regex(r"(?is)^\s*/guitnlai(?:@\w+)?(?:\s|$)"), cmd_guitnlai_text_router))
    app.add_handler(MessageHandler(filters.CaptionRegex(r"(?is)^\s*/guitnlai(?:@\w+)?(?:\s|$)"), cmd_guitnlai_text_router))
    app.add_handler(CommandHandler("setaitoken", cmd_setaitoken))
    app.add_handler(CommandHandler("clearaitoken", cmd_clearaitoken))
    app.add_handler(CommandHandler("checkaitoken", cmd_checkaitoken))
    app.add_handler(CommandHandler("aiallow", cmd_aiallow))
    app.add_handler(CommandHandler("aideny", cmd_aideny))
    app.add_handler(CommandHandler("aiquota", cmd_aiquota))
    app.add_handler(CommandHandler("zreply", cmd_zreply))
    app.add_handler(CommandHandler("loaloaall", cmd_loaloaall))
    app.add_handler(CommandHandler("quetnhanh", cmd_quetnhanh))
    app.add_handler(CommandHandler("quetfile", cmd_quetfile))
    app.add_handler(CommandHandler("quetfileepath", cmd_quetfileepath))
    app.add_handler(CommandHandler("quetfilepath", cmd_quetfilepath))
    app.add_handler(CommandHandler("quetnow", cmd_quetnow))
    app.add_handler(CommandHandler("quetuser", cmd_quetuser))

    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_scan_file_caption))
    app.add_handler(CallbackQueryHandler(handle_callback))


def main():
    if not one_file_bootstrap():
        return
    auto_restore_from_latest_backup()
    init_store()
    load_runtime_state()
    import_contacts_from_backup_dir()
    bootstrap_scan_from_cache()
    export_contacts_to_my_documents("startup")
    backup_db("startup")
    atexit.register(lambda: (persist_runtime_state(), backup_db("shutdown"), backup_session_snapshot("shutdown")))
    try:
        signal.signal(signal.SIGTERM, lambda *_: (persist_runtime_state(), backup_db("sigterm")))
        signal.signal(signal.SIGINT, lambda *_: (persist_runtime_state(), backup_db("sigint")))
    except Exception:
        pass
    slot_name = os.getenv("BOT_INSTANCE_SLOT", "").strip() or "default"
    zalo_worker_enabled = os.getenv("ZALO_WORKER_ENABLED", "1").strip().lower() not in {"0", "false", "no", "off"}
    if zalo_worker_enabled:
        set_config("zalo_worker_last", f"{int(time.time())}|started slot={slot_name}")
        threading.Thread(target=start_zalo_auto_worker, daemon=True).start()
    else:
        set_config("zalo_worker_last", f"{int(time.time())}|skipped slot={slot_name}")
    app = Application.builder().token(TOKEN_BOT).post_init(post_init_setup).build()
    if getattr(app, "job_queue", None) is not None:
        app.job_queue.run_repeating(job_scan_report_telegram, interval=SCAN_REPORT_INTERVAL_SEC, first=40)
        app.job_queue.run_repeating(job_session_backup_notify, interval=SESSION_BACKUP_INTERVAL_SEC, first=90)
        app.job_queue.run_repeating(job_process_broadcast_queue, interval=5, first=12)
    else:
        logging.warning("Khong co JobQueue (chua cai python-telegram-bot[job-queue]), bo qua scan report telegram dinh ky.")
    register_telegram_handlers(app)

    print("=========================================")
    print("BOT /nhancode88k DA KHOI CHAY")
    print(f"ADMIN_ID={ADMIN_ID}")
    print("=========================================")
    set_config("bot_polling_last", f"{int(time.time())}|start slot={slot_name}")
    try:
        try:
            asyncio.get_event_loop()
        except RuntimeError:
            asyncio.set_event_loop(asyncio.new_event_loop())
        app.run_polling()
        set_config("bot_polling_last", f"{int(time.time())}|stop slot={slot_name}")
    except Exception as exc:
        set_config(
            "bot_polling_last",
            f"{int(time.time())}|failed slot={slot_name} err={type(exc).__name__}:{str(exc)[:220]}",
        )
        raise

if __name__ == '__main__':
    main()


